from __future__ import annotations

import hashlib
import html
import csv
import ipaddress
import json
import math
import os
import re
import sqlite3
import statistics
import subprocess
import sys
import threading
import time
import uuid
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from contextlib import contextmanager
from datetime import datetime, timedelta, timezone
from functools import lru_cache
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, quote, unquote, urlparse
from urllib.request import Request, urlopen

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook, load_workbook
from openpyxl.comments import Comment
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.worksheet.table import Table, TableStyleInfo

import hierarchy_net_deposit
import rebate_churning
from ea_comment_group import (
    EaCommentGroupService,
    classify_ea_comment,
    ea_comment_identity,
    ea_match_evidence,
    ea_comment_parts,
    ea_comment_query_plan,
    ea_comment_totals,
    ea_dynamic_identity,
    ea_expert_sequence_match,
)
from signal_copy_group import SignalCopyGroupService, signal_group_totals, signal_in_identifier, signal_in_tag
from account_logs import query_account_logs


ROOT = Path(os.environ.get("K_DESK_ROOT", Path(__file__).resolve().parents[2]))
LEGACY_RISK_ROOT = Path(r"D:\risk")
OUT_DIR = Path(os.environ.get("ACCOUNT_REGISTRY_DATA_DIR", ROOT / "local_data" / "problem_account_registry"))
QUICK_ACTIONS_PATH = Path(os.environ.get("ACCOUNT_QUICK_ACTIONS_PATH", OUT_DIR / "quick_actions.json"))
DEFAULT_KLINE_OUT_DIR = LEGACY_RISK_ROOT / "output_data" if LEGACY_RISK_ROOT.exists() else ROOT / "outputs" / "kline"
KLINE_OUT_DIR = Path(os.environ.get("TRADE_KLINE_OUT_DIR", DEFAULT_KLINE_OUT_DIR))
KLINE_TIMELINE_CACHE_DIR = Path(
    os.environ.get("KDESK_KLINE_TIMELINE_CACHE_DIR", KLINE_OUT_DIR.parent / "cache" / "kline_timeline")
)
TRADE_KLINE_WEB_URL = os.environ.get("TRADE_KLINE_WEB_URL", "http://127.0.0.1:8766")
TRADE_DB_PATH = Path(os.environ.get("ACCOUNT_TRADE_DB_PATH", LEGACY_RISK_ROOT / "output_data" / "account_trade_lookup" / "trades.sqlite"))
TRADE_DB_SOURCE = os.environ.get("ACCOUNT_TRADE_DB_SOURCE", "auto").lower()
IP_HISTORY_DB_PATH = Path(os.environ.get("ACCOUNT_LOGIN_IP_DB_PATH", OUT_DIR / "account_login_ips.sqlite"))
EA_PATTERN_DB_PATH = Path(os.environ.get("ACCOUNT_EA_PATTERN_DB_PATH", OUT_DIR / "ea_comment_patterns.sqlite"))
IP_GEO_API_TEMPLATE = os.environ.get("ACCOUNT_IP_GEO_API", "https://ipwho.is/{ip}")
IP_GEO_TIMEOUT = float(os.environ.get("ACCOUNT_IP_GEO_TIMEOUT", "4"))
IP_GEO_CACHE_DAYS = int(os.environ.get("ACCOUNT_IP_GEO_CACHE_DAYS", "30"))
MYSQL_USER = os.environ.get("ACCOUNT_TRADE_MYSQL_USER", "intern")
MYSQL_PASSWORD = os.environ.get("ACCOUNT_TRADE_MYSQL_PASSWORD", "")
MYSQL_SOURCES = [
    {
        "name": "AC GB MT5",
        "host": os.environ.get("ACCOUNT_TRADE_MYSQL_AC_HOST", "rm-3nsv8k160ht47x44uio.mysql.rds.aliyuncs.com"),
        "schema": "int_sass_crm_ac_mt5_live_new",
        "table": "mt5_deals",
        "platform": "MT5",
        "server": "AC GB MT5",
        "kind": "mt5_deals",
        "default_currency": "USD",
        "crm_schema": "int_sass_crm_ac",
        "mt_server_code": "1",
        "account_route": {"schema": "int_sass_crm_ac", "mt_server_code": "1"},
        "crm_routes": [{"schema": "int_sass_crm_ac", "mt_server_code": "1"}],
    },
    {
        "name": "AC CN MT5",
        "host": os.environ.get("ACCOUNT_TRADE_MYSQL_AC_HOST", "rm-3nsv8k160ht47x44uio.mysql.rds.aliyuncs.com"),
        "schema": "sass_crm_ac_mt5_live",
        "table": "mt5_deals",
        "platform": "MT5",
        "server": "AC CN MT5",
        "kind": "mt5_deals",
        "default_currency": "USD",
        "crm_schema": "sass_crm_ac",
        "mt_server_code": "1",
        "account_route": {"schema": "sass_crm_ac", "mt_server_code": "1"},
        "crm_routes": [{"schema": "sass_crm_ac", "mt_server_code": "1"}],
    },
    {
        "name": "AC CN MT5 live3",
        "host": os.environ.get("ACCOUNT_TRADE_MYSQL_AC_HOST", "rm-3nsv8k160ht47x44uio.mysql.rds.aliyuncs.com"),
        "schema": "sass_crm_ac_mt5_live3",
        "table": "mt5_deals",
        "platform": "MT5",
        "server": "AC CN MT5 live3",
        "kind": "mt5_deals",
        "default_currency": "USD",
        "aliases": ["AC CN MT5 Live3"],
        "crm_schema": "sass_crm_ac",
        "mt_server_code": "3",
        "account_route": {"schema": "sass_crm_ac", "mt_server_code": "3"},
        "crm_routes": [{"schema": "sass_crm_ac", "mt_server_code": "3"}],
    },
    {
        "name": "AC MT4",
        "host": os.environ.get("ACCOUNT_TRADE_MYSQL_AC_HOST", "rm-3nsv8k160ht47x44uio.mysql.rds.aliyuncs.com"),
        "schema": "mt4_export_syc",
        "table": "mt4_trades",
        "platform": "MT4",
        "server": "AC CN MT4",
        "kind": "mt4_trades",
        "crm_schema": "sass_crm_ac",
        "mt_server_code": "2",
        "aliases": ["AC MT4"],
        "account_route": {"schema": "sass_crm_ac", "mt_server_code": "2"},
        "crm_routes": [{"schema": "sass_crm_ac", "mt_server_code": "2"}],
    },
    {
        "name": "AC GB MT4",
        "host": os.environ.get("ACCOUNT_TRADE_MYSQL_AC_HOST", "rm-3nsv8k160ht47x44uio.mysql.rds.aliyuncs.com"),
        "schema": "mt4_export_syc",
        "table": "mt4_trades",
        "platform": "MT4",
        "server": "AC GB MT4",
        "kind": "mt4_trades",
        "crm_schema": "int_sass_crm_ac",
        "mt_server_code": "2",
        "aliases": ["AC MT4"],
        "account_route": {"schema": "int_sass_crm_ac", "mt_server_code": "2"},
        "crm_routes": [{"schema": "int_sass_crm_ac", "mt_server_code": "2"}],
    },
    {
        "name": "DBG MT5",
        "host": os.environ.get("ACCOUNT_TRADE_MYSQL_DBG_HOST", "rm-3nspk4458ag106ugaxo.mysql.cnhk.rds.aliyuncs.com"),
        "schema": "mt5_export_new",
        "table": "mt5_deals",
        "platform": "MT5",
        "server": "DBG CN MT5",
        "kind": "mt5_deals",
        "default_currency": "USD",
        "aliases": ["DBG MT5"],
        "account_route": {"schema": "crm_cn", "mt_server_code": "4"},
        "crm_routes": [{"schema": "crm_cn", "mt_server_code": "4"}],
    },
    {
        "name": "DBG GB MT5",
        "host": os.environ.get("ACCOUNT_TRADE_MYSQL_DBG_HOST", "rm-3nspk4458ag106ugaxo.mysql.cnhk.rds.aliyuncs.com"),
        "schema": "mt5_export_new",
        "table": "mt5_deals",
        "platform": "MT5",
        "server": "DBG GB MT5",
        "kind": "mt5_deals",
        "default_currency": "USD",
        "aliases": ["DBG MT5"],
        "account_route": {"schema": "crm_vn", "mt_server_code": "2"},
        "crm_routes": [{"schema": "crm_vn", "mt_server_code": "2"}],
    },
    {
        "name": "DBG MT5 Live2",
        "host": os.environ.get("ACCOUNT_TRADE_MYSQL_DBG_HOST", "rm-3nspk4458ag106ugaxo.mysql.cnhk.rds.aliyuncs.com"),
        "schema": "crm_vn_mt5_live2",
        "table": "mt5_deals",
        "platform": "MT5",
        "server": "DBG MT5 Live2",
        "kind": "mt5_deals",
        "default_currency": "USD",
        "aliases": ["DBG MT5", "DBG GB MT5 Live2"],
        "account_route": {"schema": "crm_vn", "mt_server_code": "5"},
        "crm_routes": [{"schema": "crm_vn", "mt_server_code": "5"}],
    },
    {
        "name": "DBG MT4 CN1",
        "host": os.environ.get("ACCOUNT_TRADE_MYSQL_DBG_HOST", "rm-3nspk4458ag106ugaxo.mysql.cnhk.rds.aliyuncs.com"),
        "schema": "crm_cn_mt4_live1",
        "table": "mt4_trades",
        "platform": "MT4",
        "server": "DBG MT4 CN1",
        "kind": "mt4_trades",
        "aliases": ["DBG CN MT4 Live1"],
        "crm_schema": "crm_cn",
        "mt_server_code": "1",
        "account_route": {"schema": "crm_cn", "mt_server_code": "1"},
        "crm_routes": [{"schema": "crm_cn", "mt_server_code": "1"}],
    },
    {
        "name": "DBG MT4 CN2",
        "host": os.environ.get("ACCOUNT_TRADE_MYSQL_DBG_HOST", "rm-3nspk4458ag106ugaxo.mysql.cnhk.rds.aliyuncs.com"),
        "schema": "crm_cn_mt4_live2",
        "table": "mt4_trades",
        "platform": "MT4",
        "server": "DBG MT4 CN2",
        "kind": "mt4_trades",
        "aliases": ["DBG CN MT4 Live2"],
        "crm_schema": "crm_cn",
        "mt_server_code": "3",
        "account_route": {"schema": "crm_cn", "mt_server_code": "3"},
        "crm_routes": [{"schema": "crm_cn", "mt_server_code": "3"}],
    },
    {
        "name": "DBG MT4 VN3",
        "host": os.environ.get("ACCOUNT_TRADE_MYSQL_DBG_HOST", "rm-3nspk4458ag106ugaxo.mysql.cnhk.rds.aliyuncs.com"),
        "schema": "crm_vn_mt4_live3",
        "table": "mt4_trades",
        "platform": "MT4",
        "server": "DBG MT4 VN3",
        "kind": "mt4_trades",
        "aliases": ["DBG VN MT4 Live3"],
        "crm_schema": "crm_vn",
        "mt_server_code": "1",
        "account_route": {"schema": "crm_vn", "mt_server_code": "1"},
        "crm_routes": [{"schema": "crm_vn", "mt_server_code": "1"}],
    },
]
TRADE_KLINE_TOOL_DIR = Path(os.environ.get("TRADE_KLINE_TOOL_DIR", ROOT / "tools" / "trade_kline_tool"))
TRADE_KLINE_GENERATOR = TRADE_KLINE_TOOL_DIR / "generate_trade_kline_from_statement.py"
K_DESK_PYTHON = Path(os.environ.get("K_DESK_PYTHON", r"C:\Users\amber\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\python.exe"))
TRADE_KLINE_TERMINAL = Path(os.environ.get("TRADE_KLINE_TERMINAL", r"C:\Program Files\AC Capital Market MT5 Terminal\terminal64.exe"))
TOXIC_MT5_TERMINALS = {
    "AC GB MT5": str(TRADE_KLINE_TERMINAL),
    "AC CN MT5": str(TRADE_KLINE_TERMINAL),
    "AC MT4": str(TRADE_KLINE_TERMINAL),
    "AC CN MT4": str(TRADE_KLINE_TERMINAL),
    "AC GB MT4": str(TRADE_KLINE_TERMINAL),
}
TOXIC_MT5_QUOTE_ACCOUNTS = {
    "AC GB MT5": {"login": 11007, "server": "ACCMGlobal-Live"},
    "AC CN MT5": {"login": 11007, "server": "ACCMGlobal-Live"},
    "AC MT4": {"login": 11007, "server": "ACCMGlobal-Live"},
    "AC CN MT4": {"login": 11007, "server": "ACCMGlobal-Live"},
    "AC GB MT4": {"login": 11007, "server": "ACCMGlobal-Live"},
}
try:
    TOXIC_MT5_TERMINALS.update(json.loads(os.environ.get("TOXIC_MT5_TERMINALS_JSON", "{}")))
except (TypeError, ValueError, json.JSONDecodeError):
    pass
try:
    TOXIC_MT5_QUOTE_ACCOUNTS.update(json.loads(os.environ.get("TOXIC_MT5_QUOTE_ACCOUNTS_JSON", "{}")))
except (TypeError, ValueError, json.JSONDecodeError):
    pass
TRADE_KLINE_PYDEPS = Path(os.environ.get("TRADE_KLINE_PYDEPS", LEGACY_RISK_ROOT / "pydeps" if (LEGACY_RISK_ROOT / "pydeps").exists() else ROOT / "pydeps"))
for dependency_path in (ROOT / "pydeps", TRADE_KLINE_PYDEPS):
    if dependency_path.exists() and str(dependency_path) not in sys.path:
        sys.path.insert(0, str(dependency_path))
SOURCE_TXT = Path(os.environ.get("ACCOUNT_REGISTRY_SOURCE_TXT", ROOT / "local_data" / "source_notes.txt"))
WORKBOOK_PATH = OUT_DIR / "problematic_accounts.xlsx"
KLINE_JOBS: dict[str, dict] = {}
KLINE_JOBS_LOCK = threading.Lock()
INLINE_KLINE_LOCK = threading.Lock()
TOXIC_JOBS: dict[str, dict] = {}
TOXIC_JOBS_LOCK = threading.Lock()
PUSH_DISCOVERY_JOBS: dict[str, dict] = {}
PUSH_DISCOVERY_JOBS_LOCK = threading.Lock()
PUSH_DISCOVERY_RUN_LOCK = threading.Lock()
TOXIC_RESULT_CACHE: dict[str, dict] = {}
TOXIC_TICK_LOCK = threading.Lock()
IP_HISTORY_LOCK = threading.Lock()
QUICK_ACTIONS_LOCK = threading.Lock()
ACCOUNT_QUERY_CACHE_TTL = float(os.environ.get("ACCOUNT_QUERY_CACHE_TTL", "30"))
ACCOUNT_QUERY_CACHE: dict[tuple, tuple[float, object]] = {}
ACCOUNT_QUERY_CACHE_LOCK = threading.Lock()
HOST = "127.0.0.1"
PORT = int(os.environ.get("ACCOUNT_REGISTRY_PORT", "8776"))
if ROOT.name == "K_desk_ai_dev" and "ACCOUNT_REGISTRY_PORT" not in os.environ:
    PORT = 8777
PUSH_DISCOVERY_SCRIPT = ROOT / "scripts" / "run_platform_push_discovery.py"

TOXIC_CHECK_TYPES = [
    {"id": "market_pushing", "label": "推盘", "requiresTick": True},
    {"id": "quote_latency_arbitrage", "label": "报价延迟套利", "requiresTick": True},
    {"id": "cross_platform_spread_arbitrage", "label": "跨平台点差套利", "requiresTick": True},
    {"id": "rebate_churning", "label": "刷返佣", "requiresTick": False},
    {"id": "bonus_arbitrage", "label": "赠金套利", "requiresTick": False},
    {"id": "short_close_trading", "label": "短平交易", "requiresTick": False},
    {"id": "internal_lock_arbitrage", "label": "平台内多账户对锁", "requiresTick": False},
    {"id": "high_leverage_lock_arbitrage", "label": "高杠杆锁仓套利", "requiresTick": False},
    {"id": "weekend_gap_trading", "label": "周末跳空交易", "requiresTick": False},
    {"id": "open_betting", "label": "赌开盘", "requiresTick": False},
    {"id": "news_event_betting", "label": "新闻 / 高波动赌博", "requiresTick": True},
]
TOXIC_CHECK_TYPE_MAP = {item["id"]: item for item in TOXIC_CHECK_TYPES}

SHEET_ACCOUNTS = "问题账户"
SHEET_RAW = "原始记录"
SHEET_HELP = "字段说明"
SHEET_HISTORY = "修改历史"

HEADERS = [
    "记录ID",
    "账号",
    "记录类型",
    "关联账号/主体",
    "建议动作",
    "当前分组",
    "风险标签",
    "风险/问题备注",
    "原始记录",
    "加入时间",
    "修改时间",
    "状态",
    "处理人/来源",
    "AI风险等级",
    "AI备注",
    "AI分析时间",
    "AI证据图表",
]

LEGACY_HEADER_ALIASES = {
    "加入时间": ["加入时间", "首次录入日期"],
    "修改时间": ["修改时间", "更新时间"],
}

HISTORY_HEADERS = [
    "历史ID",
    "记录ID",
    "账号",
    "操作",
    "修改时间",
    "修改字段",
    "修改前JSON",
    "修改后JSON",
    "处理人/来源",
]

EDITABLE_FIELDS = [
    "账号",
    "记录类型",
    "关联账号/主体",
    "建议动作",
    "当前分组",
    "风险标签",
    "风险/问题备注",
    "原始记录",
    "状态",
    "处理人/来源",
]

ACTION_CHOICES = [
    "",
    "B",
    "M",
    "M观察",
    "P",
    "P观察",
    "T",
    "A",
    "A/TA",
    "B-M",
    "B-P",
    "M-P",
    "P->A/T",
    "限制出金",
    "自定义",
    "待定",
]
DEFAULT_QUICK_ACTIONS = [value for value in ACTION_CHOICES if value]
PROTECTED_QUICK_ACTIONS = {"自定义"}
STATUS_CHOICES = ["待复核", "观察中", "已确认", "已关闭"]
TYPE_CHOICES = ["账户", "IB/组", "其他"]
JOURNAL_DIR = Path(r"C:\Users\amber\Downloads")


def now_text() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def normalize_datetime(value: object, default_time: str = "12:00:00") -> str:
    text = normalize_text(value)
    if not text:
        return ""
    text = text.replace("/", "-")
    if re.fullmatch(r"\d{8}", text):
        return f"{text[:4]}-{text[4:6]}-{text[6:]} {default_time}"
    if re.fullmatch(r"\d{4}-\d{1,2}-\d{1,2}", text):
        dt = datetime.strptime(text, "%Y-%m-%d")
        return dt.strftime(f"%Y-%m-%d {default_time}")
    if re.fullmatch(r"\d{4}-\d{1,2}-\d{1,2}\s+\d{1,2}:\d{1,2}", text):
        dt = datetime.strptime(text, "%Y-%m-%d %H:%M")
        return dt.strftime("%Y-%m-%d %H:%M:%S")
    if re.fullmatch(r"\d{4}-\d{1,2}-\d{1,2}\s+\d{1,2}:\d{1,2}:\d{1,2}", text):
        dt = datetime.strptime(text, "%Y-%m-%d %H:%M:%S")
        return dt.strftime("%Y-%m-%d %H:%M:%S")
    return text


def account_join_times_from_journals() -> dict[str, str]:
    mapping: dict[str, str] = {}
    group_mapping: dict[str, str] = {}
    for path in sorted(JOURNAL_DIR.glob("journal_202606*.md")):
        match = re.search(r"(\d{8})", path.name)
        if not match:
            continue
        joined_at = normalize_datetime(match.group(1))
        text = path.read_text(encoding="utf-8", errors="ignore")
        for line in text.splitlines():
            if line.startswith("[image") or "base64," in line:
                continue
            if line.strip().lower().startswith("ib "):
                group_mapping.setdefault("IB/组", joined_at)
            for account_id in re.findall(r"\d{5,10}", line):
                mapping.setdefault(account_id, joined_at)

    # These source rows say "同上" in the journal and do not repeat every account id.
    for account_id in ("241002225", "5002797", "634847"):
        mapping.setdefault(account_id, "2026-06-22 12:00:00")
    mapping["5006543"] = "2026-06-23 10:30:00"
    mapping.update({f"__GROUP__{key}": value for key, value in group_mapping.items()})
    return mapping


def default_join_time(record: dict[str, str], journal_times: dict[str, str] | None = None) -> str:
    journal_times = journal_times or account_join_times_from_journals()
    account_id = normalize_text(record.get("账号"))
    if account_id and account_id in journal_times:
        return journal_times[account_id]
    if record.get("记录类型") == "IB/组" and "__GROUP__IB/组" in journal_times:
        return journal_times["__GROUP__IB/组"]
    existing = normalize_text(record.get("加入时间") or record.get("首次录入日期"))
    if existing and existing != "2026-06-23 12:00:00":
        return normalize_datetime(existing)
    raw = normalize_text(record.get("原始记录"))
    for account_id in re.findall(r"\b\d{5,10}\b", raw):
        if account_id in journal_times:
            return journal_times[account_id]
    return "2026-06-16 12:00:00"


def read_source_text() -> str:
    try:
        data = SOURCE_TXT.read_bytes()
    except FileNotFoundError:
        # Source notes are optional after the ledger workbook has been initialized.
        # A missing notes file must not make an account-detail read unavailable.
        return ""
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def normalize_text(value: object) -> str:
    return str(value or "").strip()


def payload_bool(value: object, default: bool = False) -> bool:
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.strip().lower() in {"true", "1", "yes", "on"}
    return bool(value)


def uniq_join(values: list[str], sep: str = "；") -> str:
    seen: set[str] = set()
    out: list[str] = []
    for value in values:
        text = normalize_text(value)
        if text and text not in seen:
            seen.add(text)
            out.append(text)
    return sep.join(out)


def make_record_id(account_id: str, seed: str = "") -> str:
    account_id = normalize_text(account_id)
    if account_id:
        safe_id = re.sub(r"[^0-9A-Za-z_-]+", "_", account_id).strip("_")
        if re.fullmatch(r"[0-9A-Za-z_-]+", account_id) and safe_id:
            return f"ACC-{safe_id}"
        digest = hashlib.sha1(account_id.encode("utf-8")).hexdigest()[:10]
        return f"ACC-{safe_id or 'TEXT'}-{digest}"
    digest = hashlib.sha1(seed.encode("utf-8")).hexdigest()[:10]
    return f"REC-{digest}"


def record_sort_key(record: dict[str, str]) -> tuple[int, int, str]:
    account_id = normalize_text(record.get("账号"))
    if record.get("记录类型") != "账户":
        return (1, 10**18, account_id)
    if account_id.isdigit():
        return (0, int(account_id), account_id)
    return (0, 10**18, account_id)


def derive_action(note: str) -> str:
    text = note.lower()
    if "限制出金" in note:
        return "限制出金"
    if "直接转T" in note or "转T" in note:
        return "P->A/T" if "放P" in note or "抛A" in note else "T"
    if "违规放T" in note or "放T" in note:
        return "T"
    if "b-p" in text or "B-P" in note:
        return "B-P"
    if "b-m" in text or "B-M" in note:
        return "B-M"
    if "m-p" in text or "M-P" in note or "放M-P" in note:
        return "M-P"
    if "放p观察" in text or "放P观察" in note:
        return "P观察"
    if "放p" in text or "放P" in note:
        return "P"
    if "放m观察" in text or "放M观察" in note:
        return "M观察"
    if "放m" in text or "放M" in note:
        return "M"
    return "待定"


def derive_group(action: str) -> str:
    if action in {"M", "M观察", "B-M"}:
        return "M"
    if action in {"P", "P观察", "B-P", "M-P", "P->A/T"}:
        return "P"
    if action == "T":
        return "T"
    if action == "限制出金":
        return "限制出金"
    return ""


def derive_tags(note: str, record_type: str) -> str:
    tag_rules = [
        ("高频", "高频"),
        ("短平", "短平"),
        ("短线", "短线"),
        ("EA", "EA"),
        ("长持仓", "长持仓"),
        ("长时间持仓", "长持仓"),
        ("大手数", "大手数"),
        ("小手数", "小手数"),
        ("逆势加仓", "逆势加仓"),
        ("稳定盈利", "稳定盈利"),
        ("盈利能力强", "盈利能力强"),
        ("抗损", "抗损"),
        ("爆仓", "爆仓风险"),
        ("同名账户", "同名账户"),
        ("跟单", "跟单/同步"),
        ("同步", "跟单/同步"),
        ("套赠金", "疑似套赠金"),
        ("credit", "Credit"),
        ("Credit", "Credit"),
        ("满仓", "满仓/过周末"),
        ("过周末", "满仓/过周末"),
        ("违规", "违规"),
        ("手数变大", "手数变大"),
        ("注册", "新注册"),
        ("没活跃", "不活跃"),
        ("IB", "IB关联"),
        ("ib", "IB关联"),
        ("穿仓", "穿仓"),
        ("菲律宾ip", "菲律宾IP"),
    ]
    tags = [label for keyword, label in tag_rules if keyword in note]
    if record_type == "IB/组":
        tags.append("IB关联")
    return uniq_join(tags)


def parse_source_records() -> list[dict[str, str]]:
    text = read_source_text()
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    records_by_id: dict[str, dict[str, str]] = {}
    last_context = ""

    for index, line in enumerate(lines, start=1):
        record_type = "账户"
        subject = ""
        m = re.match(r"^\s*((?:\d{5,10}\s*)+)(.*)$", line)
        if m:
            account_ids = re.findall(r"\d{5,10}", m.group(1))
            note = m.group(2).strip() or line
        else:
            account_ids = [""]
            note = line
            record_type = "IB/组" if line.lower().startswith("ib ") else "其他"
            ib_match = re.match(r"^(IB\s+.+?)\s+下的", line, flags=re.IGNORECASE)
            subject = ib_match.group(1) if ib_match else ""

        if note == "同上" and last_context:
            note = f"同上：{last_context}"
        elif note:
            last_context = note

        for account_id in account_ids:
            related_ids = [
                value for value in re.findall(r"\d{5,10}", note) if value != account_id
            ]
            related = uniq_join(related_ids)
            if subject:
                related = uniq_join([subject, related])
            action = derive_action(note)
            group = derive_group(action)
            tags = derive_tags(note, record_type)
            record_id = make_record_id(account_id, f"{index}:{line}")

            if record_id not in records_by_id:
                records_by_id[record_id] = {
                    "记录ID": record_id,
                    "账号": account_id,
                    "记录类型": record_type,
                    "关联账号/主体": related,
                    "建议动作": action,
                    "当前分组": group,
                    "风险标签": tags,
                    "风险/问题备注": note,
                    "原始记录": line,
                    "加入时间": datetime.now().strftime("%Y-%m-%d 12:00:00"),
                    "修改时间": now_text(),
                    "状态": "待复核",
                    "处理人/来源": SOURCE_TXT.name,
                }
            else:
                rec = records_by_id[record_id]
                rec["关联账号/主体"] = uniq_join([rec["关联账号/主体"], related])
                rec["风险标签"] = uniq_join([rec["风险标签"], tags])
                rec["风险/问题备注"] = uniq_join([rec["风险/问题备注"], note], "\n")
                rec["原始记录"] = uniq_join([rec["原始记录"], line], "\n")
                if rec["建议动作"] in {"", "待定"} and action != "待定":
                    rec["建议动作"] = action
                    rec["当前分组"] = group

    return sorted(
        records_by_id.values(),
        key=record_sort_key,
    )


def ensure_output_dir() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)


def style_accounts_sheet(ws) -> None:
    ws.freeze_panes = "B2"
    ws.sheet_view.showGridLines = False
    last_col = ws.cell(1, len(HEADERS)).column_letter
    ws.auto_filter.ref = f"A1:{last_col}{max(ws.max_row, 2)}"

    header_fill = PatternFill("solid", fgColor="263238")
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(style="thin", color="D7DEE2")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = Border(bottom=Side(style="medium", color="90A4AE"))

    widths = {
        "A": 16,
        "B": 14,
        "C": 12,
        "D": 24,
        "E": 14,
        "F": 12,
        "G": 30,
        "H": 48,
        "I": 52,
        "J": 20,
        "K": 20,
        "L": 12,
        "M": 20,
    }
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=len(HEADERS)):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=cell.column in {7, 8, 9})
            cell.border = Border(bottom=thin)
        row[0].font = Font(color="607D8B")
        row[1].number_format = "@"

    for row_idx in range(2, ws.max_row + 1):
        ws.row_dimensions[row_idx].height = 42

    add_list_validation(ws, "C", TYPE_CHOICES)
    add_list_validation(ws, "E", ACTION_CHOICES)
    add_list_validation(ws, "L", STATUS_CHOICES)

    ref = f"A1:{last_col}{max(ws.max_row, 2)}"
    if "ProblemAccounts" not in ws.tables:
        table = Table(displayName="ProblemAccounts", ref=ref)
        table.tableStyleInfo = TableStyleInfo(
            name="TableStyleMedium2",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False,
        )
        ws.add_table(table)
    else:
        ws.tables["ProblemAccounts"].ref = ref


def add_list_validation(ws, column: str, values: list[str]) -> None:
    formula = '"' + ",".join(values) + '"'
    dv = DataValidation(type="list", formula1=formula, allow_blank=True)
    ws.add_data_validation(dv)
    dv.add(f"{column}2:{column}5000")


def style_history_sheet(ws) -> None:
    ws.freeze_panes = "A2"
    ws.sheet_view.showGridLines = False
    header_fill = PatternFill("solid", fgColor="37474F")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    widths = {"A": 18, "B": 16, "C": 14, "D": 12, "E": 20, "F": 36, "G": 70, "H": 70, "I": 20}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=cell.column in {6, 7, 8})


def write_workbook(
    records: list[dict[str, str]],
    raw_lines: list[str] | None = None,
    history_rows: list[dict[str, str]] | None = None,
) -> None:
    ensure_output_dir()
    wb = Workbook()
    ws = wb.active
    ws.title = SHEET_ACCOUNTS
    ws.append(HEADERS)
    for record in records:
        ws.append([record.get(header, "") for header in HEADERS])
    style_accounts_sheet(ws)

    raw = wb.create_sheet(SHEET_RAW)
    raw.append(["序号", "原始记录"])
    for idx, line in enumerate(raw_lines or [], start=1):
        raw.append([idx, line])
    raw.column_dimensions["A"].width = 10
    raw.column_dimensions["B"].width = 100
    raw.freeze_panes = "A2"
    raw.sheet_view.showGridLines = False
    for cell in raw[1]:
        cell.fill = PatternFill("solid", fgColor="455A64")
        cell.font = Font(color="FFFFFF", bold=True)
    for row in raw.iter_rows(min_row=2, max_row=raw.max_row):
        row[1].alignment = Alignment(wrap_text=True, vertical="top")

    help_ws = wb.create_sheet(SHEET_HELP)
    help_rows = [
        ("字段", "说明"),
        ("记录ID", "网页增删改查使用的唯一键。账号记录使用 ACC-账号。"),
        ("账号", "MT4/MT5 账号，按文本原始记录拆分后去重。"),
        ("建议动作", "从原始备注初步提取，可在网页或 Excel 中调整。"),
        ("风险标签", "从备注自动抽取的可筛选标签，后续可手动增删。"),
        ("原始记录", "保留来源文本，便于回溯。"),
    ]
    for row in help_rows:
        help_ws.append(row)
    help_ws.column_dimensions["A"].width = 18
    help_ws.column_dimensions["B"].width = 80
    for cell in help_ws[1]:
        cell.fill = PatternFill("solid", fgColor="00695C")
        cell.font = Font(color="FFFFFF", bold=True)
    help_ws["B2"].comment = Comment(
        "本文件只用于本地维护已查验的问题账户记录，不执行任何 MT4/MT5 Manager 修改操作。",
        "Codex",
    )

    history = wb.create_sheet(SHEET_HISTORY)
    history.append(HISTORY_HEADERS)
    for row in history_rows or []:
        history.append([row.get(header, "") for header in HISTORY_HEADERS])
    style_history_sheet(history)

    wb.save(WORKBOOK_PATH)


def init_workbook(force: bool = False) -> None:
    if WORKBOOK_PATH.exists() and not force:
        migrate_workbook()
        return
    text = read_source_text()
    raw_lines = [line.strip() for line in text.splitlines() if line.strip()]
    write_workbook(parse_source_records(), raw_lines)
    migrate_workbook()


def read_history_rows(wb=None) -> list[dict[str, str]]:
    close_after = False
    if wb is None:
        if not WORKBOOK_PATH.exists():
            return []
        wb = load_workbook(WORKBOOK_PATH, read_only=True, data_only=True)
        close_after = True
    if SHEET_HISTORY not in wb.sheetnames:
        return []
    ws = wb[SHEET_HISTORY]
    headers = [normalize_text(cell.value) for cell in ws[1]]
    rows: list[dict[str, str]] = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        item = {headers[i]: normalize_text(row[i]) for i in range(min(len(headers), len(row)))}
        for header in HISTORY_HEADERS:
            item.setdefault(header, "")
        rows.append(item)
    if close_after:
        wb.close()
    return rows


def history_changed_fields(before: dict[str, str], after: dict[str, str]) -> str:
    changed = []
    for header in HEADERS:
        if header == "修改时间":
            continue
        if normalize_text(before.get(header)) != normalize_text(after.get(header)):
            changed.append(header)
    return "；".join(changed)


def make_history_row(
    before: dict[str, str],
    after: dict[str, str] | None,
    operation: str,
    changed_fields: str | None = None,
) -> dict[str, str]:
    after = after or {}
    stamp = now_text()
    return {
        "历史ID": f"HIS-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}",
        "记录ID": before.get("记录ID") or after.get("记录ID", ""),
        "账号": before.get("账号") or after.get("账号", ""),
        "操作": operation,
        "修改时间": stamp,
        "修改字段": changed_fields if changed_fields is not None else history_changed_fields(before, after),
        "修改前JSON": json.dumps(before, ensure_ascii=False, sort_keys=True),
        "修改后JSON": json.dumps(after, ensure_ascii=False, sort_keys=True) if after else "",
        "处理人/来源": after.get("处理人/来源") or before.get("处理人/来源", ""),
    }


def migrate_workbook() -> None:
    if not WORKBOOK_PATH.exists():
        return
    wb = load_workbook(WORKBOOK_PATH)
    if SHEET_ACCOUNTS not in wb.sheetnames:
        return
    ws = wb[SHEET_ACCOUNTS]
    current_headers = [normalize_text(cell.value) for cell in ws[1]]
    needs_migration = current_headers != HEADERS or SHEET_HISTORY not in wb.sheetnames
    if not needs_migration:
        return

    header_index = {header: idx for idx, header in enumerate(current_headers)}
    journal_times = account_join_times_from_journals()
    records: list[dict[str, str]] = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        rec: dict[str, str] = {}
        for header in HEADERS:
            value = ""
            for source_header in LEGACY_HEADER_ALIASES.get(header, [header]):
                if source_header in header_index and header_index[source_header] < len(row):
                    value = normalize_text(row[header_index[source_header]])
                    if value:
                        break
            rec[header] = value
        if not rec["记录ID"]:
            rec["记录ID"] = make_record_id(rec["账号"], json.dumps(rec, ensure_ascii=False))
        rec["加入时间"] = default_join_time(rec, journal_times)
        rec["修改时间"] = normalize_datetime(rec["修改时间"]) or rec["加入时间"]
        records.append(rec)

    raw_lines: list[str] = []
    if SHEET_RAW in wb.sheetnames:
        raw_ws = wb[SHEET_RAW]
        for row in raw_ws.iter_rows(min_row=2, values_only=True):
            if len(row) > 1 and row[1]:
                raw_lines.append(str(row[1]))
    history_rows = read_history_rows(wb)
    write_workbook(records, raw_lines, history_rows)


def normalize_initial_record_times() -> None:
    if not WORKBOOK_PATH.exists():
        return
    records = load_records()
    history_rows = read_history_rows()
    changed_ids = {row["记录ID"] for row in history_rows if row["操作"] == "修改"}
    journal_times = account_join_times_from_journals()
    changed = False
    for record in records:
        joined_at = default_join_time(record, journal_times)
        if record["加入时间"] != joined_at:
            record["加入时间"] = joined_at
            changed = True
        if record["记录ID"] not in changed_ids and record["修改时间"] != record["加入时间"]:
            record["修改时间"] = record["加入时间"]
            changed = True
    if changed:
        save_records(records)


def load_records() -> list[dict[str, str]]:
    init_workbook()
    wb = load_workbook(WORKBOOK_PATH)
    ws = wb[SHEET_ACCOUNTS]
    headers = [normalize_text(cell.value) for cell in ws[1]]
    header_index = {header: idx for idx, header in enumerate(headers)}
    records: list[dict[str, str]] = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        rec = {}
        for header in HEADERS:
            value = ""
            for source_header in LEGACY_HEADER_ALIASES.get(header, [header]):
                idx = header_index.get(source_header)
                if idx is not None and idx < len(row):
                    value = normalize_text(row[idx])
                    if value:
                        break
            rec[header] = value
        if not rec["记录ID"]:
            rec["记录ID"] = make_record_id(rec["账号"], json.dumps(rec, ensure_ascii=False))
        rec["加入时间"] = normalize_datetime(rec["加入时间"])
        rec["修改时间"] = normalize_datetime(rec["修改时间"]) or rec["加入时间"]
        records.append(rec)
    wb.close()
    return records


def save_records(records: list[dict[str, str]], extra_history: list[dict[str, str]] | None = None) -> None:
    raw_lines: list[str] = []
    history_rows: list[dict[str, str]] = []
    if WORKBOOK_PATH.exists():
        wb = load_workbook(WORKBOOK_PATH, read_only=True, data_only=True)
        if SHEET_RAW in wb.sheetnames:
            raw_ws = wb[SHEET_RAW]
            for row in raw_ws.iter_rows(min_row=2, values_only=True):
                if len(row) > 1 and row[1]:
                    raw_lines.append(str(row[1]))
        history_rows = read_history_rows(wb)
        wb.close()
    history_rows.extend(extra_history or [])
    write_workbook(records, raw_lines, history_rows)


def summarize(records: list[dict[str, str]]) -> dict:
    actions: dict[str, int] = {}
    statuses: dict[str, int] = {}
    for record in records:
        actions[record["建议动作"] or "未填"] = actions.get(record["建议动作"] or "未填", 0) + 1
        statuses[record["状态"] or "未填"] = statuses.get(record["状态"] or "未填", 0) + 1
    return {
        "total": len(records),
        "accountRecords": sum(1 for r in records if r["账号"]),
        "groupRecords": sum(1 for r in records if not r["账号"]),
        "actions": actions,
        "statuses": statuses,
        "workbook": str(WORKBOOK_PATH),
        "updatedAt": now_text(),
    }


@lru_cache(maxsize=200000)
def parse_trade_time(value: str | None) -> datetime | None:
    text = normalize_text(value)
    if not text:
        return None
    text = text.replace("T", " ")
    try:
        return datetime.fromisoformat(text)
    except ValueError:
        pass
    for fmt, size in (("%Y-%m-%d %H:%M:%S", 19), ("%Y-%m-%d %H:%M", 16), ("%Y-%m-%d", 10)):
        try:
            return datetime.strptime(text[:size], fmt)
        except ValueError:
            continue
    return None


def trade_time_text(value: datetime | None) -> str:
    return value.strftime("%Y-%m-%d %H:%M:%S") if value else ""


def display_unknown(value: str) -> str:
    text = normalize_text(value)
    if not text or text == "未指定" or "δ" in text:
        return "未指定"
    return text


def trade_db_connect() -> sqlite3.Connection:
    if not TRADE_DB_PATH.exists():
        raise FileNotFoundError(f"交易数据库不存在：{TRADE_DB_PATH}")
    conn = sqlite3.connect(f"file:{TRADE_DB_PATH}?mode=ro", uri=True, timeout=10)
    conn.row_factory = sqlite3.Row
    return conn


def mysql_trade_connect(source: dict, *, connect_timeout: int | None = None, read_timeout: int | None = None):
    if not MYSQL_PASSWORD:
        raise RuntimeError("远程交易数据库未配置密码：请设置 ACCOUNT_TRADE_MYSQL_PASSWORD")
    try:
        import pymysql
    except ImportError as exc:
        raise RuntimeError("缺少 pymysql，无法连接远程 MySQL 交易数据库") from exc
    try:
        effective_connect_timeout = max(int(connect_timeout if connect_timeout is not None else 10), 1)
    except (TypeError, ValueError):
        effective_connect_timeout = 10
    try:
        effective_read_timeout = max(int(read_timeout if read_timeout is not None else source.get("read_timeout", 90)), 1)
    except (TypeError, ValueError):
        effective_read_timeout = 90
    return pymysql.connect(
        host=source["host"],
        port=int(source.get("port", 3306)),
        user=MYSQL_USER,
        password=MYSQL_PASSWORD,
        connect_timeout=effective_connect_timeout,
        read_timeout=effective_read_timeout,
        charset="utf8mb4",
        cursorclass=pymysql.cursors.DictCursor,
    )


def account_cache_get(key: tuple):
    if ACCOUNT_QUERY_CACHE_TTL <= 0:
        return None
    with ACCOUNT_QUERY_CACHE_LOCK:
        cached = ACCOUNT_QUERY_CACHE.get(key)
        if not cached:
            return None
        saved_at, value = cached
        if time.monotonic() - saved_at > ACCOUNT_QUERY_CACHE_TTL:
            ACCOUNT_QUERY_CACHE.pop(key, None)
            return None
        return value


def account_cache_set(key: tuple, value):
    if ACCOUNT_QUERY_CACHE_TTL <= 0:
        return value
    with ACCOUNT_QUERY_CACHE_LOCK:
        if len(ACCOUNT_QUERY_CACHE) >= 256:
            oldest = min(ACCOUNT_QUERY_CACHE, key=lambda item: ACCOUNT_QUERY_CACHE[item][0])
            ACCOUNT_QUERY_CACHE.pop(oldest, None)
        ACCOUNT_QUERY_CACHE[key] = (time.monotonic(), value)
    return value


def normalize_mt5_volume(value: object) -> float:
    try:
        return round(float(value or 0) / 10000, 6)
    except (TypeError, ValueError):
        return 0.0


def normalize_mt4_volume(value: object) -> float:
    try:
        return round(float(value or 0) / 100, 6)
    except (TypeError, ValueError):
        return 0.0


MT5_DEAL_REASON_LABELS = {
    0: "Client", 1: "Expert", 2: "Dealer", 3: "Stop Loss", 4: "Take Profit",
    5: "Stop Out", 6: "Rollover", 7: "Variation Margin", 8: "Gateway", 9: "Signal",
    10: "Settlement", 11: "Transfer", 12: "Synchronization", 13: "External Service",
    14: "Migration", 15: "Mobile", 16: "Web", 17: "Split", 18: "Corporate Action",
}
MT4_TRADE_REASON_LABELS = {
    0: "Client", 1: "Expert", 2: "Dealer", 3: "Signal", 4: "Gateway",
    5: "Mobile", 6: "Web", 7: "API",
}


def trade_reason_label(platform: object, value: object) -> str:
    text = normalize_text(value)
    if not text:
        return ""
    try:
        code = int(value)
    except (TypeError, ValueError):
        return text
    labels = MT5_DEAL_REASON_LABELS if normalize_text(platform).upper() == "MT5" else MT4_TRADE_REASON_LABELS
    return labels.get(code, f"Reason {code}")


def combined_trade_comment(*values: object) -> str:
    comments = []
    for value in values:
        comment = normalize_text(value)
        if comment and comment not in comments:
            comments.append(comment)
    return " / ".join(comments)


def is_ea_trade(row: dict) -> bool:
    reason = normalize_text(row.get("reason")).lower()
    expert_id = normalize_text(row.get("expert_id"))
    comment = normalize_text(row.get("comment"))
    ea_hint = (
        reason == "expert"
        or mysql_int(row.get("reason_code"), -1) == 1
        or expert_id not in {"", "0"}
        or bool(re.search(r"(?i)(\bEA\b|expert|auto\s*trade|robot)", comment))
    )
    if not ea_hint:
        return False
    source_comment = row.get("open_comment") or comment
    if not normalize_text(source_comment):
        return True
    classification = classify_ea_comment(source_comment, ea_hint=True)
    if classification["classification"] == "possible_copy_route":
        return False
    if classification["classification"] == "system_excluded":
        return normalize_text(source_comment).casefold() in {"ea", "expert", "robot", "auto", "auto trade", "autotrade"}
    return True


def is_copy_trade(row: dict) -> bool:
    comment = normalize_text(row.get("comment"))
    reason = normalize_text(row.get("reason")).lower()
    platform = normalize_text(row.get("platform")).upper()
    reason_code = mysql_int(row.get("reason_code"), -1)
    signal_reason = reason in {"signal", "synchronization"} or (platform == "MT5" and reason_code in {9, 12}) or (platform == "MT4" and reason_code == 3)
    explicit_comment = bool(
        re.search(r"(?i)CPT-[A-Z0-9]+#\d+", comment)
        or signal_in_identifier(comment)
        or re.search(r"(?i)\b(copy\s*trade|trade\s*copier|copier|social\s*trading)\b", comment)
    )
    return bool(signal_reason or explicit_comment)


def copy_trade_order_ids(row: dict) -> list[str]:
    source_comment = normalize_text(row.get("open_comment"))
    if not source_comment:
        source_comment = normalize_text(row.get("comment")).split(" / ", 1)[0]
    return list(dict.fromkeys(re.findall(r"(?i)CPT-[A-Z0-9]+#(\d+)", source_comment)))


def copy_trade_channels(row: dict) -> list[str]:
    source_comment = normalize_text(row.get("open_comment"))
    if not source_comment:
        source_comment = normalize_text(row.get("comment")).split(" / ", 1)[0]
    return list(dict.fromkeys(
        value.upper() for value in re.findall(r"(?i)(CPT-[A-Z0-9]+)#\d+", source_comment)
    ))


def mysql_datetime_text(value: object) -> str:
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M:%S")
    return normalize_text(value)


def mysql_int(value: object, default: int = 0) -> int:
    if value is None or value == "":
        return default
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def source_allowed(source: dict, platform: str = "", server: str = "") -> bool:
    platform = normalize_text(platform).upper()
    server = normalize_text(server)
    if platform and platform != normalize_text(source.get("platform")).upper():
        return False
    source_names = {
        normalize_text(source.get("server")),
        normalize_text(source.get("name")),
        *(normalize_text(alias) for alias in source.get("aliases", []) if normalize_text(alias)),
    }
    if server and server not in source_names:
        return False
    return True


def source_physical_identity(source: dict) -> tuple[str, str, str, str, str]:
    """Identify the physical read-only export behind one or more logical CRM routes."""
    return (
        normalize_text(source.get("host")).casefold(),
        normalize_text(source.get("schema")).casefold(),
        normalize_text(source.get("table")).casefold(),
        normalize_text(source.get("kind")).casefold(),
        normalize_text(source.get("platform")).upper(),
    )


def source_trade_user_exists(cur, source: dict, account: str) -> bool:
    if source.get("kind") == "mt5_deals":
        sql = f"select Login from `{source['schema']}`.`mt5_users_view` where Login = %s limit 1"
    else:
        sql = f"select LOGIN from `{source['schema']}`.`mt4_users_view` where LOGIN = %s limit 1"
    cur.execute(sql, (int(account),))
    return bool(cur.fetchone())


def source_account_route_status(cur, source: dict, account: str) -> str:
    """Validate a logical CRM route, with a fail-closed unique trade-user fallback.

    Some newly-created accounts arrive in the trade export before CRM mapping.  The fallback is
    intentionally restricted to the first logical route for one physical source and only when no
    other independent source on the same DB host/platform has the same Login.
    """
    route = source.get("account_route")
    if not isinstance(route, dict):
        return "not_routed"
    schema = normalize_text(route.get("schema"))
    server_code = normalize_text(route.get("mt_server_code"))
    if not schema or not server_code:
        return "not_routed"
    source_key = source_physical_identity(source)
    registry_key = tuple(
        (source_physical_identity(item), normalize_text(item.get("name")))
        for item in MYSQL_SOURCES
    )
    cache_key = ("source-route-status", registry_key, source_key, normalize_text(source.get("name")), normalize_text(account))
    cached = account_cache_get(cache_key)
    if cached is not None:
        return cached
    cur.execute(
        f"select mt_login from `{schema}`.`mt_users_account` "
        "where mt_login = %s and mt_server_code = %s limit 1",
        (int(account), server_code),
    )
    if cur.fetchone():
        return account_cache_set(cache_key, "crm_confirmed")

    # A physical export can back several logical CRM routes.  Without CRM evidence, only its
    # configured canonical route may represent a unique trade-user fallback.
    shared_sources = [item for item in MYSQL_SOURCES if source_physical_identity(item) == source_key]
    canonical_source = shared_sources[0] if shared_sources else source
    if canonical_source is not source and normalize_text(canonical_source.get("name")) != normalize_text(source.get("name")):
        return account_cache_set(cache_key, "shared_physical_source")
    try:
        if not source_trade_user_exists(cur, source, account):
            return account_cache_set(cache_key, "missing")
        for candidate in MYSQL_SOURCES:
            if source_physical_identity(candidate) == source_key:
                continue
            if normalize_text(candidate.get("host")).casefold() != normalize_text(source.get("host")).casefold():
                continue
            if normalize_text(candidate.get("platform")).upper() != normalize_text(source.get("platform")).upper():
                continue
            if source_trade_user_exists(cur, candidate, account):
                return account_cache_set(cache_key, "ambiguous_trade_user_fallback")
    except Exception:
        # We cannot prove uniqueness when a compatible users view is unavailable.
        return account_cache_set(cache_key, "trade_user_fallback_unavailable")
    return account_cache_set(cache_key, "unique_trade_user_fallback")


def source_account_exists(cur, source: dict, account: str) -> bool:
    return source_account_route_status(cur, source, account) in {
        "not_routed", "crm_confirmed", "unique_trade_user_fallback",
    }


def query_mysql_account_lookup_source(source: dict, account: str) -> dict | None:
    if source.get("kind") == "mt5_deals":
        sql = f"""
            select count(*) as RawRows,
                   count(distinct nullif(PositionID, 0)) as OrderCount,
                   min(Time) as FirstTime, max(Time) as LastTime,
                   group_concat(distinct Symbol order by Symbol separator ',') as Symbols
            from `{source['schema']}`.`{source['table']}`
            where Login = %s and Action in (0, 1)
        """
    else:
        sql = f"""
            select count(*) as RawRows, count(*) as OrderCount,
                   min(OPEN_TIME) as FirstTime, max(CLOSE_TIME) as LastTime,
                   group_concat(distinct SYMBOL order by SYMBOL separator ',') as Symbols
            from `{source['schema']}`.`{source['table']}`
            where LOGIN = %s and CMD in (0, 1)
        """
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            if not source_account_exists(cur, source, account):
                return None
            route_validation = source_account_route_status(cur, source, account)
            cur.execute(sql, (int(account),))
            row = cur.fetchone() or {}
            account_meta = account_money_meta(source_name=source.get("name"))
            if mysql_int(row.get("RawRows")):
                if source.get("kind") == "mt5_deals":
                    account_meta = query_mysql_mt5_account_meta(cur, source, account)
                else:
                    try:
                        cur.execute(
                            f"select CURRENCY, `GROUP` as AccountGroup from `{source['schema']}`.`mt4_users_view` "
                            "where LOGIN = %s limit 1",
                            (int(account),),
                        )
                        meta_row = cur.fetchone() or {}
                        if meta_row:
                            account_meta = account_money_meta(
                                meta_row.get("CURRENCY"), meta_row.get("AccountGroup"),
                                source.get("name"), "mt4_users_view",
                            )
                    except Exception:
                        pass
    if not mysql_int(row.get("RawRows")):
        source_label = " / ".join(
            item for item in (normalize_text(source.get("platform")), normalize_text(source.get("server"))) if item
        )
        return {
            "exists": False,
            "dbSource": "mysql",
            "account": account,
            "orderCount": 0,
            "chartableOrderCount": 0,
            "firstTime": "",
            "lastTime": "",
            "platforms": [{"value": source.get("platform", ""), "label": source.get("platform", "")}],
            "servers": [{"value": source.get("server", ""), "label": source.get("server", "")}],
            "symbols": [],
            "latestSource": {"platform": source.get("platform", ""), "server": source.get("server", "")},
            "accountMeta": account_meta,
            "routeValidation": route_validation,
            "error": f"已确认 {source_label or '账户来源'}，账户暂未做单",
            "refreshedAt": now_text(),
        }
    symbols = [item for item in normalize_text(row.get("Symbols")).split(",") if item]
    return {
        "exists": True,
        "dbSource": "mysql",
        "account": account,
        "orderCount": mysql_int(row.get("OrderCount")),
        "chartableOrderCount": mysql_int(row.get("OrderCount")),
        "firstTime": mysql_datetime_text(row.get("FirstTime")),
        "lastTime": mysql_datetime_text(row.get("LastTime")),
        "platforms": [{"value": source.get("platform", ""), "label": source.get("platform", "")}],
        "servers": [{"value": source.get("server", ""), "label": source.get("server", "")}],
        "symbols": symbols,
        "latestSource": {"platform": source.get("platform", ""), "server": source.get("server", "")},
        "accountMeta": account_meta,
        "routeValidation": route_validation,
        "refreshedAt": now_text(),
    }


def query_sqlite_account_lookup(account: str) -> list[dict]:
    if not TRADE_DB_PATH.exists():
        return []
    sql = """
        select platform, server, count(*) as OrderCount,
               min(coalesce(open_time, close_time)) as FirstTime,
               max(coalesce(close_time, open_time)) as LastTime,
               group_concat(distinct symbol) as Symbols
        from trades where account = ?
        group by platform, server
        order by platform, server
    """
    with trade_db_connect() as conn:
        rows = conn.execute(sql, (account,)).fetchall()
    return [{
        "exists": True,
        "dbSource": "sqlite",
        "account": account,
        "orderCount": mysql_int(row["OrderCount"]),
        "chartableOrderCount": mysql_int(row["OrderCount"]),
        "firstTime": normalize_text(row["FirstTime"]),
        "lastTime": normalize_text(row["LastTime"]),
        "platforms": [{"value": normalize_text(row["platform"]), "label": display_unknown(row["platform"])}],
        "servers": [{"value": normalize_text(row["server"]), "label": display_unknown(row["server"])}],
        "symbols": [item for item in normalize_text(row["Symbols"]).split(",") if item],
        "latestSource": {"platform": normalize_text(row["platform"]), "server": display_unknown(row["server"])},
        "accountMeta": account_money_meta(),
        "refreshedAt": now_text(),
    } for row in rows]


def account_lookup_databases(account: str) -> list[dict]:
    account = normalize_text(account)
    if not re.fullmatch(r"\d+", account):
        return []
    cache_key = ("lookup", account)
    cached = account_cache_get(cache_key)
    if cached is not None:
        return cached
    matches: list[dict] = []
    errors: list[str] = []
    if TRADE_DB_SOURCE in {"mysql", "auto"}:
        with ThreadPoolExecutor(max_workers=min(8, len(MYSQL_SOURCES)), thread_name_prefix="lookup-source") as executor:
            futures = {executor.submit(query_mysql_account_lookup_source, source, account): source for source in MYSQL_SOURCES}
            for future in as_completed(futures):
                try:
                    result = future.result()
                    if result:
                        matches.append(result)
                except Exception as exc:
                    errors.append(f"{futures[future].get('name', '交易源')}: {exc}")
                    continue
    if not matches and TRADE_DB_SOURCE in {"sqlite", "auto"}:
        matches = query_sqlite_account_lookup(account)
    if not matches and errors:
        matches = [{
            "exists": False,
            "queryFailed": True,
            "dbSource": "mysql",
            "account": account,
            "orderCount": 0,
            "chartableOrderCount": 0,
            "symbols": [],
            "latestSource": {},
            "accountMeta": account_money_meta(),
            "error": "账号查询失败，未能确认交易记录；请稍后重试",
            "errorDetails": errors[:3],
            "refreshedAt": now_text(),
        }]
    source_order = {(source["platform"], source["server"]): index for index, source in enumerate(MYSQL_SOURCES)}
    matches.sort(key=lambda item: (
        not bool(item.get("exists")),
        source_order.get((
            (item.get("latestSource") or {}).get("platform", ""),
            (item.get("latestSource") or {}).get("server", ""),
        ), len(source_order)),
    ))
    return account_cache_set(cache_key, matches)


def group_indicates_cent_account(group: object) -> bool:
    group_parts = {part for part in re.split(r"[\\/_\-.]+", normalize_text(group).upper()) if part}
    return bool(group_parts.intersection({"CENT", "USC"}))


def mt5_group_currency(group: object, default_currency: object = "USD") -> str:
    group_parts = [part for part in re.split(r"[\\/_\-.]+", normalize_text(group).upper()) if part]
    if any(part in {"CENT", "USC"} for part in group_parts):
        return "USC"
    supported = {"USD", "EUR", "GBP", "AUD", "JPY", "CAD", "CHF", "NZD", "SGD", "HKD"}
    return next((part for part in group_parts if part in supported), normalize_text(default_currency).upper() or "USD")


def account_money_meta(
    currency: object = "",
    group: object = "",
    source_name: object = "",
    currency_source: object = "",
) -> dict:
    account_currency = normalize_text(currency).upper()
    account_group = normalize_text(group)
    group_is_cent = group_indicates_cent_account(account_group)
    is_cent = account_currency == "USC" or group_is_cent
    if is_cent:
        account_currency = "USC"
    detected_source = normalize_text(currency_source)
    if not detected_source:
        detected_source = "currency" if normalize_text(currency) else "group" if group_is_cent else ""
    return {
        "currency": account_currency,
        "displayCurrency": "USD" if is_cent else account_currency,
        "moneyScale": 0.01 if is_cent else 1.0,
        "isCentAccount": is_cent,
        "group": account_group,
        "source": normalize_text(source_name),
        "currencySource": detected_source,
    }


def query_mysql_mt5_account_meta(cur, source: dict, account: str) -> dict:
    default_currency = normalize_text(source.get("default_currency")).upper() or "USD"
    try:
        cur.execute(
            f"select `Group` as AccountGroup from `{source['schema']}`.`mt5_users_view` where Login = %s limit 1",
            (int(account),),
        )
        row = cur.fetchone() or {}
        if row:
            account_group = row.get("AccountGroup")
            return account_money_meta(
                mt5_group_currency(account_group, default_currency),
                account_group,
                source.get("name"),
                "mt5_users_view.Group",
            )
    except Exception:
        # Some export schemas contain deals before a compatible users view is available.
        pass
    return account_money_meta(
        default_currency,
        source_name=source.get("name"),
        currency_source="source.default_currency",
    )


def mt5_deals_to_trades(deals: list[dict], source: dict, account: str, account_meta: dict | None = None) -> list[dict]:
    account_meta = account_meta or account_money_meta(source_name=source.get("name"))
    money_scale = numeric_value(account_meta.get("moneyScale")) or 1.0
    grouped: dict[str, list[dict]] = defaultdict(list)
    for deal in deals:
        position_id = normalize_text(deal.get("PositionID"))
        if not position_id or position_id == "0":
            continue
        grouped[position_id].append(deal)

    out: list[dict] = []
    for position_id, items in grouped.items():
        items.sort(key=lambda row: (mysql_datetime_text(row.get("TimeMsc") or row.get("Time")), mysql_int(row.get("Deal"))))
        opens = [row for row in items if mysql_int(row.get("Entry"), -1) == 0 and mysql_int(row.get("Action"), -1) in (0, 1)]
        closes = [row for row in items if mysql_int(row.get("Entry"), -1) == 1 and mysql_int(row.get("Action"), -1) in (0, 1)]
        if not opens or not closes:
            reversal = next(
                (row for row in items if mysql_int(row.get("Entry"), -1) in (2, 3) and mysql_int(row.get("Action"), -1) in (0, 1)),
                None,
            )
            if reversal:
                event_time = mysql_datetime_text(reversal.get("Time"))
                event_time_msc = mysql_datetime_text(reversal.get("TimeMsc") or reversal.get("Time"))
                action = mysql_int(reversal.get("Action"))
                out.append({
                    "id": normalize_text(reversal.get("Deal") or position_id),
                    "source_id": normalize_text(source.get("name")),
                    "data_source": "mysql",
                    "platform": source.get("platform", "MT5"), "server": source.get("server", source.get("name", "")),
                    "account": account, "account_currency": account_meta.get("currency", ""),
                    "display_currency": account_meta.get("displayCurrency", ""), "money_scale": money_scale,
                    "is_cent_account": bool(account_meta.get("isCentAccount")),
                    "currency_source": normalize_text(account_meta.get("currencySource")),
                    "ticket": normalize_text(reversal.get("PositionID") or reversal.get("Order") or reversal.get("Deal")),
                    "open_time": event_time, "close_time": event_time,
                    "open_time_msc": event_time_msc, "close_time_msc": event_time_msc,
                    "type": "buy" if action == 0 else "sell",
                    "volume": normalize_mt5_volume(reversal.get("VolumeClosed") or reversal.get("Volume")),
                    "symbol": normalize_text(reversal.get("Symbol")),
                    "open_price": reversal.get("Price") or 0, "close_price": reversal.get("Price") or 0,
                    "commission": numeric_value(reversal.get("Commission")) * money_scale,
                    "fee": numeric_value(reversal.get("Fee")) * money_scale, "taxes": 0,
                    "swap": numeric_value(reversal.get("Storage")) * money_scale,
                    "profit": numeric_value(reversal.get("Profit")) * money_scale,
                    "sl": reversal.get("PriceSL") or "", "tp": reversal.get("PriceTP") or "",
                    "reason": trade_reason_label("MT5", reversal.get("Reason")),
                    "reason_code": mysql_int(reversal.get("Reason"), -1),
                    "open_comment": normalize_text(reversal.get("Comment")), "comment": normalize_text(reversal.get("Comment")),
                    "expert_id": normalize_text(reversal.get("ExpertID")),
                    "tick_value": numeric_value(reversal.get("TickValue")), "tick_size": numeric_value(reversal.get("TickSize")),
                    "contract_size": numeric_value(reversal.get("ContractSize")),
                    "market_bid": numeric_value(reversal.get("MarketBid")), "market_ask": numeric_value(reversal.get("MarketAsk")),
                    "price_gateway": numeric_value(reversal.get("PriceGateway")), "holding_seconds": 0,
                    "raw_json": json.dumps({"source": source.get("name"), "reversal": reversal}, ensure_ascii=False, default=str),
                })
            continue
        open_row = opens[0]
        for close_row in closes:
            open_time = mysql_datetime_text(open_row.get("Time"))
            close_time = mysql_datetime_text(close_row.get("Time"))
            open_time_msc = mysql_datetime_text(open_row.get("TimeMsc") or open_row.get("Time"))
            close_time_msc = mysql_datetime_text(close_row.get("TimeMsc") or close_row.get("Time"))
            open_dt = parse_trade_time(open_time_msc)
            close_dt = parse_trade_time(close_time_msc)
            action = mysql_int(open_row.get("Action"))
            out.append({
                "id": normalize_text(close_row.get("Deal") or position_id),
                "source_id": normalize_text(source.get("name")),
                "data_source": "mysql",
                "platform": source.get("platform", "MT5"),
                "server": source.get("server", source.get("name", "")),
                "account": account,
                "account_currency": account_meta.get("currency", ""),
                "display_currency": account_meta.get("displayCurrency", ""),
                "money_scale": money_scale,
                "is_cent_account": bool(account_meta.get("isCentAccount")),
                "currency_source": normalize_text(account_meta.get("currencySource")),
                "ticket": normalize_text(open_row.get("PositionID") or open_row.get("Order") or open_row.get("Deal")),
                "open_time": open_time,
                "close_time": close_time,
                "open_time_msc": open_time_msc,
                "close_time_msc": close_time_msc,
                "type": "buy" if action == 0 else "sell",
                "volume": normalize_mt5_volume(close_row.get("VolumeClosed") or close_row.get("Volume") or open_row.get("Volume")),
                "symbol": normalize_text(close_row.get("Symbol") or open_row.get("Symbol")),
                "open_price": open_row.get("Price") or 0,
                "close_price": close_row.get("Price") or 0,
                "commission": (float(open_row.get("Commission") or 0) + float(close_row.get("Commission") or 0)) * money_scale,
                "fee": (float(open_row.get("Fee") or 0) + float(close_row.get("Fee") or 0)) * money_scale,
                "taxes": 0,
                "swap": (float(open_row.get("Storage") or 0) + float(close_row.get("Storage") or 0)) * money_scale,
                "profit": numeric_value(close_row.get("Profit")) * money_scale,
                "sl": open_row.get("PriceSL") or "",
                "tp": open_row.get("PriceTP") or "",
                "reason": trade_reason_label("MT5", open_row.get("Reason") if open_row.get("Reason") is not None else close_row.get("Reason")),
                "reason_code": mysql_int(open_row.get("Reason") if open_row.get("Reason") is not None else close_row.get("Reason"), -1),
                "open_comment": normalize_text(open_row.get("Comment")),
                "comment": combined_trade_comment(open_row.get("Comment"), close_row.get("Comment")),
                "expert_id": normalize_text(open_row.get("ExpertID") or close_row.get("ExpertID")),
                "tick_value": numeric_value(open_row.get("TickValue") or close_row.get("TickValue")),
                "tick_size": numeric_value(open_row.get("TickSize") or close_row.get("TickSize")),
                "contract_size": numeric_value(open_row.get("ContractSize") or close_row.get("ContractSize")),
                "market_bid": numeric_value(open_row.get("MarketBid")),
                "market_ask": numeric_value(open_row.get("MarketAsk")),
                "price_gateway": numeric_value(open_row.get("PriceGateway")),
                "holding_seconds": (close_dt - open_dt).total_seconds() if open_dt and close_dt else "",
                "raw_json": json.dumps({"source": source.get("name"), "open": open_row, "close": close_row}, ensure_ascii=False, default=str),
            })
    return out


def query_mysql_mt5_source(source: dict, account: str, symbol: str = "", start: str = "", end: str = "", limit: int | None = 50000) -> list[dict]:
    where = ["Login = %s", "Action in (0, 1)", "Entry in (0, 1, 2, 3)"]
    args: list[object] = [int(account)]
    if symbol:
        where.append("Symbol = %s")
        args.append(symbol)
    limit_clause = "limit %s" if limit is not None else ""
    sql = f"""
        select Deal, Login, `Order`, PositionID, Action, Entry, Reason, Time, TimeMsc,
               Symbol, Price, Volume, VolumeExt, VolumeClosed, VolumeClosedExt,
               Profit, Commission, Storage, Fee, Comment, ExpertID, PriceSL, PriceTP,
               ContractSize, TickValue, TickSize, MarketBid, MarketAsk, PriceGateway
        from `{source['schema']}`.`{source['table']}`
        where {' and '.join(where)}
        order by PositionID, Time, Deal
        {limit_clause}
    """
    if limit is not None:
        args.append(limit * 2)
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            if not source_account_exists(cur, source, account):
                return []
            cur.execute(sql, args)
            deals = cur.fetchall()
            account_meta = query_mysql_mt5_account_meta(cur, source, account) if deals else account_money_meta(source_name=source.get("name"))
    trades = mt5_deals_to_trades(deals, source, account, account_meta)
    start_dt = parse_trade_time(start)
    end_dt = parse_trade_time(end)
    if start_dt:
        trades = [row for row in trades if (parse_trade_time(row.get("close_time")) or datetime.min) >= start_dt]
    if end_dt:
        trades = [row for row in trades if (parse_trade_time(row.get("open_time")) or datetime.max) <= end_dt]
    return trades if limit is None else trades[:limit]


def query_mysql_mt4_source(
    source: dict,
    account: str,
    symbol: str = "",
    start: str = "",
    end: str = "",
    limit: int | None = 50000,
    tickets: list[int] | None = None,
) -> list[dict]:
    where = ["LOGIN = %s", "CMD in (0, 1)", "CLOSE_TIME > OPEN_TIME"]
    args: list[object] = [int(account)]
    if symbol:
        where.append("SYMBOL = %s")
        args.append(symbol)
    if start:
        where.append("CLOSE_TIME >= %s")
        args.append(start)
    if end:
        where.append("OPEN_TIME <= %s")
        args.append(end)
    if tickets:
        where.append(f"TICKET in ({','.join(['%s'] * len(tickets))})")
        args.extend(tickets)
    limit_clause = ""
    if limit is not None:
        if limit <= 0:
            return []
        limit_clause = "limit %s"
        args.append(limit)
    sql = f"""
        select TICKET, LOGIN, CMD, SYMBOL, VOLUME, OPEN_TIME, OPEN_PRICE, CLOSE_TIME,
               CLOSE_PRICE, COMMISSION, TAXES, SWAPS, PROFIT, SL, TP, REASON, MAGIC, COMMENT
        from `{source['schema']}`.`{source['table']}`
        where {' and '.join(where)}
        order by OPEN_TIME, TICKET
        {limit_clause}
    """
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            if not source_account_exists(cur, source, account):
                return []
            cur.execute(sql, args)
            rows = cur.fetchall()
            account_meta = account_money_meta(source_name=source.get("name"))
            if rows:
                try:
                    cur.execute(
                        f"select CURRENCY, `GROUP` as AccountGroup from `{source['schema']}`.`mt4_users_view` where LOGIN = %s limit 1",
                        (int(account),),
                    )
                    meta_row = cur.fetchone() or {}
                    if meta_row:
                        account_meta = account_money_meta(
                            meta_row.get("CURRENCY"), meta_row.get("AccountGroup"),
                            source.get("name"), "mt4_users_view",
                        )
                except Exception:
                    pass
    money_scale = numeric_value(account_meta.get("moneyScale")) or 1.0
    out: list[dict] = []
    for row in rows:
        open_time = mysql_datetime_text(row.get("OPEN_TIME"))
        close_time = mysql_datetime_text(row.get("CLOSE_TIME"))
        open_dt = parse_trade_time(open_time)
        close_dt = parse_trade_time(close_time)
        if not open_dt or not close_dt or close_dt <= open_dt:
            continue
        out.append({
            "id": normalize_text(row.get("TICKET")),
            "source_id": normalize_text(source.get("name")),
            "data_source": "mysql",
            "platform": source.get("platform", "MT4"),
            "server": source.get("server", source.get("name", "")),
            "account": account,
            "account_currency": account_meta.get("currency", ""),
            "display_currency": account_meta.get("displayCurrency", ""),
            "money_scale": money_scale,
            "is_cent_account": bool(account_meta.get("isCentAccount")),
            "currency_source": normalize_text(account_meta.get("currencySource")),
            "ticket": normalize_text(row.get("TICKET")),
            "open_time": open_time,
            "close_time": close_time,
            "type": "buy" if mysql_int(row.get("CMD")) == 0 else "sell",
            "volume": normalize_mt4_volume(row.get("VOLUME")),
            "symbol": normalize_text(row.get("SYMBOL")),
            "open_price": row.get("OPEN_PRICE") or 0,
            "close_price": row.get("CLOSE_PRICE") or 0,
            "commission": numeric_value(row.get("COMMISSION")) * money_scale,
            "taxes": numeric_value(row.get("TAXES")) * money_scale,
            "swap": numeric_value(row.get("SWAPS")) * money_scale,
            "profit": numeric_value(row.get("PROFIT")) * money_scale,
            "sl": row.get("SL") or "",
            "tp": row.get("TP") or "",
            "reason": trade_reason_label("MT4", row.get("REASON")),
            "reason_code": mysql_int(row.get("REASON"), -1),
            "comment": normalize_text(row.get("COMMENT")),
            "expert_id": normalize_text(row.get("MAGIC")) if mysql_int(row.get("MAGIC")) else "",
            "holding_seconds": (close_dt - open_dt).total_seconds() if open_dt and close_dt else "",
        })
    return out


def query_mysql_mt4_orders_page_source(
    source: dict,
    account: str,
    page: int,
    page_size: int,
) -> tuple[int, list[dict]]:
    where = "LOGIN = %s and CMD in (0, 1) and CLOSE_TIME > OPEN_TIME"
    offset = (page - 1) * page_size
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            if not source_account_exists(cur, source, account):
                return 0, []
            cur.execute(
                f"select count(*) as Total from `{source['schema']}`.`{source['table']}` where {where}",
                (int(account),),
            )
            total = mysql_int((cur.fetchone() or {}).get("Total"))
            cur.execute(
                f"select TICKET from `{source['schema']}`.`{source['table']}` where {where} "
                "order by CLOSE_TIME desc, TICKET desc limit %s offset %s",
                (int(account), page_size, offset),
            )
            tickets = [mysql_int(row.get("TICKET")) for row in cur.fetchall() if mysql_int(row.get("TICKET"))]
    rows = query_mysql_mt4_source(source, account, limit=None, tickets=tickets) if tickets else []
    rows.sort(
        key=lambda row: (
            parse_trade_time(row.get("close_time") or row.get("open_time")) or datetime.min,
            mysql_int(row.get("ticket")),
        ),
        reverse=True,
    )
    return total, rows


def query_mysql_trades(
    account: str,
    platform: str = "",
    server: str = "",
    symbol: str = "",
    start: str = "",
    end: str = "",
    limit: int | None = 50000,
) -> list[dict]:
    account = normalize_text(account)
    if not re.fullmatch(r"\d+", account):
        return []
    rows: list[dict] = []
    errors: list[str] = []
    sources = [source for source in MYSQL_SOURCES if source_allowed(source, platform=platform, server=server)]

    def query_source(source: dict) -> list[dict]:
        if source.get("kind") == "mt5_deals":
            return query_mysql_mt5_source(source, account, symbol=symbol, start=start, end=end, limit=limit)
        return query_mysql_mt4_source(source, account, symbol=symbol, start=start, end=end, limit=limit)

    if len(sources) == 1:
        try:
            rows.extend(query_source(sources[0]))
        except Exception as exc:
            errors.append(f"{sources[0].get('name')}: {exc}")
    elif sources:
        with ThreadPoolExecutor(max_workers=min(8, len(sources)), thread_name_prefix="trade-source") as executor:
            futures = {executor.submit(query_source, source): source for source in sources}
            for future in as_completed(futures):
                source = futures[future]
                try:
                    rows.extend(future.result())
                except Exception as exc:
                    errors.append(f"{source.get('name')}: {exc}")
    if errors and not rows:
        raise RuntimeError("；".join(errors[:3]))
    rows.sort(key=lambda row: (normalize_text(row.get("open_time")), normalize_text(row.get("ticket"))))
    return rows if limit is None else rows[:limit]


def query_mysql_historical_funds_source(source: dict, account: str) -> dict | None:
    """Read the complete read-only ledger, deal and daily-anchor facts for one routed account."""
    account = normalize_text(account)
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            if not source_account_exists(cur, source, account):
                return None
            if source.get("kind") == "mt5_deals":
                cur.execute(
                    f"select Deal, Login, `Order`, PositionID, Action, Entry, Reason, Time, TimeMsc, Symbol, "
                    f"Profit, Commission, Storage, Fee, Comment "
                    f"from `{source['schema']}`.`{source['table']}` where Login = %s "
                    "order by TimeMsc, Deal",
                    (int(account),),
                )
                events = cur.fetchall()
                meta = query_mysql_mt5_account_meta(cur, source, account)
                cur.execute(
                    f"select Balance, Credit, Equity from `{source['schema']}`.`mt5_accounts` "
                    "where Login = %s limit 1",
                    (int(account),),
                )
                current_row = cur.fetchone() or {}
                current_anchor = {
                    "timestamp": now_text(),
                    "balance": current_row.get("Balance"),
                    "credit": current_row.get("Credit"),
                    "equity": current_row.get("Equity"),
                } if current_row else None
                # This MT5 daily view has no Login index. A historical filter therefore becomes a
                # full-view scan and can block the account page for the database read timeout.
                # Use the indexed deal ledger plus the current account row; do not query the view.
                anchors = []
                daily_anchor_available = False
                daily_anchor_reason = (
                    "MT5 日快照视图未按账号索引，未执行全表查询；"
                    "余额和 Credit 已按当前账户状态回放，历史权益快照不可用"
                )
            else:
                cur.execute(
                    f"select TICKET, LOGIN, CMD, SYMBOL, VOLUME, OPEN_TIME, OPEN_PRICE, CLOSE_TIME, "
                    f"CLOSE_PRICE, COMMISSION, TAXES, SWAPS, PROFIT, COMMENT, REASON, MAGIC "
                    f"from `{source['schema']}`.`{source['table']}` where LOGIN = %s "
                    "order by OPEN_TIME, TICKET",
                    (int(account),),
                )
                events = cur.fetchall()
                cur.execute(
                    f"select TIME, BALANCE, CREDIT, DEPOSIT, EQUITY from `{source['schema']}`.`mt4_daily` "
                    "where LOGIN = %s order by TIME",
                    (int(account),),
                )
                anchors = cur.fetchall()
                meta = account_money_meta(source_name=source.get("name"))
                try:
                    cur.execute(
                        f"select CURRENCY, `GROUP` as AccountGroup from `{source['schema']}`.`mt4_users_view` "
                        "where LOGIN = %s limit 1",
                        (int(account),),
                    )
                    meta_row = cur.fetchone() or {}
                    if meta_row:
                        meta = account_money_meta(
                            meta_row.get("CURRENCY"), meta_row.get("AccountGroup"), source.get("name"),
                            "mt4_users_view",
                        )
                except Exception:
                    pass
                current_anchor = None
                daily_anchor_available = True
                daily_anchor_reason = ""
    return {
        "account": account,
        "platform": source.get("platform", ""),
        "server": source.get("server", source.get("name", "")),
        "source": source.get("name", ""),
        "currency": meta.get("displayCurrency") or meta.get("currency") or "USD",
        "moneyScale": numeric_value(meta.get("moneyScale")) or 1.0,
        "events": [dict(row) for row in events],
        "anchors": [dict(row) for row in anchors],
        "currentAnchor": current_anchor,
        "coverage": {
            "eventRows": len(events),
            "dailyAnchors": len(anchors),
            "dailyAnchorsAvailable": daily_anchor_available,
            "dailyAnchorReason": daily_anchor_reason,
            "completeHistory": True,
        },
    }


def account_historical_funds_source_payload(login: str, filters: dict | None = None) -> dict:
    login = normalize_text(login)
    filters = filters or {}
    platform = normalize_text(filters.get("platform")).upper()
    server = normalize_text(filters.get("server"))
    if not re.fullmatch(r"\d+", login):
        raise ValueError("账号格式无效")
    sources = [source for source in MYSQL_SOURCES if source_allowed(source, platform=platform, server=server)]
    matches: list[dict] = []
    errors: list[str] = []
    for source in sources:
        try:
            result = query_mysql_historical_funds_source(source, login)
            if result:
                matches.append(result)
        except Exception as exc:
            errors.append(f"{source.get('name')}: {exc}")
    if len(matches) > 1:
        return {
            "available": False,
            "account": login,
            "reason": "账号匹配多个服务器，请先选择平台和服务器",
            "sources": [{"platform": item["platform"], "server": item["server"]} for item in matches],
        }
    if not matches:
        if errors:
            raise RuntimeError("；".join(errors[:3]))
        return {"available": False, "account": login, "reason": "数据库没有找到该账号的历史资金数据"}
    return {"available": True, **matches[0]}


def query_mysql_trade_cost_source(
    source: dict,
    account: str,
    symbol: str = "",
    start: str = "",
    end: str = "",
) -> dict | None:
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            if not source_account_exists(cur, source, account):
                return None
            if source.get("kind") == "mt5_deals":
                where = ["Login = %s", "Action in (0, 1)"]
                args: list[object] = [int(account)]
                if symbol:
                    where.append("Symbol = %s")
                    args.append(symbol)
                if start:
                    where.append("Time >= %s")
                    args.append(start)
                if end:
                    where.append("Time <= %s")
                    args.append(end)
                cur.execute(
                    f"select count(*) as RowCount, sum(Commission + Fee) as Commission, sum(Storage) as Swap from `{source['schema']}`.`{source['table']}` where {' and '.join(where)}",
                    args,
                )
                row = cur.fetchone() or {}
                cur.execute(
                    f"select Symbol, sum(Commission + Fee) as Commission, sum(Storage) as Swap from `{source['schema']}`.`{source['table']}` where {' and '.join(where)} group by Symbol",
                    args,
                )
                raw_symbol_costs = cur.fetchall()
                meta = query_mysql_mt5_account_meta(cur, source, account) if mysql_int(row.get("RowCount")) else account_money_meta(source_name=source.get("name"))
                taxes = 0.0
            else:
                where = ["LOGIN = %s", "CMD in (0, 1)"]
                args = [int(account)]
                if symbol:
                    where.append("SYMBOL = %s")
                    args.append(symbol)
                if start:
                    where.append("CLOSE_TIME >= %s")
                    args.append(start)
                if end:
                    where.append("OPEN_TIME <= %s")
                    args.append(end)
                cur.execute(
                    f"select count(*) as RowCount, sum(COMMISSION) as Commission, sum(SWAPS) as Swap, sum(TAXES) as Taxes from `{source['schema']}`.`{source['table']}` where {' and '.join(where)}",
                    args,
                )
                row = cur.fetchone() or {}
                cur.execute(
                    f"select SYMBOL as Symbol, sum(COMMISSION) as Commission, sum(SWAPS) as Swap, sum(TAXES) as Taxes from `{source['schema']}`.`{source['table']}` where {' and '.join(where)} group by SYMBOL",
                    args,
                )
                raw_symbol_costs = cur.fetchall()
                meta = account_money_meta(source_name=source.get("name"))
                if mysql_int(row.get("RowCount")):
                    try:
                        cur.execute(
                            f"select CURRENCY, `GROUP` as AccountGroup from `{source['schema']}`.`mt4_users_view` where LOGIN = %s limit 1",
                            (int(account),),
                        )
                        meta_row = cur.fetchone() or {}
                        if meta_row:
                            meta = account_money_meta(meta_row.get("CURRENCY"), meta_row.get("AccountGroup"), source.get("name"))
                    except Exception:
                        pass
                taxes = numeric_value(row.get("Taxes"))

    if not mysql_int(row.get("RowCount")):
        return None
    scale = numeric_value(meta.get("moneyScale")) or 1.0
    by_symbol = {}
    for item in raw_symbol_costs:
        item_symbol = normalize_text(item.get("Symbol")) or "未指定"
        by_symbol[item_symbol] = {
            "commission": numeric_value(item.get("Commission")) * scale,
            "swap": numeric_value(item.get("Swap")) * scale,
            "taxes": numeric_value(item.get("Taxes")) * scale,
        }
    return {
        "source": {
            "platform": source.get("platform", ""),
            "server": source.get("server", source.get("name", "")),
            "currency": meta.get("currency", ""),
            "commission": rounded(numeric_value(row.get("Commission")) * scale),
            "swap": rounded(numeric_value(row.get("Swap")) * scale),
            "taxes": rounded(taxes * scale),
        },
        "bySymbol": by_symbol,
    }


def _query_mysql_trade_costs_uncached(
    account: str,
    platform: str = "",
    server: str = "",
    symbol: str = "",
    start: str = "",
    end: str = "",
) -> dict:
    account = normalize_text(account)
    summary = {
        "commission": 0.0,
        "swap": 0.0,
        "taxes": 0.0,
        "complete": True,
        "includesOpenTradeFees": True,
        "sources": [],
        "bySymbol": {},
        "errors": [],
    }
    if not re.fullmatch(r"\d+", account):
        return summary
    sources = [source for source in MYSQL_SOURCES if source_allowed(source, platform=platform, server=server)]
    results: list[dict] = []
    if len(sources) == 1:
        try:
            result = query_mysql_trade_cost_source(sources[0], account, symbol=symbol, start=start, end=end)
            if result:
                results.append(result)
        except Exception as exc:
            summary["complete"] = False
            summary["errors"].append(f"{sources[0].get('name')}: {exc}")
    elif sources:
        with ThreadPoolExecutor(max_workers=min(8, len(sources)), thread_name_prefix="cost-source") as executor:
            futures = {
                executor.submit(query_mysql_trade_cost_source, source, account, symbol, start, end): source
                for source in sources
            }
            for future in as_completed(futures):
                source = futures[future]
                try:
                    result = future.result()
                    if result:
                        results.append(result)
                except Exception as exc:
                    summary["complete"] = False
                    summary["errors"].append(f"{source.get('name')}: {exc}")

    for result in results:
        source_costs = result["source"]
        for key in ("commission", "swap", "taxes"):
            summary[key] += numeric_value(source_costs.get(key))
        summary["sources"].append(source_costs)
        for item_symbol, item_costs in result["bySymbol"].items():
            target = summary["bySymbol"].setdefault(item_symbol, {"commission": 0.0, "swap": 0.0, "taxes": 0.0})
            for key in ("commission", "swap", "taxes"):
                target[key] += numeric_value(item_costs.get(key))

    source_order = {normalize_text(source.get("server") or source.get("name")): index for index, source in enumerate(MYSQL_SOURCES)}
    summary["sources"].sort(key=lambda item: source_order.get(normalize_text(item.get("server")), len(source_order)))
    for key in ("commission", "swap", "taxes"):
        summary[key] = rounded(summary[key])
    for values in summary["bySymbol"].values():
        for key in ("commission", "swap", "taxes"):
            values[key] = rounded(values[key])
    return summary


def query_mysql_trade_costs(
    account: str,
    platform: str = "",
    server: str = "",
    symbol: str = "",
    start: str = "",
    end: str = "",
) -> dict:
    cache_key = ("costs", normalize_text(account), platform, server, symbol, start, end)
    cached = account_cache_get(cache_key)
    if cached is not None:
        return cached
    return account_cache_set(cache_key, _query_mysql_trade_costs_uncached(account, platform, server, symbol, start, end))


def clean_trade_type(value: str) -> str:
    text = normalize_text(value).lower()
    if text.startswith("buy"):
        return "buy"
    if text.startswith("sell"):
        return "sell"
    return text


def is_chartable_trade(row: dict) -> bool:
    return clean_trade_type(row.get("type", "")) in {"buy", "sell"} and bool(row.get("open_time")) and bool(row.get("close_time"))


def recent_chartable_kline_trades(rows: list[dict], limit: int) -> list[dict]:
    """Keep the newest completed trade window while retaining chronological chart input."""
    if limit <= 0:
        return rows
    chartable = [row for row in rows if is_chartable_trade(row)]
    chartable.sort(
        key=lambda row: (
            parse_trade_time(row.get("close_time") or row.get("open_time")) or datetime.min,
            normalize_text(row.get("ticket")),
        ),
        reverse=True,
    )
    selected = chartable[:limit]
    selected.sort(
        key=lambda row: (
            parse_trade_time(row.get("open_time") or row.get("close_time")) or datetime.min,
            normalize_text(row.get("ticket")),
        ),
    )
    return selected


def account_inline_kline_html(account: str, filters: dict | None = None) -> str:
    """Build the bounded account-detail chart without creating a K-line job.

    The manual K-line flow deliberately remains a durable :8766 task that
    writes artifacts.  This page-only path reads the chosen trading source and
    the configured quote Terminal, uses its M1 cache when available, then
    returns a self-contained Lightweight Charts document for an in-page iframe.
    It never writes trades, generated HTML, jobs or remote trading state.
    """
    filters = filters or {}
    account = normalize_text(account)
    platform = normalize_text(filters.get("platform")).upper()
    server = normalize_text(filters.get("server"))
    try:
        recent_orders = int(filters.get("recentOrders") or 300)
    except (TypeError, ValueError) as exc:
        raise ValueError("最近订单数量必须是整数") from exc
    recent_orders = max(1, min(300, recent_orders))
    if not account:
        raise ValueError("账号不能为空")
    if not platform or not server:
        raise ValueError("请先选择平台和服务器")

    rows = recent_chartable_kline_trades(
        query_db_trades(account, platform=platform, server=server, limit=50000),
        recent_orders,
    )
    if not rows:
        raise RuntimeError("账户暂未找到可展示的已平仓买卖订单")
    newest = max(
        rows,
        key=lambda row: (
            normalize_text(row.get("close_time") or row.get("open_time")),
            normalize_text(row.get("ticket")),
        ),
    )
    cache_key = (
        "inline-kline-html",
        account,
        platform,
        server,
        recent_orders,
        normalize_text(newest.get("ticket")),
        normalize_text(newest.get("close_time") or newest.get("open_time")),
    )
    cached = account_cache_get(cache_key)
    if cached is not None:
        return cached

    # MT5's local terminal IPC is process-global.  Serialise only this bounded
    # reader so an account page cannot race a quote calibration on another tab.
    with INLINE_KLINE_LOCK:
        cached = account_cache_get(cache_key)
        if cached is not None:
            return cached

        tool_dir = TRADE_KLINE_TOOL_DIR
        if not (tool_dir / "generate_trade_kline_from_statement.py").is_file():
            raise RuntimeError("未配置账户详情直出 K 线报价工具")
        if str(tool_dir) not in sys.path:
            sys.path.insert(0, str(tool_dir))
        try:
            import pandas as pd
            import generate_trade_kline_from_statement as quote_generator
            from kdesk.infrastructure.quote_sources import QuoteSourceRegistry
            from lightweight_trade_kline import build_lightweight_html
        except Exception as exc:  # pragma: no cover - deployment dependency guard
            raise RuntimeError(f"账户详情直出 K 线组件不可用：{exc}") from exc
        if quote_generator.mt5 is None:
            raise RuntimeError("账户详情直出 K 线缺少 MT5 只读报价组件")

        canonical = [canonical_trade_row(row) for row in rows]
        trades = pd.DataFrame(canonical)
        for column in ("Open Time", "Close Time"):
            trades[column] = pd.to_datetime(trades[column], errors="coerce")
        trades = trades.dropna(subset=["Open Time", "Close Time", "Open Price", "Close Price"])
        if trades.empty:
            raise RuntimeError("账户订单缺少完整开平仓时间或价格，无法展示 K 线")
        trades = trades.sort_values("Open Time").reset_index(drop=True)
        stem = f"inline_{safe_stem_text(account)}_{recent_orders}"
        # The page never creates an output HTML or a durable job, but it may
        # refresh the bounded, reusable local M1 quote cache used for display.
        KLINE_OUT_DIR.mkdir(parents=True, exist_ok=True)
        quote_generator.OUT_DIR = KLINE_OUT_DIR
        registry = QuoteSourceRegistry.load(
            str(TRADE_KLINE_TERMINAL),
            os.environ.get("KDESK_KLINE_QUOTE_SOURCES") or None,
        )
        bars_by_symbol: dict[str, object] = {}
        mapping_by_symbol: dict[str, dict] = {}
        errors: list[str] = []
        for report_symbol, group in trades.groupby("Item", sort=True):
            group = group.sort_values("Open Time").reset_index(drop=True)
            source_candidates = registry.candidates(platform, server)
            if not source_candidates:
                errors.append(f"{report_symbol}: 未配置 {platform} / {server} 的同源报价")
                continue
            attempted: list[str] = []
            for provider, is_fallback in source_candidates:
                attempted.append(provider.id)
                if not quote_generator.mt5.initialize(path=str(provider.terminal), timeout=6000):
                    errors.append(
                        f"{report_symbol} / {provider.id}: M1 报价终端连接失败：{quote_generator.mt5.last_error()}"
                    )
                    continue
                try:
                    terminal_server = str(getattr(quote_generator.mt5.account_info(), "server", "") or "")
                    mapping, _align, sample = quote_generator.choose_by_m1_envelope(
                        str(report_symbol),
                        group,
                        aliases=provider.aliases,
                        fallback_source=is_fallback,
                        allowed_hour_offsets=provider.allowed_hour_offsets,
                    )
                    correction = provider.price_corrections.get(str(report_symbol))
                    if correction is None:
                        correction = provider.price_corrections.get(quote_generator.canonical_symbol(str(report_symbol)), 0.0)
                    mapping.update(
                        {
                            "provider": provider.id,
                            "provider_server": terminal_server,
                            "requested_server": server,
                            "tried_quote_sources": list(attempted),
                            "configured_price_correction": float(correction or 0.0),
                            "render_mode": "inline-direct",
                        }
                    )
                    bars = quote_generator.load_or_fetch_bars(
                        stem,
                        str(report_symbol),
                        mapping["mt5_symbol"],
                        mapping["time_mode"],
                        mapping["hour_delta"],
                        group,
                        provider_id=provider.id,
                    )
                    bars_by_symbol[str(report_symbol)] = quote_generator.apply_display_price_alignment(
                        str(report_symbol), bars, sample, mapping
                    )
                    mapping_by_symbol[str(report_symbol)] = mapping
                    break
                except Exception as exc:
                    errors.append(f"{report_symbol} / {provider.id}: {exc}")
                finally:
                    quote_generator.mt5.shutdown()

        if not bars_by_symbol:
            detail = "；".join(errors[-3:]) or "同源 M1 报价不可用"
            raise RuntimeError(f"账户详情 K 线暂时无法加载：{detail}")
        if errors:
            for report_symbol in set(trades["Item"].astype(str)) - set(bars_by_symbol):
                mapping_by_symbol.setdefault(
                    report_symbol,
                    {"report_symbol": report_symbol, "validation_status": "rejected", "failure": {"reason": "报价不可用"}},
                )
        return account_cache_set(
            cache_key,
            build_lightweight_html(account, stem, trades, bars_by_symbol, mapping_by_symbol),
        )


def _query_db_trades_uncached(
    account: str,
    platform: str = "",
    server: str = "",
    symbol: str = "",
    start: str = "",
    end: str = "",
    limit: int | None = 20000,
) -> list[dict]:
    account = normalize_text(account)
    if not account:
        return []
    if TRADE_DB_SOURCE in {"mysql", "auto"}:
        mysql_rows = query_mysql_trades(
            account,
            platform=platform,
            server=server,
            symbol=symbol,
            start=start,
            end=end,
            limit=limit,
        )
        if mysql_rows or TRADE_DB_SOURCE == "mysql":
            return mysql_rows
    where = ["account = ?"]
    args: list[object] = [account]
    if platform:
        where.append("platform = ?")
        args.append(platform)
    if server:
        where.append("server = ?")
        args.append(server)
    if symbol:
        where.append("symbol = ?")
        args.append(symbol)
    if start:
        where.append("coalesce(close_time, open_time) >= ?")
        args.append(start)
    if end:
        where.append("coalesce(open_time, close_time) <= ?")
        args.append(end)
    clause = " and ".join(where)
    sql = f"""
        select id, source_id, platform, server, account, ticket, open_time, close_time, type,
               volume, symbol, open_price, close_price, commission, taxes, swap, profit,
               holding_seconds, raw_json
        from trades
        where {clause}
        order by coalesce(open_time, close_time), id
        limit ?
    """
    sqlite_limit = int(limit) if limit is not None else 1000000
    args.append(sqlite_limit)
    with trade_db_connect() as conn:
        rows = [dict(row) for row in conn.execute(sql, args).fetchall()]
    seen: set[tuple] = set()
    deduped: list[dict] = []
    for row in rows:
        key = (
            row.get("platform"),
            row.get("server"),
            row.get("account"),
            row.get("ticket"),
            row.get("open_time"),
            row.get("close_time"),
            row.get("symbol"),
            row.get("type"),
            row.get("volume"),
        )
        if key in seen:
            continue
        seen.add(key)
        deduped.append(row)
    return deduped


def query_db_trades(
    account: str,
    platform: str = "",
    server: str = "",
    symbol: str = "",
    start: str = "",
    end: str = "",
    limit: int | None = 20000,
) -> list[dict]:
    cache_key = ("trades", normalize_text(account), platform, server, symbol, start, end, limit)
    cached = account_cache_get(cache_key)
    if cached is not None:
        return cached
    rows = _query_db_trades_uncached(account, platform, server, symbol, start, end, limit)
    return account_cache_set(cache_key, rows)


def account_history_limit(platform: object) -> int | None:
    """MT4 detail analytics must not silently truncate high-volume account history."""
    return None if normalize_text(platform).upper() == "MT4" else 50000


def account_trade_analysis(login: str, platform: str = "", server: str = "") -> dict:
    """Share the expensive full-history read and metric calculation across account panels."""
    login = normalize_text(login)
    platform = normalize_text(platform).upper()
    server = normalize_text(server)
    cache_key = ("trade-analysis", login, platform, server)
    cached = account_cache_get(cache_key)
    if cached is not None:
        return cached
    scope = {"platform": platform, "server": server}
    history_limit = account_history_limit(platform)
    with ThreadPoolExecutor(max_workers=2, thread_name_prefix="account-analysis") as executor:
        rows_future = executor.submit(query_db_trades, login, limit=history_limit, **scope)
        costs_future = executor.submit(query_mysql_trade_costs, login, **scope) if TRADE_DB_SOURCE in {"mysql", "auto"} else None
        rows = rows_future.result()
        uses_mysql = any(row.get("data_source") == "mysql" for row in rows)
        costs = costs_future.result() if uses_mysql and costs_future else None
    return account_cache_set(cache_key, {
        "rows": rows,
        "costs": costs,
        "metrics": trade_metrics(rows, costs),
        "usesMysql": uses_mysql,
        "historyLimit": history_limit,
    })


def query_copy_origin_source(source: dict, order_ids: list[str], batch_size: int = 500) -> list[dict]:
    numeric_ids = sorted({int(order_id) for order_id in order_ids if normalize_text(order_id).isdigit()})
    if not numeric_ids:
        return []
    requested = {str(value) for value in numeric_ids}
    raw_rows: list[dict] = []
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            for offset in range(0, len(numeric_ids), batch_size):
                batch = numeric_ids[offset:offset + batch_size]
                placeholders = ",".join(["%s"] * len(batch))
                if source.get("kind") == "mt5_deals":
                    sql = f"""
                        select Deal, Login, `Order`, PositionID, Time, Symbol, Comment
                        from `{source['schema']}`.`{source['table']}`
                        where PositionID in ({placeholders})
                        order by Time, Deal
                    """
                else:
                    sql = f"""
                        select TICKET, LOGIN, OPEN_TIME, SYMBOL, COMMENT
                        from `{source['schema']}`.`{source['table']}`
                        where TICKET in ({placeholders})
                        order by OPEN_TIME, TICKET
                    """
                cur.execute(sql, tuple(batch))
                raw_rows.extend(cur.fetchall())
    results = []
    for row in raw_rows:
        if source.get("kind") == "mt5_deals":
            candidate_ids = {normalize_text(row.get("PositionID"))}
            account = normalize_text(row.get("Login"))
            event_time = mysql_datetime_text(row.get("Time"))
            ticket = normalize_text(row.get("Order") or row.get("PositionID") or row.get("Deal"))
        else:
            candidate_ids = {normalize_text(row.get("TICKET"))}
            account = normalize_text(row.get("LOGIN"))
            event_time = mysql_datetime_text(row.get("OPEN_TIME"))
            ticket = normalize_text(row.get("TICKET"))
        matched_ids = sorted(requested & candidate_ids)
        if not matched_ids:
            continue
        results.append({
            "account": account,
            "platform": normalize_text(source.get("platform")),
            "server": normalize_text(source.get("server") or source.get("name")),
            "ticket": ticket,
            "matchedOrderIds": matched_ids,
            "time": event_time,
            "symbol": normalize_text(row.get("Symbol") or row.get("SYMBOL")),
            "comment": normalize_text(row.get("Comment") or row.get("COMMENT")),
        })
    return results


def _copy_source_windows(
    source_orders: list[dict],
    limit: int | None = None,
) -> tuple[list[tuple[datetime, datetime, set[str]]], bool]:
    valid_orders: list[tuple[datetime, str]] = []
    for order in source_orders:
        order_id = normalize_text(order.get("orderId"))
        event_time = parse_trade_time(order.get("time"))
        if order_id.isdigit() and event_time:
            valid_orders.append((event_time, order_id))
    valid_orders.sort(key=lambda item: (item[0], int(item[1])))
    truncated = limit is not None and len(valid_orders) > limit
    selected_orders = valid_orders if limit is None else valid_orders[:limit]
    grouped: dict[str, list[tuple[datetime, str]]] = defaultdict(list)
    for event_time, order_id in selected_orders:
        grouped[event_time.strftime("%Y-%m-%d %H")].append((event_time, order_id))
    windows = []
    for values in grouped.values():
        times = [item[0] for item in values]
        windows.append((min(times) - timedelta(minutes=5), max(times) + timedelta(minutes=5), {item[1] for item in values}))
    windows.sort(key=lambda item: item[0])
    return windows, truncated


def _copy_follower_money_meta(cur, source: dict, accounts: set[str]) -> dict[str, dict]:
    metadata: dict[str, dict] = {}
    numeric_accounts = sorted({int(account) for account in accounts if normalize_text(account).isdigit()})
    if source.get("kind") == "mt5_deals":
        default_currency = normalize_text(source.get("default_currency")).upper() or "USD"
        for offset in range(0, len(numeric_accounts), 500):
            batch = numeric_accounts[offset:offset + 500]
            placeholders = ",".join(["%s"] * len(batch))
            cur.execute(
                f"select Login, `Group` as AccountGroup from `{source['schema']}`.`mt5_users_view` "
                f"where Login in ({placeholders})",
                tuple(batch),
            )
            for row in cur.fetchall():
                account = normalize_text(row.get("Login"))
                account_group = row.get("AccountGroup")
                metadata[account] = account_money_meta(
                    mt5_group_currency(account_group, default_currency),
                    account_group,
                    source.get("name"),
                    "mt5_users_view.Group",
                )
        for account in numeric_accounts:
            metadata.setdefault(
                str(account),
                account_money_meta(
                    default_currency,
                    source_name=source.get("name"),
                    currency_source="source.default_currency",
                ),
            )
        return metadata
    for offset in range(0, len(numeric_accounts), 300):
        batch = numeric_accounts[offset:offset + 300]
        placeholders = ",".join(["%s"] * len(batch))
        cur.execute(
            f"select LOGIN, CURRENCY, `GROUP` as AccountGroup from `{source['schema']}`.`mt4_users_view` "
            f"where LOGIN in ({placeholders})",
            tuple(batch),
        )
        for row in cur.fetchall():
            account = normalize_text(row.get("LOGIN"))
            metadata[account] = account_money_meta(row.get("CURRENCY"), row.get("AccountGroup"), source.get("name"))
    return metadata


def query_copy_followers_source(
    source: dict,
    source_orders: list[dict],
    *,
    copy_channels: list[str] | None = None,
    start: str = "",
    end: str = "",
    comment_batch_size: int = 500,
    position_batch_size: int = 5000,
) -> dict:
    range_start = parse_trade_time(start)
    range_end = parse_trade_time(end)
    requested_ids = {
        normalize_text(order.get("orderId"))
        for order in source_orders
        if normalize_text(order.get("orderId")).isdigit()
    }
    channels = sorted({
        normalize_text(channel).upper()
        for channel in (copy_channels or [])
        if re.fullmatch(r"(?i)CPT-[A-Z0-9]+", normalize_text(channel))
    })
    if not requested_ids:
        return {
            "rows": [], "sourceOrdersScanned": 0, "sourceOrdersTruncated": False,
            "candidateRowsTruncated": False, "queryStrategy": "exact-comment",
        }
    matches: dict[tuple[str, str], dict] = {}
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            if source.get("kind") == "mt5_deals" and channels:
                exact_comments = sorted(
                    f"{channel}#{order_id}"
                    for channel in channels
                    for order_id in requested_ids
                )
                candidate_rows: list[dict] = []
                for offset in range(0, len(exact_comments), comment_batch_size):
                    batch = exact_comments[offset:offset + comment_batch_size]
                    placeholders = ",".join(["%s"] * len(batch))
                    time_clause = ""
                    time_args: list[object] = []
                    if start:
                        time_clause += " and Time >= %s"
                        time_args.append(start)
                    if end:
                        time_clause += " and Time <= %s"
                        time_args.append(end)
                    sql = f"""
                        select Deal, Login, `Order`, PositionID, Time, TimeMsc, Action, Entry,
                               Symbol, Volume, VolumeExt, Profit, Commission, Storage, Fee, Comment
                        from `{source['schema']}`.`{source['table']}`
                        where Comment in ({placeholders}) and Action in (0,1) and Entry in (0,2){time_clause}
                        order by TimeMsc, Deal
                    """
                    cur.execute(sql, (*batch, *time_args))
                    candidate_rows.extend(cur.fetchall())
            else:
                candidate_rows = []
                windows, _ = _copy_source_windows(source_orders)
                for source_start, source_end, window_ids in windows:
                    window_start = max(source_start, range_start) if range_start else source_start
                    window_end = min(source_end, range_end) if range_end else source_end
                    if window_start > window_end:
                        continue
                    if channels:
                        comments = sorted(
                            f"{channel}#{order_id}"
                            for channel in channels
                            for order_id in window_ids
                        )
                        comment_clause = f"COMMENT in ({','.join(['%s'] * len(comments))})"
                        comment_args: tuple[object, ...] = tuple(comments)
                    else:
                        comment_clause = "COMMENT like %s"
                        comment_args = ("%CPT-%",)
                    sql = f"""
                        select LOGIN, TICKET, OPEN_TIME, CLOSE_TIME, SYMBOL, VOLUME,
                               PROFIT, COMMISSION, SWAPS, TAXES, COMMENT
                        from `{source['schema']}`.`{source['table']}`
                        where OPEN_TIME >= %s and OPEN_TIME <= %s and CMD in (0,1) and {comment_clause}
                        order by OPEN_TIME, TICKET
                    """
                    cur.execute(sql, (
                        window_start.strftime("%Y-%m-%d %H:%M:%S"),
                        window_end.strftime("%Y-%m-%d %H:%M:%S"),
                        *comment_args,
                    ))
                    candidate_rows.extend(cur.fetchall())

            for row in candidate_rows:
                comment = normalize_text(row.get("Comment") or row.get("COMMENT"))
                matched_ids = sorted(requested_ids.intersection(copy_trade_order_ids({"comment": comment})), key=int)
                if not matched_ids:
                    continue
                account = normalize_text(row.get("Login") or row.get("LOGIN"))
                ticket = normalize_text(row.get("PositionID") or row.get("TICKET"))
                if not account or not ticket:
                    continue
                key = (account, ticket)
                item = matches.setdefault(key, {"seed": row, "matchedSourceOrderIds": set()})
                item["matchedSourceOrderIds"].update(matched_ids)

            accounts = {key[0] for key in matches}
            metadata = _copy_follower_money_meta(cur, source, accounts)
            result_rows: list[dict] = []
            if source.get("kind") == "mt5_deals":
                position_ids = sorted({int(key[1]) for key in matches if key[1].isdigit()})
                for offset in range(0, len(position_ids), position_batch_size):
                    batch = position_ids[offset:offset + position_batch_size]
                    placeholders = ",".join(["%s"] * len(batch))
                    cur.execute(
                        f"""
                            select Login, PositionID, min(Symbol) as Symbol,
                                   min(case when Entry in (0,2) then Time end) as OpenTime,
                                   max(case when Entry in (1,2,3) then Time end) as CloseTime,
                                   sum(case when Entry in (0,2) then Volume else 0 end) as OpenVolume,
                                   sum(Profit) as GrossProfit,
                                   sum(Commission) + sum(Fee) as CommissionFee,
                                   sum(Storage) as Swap
                            from `{source['schema']}`.`{source['table']}`
                            where PositionID in ({placeholders}) and Action in (0,1)
                            group by Login, PositionID
                        """,
                        tuple(batch),
                    )
                    for aggregate in cur.fetchall():
                        account = normalize_text(aggregate.get("Login"))
                        position = normalize_text(aggregate.get("PositionID"))
                        match = matches.get((account, position))
                        if not match:
                            continue
                        meta = metadata.get(account) or account_money_meta(source_name=source.get("name"))
                        scale = numeric_value(meta.get("moneyScale")) or 1.0
                        gross = numeric_value(aggregate.get("GrossProfit")) * scale
                        commission = numeric_value(aggregate.get("CommissionFee")) * scale
                        swap = numeric_value(aggregate.get("Swap")) * scale
                        symbol = normalize_text(aggregate.get("Symbol") or match["seed"].get("Symbol"))
                        result_rows.append({
                            "account": account, "platform": normalize_text(source.get("platform")),
                            "server": normalize_text(source.get("server") or source.get("name")),
                            "orders": 1, "ticket": position, "tickets": [position],
                            "matchedSourceOrderIds": sorted(match["matchedSourceOrderIds"], key=int),
                            "openTime": mysql_datetime_text(aggregate.get("OpenTime") or match["seed"].get("Time")),
                            "closeTime": mysql_datetime_text(aggregate.get("CloseTime")),
                            "symbol": symbol, "symbols": [symbol] if symbol else [],
                            "volume": rounded(normalize_mt5_volume(aggregate.get("OpenVolume")), 4),
                            "grossProfit": rounded(gross, 6),
                            "commission": rounded(commission, 6), "swap": rounded(swap, 6), "taxes": 0.0,
                            "netProfit": rounded(gross + commission + swap, 6),
                            "currency": normalize_text(meta.get("currency")), "displayCurrency": normalize_text(meta.get("displayCurrency")),
                            "isCentAccount": bool(meta.get("isCentAccount")),
                        })
            else:
                for key, match in matches.items():
                    row = match["seed"]
                    account, ticket = key
                    meta = metadata.get(account) or account_money_meta(source_name=source.get("name"))
                    scale = numeric_value(meta.get("moneyScale")) or 1.0
                    gross = numeric_value(row.get("PROFIT")) * scale
                    commission = numeric_value(row.get("COMMISSION")) * scale
                    swap = numeric_value(row.get("SWAPS")) * scale
                    taxes = numeric_value(row.get("TAXES")) * scale
                    result_rows.append({
                        "account": account, "platform": normalize_text(source.get("platform")),
                        "server": normalize_text(source.get("server") or source.get("name")), "ticket": ticket,
                        "matchedSourceOrderIds": sorted(match["matchedSourceOrderIds"], key=int),
                        "openTime": mysql_datetime_text(row.get("OPEN_TIME")), "closeTime": mysql_datetime_text(row.get("CLOSE_TIME")),
                        "symbol": normalize_text(row.get("SYMBOL")), "volume": rounded(normalize_mt4_volume(row.get("VOLUME")), 4),
                        "grossProfit": rounded(gross, 6), "commission": rounded(commission, 6), "swap": rounded(swap, 6),
                        "taxes": rounded(taxes, 6), "netProfit": rounded(gross + commission + swap + taxes, 6),
                        "currency": normalize_text(meta.get("currency")), "displayCurrency": normalize_text(meta.get("displayCurrency")),
                        "isCentAccount": bool(meta.get("isCentAccount")),
                    })
    if range_start or range_end:
        result_rows = [
            row for row in result_rows
            if (opened := parse_trade_time(row.get("openTime")))
            and (not range_start or opened >= range_start)
            and (not range_end or opened <= range_end)
        ]
    result_rows.sort(key=lambda row: (
        row["account"],
        row.get("openTime") or "",
        row.get("ticket") or next(iter(row.get("tickets") or []), ""),
    ))
    return {
        "rows": result_rows,
        "sourceOrdersScanned": len(requested_ids),
        "sourceOrdersTruncated": False,
        "candidateRowsTruncated": False,
        "queryStrategy": "exact-comment" if channels else "time-window-fallback",
    }


def summarize_copy_followers(rows: list[dict], current_account: str = "") -> tuple[list[dict], dict]:
    grouped: dict[tuple[str, str, str], dict] = {}
    for row in rows:
        key = (normalize_text(row.get("account")), normalize_text(row.get("platform")), normalize_text(row.get("server")))
        if not key[0]:
            continue
        item = grouped.setdefault(key, {
            "account": key[0], "platform": key[1], "server": key[2], "orders": 0, "volume": 0.0,
            "grossProfit": 0.0, "commission": 0.0, "swap": 0.0, "taxes": 0.0, "netProfit": 0.0,
            "sourceOrderIds": set(), "symbols": set(), "tickets": [], "firstTime": "", "lastTime": "",
            "currency": normalize_text(row.get("currency")), "displayCurrency": normalize_text(row.get("displayCurrency")),
            "isCentAccount": bool(row.get("isCentAccount")),
        })
        item["orders"] += max(1, mysql_int(row.get("orders")))
        for field in ("volume", "grossProfit", "commission", "swap", "taxes", "netProfit"):
            item[field] += numeric_value(row.get(field))
        item["sourceOrderIds"].update(normalize_text(value) for value in row.get("matchedSourceOrderIds", []) if normalize_text(value))
        row_symbols = row.get("symbols") or [row.get("symbol")]
        item["symbols"].update(normalize_text(value) for value in row_symbols if normalize_text(value))
        row_tickets = row.get("tickets") or [row.get("ticket")]
        for ticket in row_tickets:
            if normalize_text(ticket) and len(item["tickets"]) < 12:
                item["tickets"].append(normalize_text(ticket))
        opened = normalize_text(row.get("firstTime") or row.get("openTime"))
        closed = normalize_text(row.get("lastTime") or row.get("closeTime")) or opened
        if opened:
            item["firstTime"] = min(filter(None, [item["firstTime"], opened]), default=opened)
        if closed:
            item["lastTime"] = max(item["lastTime"], closed)
    followers = []
    for item in grouped.values():
        for field in ("volume", "grossProfit", "commission", "swap", "taxes", "netProfit"):
            item[field] = rounded(item[field], 4 if field == "volume" else 2)
        item["sourceOrderIds"] = sorted(item["sourceOrderIds"], key=lambda value: int(value) if value.isdigit() else value)
        item["matchedSourceOrders"] = len(item["sourceOrderIds"])
        item["symbols"] = sorted(item["symbols"])
        item["isCurrentAccount"] = item["account"] == normalize_text(current_account)
        followers.append(item)
    followers.sort(key=lambda item: (not item["isCurrentAccount"], -item["orders"], item["account"]))
    currencies = sorted({normalize_text(item.get("displayCurrency") or item.get("currency")) for item in followers if normalize_text(item.get("displayCurrency") or item.get("currency"))})
    summary = {
        "accounts": len(followers),
        "profitableAccounts": sum(1 for item in followers if numeric_value(item.get("netProfit")) > 0),
        "losingAccounts": sum(1 for item in followers if numeric_value(item.get("netProfit")) < 0),
        "orders": sum(mysql_int(item.get("orders")) for item in followers),
        "volume": rounded(sum(numeric_value(item.get("volume")) for item in followers), 4),
        "grossProfit": rounded(sum(numeric_value(item.get("grossProfit")) for item in followers)),
        "commission": rounded(sum(numeric_value(item.get("commission")) for item in followers)),
        "swap": rounded(sum(numeric_value(item.get("swap")) for item in followers)),
        "taxes": rounded(sum(numeric_value(item.get("taxes")) for item in followers)),
        "netProfit": rounded(sum(numeric_value(item.get("netProfit")) for item in followers)),
        "currencies": currencies,
        "currency": currencies[0] if len(currencies) == 1 else "多币种" if currencies else "",
    }
    return followers, summary


def account_copy_origins_payload(login: str, filters: dict | None = None) -> dict:
    filters = filters or {}
    relationship_mode = normalize_text(filters.get("_relationship")) == "1"
    login = normalize_text(login)
    if not login or not re.fullmatch(r"\d+", login):
        raise ValueError("账号格式无效")
    platform = normalize_text(filters.get("platform")).upper()
    server = normalize_text(filters.get("server"))
    start = normalize_text(filters.get("start"))
    end = normalize_text(filters.get("end"))
    scope_start = parse_trade_time(start)
    scope_end = parse_trade_time(end)
    if start and end and scope_start and scope_end and scope_start > scope_end:
        raise ValueError("跟单开始时间不能晚于结束时间")
    cache_key = (
        "copy-origins",
        login,
        *(normalize_text(filters.get(key)) for key in ("platform", "server", "symbol", "start", "end")),
    )
    if not relationship_mode:
        cached = account_cache_get(cache_key)
        if cached is not None:
            return cached
    rows = query_db_trades(login, platform=platform, server=server, start=start, end=end, limit=50000)
    # A broad overlap predicate can fill the safety page with older positions before the selected
    # opening-time window. Retry without the analytical cap only when that page is saturated.
    if (scope_start or scope_end) and len(rows) >= 50000:
        rows = query_db_trades(login, platform=platform, server=server, start=start, end=end, limit=None)
    if scope_start or scope_end:
        rows = [
            row for row in rows
            if (opened := parse_trade_time(row.get("open_time")))
            and (not scope_start or opened >= scope_start)
            and (not scope_end or opened <= scope_end)
        ]
    order_ids = list(dict.fromkeys(
        order_id
        for row in rows
        for order_id in copy_trade_order_ids(row)
    ))
    if not order_ids:
        payload = {
            "ok": True, "account": login, "detected": False, "orderIds": [], "origins": [], "primaryOrigin": None,
        }
        return payload if relationship_mode else account_cache_set(cache_key, payload)
    sources = [source for source in MYSQL_SOURCES if source_allowed(source, platform=platform, server=server)]
    candidates: list[dict] = []
    errors: list[str] = []
    with ThreadPoolExecutor(max_workers=min(8, len(sources) or 1), thread_name_prefix="copy-origin") as executor:
        futures = {executor.submit(query_copy_origin_source, source, order_ids): source for source in sources}
        for future in as_completed(futures):
            source = futures[future]
            try:
                candidates.extend(future.result())
            except Exception as exc:
                errors.append(f"{source.get('name')}: {exc}")
    candidates = [row for row in candidates if row.get("account") and row.get("account") != login]
    grouped: dict[tuple[str, str, str], dict] = {}
    for row in candidates:
        key = (row["account"], row["platform"], row["server"])
        item = grouped.setdefault(key, {
            "account": row["account"], "platform": row["platform"], "server": row["server"],
            "matchedOrderIds": set(), "symbols": set(), "sourceOrders": {}, "firstTime": "", "lastTime": "",
        })
        item["matchedOrderIds"].update(row["matchedOrderIds"])
        if row.get("symbol"):
            item["symbols"].add(row["symbol"])
        event_time = normalize_text(row.get("time"))
        if event_time:
            item["firstTime"] = min(filter(None, [item["firstTime"], event_time]), default=event_time)
            item["lastTime"] = max(item["lastTime"], event_time)
        for order_id in row["matchedOrderIds"]:
            detail = item["sourceOrders"].setdefault(order_id, {
                "orderId": order_id, "ticket": normalize_text(row.get("ticket")),
                "time": event_time, "symbol": normalize_text(row.get("symbol")),
            })
            if event_time and (not detail.get("time") or event_time < detail["time"]):
                detail.update({"ticket": normalize_text(row.get("ticket")), "time": event_time, "symbol": normalize_text(row.get("symbol"))})
    origins = []
    for item in grouped.values():
        matched = sorted(item.pop("matchedOrderIds"), key=lambda value: int(value))
        symbols = sorted(item.pop("symbols"))
        source_orders = sorted(item.pop("sourceOrders").values(), key=lambda row: (row.get("time") or "", int(row["orderId"])))
        origins.append({
            **item,
            "matchedOrders": len(matched),
            "matchedOrderIds": matched,
            "sampleOrderIds": matched[:12],
            "sourceOrders": source_orders,
            "symbols": symbols,
        })
    origins.sort(key=lambda item: (-item["matchedOrders"], item["firstTime"], item["account"]))
    copy_rows = [row for row in rows if copy_trade_order_ids(row)]
    total_orders = len(rows)
    total_volume = sum(numeric_value(row.get("volume")) for row in rows if is_chartable_trade(row))
    copy_volume = sum(numeric_value(row.get("volume")) for row in copy_rows if is_chartable_trade(row))
    assigned: dict[int, list[dict]] = defaultdict(list)
    unresolved_rows: list[dict] = []
    for row in copy_rows:
        row_order_ids = set(copy_trade_order_ids(row))
        origin_index = next(
            (index for index, origin in enumerate(origins) if row_order_ids.intersection(origin.get("matchedOrderIds", []))),
            None,
        )
        if origin_index is None:
            unresolved_rows.append(row)
        else:
            assigned[origin_index].append(row)
    follower_requests: list[tuple[dict, dict, list[str]]] = []
    for index, origin in enumerate(origins):
        source_rows = assigned.get(index, [])
        copy_channels = sorted({
            channel
            for row in source_rows
            for channel in copy_trade_channels(row)
        })
        stats = automation_stats(source_rows, total_orders, total_volume)
        source_volume = numeric_value(stats.get("volume"))
        origin.update({
            **stats,
            "copyChannels": copy_channels,
            "copyOrderRatio": rounded(len(source_rows) / len(copy_rows) * 100) if copy_rows else 0,
            "copyVolumeRatio": rounded(source_volume / copy_volume * 100) if copy_volume else 0,
        })
        source = next((item for item in sources if normalize_text(item.get("platform")) == origin["platform"] and normalize_text(item.get("server") or item.get("name")) == origin["server"]), None)
        if source:
            follower_requests.append((origin, source, copy_channels))
    if follower_requests:
        follower_scope = {key: value for key, value in {"start": start, "end": end}.items() if value}
        with ThreadPoolExecutor(
            max_workers=min(4, len(follower_requests)),
            thread_name_prefix="copy-followers",
        ) as executor:
            futures = {
                executor.submit(
                    query_copy_followers_source,
                    source,
                    origin.get("sourceOrders", []),
                    copy_channels=copy_channels,
                    **follower_scope,
                ): (origin, source)
                for origin, source, copy_channels in follower_requests
            }
            for future in as_completed(futures):
                origin, _source = futures[future]
                try:
                    discovery = future.result()
                    followers, follower_summary = summarize_copy_followers(discovery.get("rows", []), login)
                    origin.update({
                        "followers": followers,
                        "followerOrders": discovery.get("rows", []),
                        "followerSummary": follower_summary,
                        "followerDiscovery": {key: value for key, value in discovery.items() if key != "rows"},
                    })
                except Exception as exc:
                    origin.update({"followers": [], "followerOrders": [], "followerSummary": {}, "followerDiscovery": {"error": str(exc)}})
                    errors.append(f"{origin['server']} 跟单人员查询: {exc}")
    payload = {
        "ok": True,
        "account": login,
        "detected": True,
        "totalOrders": total_orders,
        "totalVolume": rounded(total_volume, 4),
        "copyOrders": len(copy_rows),
        "copyVolume": rounded(copy_volume, 4),
        "mappedCopyOrders": sum(len(source_rows) for source_rows in assigned.values()),
        "unmappedCopyOrders": len(unresolved_rows),
        "searchedOrders": len(order_ids),
        "matchedOrders": sum(item["matchedOrders"] for item in origins),
        "orderIds": order_ids[:30],
        "origins": origins,
        "primaryOrigin": origins[0] if origins else None,
        "errors": errors[:5],
        "refreshedAt": now_text(),
    }
    return payload if relationship_mode else account_cache_set(cache_key, payload)






















def account_copy_group_profit_payload(login: str, filters: dict | None = None) -> dict:
    filters = filters or {}
    relationship_mode = normalize_text(filters.get("_relationship")) == "1"
    cache_key = (
        "copy-group-profit",
        normalize_text(login),
        *(normalize_text(filters.get(key)) for key in ("platform", "server", "symbol", "start", "end")),
    )
    if not relationship_mode:
        cached = account_cache_get(cache_key)
        if cached is not None:
            return cached
    service = SignalCopyGroupService(sys.modules[__name__])
    payload = service.payload(login, filters)
    return payload if relationship_mode else account_cache_set(cache_key, payload)


def account_ea_comment_profit_payload(login: str, filters: dict | None = None) -> dict:
    filters = filters or {}
    relationship_mode = normalize_text(filters.get("_relationship")) == "1"
    cache_key = (
        "ea-comment-profit",
        normalize_text(login),
        *(normalize_text(filters.get(key)) for key in ("platform", "server", "symbol", "start", "end")),
    )
    if not relationship_mode:
        cached = account_cache_get(cache_key)
        if cached is not None:
            return cached
    service = EaCommentGroupService(sys.modules[__name__])
    payload = service.payload(login, filters)
    return payload if relationship_mode else account_cache_set(cache_key, payload)


def trade_pnl_values(row: dict) -> tuple[float, float]:
    """Return gross and trade-level net P/L for one exported order row."""
    gross = numeric_value(row.get("profit"))
    net = gross + sum(numeric_value(row.get(field)) for field in ("commission", "fee", "swap", "taxes"))
    return gross, net


def automation_stats(rows: list[dict], total_orders: int, total_volume: float) -> dict:
    volumes = [numeric_value(row.get("volume")) for row in rows if is_chartable_trade(row)]
    gross = sum(trade_pnl_values(row)[0] for row in rows)
    net = sum(trade_pnl_values(row)[1] for row in rows)
    volume = sum(volumes)
    order_count = len(rows)
    return {
        "orders": order_count,
        "orderRatio": rounded(order_count / total_orders * 100) if total_orders else 0,
        "volume": rounded(volume, 4),
        "volumeRatio": rounded(volume / total_volume * 100) if total_volume else 0,
        "grossProfit": rounded(gross),
        "netProfit": rounded(net),
    }


def account_signal_copy_seeds(login: str, platform: str = "", server: str = "") -> tuple[list[dict], list[str]]:
    cache_key = ("signal-copy-seeds", normalize_text(login), normalize_text(platform).upper(), normalize_text(server))
    cached = account_cache_get(cache_key)
    if cached is not None:
        return cached
    sources = [source for source in MYSQL_SOURCES if source_allowed(source, platform=platform, server=server)]
    service = SignalCopyGroupService(sys.modules[__name__])
    seeds: list[dict] = []
    errors: list[str] = []
    with ThreadPoolExecutor(max_workers=min(4, len(sources) or 1), thread_name_prefix="signal-copy-seed") as executor:
        futures = {executor.submit(service.query_seed, source, login): source for source in sources}
        for future in as_completed(futures):
            source = futures[future]
            try:
                seed = future.result()
                if seed:
                    seeds.append(seed)
            except Exception as exc:
                errors.append(f"{source.get('name')}: {exc}")
    seeds.sort(key=lambda item: (item.get("platform", ""), item.get("server", ""), item.get("signalId", "")))
    return account_cache_set(cache_key, (seeds, errors[:5]))


def account_automation_payload(login: str, filters: dict | None = None) -> dict:
    """Summarize copy-trading and EA activity without exposing raw order comments."""
    filters = filters or {}
    login = normalize_text(login)
    if not login or not re.fullmatch(r"\d+", login):
        raise ValueError("账号格式无效")
    active_filters = {key: normalize_text(filters.get(key)) for key in ("platform", "server", "symbol", "start", "end")}
    cache_key = ("automation-analysis", login, *(active_filters[key] for key in active_filters))
    cached = account_cache_get(cache_key)
    if cached is not None:
        return {"ok": True, **cached}

    rows = query_db_trades(login, limit=account_history_limit(active_filters["platform"]), **active_filters)
    total_orders = len(rows)
    total_volume = sum(numeric_value(row.get("volume")) for row in rows if is_chartable_trade(row))
    signal_seeds, signal_errors = account_signal_copy_seeds(
        login,
        platform=active_filters["platform"],
        server=active_filters["server"],
    )
    signal_seed_map = {
        (normalize_text(seed.get("platform")).upper(), normalize_text(seed.get("server"))): seed
        for seed in signal_seeds
    }

    def signal_seed_for_row(row: dict) -> dict | None:
        return signal_seed_map.get((normalize_text(row.get("platform")).upper(), normalize_text(row.get("server"))))

    copy_rows = [row for row in rows if is_copy_trade(row) or signal_seed_for_row(row)]
    copy_row_ids = {id(row) for row in copy_rows}
    # Signal accounts often store source ticket numbers in MT4 Magic. Those are
    # copy identifiers, not EA IDs, and must not inflate the EA group count.
    ea_rows = [row for row in rows if id(row) not in copy_row_ids and is_ea_trade(row)]

    copy_origins: list[dict] = []
    origin_errors: list[str] = list(signal_errors)
    if copy_rows:
        # Origin lookup is already cached and restricted to the selected platform/server.
        origin_payload = account_copy_origins_payload(
            login,
            {key: active_filters[key] for key in ("platform", "server")},
        )
        origin_errors.extend(origin_payload.get("errors", []))
        origins = origin_payload.get("origins", [])
        assigned: dict[int, list[dict]] = defaultdict(list)
        unassigned: list[dict] = []
        for row in copy_rows:
            ids = set(copy_trade_order_ids(row))
            match_index = next(
                (index for index, origin in enumerate(origins) if ids.intersection(origin.get("matchedOrderIds", []))),
                None,
            )
            if match_index is None:
                unassigned.append(row)
            else:
                assigned[match_index].append(row)
        for index, origin in enumerate(origins):
            source_rows = assigned.get(index, [])
            if not source_rows:
                continue
            stats = automation_stats(source_rows, total_orders, total_volume)
            stats.update({
                "account": normalize_text(origin.get("account")),
                "platform": normalize_text(origin.get("platform")),
                "server": normalize_text(origin.get("server")),
                "symbols": sorted({normalize_text(row.get("symbol")) for row in source_rows if normalize_text(row.get("symbol"))}),
                "sampleOrderIds": origin.get("sampleOrderIds", [])[:12],
            })
            copy_origins.append(stats)
        signal_rows: dict[tuple[str, str, str], list[dict]] = defaultdict(list)
        unresolved: list[dict] = []
        for row in unassigned:
            seed = signal_seed_for_row(row)
            if seed:
                key = (normalize_text(seed.get("signalTag")), normalize_text(seed.get("platform")), normalize_text(seed.get("server")))
                signal_rows[key].append(row)
            else:
                unresolved.append(row)
        for (signal_tag, seed_platform, seed_server), group_rows in signal_rows.items():
            stats = automation_stats(group_rows, total_orders, total_volume)
            stats.update({
                "account": signal_tag or "Signal 跟单",
                "platform": seed_platform,
                "server": seed_server,
                "symbols": sorted({normalize_text(row.get("symbol")) for row in group_rows if normalize_text(row.get("symbol"))}),
                "sampleOrderIds": [],
                "sourceType": "signal",
                "unresolved": True,
            })
            copy_origins.append(stats)
        if unresolved:
            stats = automation_stats(unresolved, total_orders, total_volume)
            stats.update({
                "account": "来源未解析",
                "platform": active_filters["platform"],
                "server": active_filters["server"],
                "symbols": sorted({normalize_text(row.get("symbol")) for row in unresolved if normalize_text(row.get("symbol"))}),
                "sampleOrderIds": [],
                "unresolved": True,
            })
            copy_origins.append(stats)
    copy_origins.sort(key=lambda item: (-item["orders"], item["account"]))

    ea_groups: dict[tuple[str, str, str], list[dict]] = defaultdict(list)
    for row in ea_rows:
        expert_id = normalize_text(row.get("expert_id"))
        if not expert_id or expert_id == "0":
            expert_id = "EA（未标记ID）"
        ea_groups[(expert_id, normalize_text(row.get("platform")), normalize_text(row.get("server")))].append(row)
    ea_breakdown = []
    for (expert_id, platform, server), group_rows in ea_groups.items():
        stats = automation_stats(group_rows, total_orders, total_volume)
        stats.update({
            "expertId": expert_id,
            "platform": platform,
            "server": server,
            "symbols": sorted({normalize_text(row.get("symbol")) for row in group_rows if normalize_text(row.get("symbol"))}),
            "comments": sorted({normalize_text(row.get("comment")) for row in group_rows if normalize_text(row.get("comment"))})[:3],
        })
        ea_breakdown.append(stats)
    ea_breakdown.sort(key=lambda item: (-item["orders"], item["expertId"], item["platform"], item["server"]))

    payload = {
        "account": login,
        "totalOrders": total_orders,
        "totalVolume": rounded(total_volume, 4),
        "copy": {
            "detected": bool(copy_rows),
            **automation_stats(copy_rows, total_orders, total_volume),
            "origins": copy_origins,
            "errors": origin_errors[:5],
        },
        "ea": {
            "detected": bool(ea_rows),
            **automation_stats(ea_rows, total_orders, total_volume),
            "groups": ea_breakdown,
        },
        "refreshedAt": now_text(),
    }
    account_cache_set(cache_key, payload)
    return {"ok": True, **payload}


def trade_summary_for_account(account: str, filters: dict | None = None) -> dict:
    filters = filters or {}
    account = normalize_text(account)
    if not account:
        return {"exists": False, "error": "当前记录没有账号，无法查询数据库订单"}
    try:
        all_rows = query_db_trades(account, limit=50000)
    except Exception as exc:
        return {"exists": False, "error": str(exc), "dbSource": TRADE_DB_SOURCE, "dbPath": str(TRADE_DB_PATH)}
    chartable_all = [row for row in all_rows if is_chartable_trade(row)]
    filtered_rows = query_db_trades(
        account,
        platform=normalize_text(filters.get("platform")),
        server=normalize_text(filters.get("server")),
        symbol=normalize_text(filters.get("symbol")),
        start=normalize_text(filters.get("start")),
        end=normalize_text(filters.get("end")),
        limit=50000,
    )
    chartable_filtered = [row for row in filtered_rows if is_chartable_trade(row)]

    def range_for(rows: list[dict]) -> tuple[str, str]:
        starts = [parse_trade_time(row.get("open_time")) for row in rows]
        ends = [parse_trade_time(row.get("close_time") or row.get("open_time")) for row in rows]
        starts = [item for item in starts if item]
        ends = [item for item in ends if item]
        return trade_time_text(min(starts) if starts else None), trade_time_text(max(ends) if ends else None)

    platforms = sorted({normalize_text(row.get("platform")) for row in all_rows if normalize_text(row.get("platform"))})
    servers = sorted({normalize_text(row.get("server")) for row in all_rows if normalize_text(row.get("server"))})
    symbols = sorted({normalize_text(row.get("symbol")) for row in chartable_all if normalize_text(row.get("symbol"))})
    latest_row = max(all_rows, key=lambda row: normalize_text(row.get("open_time") or row.get("close_time"))) if all_rows else None
    first_time, last_time = range_for(chartable_all)
    filtered_first, filtered_last = range_for(chartable_filtered)
    return {
        "exists": bool(all_rows),
        "dbSource": "mysql" if any(row.get("data_source") == "mysql" for row in all_rows) else ("sqlite" if all_rows else ""),
        "searchedSources": [source["name"] for source in MYSQL_SOURCES] + ([str(TRADE_DB_PATH)] if TRADE_DB_PATH.exists() else []),
        "dbPath": str(TRADE_DB_PATH),
        "account": account,
        "totalRows": len(all_rows),
        "chartableRows": len(chartable_all),
        "filteredRows": len(filtered_rows),
        "filteredChartableRows": len(chartable_filtered),
        "firstTime": first_time,
        "lastTime": last_time,
        "filteredFirstTime": filtered_first,
        "filteredLastTime": filtered_last,
        "platforms": [{"value": item, "label": display_unknown(item)} for item in platforms],
        "servers": [{"value": item, "label": display_unknown(item)} for item in servers],
        "symbols": symbols,
        "latestSource": {
            "platform": normalize_text(latest_row.get("platform")) if latest_row else "",
            "server": normalize_text(latest_row.get("server")) if latest_row else "",
            "label": " / ".join(
                item for item in [
                    normalize_text(latest_row.get("platform")) if latest_row else "",
                    display_unknown(normalize_text(latest_row.get("server"))) if latest_row else "",
                ] if item
            ),
        },
    }


def numeric_value(value: object) -> float:
    try:
        return float(value or 0)
    except (TypeError, ValueError):
        return 0.0


def rounded(value: float, digits: int = 2) -> float:
    return round(float(value or 0), digits)


def trade_metrics(rows: list[dict], cost_summary: dict | None = None) -> dict:
    chartable = [row for row in rows if is_chartable_trade(row)]
    gross_profits = [numeric_value(row.get("profit")) for row in rows]
    net_profits = [
        numeric_value(row.get("profit"))
        + numeric_value(row.get("commission"))
        + numeric_value(row.get("fee"))
        + numeric_value(row.get("swap"))
        + numeric_value(row.get("taxes"))
        for row in rows
    ]
    volumes = [numeric_value(row.get("volume")) for row in chartable]
    holding_seconds = []
    for row in chartable:
        value = numeric_value(row.get("holding_seconds"))
        if value <= 0:
            opened = parse_trade_time(row.get("open_time"))
            closed = parse_trade_time(row.get("close_time"))
            value = (closed - opened).total_seconds() if opened and closed else 0
        if value >= 0:
            holding_seconds.append(value)

    starts = [parse_trade_time(row.get("open_time")) for row in chartable]
    ends = [parse_trade_time(row.get("close_time") or row.get("open_time")) for row in chartable]
    starts = [value for value in starts if value]
    ends = [value for value in ends if value]
    winners = sum(1 for value in net_profits if value > 0)
    losers = sum(1 for value in net_profits if value < 0)
    breakeven = len(net_profits) - winners - losers
    ea_order_count = sum(1 for row in rows if is_ea_trade(row))
    copy_order_count = sum(1 for row in rows if is_copy_trade(row))
    gross_profit = sum(gross_profits)
    closed_commission = sum(numeric_value(row.get("commission")) + numeric_value(row.get("fee")) for row in rows)
    closed_swap = sum(numeric_value(row.get("swap")) for row in rows)
    closed_taxes = sum(numeric_value(row.get("taxes")) for row in rows)
    commission = numeric_value(cost_summary.get("commission")) if cost_summary is not None else closed_commission
    swap = numeric_value(cost_summary.get("swap")) if cost_summary is not None else closed_swap
    taxes = numeric_value(cost_summary.get("taxes")) if cost_summary is not None else closed_taxes
    fees = commission + swap + taxes
    net_profit = gross_profit + fees
    closed_net_profit = sum(net_profits)

    symbol_groups: dict[str, list[dict]] = defaultdict(list)
    source_groups: dict[tuple[str, str], list[dict]] = defaultdict(list)
    for row in rows:
        symbol_groups[normalize_text(row.get("symbol")) or "未指定"].append(row)
        source_groups[(normalize_text(row.get("platform")) or "未指定", display_unknown(normalize_text(row.get("server"))))].append(row)

    by_symbol = []
    symbol_costs = (cost_summary or {}).get("bySymbol", {})
    for symbol, items in symbol_groups.items():
        item_gross_profits = [numeric_value(row.get("profit")) for row in items]
        item_profits = [
            numeric_value(row.get("profit"))
            + numeric_value(row.get("commission"))
            + numeric_value(row.get("fee"))
            + numeric_value(row.get("swap"))
            + numeric_value(row.get("taxes"))
            for row in items
        ]
        item_winners = sum(1 for value in item_profits if value > 0)
        symbol_cost = symbol_costs.get(symbol)
        symbol_net = sum(item_gross_profits) + sum(numeric_value(symbol_cost.get(key)) for key in ("commission", "swap", "taxes")) if symbol_cost else sum(item_profits)
        item_starts = [parse_trade_time(row.get("open_time")) for row in items]
        item_ends = [parse_trade_time(row.get("close_time") or row.get("open_time")) for row in items]
        item_starts = [value for value in item_starts if value]
        item_ends = [value for value in item_ends if value]
        by_symbol.append({
            "symbol": symbol,
            "orders": len(items),
            "profit": rounded(symbol_net),
            "grossProfit": rounded(sum(item_gross_profits)),
            "volume": rounded(sum(numeric_value(row.get("volume")) for row in items), 4),
            "winRate": rounded(item_winners / len(items) * 100) if items else 0,
            "firstTime": trade_time_text(min(item_starts) if item_starts else None),
            "lastTime": trade_time_text(max(item_ends) if item_ends else None),
        })
    by_symbol.sort(key=lambda item: (-item["orders"], -abs(item["profit"]), item["symbol"]))

    by_source = []
    source_costs = {
        (normalize_text(item.get("platform")), display_unknown(normalize_text(item.get("server")))): item
        for item in (cost_summary or {}).get("sources", [])
    }
    for (platform, server), items in source_groups.items():
        source_currencies = sorted({normalize_text(row.get("account_currency")).upper() for row in items if normalize_text(row.get("account_currency"))})
        source_gross = sum(numeric_value(row.get("profit")) for row in items)
        source_closed_net = sum(
            numeric_value(row.get("profit"))
            + numeric_value(row.get("commission"))
            + numeric_value(row.get("fee"))
            + numeric_value(row.get("swap"))
            + numeric_value(row.get("taxes"))
            for row in items
        )
        source_cost = source_costs.get((platform, server))
        source_net = source_gross + sum(numeric_value(source_cost.get(key)) for key in ("commission", "swap", "taxes")) if source_cost else source_closed_net
        by_source.append({
            "platform": platform,
            "server": server,
            "currency": "/".join(source_currencies),
            "orders": len(items),
            "profit": rounded(source_net),
            "grossProfit": rounded(source_gross),
        })
    by_source.sort(key=lambda item: (-item["orders"], item["platform"], item["server"]))

    active_days = {value.date() for value in starts}
    sorted_starts = sorted(starts)
    gaps = [(later - earlier).total_seconds() for earlier, later in zip(sorted_starts, sorted_starts[1:])]
    minute_counts: dict[str, int] = defaultdict(int)
    for value in starts:
        minute_counts[value.strftime("%Y-%m-%d %H:%M")] += 1

    return {
        "orderCount": len(rows),
        "chartableOrderCount": len(chartable),
        "firstTradeTime": trade_time_text(min(starts) if starts else None),
        "lastTradeTime": trade_time_text(max(ends) if ends else None),
        "symbolCount": len(symbol_groups) if rows else 0,
        "totalProfit": rounded(net_profit),
        "netProfit": rounded(net_profit),
        "grossProfit": rounded(gross_profit),
        "closedNetProfit": rounded(closed_net_profit),
        "winningOrders": winners,
        "losingOrders": losers,
        "breakevenOrders": breakeven,
        "eaOrderCount": ea_order_count,
        "hasEaTrades": ea_order_count > 0,
        "copyOrderCount": copy_order_count,
        "hasCopyTrades": copy_order_count > 0,
        "winRate": rounded(winners / len(net_profits) * 100) if net_profits else 0,
        "averageProfit": rounded(statistics.fmean(net_profits)) if net_profits else 0,
        "medianProfit": rounded(statistics.median(net_profits)) if net_profits else 0,
        "totalVolume": rounded(sum(volumes), 4),
        "averageVolume": rounded(statistics.fmean(volumes), 4) if volumes else 0,
        "maxVolume": rounded(max(volumes), 4) if volumes else 0,
        "averageHoldingSeconds": rounded(statistics.fmean(holding_seconds), 1) if holding_seconds else 0,
        "medianHoldingSeconds": rounded(statistics.median(holding_seconds), 1) if holding_seconds else 0,
        "shortHoldingRatio": rounded(sum(1 for value in holding_seconds if value <= 300) / len(holding_seconds) * 100) if holding_seconds else 0,
        "oneMinuteHoldingRatio": rounded(sum(1 for value in holding_seconds if value <= 60) / len(holding_seconds) * 100) if holding_seconds else 0,
        "commissionTotal": rounded(commission),
        "swapTotal": rounded(swap),
        "taxesTotal": rounded(taxes),
        "feesTotal": rounded(fees),
        "feeToProfitRatio": rounded(abs(fees) / abs(gross_profit) * 100) if gross_profit else None,
        "costsComplete": bool(cost_summary.get("complete", True)) if cost_summary is not None else True,
        "costsIncludeOpenTradeFees": bool(cost_summary.get("includesOpenTradeFees")) if cost_summary is not None else False,
        "activeDays": len(active_days),
        "ordersPerActiveDay": rounded(len(chartable) / len(active_days), 1) if active_days else 0,
        "averageOrderGapSeconds": rounded(statistics.fmean(gaps), 1) if gaps else 0,
        "maxOrdersInOneMinute": max(minute_counts.values(), default=0),
        "bySymbol": by_symbol,
        "bySource": by_source,
    }


def sample_chart_points(points: list[dict], limit: int = 240) -> list[dict]:
    if len(points) <= limit:
        return points
    selected = {0, len(points) - 1}
    step = (len(points) - 1) / max(limit - 1, 1)
    selected.update(min(round(index * step), len(points) - 1) for index in range(limit))
    selected.add(min(range(len(points)), key=lambda index: numeric_value(points[index].get("value"))))
    selected.add(max(range(len(points)), key=lambda index: numeric_value(points[index].get("value"))))
    return [points[index] for index in sorted(selected)]


def trade_visualizations(rows: list[dict], metrics: dict) -> dict:
    ordered = sorted(
        rows,
        key=lambda row: (
            parse_trade_time(row.get("close_time") or row.get("open_time")) or datetime.min,
            normalize_text(row.get("ticket") or row.get("id")),
        ),
    )
    cumulative = []
    daily: dict[str, float] = defaultdict(float)
    running = 0.0
    peak = 0.0
    max_drawdown = 0.0
    for index, row in enumerate(ordered, 1):
        net = sum(numeric_value(row.get(field)) for field in ("profit", "commission", "swap", "taxes"))
        running += net
        peak = max(peak, running)
        max_drawdown = max(max_drawdown, peak - running)
        trade_dt = parse_trade_time(row.get("close_time") or row.get("open_time"))
        time_text = trade_time_text(trade_dt)
        if trade_dt:
            daily[trade_dt.strftime("%Y-%m-%d")] += net
        cumulative.append({
            "index": index,
            "time": time_text,
            "value": rounded(running),
            "change": rounded(net),
        })

    target_total = numeric_value(metrics.get("netProfit"))
    fee_adjustment = target_total - running
    if ordered and abs(fee_adjustment) >= 0.005:
        running += fee_adjustment
        peak = max(peak, running)
        max_drawdown = max(max_drawdown, peak - running)
        last_dt = parse_trade_time(ordered[-1].get("close_time") or ordered[-1].get("open_time"))
        if last_dt:
            daily[last_dt.strftime("%Y-%m-%d")] += fee_adjustment
        cumulative.append({
            "index": len(cumulative) + 1,
            "time": trade_time_text(last_dt),
            "value": rounded(running),
            "change": rounded(fee_adjustment),
            "adjustment": True,
        })

    symbol_rows = list(metrics.get("bySymbol") or [])
    symbol_rows.sort(key=lambda item: (-abs(numeric_value(item.get("profit"))), normalize_text(item.get("symbol"))))
    top_symbols = [dict(item) for item in symbol_rows[:10]]
    remaining = symbol_rows[10:]
    if remaining:
        top_symbols.append({
            "symbol": "其他",
            "orders": sum(mysql_int(item.get("orders")) for item in remaining),
            "profit": rounded(sum(numeric_value(item.get("profit")) for item in remaining)),
            "volume": rounded(sum(numeric_value(item.get("volume")) for item in remaining), 4),
            "winRate": None,
        })

    frequency = riskdash_high_frequency(rows)
    return {
        "pnlSeries": sample_chart_points(cumulative),
        "dailyPnl": [{"date": day, "profit": rounded(value)} for day, value in sorted(daily.items())],
        "outcomes": {
            "winning": mysql_int(metrics.get("winningOrders")),
            "losing": mysql_int(metrics.get("losingOrders")),
            "breakeven": mysql_int(metrics.get("breakevenOrders")),
        },
        "holdingBuckets": frequency.get("buckets", []),
        "symbolPerformance": top_symbols,
        "feeBreakdown": [
            {"label": "手续费", "value": rounded(numeric_value(metrics.get("commissionTotal")))},
            {"label": "利息 / Swap", "value": rounded(numeric_value(metrics.get("swapTotal")))},
            {"label": "税费", "value": rounded(numeric_value(metrics.get("taxesTotal")))},
        ],
        "maxDrawdown": rounded(max_drawdown),
        "feeAdjustment": rounded(fee_adjustment),
        "netTotal": rounded(target_total),
    }


def account_meta_for_rows(rows: list[dict]) -> dict:
    metadata = []
    seen: set[tuple] = set()
    for row in rows:
        currency = normalize_text(row.get("account_currency")).upper()
        if not currency:
            continue
        key = (normalize_text(row.get("platform")), normalize_text(row.get("server")), currency)
        if key in seen:
            continue
        seen.add(key)
        metadata.append({
            "platform": key[0],
            "server": key[1],
            "currency": currency,
            "displayCurrency": normalize_text(row.get("display_currency")) or currency,
            "moneyScale": numeric_value(row.get("money_scale")) or 1.0,
            "isCentAccount": bool(row.get("is_cent_account")),
            "source": normalize_text(row.get("currency_source")),
        })
    currencies = sorted({item["currency"] for item in metadata})
    display_currencies = sorted({item["displayCurrency"] for item in metadata})
    single = metadata[0] if len(currencies) == 1 else {}
    return {
        "currency": currencies[0] if len(currencies) == 1 else "",
        "currencies": currencies,
        "displayCurrency": display_currencies[0] if len(display_currencies) == 1 else "",
        "displayCurrencies": display_currencies,
        "moneyScale": single.get("moneyScale", 1.0),
        "isCentAccount": bool(single.get("isCentAccount")),
        "currencyDetected": bool(metadata),
        "mixedCurrencies": len(display_currencies) > 1,
        "sources": metadata,
    }


def riskdash_volume_scale(account_meta: dict) -> float:
    # RiskDash reports cent-account exposure in standard USD-lot equivalents.
    return 0.01 if account_meta.get("isCentAccount") else 1.0


def max_concurrent_volume(rows: list[dict], volume_scale: float = 1.0) -> float:
    events: list[tuple[datetime, int, float]] = []
    for row in rows:
        opened = parse_trade_time(row.get("open_time"))
        closed = parse_trade_time(row.get("close_time"))
        volume = numeric_value(row.get("volume")) * volume_scale
        if opened and volume:
            events.append((opened, 1, volume))
        if closed and volume:
            events.append((closed, 0, -volume))
    current = 0.0
    maximum = 0.0
    for _, _, delta in sorted(events, key=lambda item: (item[0], item[1])):
        current += delta
        maximum = max(maximum, current)
    return rounded(maximum, 4)


def riskdash_high_frequency(rows: list[dict]) -> dict:
    trades = [row for row in rows if is_chartable_trade(row)]
    account_meta = account_meta_for_rows(trades)
    volume_scale = riskdash_volume_scale(account_meta)
    winning_holds = [numeric_value(row.get("holding_seconds")) for row in trades if numeric_value(row.get("profit")) > 0]
    losing_holds = [numeric_value(row.get("holding_seconds")) for row in trades if numeric_value(row.get("profit")) < 0]
    all_holds = [numeric_value(row.get("holding_seconds")) for row in trades]
    definitions = [
        ("<10s", 0, 10),
        ("10s-2min", 10, 120),
        ("2-5min", 120, 300),
        ("5-10min", 300, 600),
        (">10min", 600, float("inf")),
    ]
    buckets = []
    for label, lower, upper in definitions:
        items = [row for row in trades if lower <= numeric_value(row.get("holding_seconds")) < upper]
        winners = [row for row in items if numeric_value(row.get("profit")) > 0]
        gross_profit = sum(numeric_value(row.get("profit")) for row in items)
        positive_profit = sum(numeric_value(row.get("profit")) for row in winners)
        volume = sum(numeric_value(row.get("volume")) * volume_scale for row in items)
        buckets.append({
            "label": label,
            "orders": len(items),
            "winRate": rounded(len(winners) / len(items) * 100) if items else 0,
            "grossProfit": rounded(gross_profit),
            "volume": rounded(volume, 4),
            "averageProfitPerLot": rounded(gross_profit / volume) if volume else 0,
            "averageProfitPerOrder": rounded(gross_profit / len(items)) if items else 0,
            "averageVolume": rounded(volume / len(items), 4) if items else 0,
            "positiveProfit": rounded(positive_profit),
        })
    positive_total = sum(item["positiveProfit"] for item in buckets)
    for item in buckets:
        item["profitShare"] = rounded(item["positiveProfit"] / positive_total * 100, 0) if positive_total else 0
        item.pop("positiveProfit", None)
    return {
        "orderCount": len(trades),
        "averageHoldingMinutes": rounded(statistics.fmean(all_holds) / 60, 2) if all_holds else 0,
        "winningAverageHoldingMinutes": rounded(statistics.fmean(winning_holds) / 60, 2) if winning_holds else 0,
        "losingAverageHoldingMinutes": rounded(statistics.fmean(losing_holds) / 60, 2) if losing_holds else 0,
        "highFrequencyOrderRatio": rounded(sum(1 for value in all_holds if value < 120) / len(all_holds) * 100, 2) if all_holds else 0,
        "volumeScale": volume_scale,
        "buckets": buckets,
    }


def classify_mt5_cashflows(rows: list[dict], money_scale: float = 1.0) -> dict:
    result = {
        "netDeposit": 0.0,
        "negativeBalanceClear": 0.0,
        "compensation": 0.0,
        "reward": 0.0,
        "internalTransfer": 0.0,
        "other": 0.0,
        "depositTotal": 0.0,
        "withdrawalTotal": 0.0,
        "depositCount": 0,
        "withdrawalCount": 0,
    }
    deposit_times: list[datetime] = []
    withdrawal_times: list[datetime] = []
    for row in rows:
        action = mysql_int(row.get("Action"), -1)
        amount = numeric_value(row.get("Profit")) * money_scale
        comment = normalize_text(row.get("Comment")).upper()
        if action == 2 and comment.startswith(("DEP-", "WDR-", "CRM-DP-", "CRM-CW")):
            result["netDeposit"] += amount
            event_time = parse_trade_time(row.get("TimeMsc") or row.get("Time"))
            if comment.startswith(("WDR-", "CRM-CW")) or amount < 0:
                result["withdrawalTotal"] += abs(amount)
                result["withdrawalCount"] += 1
                if event_time:
                    withdrawal_times.append(event_time)
            else:
                result["depositTotal"] += max(amount, 0.0)
                result["depositCount"] += 1
                if event_time:
                    deposit_times.append(event_time)
        elif action == 2 and (comment.startswith("RST-") or "NEGATIVE BALANCE" in comment):
            result["negativeBalanceClear"] += amount
        elif "COMP" in comment:
            result["compensation"] += amount
        elif action == 6 or "REWARD" in comment or "BONUS" in comment:
            result["reward"] += amount
        elif action == 2 and (comment.startswith("TFM-") or comment.startswith("TFH-")):
            result["internalTransfer"] += amount
        else:
            result["other"] += amount
    return {
        **{key: rounded(value) for key, value in result.items()},
        "depositTimes": [trade_time_text(value) for value in sorted(deposit_times)],
        "withdrawalTimes": [trade_time_text(value) for value in sorted(withdrawal_times)],
    }


def classify_mt4_cashflows(rows: list[dict], money_scale: float = 1.0) -> dict:
    result = {
        "netDeposit": 0.0,
        "negativeBalanceClear": 0.0,
        "compensation": 0.0,
        "reward": 0.0,
        "internalTransfer": 0.0,
        "other": 0.0,
        "depositTotal": 0.0,
        "withdrawalTotal": 0.0,
        "depositCount": 0,
        "withdrawalCount": 0,
    }
    deposit_times: list[datetime] = []
    withdrawal_times: list[datetime] = []
    for row in rows:
        command = mysql_int(row.get("CMD"), -1)
        amount = numeric_value(row.get("PROFIT")) * money_scale
        comment = normalize_text(row.get("COMMENT")).upper()
        if "RST-" in comment or "NEGATIVE BALANCE" in comment or "ZERO BALANCE" in comment or comment.startswith("CCB-"):
            result["negativeBalanceClear"] += amount
        elif "COMP" in comment or comment.startswith("CPS_"):
            result["compensation"] += amount
        elif "REWARD" in comment or "BONUS" in comment:
            result["reward"] += amount
        elif comment.startswith(("TFM-", "TFH-")) or "INTERNAL TRANSFER" in comment:
            result["internalTransfer"] += amount
        elif command == 6:
            # MT4 balance operations include both deposits and withdrawals.
            result["netDeposit"] += amount
            event_time = parse_trade_time(row.get("OPEN_TIME") or row.get("CLOSE_TIME"))
            if amount < 0 or "WITHDRAW" in comment or comment.startswith("WDR-"):
                result["withdrawalTotal"] += abs(amount)
                result["withdrawalCount"] += 1
                if event_time:
                    withdrawal_times.append(event_time)
            else:
                result["depositTotal"] += max(amount, 0.0)
                result["depositCount"] += 1
                if event_time:
                    deposit_times.append(event_time)
        else:
            result["other"] += amount
    return {
        **{key: rounded(value) for key, value in result.items()},
        "depositTimes": [trade_time_text(value) for value in sorted(deposit_times)],
        "withdrawalTimes": [trade_time_text(value) for value in sorted(withdrawal_times)],
    }


def hierarchy_net_deposit_payload(
    target: str,
    start: str,
    end: str,
    product: str = "",
    activity_rules: object = False,
) -> dict:
    parsed = hierarchy_net_deposit.parse_query(target, start, end, product, activity_rules)
    cache_key = (
        "hierarchy-net-deposit",
        parsed["target"],
        parsed["start"].isoformat(),
        parsed["end"].isoformat(),
        parsed["product"],
        parsed["activityRules"],
    )
    cached = account_cache_get(cache_key)
    if cached is not None:
        return cached
    payload = hierarchy_net_deposit.build_payload(
        parsed["target"],
        parsed["start"].isoformat(sep=" "),
        parsed["end"].isoformat(sep=" "),
        parsed["product"],
        parsed["activityRules"],
        sources=MYSQL_SOURCES,
        connect=mysql_trade_connect,
        classify_mt5_cashflows=classify_mt5_cashflows,
        classify_mt4_cashflows=classify_mt4_cashflows,
        refreshed_at=now_text(),
    )
    return account_cache_set(cache_key, payload)


def hierarchy_products_payload() -> dict:
    cache_key = ("hierarchy-products",)
    cached = account_cache_get(cache_key)
    if cached is not None:
        return cached
    payload = hierarchy_net_deposit.list_products(MYSQL_SOURCES, mysql_trade_connect)
    return account_cache_set(cache_key, payload)


def cashflow_timing_summary(cashflows: dict, rows: list[dict]) -> dict:
    deposit_times = [parse_trade_time(value) for value in cashflows.get("depositTimes", [])]
    withdrawal_times = [parse_trade_time(value) for value in cashflows.get("withdrawalTimes", [])]
    deposit_times = [value for value in deposit_times if value]
    withdrawal_times = [value for value in withdrawal_times if value]
    open_times = [parse_trade_time(row.get("open_time_msc") or row.get("open_time")) for row in rows]
    close_times = [parse_trade_time(row.get("close_time_msc") or row.get("close_time")) for row in rows]
    open_times = [value for value in open_times if value]
    close_times = [value for value in close_times if value]
    first_trade = min(open_times) if open_times else None
    last_trade = max(close_times or open_times) if (close_times or open_times) else None
    withdrawal_after_last_trade = min((value for value in withdrawal_times if last_trade and value >= last_trade), default=None)
    first_deposit_to_trade_hours = (first_trade - min(deposit_times)).total_seconds() / 3600 if first_trade and deposit_times and first_trade >= min(deposit_times) else None
    trade_to_withdrawal_hours = (withdrawal_after_last_trade - last_trade).total_seconds() / 3600 if withdrawal_after_last_trade and last_trade else None
    return {
        "depositTotal": cashflows.get("depositTotal", 0),
        "withdrawalTotal": cashflows.get("withdrawalTotal", 0),
        "depositCount": cashflows.get("depositCount", 0),
        "withdrawalCount": cashflows.get("withdrawalCount", 0),
        "firstDepositAt": trade_time_text(min(deposit_times)) if deposit_times else "",
        "lastWithdrawalAt": trade_time_text(max(withdrawal_times)) if withdrawal_times else "",
        "firstDepositToTradeHours": rounded(first_deposit_to_trade_hours, 1) if first_deposit_to_trade_hours is not None else None,
        "lastTradeToWithdrawalHours": rounded(trade_to_withdrawal_hours, 1) if trade_to_withdrawal_hours is not None else None,
    }


def calculate_comprehensive_profit(
    closed_net_profit: float,
    rebate: float,
    holding_profit: float,
    negative_balance_clear: float = 0.0,
    compensation: float = 0.0,
    reward: float = 0.0,
) -> float:
    return rounded(
        closed_net_profit
        + rebate
        + holding_profit
        + negative_balance_clear
        + compensation
        + reward
    )


def source_crm_routes(source: dict) -> list[dict[str, str]]:
    configured = source.get("crm_routes")
    if not isinstance(configured, list):
        configured = [{
            "schema": source.get("crm_schema"),
            "mt_server_code": source.get("mt_server_code"),
        }]
    routes: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for route in configured:
        if not isinstance(route, dict):
            continue
        schema = normalize_text(route.get("schema") or route.get("crm_schema"))
        server_code = normalize_text(route.get("mt_server_code"))
        key = (schema, server_code)
        if not schema or not server_code or key in seen:
            continue
        seen.add(key)
        routes.append({"schema": schema, "mt_server_code": server_code})
    return routes


def query_account_rebate(cur, source: dict, account: str) -> tuple[int, float]:
    rebate_rows = 0
    rebate = 0.0
    for route in source_crm_routes(source):
        cur.execute(
            f"select count(*) as RebateRows, sum(rebate_amount) as RebateAmount "
            f"from `{route['schema']}`.`rebate_task_detail` "
            "where trade_mt_login = %s and mt_server_code = %s",
            (int(account), route["mt_server_code"]),
        )
        rebate_data = cur.fetchone() or {}
        rebate_rows += mysql_int(rebate_data.get("RebateRows"))
        rebate += numeric_value(rebate_data.get("RebateAmount"))
    return rebate_rows, rebate


def query_mt5_finance_panel(source: dict, account: str, rows: list[dict], metrics: dict) -> dict:
    account_meta = account_meta_for_rows(rows)
    snapshot: dict = {}
    cash_rows: list[dict] = []
    position_count = 0
    rebate = 0.0
    rebate_rows = 0
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            if not account_meta.get("currencyDetected"):
                direct_meta = query_mysql_mt5_account_meta(cur, source, account)
                account_meta = {
                    "currency": direct_meta.get("currency", ""),
                    "displayCurrency": direct_meta.get("displayCurrency", ""),
                    "moneyScale": direct_meta.get("moneyScale", 1.0),
                    "isCentAccount": direct_meta.get("isCentAccount", False),
                    "currencyDetected": bool(direct_meta.get("currency")),
                }
            cur.execute(
                f"select Login, Balance, Credit, Profit, Storage, Floating, Equity, Margin, MarginFree, MarginLevel from `{source['schema']}`.`mt5_accounts` where Login = %s",
                (int(account),),
            )
            snapshot = cur.fetchone() or {}
            cur.execute(
                f"select Action, Profit, Comment, Time, TimeMsc from `{source['schema']}`.`{source['table']}` where Login = %s and Action not in (0, 1)",
                (int(account),),
            )
            cash_rows = cur.fetchall()
            cur.execute(
                f"select count(*) as PositionCount from `{source['schema']}`.`mt5_positions` where Login = %s",
                (int(account),),
            )
            position_count = mysql_int((cur.fetchone() or {}).get("PositionCount"))
            rebate_rows, rebate = query_account_rebate(cur, source, account)

    money_scale = numeric_value(account_meta.get("moneyScale")) or 1.0
    cashflows = classify_mt5_cashflows(cash_rows, money_scale)
    cashflow_timing = cashflow_timing_summary(cashflows, rows)
    holding_profit = numeric_value(snapshot.get("Profit")) * money_scale
    closed_net_profit = numeric_value(metrics.get("netProfit"))
    comprehensive_profit = calculate_comprehensive_profit(
        closed_net_profit,
        rebate,
        holding_profit,
        cashflows["negativeBalanceClear"],
        cashflows["compensation"],
        cashflows["reward"],
    )
    net_deposit = cashflows["netDeposit"]
    return {
        "currency": account_meta.get("currency", ""),
        "displayCurrency": account_meta.get("displayCurrency", ""),
        "moneyScale": money_scale,
        "balance": rounded(numeric_value(snapshot.get("Balance")) * money_scale),
        "equity": rounded(numeric_value(snapshot.get("Equity")) * money_scale),
        "credit": rounded(numeric_value(snapshot.get("Credit")) * money_scale),
        "margin": rounded(numeric_value(snapshot.get("Margin")) * money_scale),
        "marginFree": rounded(numeric_value(snapshot.get("MarginFree")) * money_scale),
        "marginLevel": rounded(numeric_value(snapshot.get("MarginLevel")), 2),
        "leverage": rounded(numeric_value(snapshot.get("MarginLeverage")), 0),
        "holdingProfit": rounded(holding_profit),
        "grossClosedProfit": rounded(numeric_value(metrics.get("grossProfit"))),
        "tradingFees": rounded(numeric_value(metrics.get("commissionTotal"))),
        "interest": rounded(numeric_value(metrics.get("swapTotal"))),
        "taxes": rounded(numeric_value(metrics.get("taxesTotal"))),
        "closedNetProfit": rounded(closed_net_profit),
        "netDeposit": rounded(net_deposit),
        **cashflow_timing,
        "negativeBalanceClear": cashflows["negativeBalanceClear"],
        "compensation": cashflows["compensation"],
        "reward": cashflows["reward"],
        "rebate": rounded(rebate),
        "rebateRows": rebate_rows,
        "comprehensiveProfit": rounded(comprehensive_profit),
        "comprehensiveProfitRate": rounded(comprehensive_profit / net_deposit, 2) if net_deposit else None,
        "liquidationRate": None,
        "liquidationAmountRatio": None,
        "internalTransfer": cashflows["internalTransfer"],
        "unclassifiedCashflow": cashflows["other"],
        "currentPositionCount": position_count,
        "highestHoldingVolume": max_concurrent_volume(rows, riskdash_volume_scale(account_meta)),
    }


def query_mt4_finance_panel(source: dict, account: str, rows: list[dict], metrics: dict) -> dict:
    account_meta = account_meta_for_rows(rows)
    snapshot: dict = {}
    cash_rows: list[dict] = []
    position_count = 0
    rebate = 0.0
    rebate_rows = 0
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            try:
                cur.execute(
                    f"""select LOGIN, `GROUP` as AccountGroup, STATUS, LEVERAGE, BALANCE, CREDIT,
                               EQUITY, MARGIN, MARGIN_LEVEL, MARGIN_FREE, CURRENCY
                        from `{source['schema']}`.`mt4_users_view` where LOGIN = %s limit 1""",
                    (int(account),),
                )
                snapshot = cur.fetchone() or {}
            except Exception:
                snapshot = {}
            if not account_meta.get("currencyDetected") and snapshot:
                direct_meta = account_money_meta(snapshot.get("CURRENCY"), snapshot.get("AccountGroup"), source.get("name"))
                account_meta = {
                    "currency": direct_meta.get("currency", ""),
                    "displayCurrency": direct_meta.get("displayCurrency", ""),
                    "moneyScale": direct_meta.get("moneyScale", 1.0),
                    "isCentAccount": direct_meta.get("isCentAccount", False),
                    "currencyDetected": bool(direct_meta.get("currency")),
                }
            try:
                cur.execute(
                    f"select CMD, PROFIT, COMMENT, OPEN_TIME, CLOSE_TIME from `{source['schema']}`.`{source['table']}` where LOGIN = %s and CMD in (6, 7)",
                    (int(account),),
                )
                cash_rows = cur.fetchall()
            except Exception:
                cash_rows = []
            try:
                cur.execute(
                    f"select count(*) as PositionCount from `{source['schema']}`.`{source['table']}` where LOGIN = %s and CMD in (0, 1) and CLOSE_TIME <= '1971-01-01'",
                    (int(account),),
                )
                position_count = mysql_int((cur.fetchone() or {}).get("PositionCount"))
            except Exception:
                position_count = 0
            try:
                rebate_rows, rebate = query_account_rebate(cur, source, account)
            except Exception:
                rebate = 0.0
                rebate_rows = 0

    money_scale = numeric_value(account_meta.get("moneyScale")) or 1.0
    cashflows = classify_mt4_cashflows(cash_rows, money_scale)
    cashflow_timing = cashflow_timing_summary(cashflows, rows)
    balance = numeric_value(snapshot.get("BALANCE")) * money_scale
    credit = numeric_value(snapshot.get("CREDIT")) * money_scale
    equity = numeric_value(snapshot.get("EQUITY")) * money_scale
    holding_profit = equity - balance - credit
    closed_net_profit = numeric_value(metrics.get("netProfit"))
    comprehensive_profit = calculate_comprehensive_profit(
        closed_net_profit,
        rebate,
        holding_profit,
        cashflows["negativeBalanceClear"],
        cashflows["compensation"],
        cashflows["reward"],
    )
    net_deposit = cashflows["netDeposit"]
    return {
        "currency": account_meta.get("currency", ""),
        "displayCurrency": account_meta.get("displayCurrency", ""),
        "moneyScale": money_scale,
        "balance": rounded(balance),
        "equity": rounded(equity),
        "credit": rounded(credit),
        "margin": rounded(numeric_value(snapshot.get("MARGIN")) * money_scale),
        "marginFree": rounded(numeric_value(snapshot.get("MARGIN_FREE")) * money_scale),
        "marginLevel": rounded(numeric_value(snapshot.get("MARGIN_LEVEL")), 2),
        "leverage": rounded(numeric_value(snapshot.get("LEVERAGE")), 0),
        "holdingProfit": rounded(holding_profit),
        "grossClosedProfit": rounded(numeric_value(metrics.get("grossProfit"))),
        "tradingFees": rounded(numeric_value(metrics.get("commissionTotal"))),
        "interest": rounded(numeric_value(metrics.get("swapTotal"))),
        "taxes": rounded(numeric_value(metrics.get("taxesTotal"))),
        "closedNetProfit": rounded(closed_net_profit),
        "netDeposit": rounded(net_deposit),
        **cashflow_timing,
        "negativeBalanceClear": cashflows["negativeBalanceClear"],
        "compensation": cashflows["compensation"],
        "reward": cashflows["reward"],
        "rebate": rounded(rebate),
        "rebateRows": rebate_rows,
        "comprehensiveProfit": rounded(comprehensive_profit),
        "comprehensiveProfitRate": rounded(comprehensive_profit / net_deposit, 2) if net_deposit else None,
        "liquidationRate": None,
        "liquidationAmountRatio": None,
        "internalTransfer": cashflows["internalTransfer"],
        "unclassifiedCashflow": cashflows["other"],
        "currentPositionCount": position_count,
        "highestHoldingVolume": max_concurrent_volume(rows, riskdash_volume_scale(account_meta)),
        "databaseStatus": normalize_text(snapshot.get("STATUS")),
    }


def source_for_crm_route(crm_schema: object, mt_server_code: object) -> dict | None:
    route_key = (normalize_text(crm_schema), normalize_text(mt_server_code))
    for candidate in MYSQL_SOURCES:
        if route_key in {
            (route["schema"], route["mt_server_code"])
            for route in source_crm_routes(candidate)
        }:
            return candidate
    return None


def query_same_name_accounts(source: dict, account: str) -> list[dict]:
    routes = source_crm_routes(source)
    fallback = [{
        "account": normalize_text(account),
        "source": source,
        "crmSchema": normalize_text((routes or [{}])[0].get("schema")),
        "mtServerCode": normalize_text((routes or [{}])[0].get("mt_server_code")),
    }]
    if not routes:
        return fallback
    matched_accounts: list[dict] = []
    seen: set[tuple[str, str, str]] = set()
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            for route in routes:
                cur.execute(
                    f"select user_id from `{route['schema']}`.`mt_users_account` where mt_login = %s and mt_server_code = %s limit 1",
                    (int(account), route["mt_server_code"]),
                )
                row = cur.fetchone() or {}
                user_id = mysql_int(row.get("user_id"))
                if not user_id:
                    continue
                cur.execute(
                    f"select mt_login, mt_server_code from `{route['schema']}`.`mt_users_account` "
                    "where user_id = %s order by mt_server_code, mt_login",
                    (user_id,),
                )
                for item in cur.fetchall():
                    same_login = normalize_text(item.get("mt_login"))
                    server_code = normalize_text(item.get("mt_server_code"))
                    key = (route["schema"], server_code, same_login)
                    if not same_login or key in seen:
                        continue
                    seen.add(key)
                    matched_accounts.append({
                        "account": same_login,
                        "source": source_for_crm_route(route["schema"], server_code),
                        "crmSchema": route["schema"],
                        "mtServerCode": server_code,
                    })
    return matched_accounts or fallback


def query_same_name_logins(source: dict, account: str) -> list[str]:
    return list(dict.fromkeys(
        item["account"] for item in query_same_name_accounts(source, account) if item.get("account")
    ))


def account_crm_ib_relationship_payload(login: str, filters: dict | None = None) -> dict:
    """Return CRM routes plus a bounded, indexed direct-rebate branch for an IB owner."""
    filters = filters or {}
    login = normalize_text(login)
    platform = normalize_text(filters.get("platform")).upper()
    server = normalize_text(filters.get("server"))
    include_ib_aggregate = bool(filters.get("includeIbAggregate", True))
    if not login.isdigit():
        raise ValueError("账号格式无效")
    # Version the key: pre-IB-branch payloads must not suppress the new direct-rebate read.
    cache_key = ("crm-ib-relationship-v3", login, platform, server, include_ib_aggregate)
    cached = account_cache_get(cache_key)
    if cached is not None:
        return cached

    records: list[dict] = []
    seen: set[tuple[str, int]] = set()
    sources = [source for source in MYSQL_SOURCES if source_allowed(source, platform=platform, server=server)]
    for source in sources:
        with mysql_trade_connect(source) as conn:
            with conn.cursor() as cur:
                for route in source_crm_routes(source):
                    crm_schema = route["schema"]
                    server_code = route["mt_server_code"]
                    cur.execute(
                        f"select user_id from `{crm_schema}`.`mt_users_account` "
                        "where mt_login = %s and mt_server_code = %s limit 1",
                        (int(login), server_code),
                    )
                    mapping = cur.fetchone() or {}
                    crm_user_id = mysql_int(mapping.get("user_id"))
                    if not crm_user_id or (crm_schema, crm_user_id) in seen:
                        continue
                    seen.add((crm_schema, crm_user_id))
                    cur.execute(
                        f"select id, supper_id, top_ib_id from `{crm_schema}`.`sys_user_view` where id = %s limit 1",
                        (crm_user_id,),
                    )
                    crm_user = cur.fetchone() or {}
                    direct_ib_user_id = mysql_int(crm_user.get("supper_id"))
                    top_ib_user_id = mysql_int(crm_user.get("top_ib_id"))
                    direct_ib_accounts: list[dict] = []
                    direct_ib_account_truncated = False
                    if direct_ib_user_id:
                        cur.execute(
                            f"select mt_login, mt_server_code from `{crm_schema}`.`mt_users_account` "
                            "where user_id = %s order by mt_server_code, mt_login limit 21",
                            (direct_ib_user_id,),
                        )
                        direct_rows = cur.fetchall() or []
                        direct_ib_account_truncated = len(direct_rows) > 20
                        for item in direct_rows[:20]:
                            account = normalize_text(item.get("mt_login"))
                            account_server_code = normalize_text(item.get("mt_server_code"))
                            account_source = source_for_crm_route(crm_schema, account_server_code)
                            if not account or not account_source:
                                continue
                            direct_ib_accounts.append({
                                "account": account,
                                "platform": normalize_text(account_source.get("platform")),
                                "server": normalize_text(account_source.get("server") or account_source.get("name")),
                            })

                    # This is intentionally *not* a top-IB tree scan.  It asks one
                    # indexed question: which trading accounts have actually produced
                    # direct rebate records for the current CRM user?  Grouping returns
                    # a single graph edge per account rather than raw rebate history.
                    # The graph-wide account guard remains authoritative.  A much smaller
                    # branch limit keeps one IB from materialising thousands of otherwise
                    # unrelated query candidates into the 8777 process at once.
                    cur.execute(
                        f"select trade_mt_login as Account, mt_server_code as ServerCode, "
                        "count(distinct coalesce(trade_mt_deal, trade_mt_ticket)) as RebateOrderCount, "
                        "max(create_time) as LastRebateAt "
                        f"from `{crm_schema}`.`rebate_task_detail` "
                        "where rebate_ib_id = %s "
                        "group by trade_mt_login, mt_server_code "
                        # Do not order this potentially broad group.  The exact IB-ID
                        # index can stop at the 2,001st distinct account; deterministic
                        # recency ordering would force a full temporary/filesort scan.
                        "limit 151",
                        (crm_user_id,),
                    )
                    own_ib_rows = cur.fetchall() or []
                    own_ib_direct_rebate_truncated = len(own_ib_rows) > 150
                    own_ib_direct_rebate_accounts: list[dict] = []
                    for item in own_ib_rows[:150]:
                        account = normalize_text(item.get("Account"))
                        account_server_code = normalize_text(item.get("ServerCode"))
                        account_source = source_for_crm_route(crm_schema, account_server_code)
                        if not account or not account_source:
                            continue
                        own_ib_direct_rebate_accounts.append({
                            "account": account,
                            "platform": normalize_text(account_source.get("platform")),
                            "server": normalize_text(account_source.get("server") or account_source.get("name")),
                            "rebateOrderCount": mysql_int(item.get("RebateOrderCount")),
                            "lastRebateAt": normalize_text(item.get("LastRebateAt")),
                        })

                    top_ib_account_count = top_ib_client_count = 0
                    if top_ib_user_id and include_ib_aggregate:
                        cur.execute(
                            f"select count(distinct a.mt_login) as AccountCount, count(distinct a.user_id) as ClientCount "
                            f"from `{crm_schema}`.`mt_users_account` a "
                            f"inner join `{crm_schema}`.`sys_user_view` u on u.id = a.user_id "
                            "where u.top_ib_id = %s",
                            (top_ib_user_id,),
                        )
                        aggregate = cur.fetchone() or {}
                        top_ib_account_count = mysql_int(aggregate.get("AccountCount"))
                        top_ib_client_count = mysql_int(aggregate.get("ClientCount"))

                    records.append({
                        "crmSchema": crm_schema,
                        "platform": normalize_text(source.get("platform")),
                        "server": normalize_text(source.get("server") or source.get("name")),
                        "crmUserId": crm_user_id,
                        "directIbUserId": direct_ib_user_id or "",
                        "topIbUserId": top_ib_user_id or "",
                        "directIbAccounts": direct_ib_accounts,
                        "directIbAccountTruncated": direct_ib_account_truncated,
                        # A user is represented as an expandable IB node only when
                        # direct-rebate evidence actually exists; an empty indexed probe
                        # is merely a fast negative check, not an IB classification.
                        "ownIbDirectRebateChecked": bool(own_ib_rows),
                        "ownIbDirectRebateAccounts": own_ib_direct_rebate_accounts,
                        "ownIbDirectRebateTruncated": own_ib_direct_rebate_truncated,
                        "topIbAggregateAvailable": include_ib_aggregate,
                        "topIbAccountCount": top_ib_account_count,
                        "topIbClientCount": top_ib_client_count,
                    })
    return account_cache_set(cache_key, {"ok": True, "account": login, "records": records})


def query_mt5_database_statuses(source: dict, accounts: list[str]) -> dict[str, str]:
    logins = [normalize_text(account) for account in accounts if normalize_text(account).isdigit()]
    if not logins:
        return {}
    placeholders = ",".join(["%s"] * len(logins))
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            cur.execute(
                f"select Login, Status from `{source['schema']}`.`mt5_users_view` where Login in ({placeholders})",
                tuple(int(login) for login in logins),
            )
            return {
                normalize_text(row.get("Login")): normalize_text(row.get("Status"))
                for row in cur.fetchall()
            }


def query_mt4_database_statuses(source: dict, accounts: list[str]) -> dict[str, str]:
    logins = [normalize_text(account) for account in accounts if normalize_text(account).isdigit()]
    if not logins:
        return {}
    placeholders = ",".join(["%s"] * len(logins))
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            try:
                cur.execute(
                    f"select LOGIN, STATUS from `{source['schema']}`.`mt4_users_view` where LOGIN in ({placeholders})",
                    tuple(int(login) for login in logins),
                )
                return {
                    normalize_text(row.get("LOGIN")): normalize_text(row.get("STATUS"))
                    for row in cur.fetchall()
                }
            except Exception:
                return {}


def build_riskdash_panels(login: str, all_rows: list[dict], all_metrics: dict, filtered_rows: list[dict]) -> dict:
    servers = {normalize_text(row.get("server")) for row in all_rows if normalize_text(row.get("server"))}
    platforms = {normalize_text(row.get("platform")) for row in all_rows if normalize_text(row.get("platform"))}
    if len(servers) != 1 or len(platforms) != 1 or not platforms.issubset({"MT4", "MT5"}):
        return {
            "available": False,
            "reason": "请选择单一平台和服务器后查看风控面板" if all_rows else "账户暂未做单",
        }
    server = next(iter(servers))
    platform = next(iter(platforms))
    expected_kind = "mt5_deals" if platform == "MT5" else "mt4_trades"
    source = next((item for item in MYSQL_SOURCES if item.get("server") == server and item.get("kind") == expected_kind), None)
    if not source:
        return {"available": False, "reason": "当前服务器尚未配置风控面板数据源"}

    finance_query = query_mt5_finance_panel if platform == "MT5" else query_mt4_finance_panel
    finance = finance_query(source, login, all_rows, all_metrics)
    records = load_records()
    same_name_accounts = query_same_name_accounts(source, login)
    status_groups: dict[tuple[str, str], dict] = {}
    for item in same_name_accounts:
        item_source = item.get("source")
        if not item_source:
            continue
        status_key = (normalize_text(item_source.get("platform")), normalize_text(item_source.get("server")))
        group = status_groups.setdefault(status_key, {"source": item_source, "accounts": []})
        group["accounts"].append(normalize_text(item.get("account")))
    database_statuses: dict[tuple[str, str, str], str] = {}
    for status_key, group in status_groups.items():
        item_source = group["source"]
        status_query = query_mt5_database_statuses if item_source.get("kind") == "mt5_deals" else query_mt4_database_statuses
        try:
            statuses = status_query(item_source, group["accounts"])
        except Exception:
            statuses = {}
        for same_login, status in statuses.items():
            database_statuses[(*status_key, same_login)] = status

    def item_key(item: dict) -> tuple[str, str, str]:
        item_source = item.get("source") or {}
        return (
            normalize_text(item_source.get("platform")),
            normalize_text(item_source.get("server")),
            normalize_text(item.get("account")),
        )

    def empty_same_name_row(item: dict) -> dict:
        same_login = normalize_text(item.get("account"))
        item_source = item.get("source") or {}
        item_platform = normalize_text(item_source.get("platform"))
        item_server = normalize_text(item_source.get("server"))
        if not item_server:
            item_server = f"{normalize_text(item.get('crmSchema'))} code {normalize_text(item.get('mtServerCode'))}".strip()
        ledger = ledger_record_for_login(same_login, records)
        return {
            "server": item_server, "platform": item_platform, "account": same_login, "currency": "",
            "balance": 0, "equity": 0, "netDeposit": 0, "holdingProfit": 0,
            "closedNetProfit": 0, "adjustments": 0, "rebate": 0, "comprehensiveProfit": 0,
            "highestHoldingVolume": 0,
            "databaseStatus": "未配置数据源" if not item.get("source") else database_statuses.get(item_key(item), ""),
            "localStatus": normalize_text(ledger.get("建议动作")) if ledger else "",
        }

    def build_same_name_row(item: dict) -> dict:
        same_login = normalize_text(item.get("account"))
        item_source = item.get("source")
        if not item_source:
            return empty_same_name_row(item)
        item_platform = normalize_text(item_source.get("platform"))
        item_server = normalize_text(item_source.get("server"))
        is_current = same_login == normalize_text(login) and item_platform == platform and item_server == server
        item_finance_query = query_mt5_finance_panel if item_source.get("kind") == "mt5_deals" else query_mt4_finance_panel
        item_trade_query = query_mysql_mt5_source if item_source.get("kind") == "mt5_deals" else query_mysql_mt4_source
        if is_current:
            trade_rows = all_rows
            metrics = all_metrics
            account_finance = finance
        else:
            peer_limit = None if item_source.get("kind") == "mt4_trades" else 50000
            trade_rows = item_trade_query(item_source, same_login, limit=peer_limit)
            costs = query_mysql_trade_costs(same_login, platform=item_platform, server=item_server)
            metrics = trade_metrics(trade_rows, costs)
            account_finance = item_finance_query(item_source, same_login, trade_rows, metrics)
        ledger = ledger_record_for_login(same_login, records)
        return {
            "server": item_server,
            "platform": item_platform,
            "account": same_login,
            "currency": account_finance.get("currency", ""),
            "balance": account_finance["balance"],
            "equity": account_finance["equity"],
            "netDeposit": account_finance["netDeposit"],
            "holdingProfit": account_finance["holdingProfit"],
            "closedNetProfit": account_finance["closedNetProfit"],
            "adjustments": rounded(account_finance["negativeBalanceClear"] + account_finance["compensation"] + account_finance["reward"]),
            "rebate": account_finance["rebate"],
            "comprehensiveProfit": account_finance["comprehensiveProfit"],
            "highestHoldingVolume": account_finance["highestHoldingVolume"],
            "databaseStatus": database_statuses.get(item_key(item), ""),
            "localStatus": normalize_text(ledger.get("建议动作")) if ledger else "",
        }

    same_name_by_key: dict[tuple[str, str, str], dict] = {}
    current_items = [item for item in same_name_accounts if item_key(item) == (platform, server, normalize_text(login))]
    for item in current_items:
        same_name_by_key[item_key(item)] = build_same_name_row(item)
    other_items = [item for item in same_name_accounts if item_key(item) != (platform, server, normalize_text(login))]
    if len(other_items) == 1:
        item = other_items[0]
        try:
            same_name_by_key[item_key(item)] = build_same_name_row(item)
        except Exception:
            same_name_by_key[item_key(item)] = empty_same_name_row(item)
    elif other_items:
        with ThreadPoolExecutor(max_workers=min(8, len(other_items)), thread_name_prefix="same-name") as executor:
            futures = {executor.submit(build_same_name_row, item): item for item in other_items}
            for future in as_completed(futures):
                item = futures[future]
                try:
                    same_name_by_key[item_key(item)] = future.result()
                except Exception:
                    same_name_by_key[item_key(item)] = empty_same_name_row(item)
    same_name_rows = [same_name_by_key[item_key(item)] for item in same_name_accounts if item_key(item) in same_name_by_key]
    total_fields = ("balance", "equity", "netDeposit", "holdingProfit", "closedNetProfit", "adjustments", "rebate", "comprehensiveProfit")
    totals = {field: rounded(sum(numeric_value(row.get(field)) for row in same_name_rows)) for field in total_fields}
    return {
        "available": True,
        "server": server,
        "finance": finance,
        "highFrequency": riskdash_high_frequency(filtered_rows),
        "sameName": same_name_rows,
        "sameNameTotals": totals,
    }


def ledger_record_for_login(login: str, records: list[dict[str, str]] | None = None) -> dict[str, str] | None:
    login = normalize_text(login)
    records = records if records is not None else load_records()
    return next((record for record in records if normalize_text(record.get("账号")) == login), None)


PUBLIC_LEDGER_FIELDS = (
    "记录ID", "账号", "记录类型", "建议动作", "当前分组", "风险标签", "风险/问题备注",
    "加入时间", "修改时间", "状态", "处理人/来源", "AI风险等级", "AI备注", "AI分析时间", "AI证据图表",
)


def public_ledger_record(record: dict[str, str] | None) -> dict[str, str] | None:
    if not record:
        return None
    return {field: normalize_text(record.get(field)) for field in PUBLIC_LEDGER_FIELDS}


def public_history_row(row: dict[str, str]) -> dict[str, str]:
    return {
        field: normalize_text(row.get(field))
        for field in ("历史ID", "记录ID", "账号", "操作", "修改时间", "修改字段", "处理人/来源")
    }


def _account_database_detail_uncached(login: str, filters: dict | None = None) -> dict:
    filters = filters or {}
    login = normalize_text(login)
    if not login:
        return {"exists": False, "error": "请输入账号", "metrics": trade_metrics([])}
    active_filters = {key: normalize_text(filters.get(key)) for key in ("platform", "server", "symbol", "start", "end")}
    source_scope = {key: active_filters[key] for key in ("platform", "server")}
    detail_filters = {key: active_filters[key] for key in ("symbol", "start", "end")}
    if not any(source_scope.values()):
        candidates = account_lookup_databases(login)
        if len(candidates) > 1:
            return {
                "exists": False,
                "account": login,
                "requiresSourceSelection": True,
                "sourceCandidates": [
                    {
                        "platform": normalize_text((item.get("latestSource") or {}).get("platform")),
                        "server": normalize_text((item.get("latestSource") or {}).get("server")),
                        "orderCount": mysql_int(item.get("orderCount")),
                    }
                    for item in candidates
                ],
                "platforms": [],
                "servers": [],
                "metrics": trade_metrics([]),
                "error": "账号同时存在多个平台/服务器，请先选择查看来源",
                "refreshedAt": now_text(),
            }
    try:
        analysis = account_trade_analysis(login, **source_scope)
        history_limit = analysis["historyLimit"]
        all_rows = analysis["rows"]
        all_costs = analysis["costs"]
        uses_mysql = analysis["usesMysql"]
        filtered_rows = query_db_trades(login, limit=history_limit, **active_filters) if any(detail_filters.values()) else all_rows
    except Exception as exc:
        return {
            "exists": False,
            "error": str(exc),
            "dbSource": TRADE_DB_SOURCE,
            "metrics": trade_metrics([]),
            "refreshedAt": now_text(),
        }

    filtered_costs = query_mysql_trade_costs(login, **active_filters) if uses_mysql and any(detail_filters.values()) else all_costs
    all_metrics = analysis["metrics"]
    filtered_metrics = trade_metrics(filtered_rows, filtered_costs) if filtered_rows is not all_rows else all_metrics
    risk_panels = {"available": False, "loading": True, "reason": "正在加载风控面板..."}
    platforms = sorted({normalize_text(row.get("platform")) for row in all_rows if normalize_text(row.get("platform"))})
    servers = sorted({normalize_text(row.get("server")) for row in all_rows if normalize_text(row.get("server"))})
    symbols = sorted({normalize_text(row.get("symbol")) for row in all_rows if normalize_text(row.get("symbol"))})
    latest_row = max(all_rows, key=lambda row: normalize_text(row.get("close_time") or row.get("open_time"))) if all_rows else None
    known_source = None
    if not all_rows:
        known_source = next((
            item for item in account_lookup_databases(login)
            if source_allowed(item.get("latestSource") or {}, platform=source_scope["platform"], server=source_scope["server"])
        ), None)
    latest_source = (known_source or {}).get("latestSource") or {
        "platform": normalize_text(latest_row.get("platform")) if latest_row else "",
        "server": display_unknown(normalize_text(latest_row.get("server"))) if latest_row else "",
    }
    platforms = platforms or ([normalize_text(latest_source.get("platform"))] if normalize_text(latest_source.get("platform")) else [])
    servers = servers or ([normalize_text(latest_source.get("server"))] if normalize_text(latest_source.get("server")) else [])
    return {
        "exists": bool(all_rows),
        "account": login,
        "dbSource": "mysql" if uses_mysql else ("sqlite" if all_rows else normalize_text((known_source or {}).get("dbSource"))),
        "searchedSources": [source["name"] for source in MYSQL_SOURCES] + ([str(TRADE_DB_PATH)] if TRADE_DB_PATH.exists() else []),
        "orderCount": len(all_rows),
        "chartableOrderCount": all_metrics["chartableOrderCount"],
        "firstTime": all_metrics["firstTradeTime"],
        "lastTime": all_metrics["lastTradeTime"],
        "platforms": [{"value": value, "label": display_unknown(value)} for value in platforms],
        "servers": [{"value": value, "label": display_unknown(value)} for value in servers],
        "symbols": symbols,
        "latestSource": latest_source,
        "accountMeta": account_meta_for_rows(all_rows) if all_rows else (known_source or {}).get("accountMeta", account_money_meta()),
        "metrics": filtered_metrics,
        "allMetrics": all_metrics,
        "visualizations": trade_visualizations(filtered_rows, filtered_metrics),
        "riskPanels": risk_panels,
        "filters": active_filters,
        "error": normalize_text((known_source or {}).get("error")) if not all_rows else "",
        "refreshedAt": now_text(),
    }


def account_database_detail(login: str, filters: dict | None = None) -> dict:
    filters = filters or {}
    cache_key = (
        "detail",
        normalize_text(login),
        *(normalize_text(filters.get(key)) for key in ("platform", "server", "symbol", "start", "end")),
    )
    cached = account_cache_get(cache_key)
    if cached is not None:
        return cached
    return account_cache_set(cache_key, _account_database_detail_uncached(login, filters))


def account_risk_panels_payload(login: str, filters: dict | None = None) -> dict:
    filters = filters or {}
    active_filters = {key: normalize_text(filters.get(key)) for key in ("platform", "server", "symbol", "start", "end")}
    source_scope = {key: active_filters[key] for key in ("platform", "server")}
    cache_key = (
        "risk-panels", normalize_text(login),
        *(active_filters[key] for key in ("platform", "server", "symbol", "start", "end")),
    )
    cached = account_cache_get(cache_key)
    if cached is not None:
        return {"ok": True, "riskPanels": cached}
    try:
        analysis = account_trade_analysis(login, **source_scope)
        history_limit = analysis["historyLimit"]
        all_rows = analysis["rows"]
        if not all_rows:
            raise RuntimeError("账户暂未做单")
        filtered_rows = query_db_trades(login, limit=history_limit, **active_filters) if any(active_filters[key] for key in ("symbol", "start", "end")) else all_rows
        panels = build_riskdash_panels(login, all_rows, analysis["metrics"], filtered_rows)
    except Exception as exc:
        panels = {"available": False, "reason": str(exc)}
    account_cache_set(cache_key, panels)
    return {"ok": True, "riskPanels": panels}


def account_relationship_core_payload(login: str, filters: dict | None = None) -> dict:
    """Return only CRM account mapping facts needed by relationship expansion.

    This deliberately avoids ``account_trade_analysis`` and its full order-history cache.
    The full dashboard keeps using ``account_risk_panels_payload``; graph expansion only
    needs account identifiers, platform, and server for a same-CRM-user edge.
    """
    filters = filters or {}
    login = normalize_text(login)
    platform = normalize_text(filters.get("platform")).upper()
    server = normalize_text(filters.get("server"))
    if not login or not re.fullmatch(r"\d+", login):
        raise ValueError("账号格式无效")
    source = next((
        item for item in MYSQL_SOURCES
        if source_allowed(item, platform=platform, server=server)
    ), None)
    if not source:
        return {"ok": True, "riskPanels": {"available": False, "reason": "当前服务器尚未配置 CRM 关系数据源"}}
    source_rows = query_same_name_accounts(source, login)
    accounts = [login, *(normalize_text(item.get("account")) for item in source_rows)]
    try:
        status_query = query_mt5_database_statuses if source.get("kind") == "mt5_deals" else query_mt4_database_statuses
        database_statuses = status_query(source, accounts)
    except Exception:
        database_statuses = {}
    rows = []
    for item in source_rows:
        item_source = item.get("source") or {}
        account = normalize_text(item.get("account"))
        if not account:
            continue
        rows.append({
            "account": account,
            "platform": normalize_text(item_source.get("platform")),
            "server": normalize_text(item_source.get("server")),
            "databaseStatus": database_statuses.get(account, ""),
        })
    return {
        "ok": True,
        "riskPanels": {
            "available": True,
            "databaseStatus": database_statuses.get(login, ""),
            "sameName": rows,
        },
    }


def account_lookup_finance_payload(login: str, platform: str = "", server: str = "") -> dict:
    login = normalize_text(login)
    platform = normalize_text(platform).upper()
    server = normalize_text(server)
    if not login or not re.fullmatch(r"\d+", login):
        raise ValueError("账号格式无效")
    cache_key = ("lookup-finance", login, platform, server)
    cached = account_cache_get(cache_key)
    if cached is not None:
        return {"ok": True, **cached}
    analysis = account_trade_analysis(login, platform=platform, server=server)
    rows = analysis["rows"]
    row_servers = {normalize_text(row.get("server")) for row in rows if normalize_text(row.get("server"))}
    resolved_server = next(iter(row_servers)) if len(row_servers) == 1 else server
    source = next((
        item for item in MYSQL_SOURCES
        if source_allowed(item, platform=platform, server=resolved_server)
    ), None)
    if not source:
        raise ValueError("当前平台或服务器未配置")
    metrics = analysis["metrics"]
    if platform == "MT5":
        finance = query_mt5_finance_panel(source, login, rows, metrics)
        database_status = query_mt5_database_statuses(source, [login]).get(login, "")
    else:
        finance = query_mt4_finance_panel(source, login, rows, metrics)
        database_status = query_mt4_database_statuses(source, [login]).get(login, "")
    ledger = ledger_record_for_login(login)
    payload = {
        "account": login,
        "platform": platform,
        "server": server,
        "databaseStatus": database_status,
        "localStatus": normalize_text(ledger.get("建议动作")) if ledger else "",
        "workflowStatus": normalize_text(ledger.get("状态")) if ledger else "",
        "comprehensiveProfit": finance.get("comprehensiveProfit", 0),
        "currency": finance.get("displayCurrency") or finance.get("currency") or "",
        "refreshedAt": now_text(),
    }
    account_cache_set(cache_key, payload)
    return {"ok": True, **payload}


def account_detail_payload(login: str, filters: dict | None = None) -> dict:
    records = load_records()
    record = ledger_record_for_login(login, records)
    database = account_database_detail(login, filters)
    charts = [chart for chart in scan_chart_files(records) if normalize_text(chart.get("account")) == normalize_text(login)]
    history = []
    if record:
        history = [row for row in read_history_rows() if row["记录ID"] == record["记录ID"]]
        history.sort(key=lambda row: row["修改时间"], reverse=True)
    return {
        "ok": True,
        "account": normalize_text(login),
        "marked": bool(record),
        "record": public_ledger_record(record),
        "database": database,
        "charts": charts,
        "history": [public_history_row(row) for row in history],
        "actions": load_quick_actions(),
        "statuses": STATUS_CHOICES,
    }


def account_ledger_payload(login: str) -> dict:
    login = normalize_text(login)
    records = load_records()
    record = ledger_record_for_login(login, records)
    return {
        "ok": True,
        "account": login,
        "marked": bool(record),
        "record": public_ledger_record(record),
        "actions": load_quick_actions(),
        "statuses": STATUS_CHOICES,
    }


@contextmanager
def login_ip_db_connect():
    IP_HISTORY_DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(IP_HISTORY_DB_PATH, timeout=10)
    conn.row_factory = sqlite3.Row
    conn.execute("pragma journal_mode = wal")
    conn.execute("""
        create table if not exists account_login_ips (
            account text not null,
            platform text not null,
            server text not null,
            ip text not null,
            first_seen_at text not null,
            last_seen_at text not null,
            last_access_at text not null default '',
            primary key (account, platform, server, ip)
        )
    """)
    conn.execute("""
        create table if not exists ip_geo_cache (
            ip text primary key,
            country text not null default '',
            region text not null default '',
            city text not null default '',
            isp text not null default '',
            asn text not null default '',
            latitude real,
            longitude real,
            status text not null default '',
            checked_at text not null
        )
    """)
    try:
        yield conn
        conn.commit()
    finally:
        conn.close()


def normalize_ip_address(value: object) -> tuple[str, str]:
    text = normalize_text(value).strip()
    if not text:
        return "", "missing"
    if "," in text:
        text = text.split(",", 1)[0].strip()
    try:
        address = ipaddress.ip_address(text)
    except ValueError:
        return text[:128], "invalid"
    return str(address), "public" if address.is_global else "private"


def public_geo_record(row: dict | sqlite3.Row | None) -> dict:
    if not row:
        return {"status": "unavailable", "country": "", "region": "", "city": "", "isp": "", "asn": ""}
    return {
        "status": normalize_text(row["status"]),
        "country": normalize_text(row["country"]),
        "region": normalize_text(row["region"]),
        "city": normalize_text(row["city"]),
        "isp": normalize_text(row["isp"]),
        "asn": normalize_text(row["asn"]),
        "latitude": row["latitude"],
        "longitude": row["longitude"],
        "checkedAt": normalize_text(row["checked_at"]),
    }


def geo_cache_is_fresh(row: sqlite3.Row | None) -> bool:
    if not row:
        return False
    checked = parse_trade_time(row["checked_at"])
    if not checked:
        return False
    ttl = timedelta(hours=1) if normalize_text(row["status"]) == "unavailable" else timedelta(days=IP_GEO_CACHE_DAYS)
    return datetime.now() - checked <= ttl


def query_public_ip_geo(ip: str) -> dict:
    url = IP_GEO_API_TEMPLATE.format(ip=quote(ip, safe=""))
    request = Request(url, headers={"User-Agent": "AccountRegistry/1.0", "Accept": "application/json"})
    with urlopen(request, timeout=IP_GEO_TIMEOUT) as response:
        payload = json.loads(response.read().decode("utf-8", errors="replace"))
    if payload.get("success") is False:
        raise RuntimeError(normalize_text(payload.get("message")) or "IP 查询失败")
    connection = payload.get("connection") if isinstance(payload.get("connection"), dict) else {}
    return {
        "country": normalize_text(payload.get("country")),
        "region": normalize_text(payload.get("region")),
        "city": normalize_text(payload.get("city")),
        "isp": normalize_text(connection.get("isp") or connection.get("org")),
        "asn": normalize_text(connection.get("asn")),
        "latitude": payload.get("latitude"),
        "longitude": payload.get("longitude"),
        "status": "ok",
    }


def cached_ip_geo(ip: str, address_kind: str = "public") -> dict:
    with IP_HISTORY_LOCK:
        with login_ip_db_connect() as conn:
            cached = conn.execute("select * from ip_geo_cache where ip = ?", (ip,)).fetchone()
            if geo_cache_is_fresh(cached):
                return public_geo_record(cached)

    if address_kind != "public":
        geo = {
            "country": "",
            "region": "",
            "city": "",
            "isp": "内网或保留地址" if address_kind == "private" else "IP 格式无效",
            "asn": "",
            "latitude": None,
            "longitude": None,
            "status": address_kind,
        }
    else:
        try:
            geo = query_public_ip_geo(ip)
        except Exception:
            geo = {"country": "", "region": "", "city": "", "isp": "", "asn": "", "latitude": None, "longitude": None, "status": "unavailable"}

    checked_at = now_text()
    with IP_HISTORY_LOCK:
        with login_ip_db_connect() as conn:
            conn.execute("""
                insert into ip_geo_cache (ip, country, region, city, isp, asn, latitude, longitude, status, checked_at)
                values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                on conflict(ip) do update set
                    country=excluded.country, region=excluded.region, city=excluded.city,
                    isp=excluded.isp, asn=excluded.asn, latitude=excluded.latitude,
                    longitude=excluded.longitude, status=excluded.status, checked_at=excluded.checked_at
            """, (
                ip, geo["country"], geo["region"], geo["city"], geo["isp"], geo["asn"],
                geo["latitude"], geo["longitude"], geo["status"], checked_at,
            ))
            saved = conn.execute("select * from ip_geo_cache where ip = ?", (ip,)).fetchone()
    return public_geo_record(saved)


def record_login_ip_observation(account: str, platform: str, server: str, ip: str, last_access_at: str = "") -> None:
    observed_at = now_text()
    with IP_HISTORY_LOCK:
        with login_ip_db_connect() as conn:
            conn.execute("""
                insert into account_login_ips (account, platform, server, ip, first_seen_at, last_seen_at, last_access_at)
                values (?, ?, ?, ?, ?, ?, ?)
                on conflict(account, platform, server, ip) do update set
                    last_seen_at=excluded.last_seen_at,
                    last_access_at=case when excluded.last_access_at <> '' then excluded.last_access_at else account_login_ips.last_access_at end
            """, (account, platform, server, ip, observed_at, observed_at, last_access_at))


def read_login_ip_history(account: str) -> list[dict]:
    with IP_HISTORY_LOCK:
        with login_ip_db_connect() as conn:
            rows = conn.execute("""
                select h.*, g.country, g.region, g.city, g.isp, g.asn, g.latitude, g.longitude,
                       g.status as geo_status, g.checked_at as geo_checked_at
                from account_login_ips h
                left join ip_geo_cache g on g.ip = h.ip
                where h.account = ?
                order by h.last_seen_at desc, h.server, h.ip
            """, (account,)).fetchall()
    return [
        {
            "platform": normalize_text(row["platform"]),
            "server": normalize_text(row["server"]),
            "ip": normalize_text(row["ip"]),
            "firstSeenAt": normalize_text(row["first_seen_at"]),
            "lastSeenAt": normalize_text(row["last_seen_at"]),
            "lastAccessAt": normalize_text(row["last_access_at"]),
            "geo": {
                "status": normalize_text(row["geo_status"]) or "unavailable",
                "country": normalize_text(row["country"]),
                "region": normalize_text(row["region"]),
                "city": normalize_text(row["city"]),
                "isp": normalize_text(row["isp"]),
                "asn": normalize_text(row["asn"]),
                "latitude": row["latitude"],
                "longitude": row["longitude"],
                "checkedAt": normalize_text(row["geo_checked_at"]),
            },
        }
        for row in rows
    ]


def account_login_ips_payload(login: str) -> dict:
    login = normalize_text(login)
    if not login or len(login) > 64 or not re.fullmatch(r"[0-9A-Za-z_.-]+", login):
        raise ValueError("账号格式无效")
    sources = []
    current_ips: dict[str, str] = {}
    for source in MYSQL_SOURCES:
        platform = normalize_text(source.get("platform"))
        server = normalize_text(source.get("server") or source.get("name"))
        kind = normalize_text(source.get("kind"))
        if not login.isdigit():
            sources.append({"platform": platform, "server": server, "available": False, "reason": "当前数据源仅支持数字账号"})
            continue
        if kind == "mt4_trades":
            try:
                with mysql_trade_connect(source) as conn:
                    with conn.cursor() as cur:
                        cur.execute(
                            f"select LOGIN, LASTDATE from `{source['schema']}`.`mt4_users_view` where LOGIN = %s limit 1",
                            (int(login),),
                        )
                        row = cur.fetchone() or {}
            except Exception:
                sources.append({"platform": platform, "server": server, "available": False, "accountExists": False, "reason": "MT4 账户数据源暂时不可用"})
                continue
            if not row:
                sources.append({"platform": platform, "server": server, "available": False, "accountExists": False, "reason": "账号不在该数据源"})
                continue
            last_access = trade_time_text(parse_trade_time(row.get("LASTDATE"))) or normalize_text(row.get("LASTDATE"))
            sources.append({
                "platform": platform,
                "server": server,
                "available": False,
                "accountExists": True,
                "lastAccessAt": last_access,
                "reason": "MT4 当前导出库未包含登录 IP 字段",
            })
            continue
        if kind != "mt5_deals":
            sources.append({"platform": platform, "server": server, "available": False, "accountExists": False, "reason": "当前数据源没有登录 IP 字段"})
            continue
        try:
            with mysql_trade_connect(source) as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        f"select Login, LastIP, LastAccess from `{source['schema']}`.`mt5_users_view` where Login = %s limit 1",
                        (int(login),),
                    )
                    row = cur.fetchone() or {}
        except Exception:
            sources.append({"platform": platform, "server": server, "available": False, "reason": "登录 IP 数据源暂时不可用"})
            continue
        if not row:
            sources.append({"platform": platform, "server": server, "available": False, "reason": "账号不在该数据源"})
            continue
        ip, address_kind = normalize_ip_address(row.get("LastIP"))
        last_access = trade_time_text(parse_trade_time(row.get("LastAccess"))) or normalize_text(row.get("LastAccess"))
        if ip:
            record_login_ip_observation(login, platform, server, ip, last_access)
            current_ips[ip] = address_kind
        sources.append({
            "platform": platform,
            "server": server,
            "available": bool(ip),
            "accountExists": True,
            "ip": ip,
            "lastAccessAt": last_access,
            "reason": "" if ip else "数据库暂无登录 IP",
        })

    for ip, address_kind in current_ips.items():
        cached_ip_geo(ip, address_kind)
    records = read_login_ip_history(login)
    observed_since = min((row["firstSeenAt"] for row in records if row.get("firstSeenAt")), default="")
    return {
        "ok": True,
        "account": login,
        "coverage": {
            "type": "observed_since",
            "observedSince": observed_since,
            "notice": "MT5 数据库仅提供最后登录 IP，本地历史自功能上线后开始累计；MT4 当前导出未包含 IP，只能显示最近登录时间。",
        },
        "geoProvider": "ipwho.is",
        "externalLookup": True,
        "sources": sources,
        "records": records,
        "refreshedAt": now_text(),
    }


def account_shared_last_ip_payload(login: str, filters: dict | None = None) -> dict:
    """Read-only same-server current-LastIP relation for graph expansion."""
    filters = filters or {}
    login = normalize_text(login)
    platform = normalize_text(filters.get("platform")).upper()
    server = normalize_text(filters.get("server"))
    if not login.isdigit() or platform != "MT5" or not server:
        return {"ok": True, "peers": [], "coverage": [{"source": "sharedLastIp", "status": "skipped", "reason": "仅支持已路由 MT5 账户"}]}
    source = next((item for item in MYSQL_SOURCES if source_allowed(item, platform=platform, server=server) and item.get("kind") == "mt5_deals"), None)
    if not source:
        return {"ok": True, "peers": [], "coverage": [{"source": "sharedLastIp", "status": "skipped", "reason": "当前服务器没有 MT5 LastIP 数据源"}]}
    try:
        requested_timeout = float(filters.get("relationshipQueryTimeoutSeconds") or 3)
    except (TypeError, ValueError):
        requested_timeout = 3
    query_timeout_seconds = min(max(int(requested_timeout + 0.999), 1), 3)
    try:
        with mysql_trade_connect(source, connect_timeout=query_timeout_seconds, read_timeout=query_timeout_seconds) as conn:
            with conn.cursor() as cur:
                cur.execute(f"SELECT LastIP FROM `{source['schema']}`.mt5_users_view WHERE Login = %s LIMIT 1", (int(login),))
                target = cur.fetchone() or {}
                ip, _kind = normalize_ip_address(target.get("LastIP"))
                if not ip:
                    return {"ok": True, "peers": [], "coverage": [{"source": "sharedLastIp", "status": "available", "reason": "当前账号没有 LastIP"}]}
                cur.execute(
                    f"SELECT Login, LastIP, LastAccess FROM `{source['schema']}`.mt5_users_view WHERE LastIP = %s AND Login <> %s LIMIT 200",
                    (ip, int(login)),
                )
                rows = cur.fetchall() or []
    except Exception as exc:
        return {"ok": True, "peers": [], "coverage": [{"source": "sharedLastIp", "status": "failed", "reason": str(exc)}]}
    return {
        "ok": True,
        "peers": [{"account": str(row.get("Login")), "platform": "MT5", "server": server, "ip": ip, "lastAccessAt": normalize_text(row.get("LastAccess"))} for row in rows if row.get("Login")],
        "coverage": [{"source": "sharedLastIp", "status": "available", "reason": ""}],
    }


def public_trade_order(row: dict) -> dict:
    profit = numeric_value(row.get("profit"))
    commission = numeric_value(row.get("commission")) + numeric_value(row.get("fee"))
    swap = numeric_value(row.get("swap"))
    taxes = numeric_value(row.get("taxes"))
    return {
        "ticket": normalize_text(row.get("ticket")),
        "platform": normalize_text(row.get("platform")),
        "server": display_unknown(normalize_text(row.get("server"))),
        "symbol": normalize_text(row.get("symbol")),
        "type": clean_trade_type(row.get("type", "")),
        "reason": normalize_text(row.get("reason")),
        "comment": normalize_text(row.get("comment")),
        "isCopyTrade": is_copy_trade(row),
        "expertId": normalize_text(row.get("expert_id")),
        "openTime": trade_time_text(parse_trade_time(row.get("open_time"))),
        "closeTime": trade_time_text(parse_trade_time(row.get("close_time"))),
        "holdingSeconds": rounded(numeric_value(row.get("holding_seconds")), 1),
        "volume": rounded(numeric_value(row.get("volume")), 4),
        "profit": rounded(profit),
        "commission": rounded(commission),
        "swap": rounded(swap),
        "taxes": rounded(taxes),
        "netProfit": rounded(profit + commission + swap + taxes),
        "currency": normalize_text(row.get("account_currency")),
        "displayCurrency": normalize_text(row.get("display_currency")),
        "isCentAccount": bool(row.get("is_cent_account")),
    }


def account_orders_payload(
    login: str,
    page: int = 1,
    page_size: int = 100,
    platform: str = "",
    server: str = "",
) -> dict:
    login = normalize_text(login)
    if not login or len(login) > 64 or not re.fullmatch(r"[0-9A-Za-z_.-]+", login):
        raise ValueError("账号格式无效")
    page = max(int(page or 1), 1)
    page_size = min(max(int(page_size or 100), 20), 200)
    platform = normalize_text(platform).upper()
    server = normalize_text(server)
    sources = [source for source in MYSQL_SOURCES if source_allowed(source, platform=platform, server=server)]
    if platform == "MT4" and len(sources) == 1 and sources[0].get("kind") == "mt4_trades":
        cache_key = ("orders-page", login, platform, server, page, page_size)
        cached = account_cache_get(cache_key)
        if cached is not None:
            return cached
        total, rows = query_mysql_mt4_orders_page_source(sources[0], login, page, page_size)
        payload = {
            "ok": True,
            "account": login,
            "page": page,
            "pageSize": page_size,
            "total": total,
            "pages": max((total + page_size - 1) // page_size, 1),
            "truncated": False,
            "orders": [public_trade_order(row) for row in rows],
        }
        return account_cache_set(cache_key, payload)
    rows = query_db_trades(login, limit=50000)
    rows.sort(
        key=lambda row: parse_trade_time(row.get("close_time") or row.get("open_time")) or datetime.min,
        reverse=True,
    )
    total = len(rows)
    start = (page - 1) * page_size
    return {
        "ok": True,
        "account": login,
        "page": page,
        "pageSize": page_size,
        "total": total,
        "pages": max((total + page_size - 1) // page_size, 1),
        "truncated": total >= 50000,
        "orders": [public_trade_order(row) for row in rows[start:start + page_size]],
    }


def toxic_ramp(value: object, low: float, high: float, reverse: bool = False) -> float:
    number = numeric_value(value)
    if high <= low:
        return 0.0
    score = max(0.0, min(1.0, (number - low) / (high - low)))
    return (1.0 - score if reverse else score) * 100.0


def market_pushing_consistency_bonus(
    structure_score: object,
    tick_score: object,
    counterevidence_deduction: object,
) -> float:
    """Reward strong independent structure and Tick evidence without making either a floor."""
    structure = numeric_value(structure_score)
    tick = numeric_value(tick_score)
    counterevidence = numeric_value(counterevidence_deduction)
    if structure < 35 or tick < 15 or counterevidence > 5:
        return 0.0
    return min(12.0, 5.0 + (structure - 35.0) * 1.5 + (tick - 15.0) * 0.5)


def toxic_level(score: object) -> str:
    value = numeric_value(score)
    if value >= 90:
        return "严重形态"
    if value >= 75:
        return "高危形态"
    if value >= 60:
        return "预警"
    if value >= 40:
        return "关注"
    return "无明显风险"


def toxic_trade_holding(row: dict) -> float:
    value = numeric_value(row.get("holding_seconds"))
    if value > 0:
        return value
    opened = parse_trade_time(row.get("open_time_msc") or row.get("open_time"))
    closed = parse_trade_time(row.get("close_time_msc") or row.get("close_time"))
    return max((closed - opened).total_seconds(), 0.0) if opened and closed else 0.0


def toxic_trade_net(row: dict) -> float:
    return sum(numeric_value(row.get(key)) for key in ("profit", "commission", "fee", "swap", "taxes"))


def toxic_volume_ratio(rows: list[dict], predicate) -> float:
    total = sum(max(numeric_value(row.get("volume")), 0.0) for row in rows)
    if not total:
        return 0.0
    selected = sum(max(numeric_value(row.get("volume")), 0.0) for row in rows if predicate(row))
    return selected / total


def toxic_burst(rows: list[dict], seconds: int) -> dict:
    timed = []
    for row in rows:
        opened = parse_trade_time(row.get("open_time_msc") or row.get("open_time"))
        if opened and clean_trade_type(row.get("type", "")) in {"buy", "sell"}:
            timed.append((opened, row))
    timed.sort(key=lambda item: item[0])
    best = {"count": 0, "volume": 0.0, "direction": 0.0, "start": "", "end": "", "tickets": []}
    left = 0
    buy_volume = sell_volume = volume = 0.0
    for right, (opened, row) in enumerate(timed):
        item_volume = max(numeric_value(row.get("volume")), 0.0)
        volume += item_volume
        if clean_trade_type(row.get("type", "")) == "buy":
            buy_volume += item_volume
        else:
            sell_volume += item_volume
        while left <= right and (opened - timed[left][0]).total_seconds() > seconds:
            old = timed[left][1]
            old_volume = max(numeric_value(old.get("volume")), 0.0)
            volume -= old_volume
            if clean_trade_type(old.get("type", "")) == "buy":
                buy_volume -= old_volume
            else:
                sell_volume -= old_volume
            left += 1
        count = right - left + 1
        direction = max(buy_volume, sell_volume) / volume if volume else 0.0
        if (volume, count) > (best["volume"], best["count"]):
            window_rows = [item[1] for item in timed[left:right + 1]]
            best = {
                "count": count,
                "volume": rounded(volume, 4),
                "direction": rounded(direction * 100, 1),
                "start": trade_time_text(timed[left][0]),
                "end": trade_time_text(opened),
                "tickets": [normalize_text(item.get("ticket")) for item in window_rows[:12]],
            }
    return best


def toxic_best_opposite_pair(rows: list[dict]) -> dict:
    by_symbol: dict[str, list[dict]] = defaultdict(list)
    for row in rows:
        if clean_trade_type(row.get("type", "")) in {"buy", "sell"}:
            by_symbol[normalize_text(row.get("symbol"))].append(row)
    best = {"found": False, "match": 0.0, "seconds": None, "volume": 0.0, "overlap": 0.0, "tickets": [], "symbol": ""}
    for symbol, items in by_symbol.items():
        items.sort(key=lambda row: parse_trade_time(row.get("open_time_msc") or row.get("open_time")) or datetime.min)
        for index, first in enumerate(items):
            first_open = parse_trade_time(first.get("open_time_msc") or first.get("open_time"))
            if not first_open:
                continue
            for second in items[index + 1:]:
                second_open = parse_trade_time(second.get("open_time_msc") or second.get("open_time"))
                if not second_open:
                    continue
                gap = abs((second_open - first_open).total_seconds())
                if gap > 60:
                    break
                if clean_trade_type(first.get("type", "")) == clean_trade_type(second.get("type", "")):
                    continue
                first_volume = max(numeric_value(first.get("volume")), 0.0)
                second_volume = max(numeric_value(second.get("volume")), 0.0)
                match = min(first_volume, second_volume) / max(first_volume, second_volume, 1e-9)
                first_close = parse_trade_time(first.get("close_time_msc") or first.get("close_time"))
                second_close = parse_trade_time(second.get("close_time_msc") or second.get("close_time"))
                overlap = 0.0
                if first_close and second_close:
                    overlap_start = max(first_open, second_open)
                    overlap_end = min(first_close, second_close)
                    shorter = min((first_close - first_open).total_seconds(), (second_close - second_open).total_seconds())
                    overlap = max((overlap_end - overlap_start).total_seconds(), 0.0) / max(shorter, 1e-9)
                strength = match * (first_volume + second_volume)
                current_strength = numeric_value(best.get("match")) / 100 * numeric_value(best.get("volume"))
                if strength > current_strength:
                    best = {
                        "found": True,
                        "match": rounded(match * 100, 1),
                        "seconds": rounded(gap, 3),
                        "volume": rounded(first_volume + second_volume, 4),
                        "overlap": rounded(overlap * 100, 1),
                        "tickets": [normalize_text(first.get("ticket")), normalize_text(second.get("ticket"))],
                        "symbol": symbol,
                    }
    return best


def toxic_base_symbol(value: object) -> str:
    return normalize_text(value).upper().split(".")[0].replace("ROLL", "")


def toxic_filter_push_orders(rows: list[dict]) -> dict:
    tradable = [
        row for row in rows
        if clean_trade_type(row.get("type", "")) in {"buy", "sell"}
        and parse_trade_time(row.get("open_time_msc") or row.get("open_time"))
        and numeric_value(row.get("volume")) > 0
    ]
    positive_volumes = sorted(max(numeric_value(row.get("volume")), 0.0) for row in tradable)
    median_volume = statistics.median(positive_volumes) if positive_volumes else 0.0
    filtered_ids: set[int] = set()
    volume_cutoffs: dict[str, float] = {}
    by_symbol: dict[str, list[dict]] = defaultdict(list)
    for row in tradable:
        by_symbol[toxic_base_symbol(row.get("symbol")) or normalize_text(row.get("symbol")) or "-"].append(row)
    for symbol, symbol_rows in by_symbol.items():
        if len(symbol_rows) < 5:
            filtered_ids.update(id(row) for row in symbol_rows)
            volume_cutoffs[symbol] = 0.0
            continue
        ranked = sorted(symbol_rows, key=lambda row: numeric_value(row.get("volume")), reverse=True)
        total_symbol_volume = sum(max(numeric_value(row.get("volume")), 0.0) for row in ranked)
        target_volume = total_symbol_volume * 0.95
        running_volume = 0.0
        cutoff = numeric_value(ranked[-1].get("volume"))
        for index, row in enumerate(ranked):
            running_volume += max(numeric_value(row.get("volume")), 0.0)
            cutoff = numeric_value(row.get("volume"))
            if running_volume >= target_volume and index + 1 >= 5:
                break
        selected = [row for row in symbol_rows if numeric_value(row.get("volume")) >= cutoff]
        if len(selected) < 5:
            selected = ranked[:5]
            cutoff = numeric_value(selected[-1].get("volume"))
        filtered_ids.update(id(row) for row in selected)
        volume_cutoffs[symbol] = rounded(cutoff, 4)
    filtered = [row for row in tradable if id(row) in filtered_ids]
    excluded = [row for row in tradable if id(row) not in filtered_ids]
    raw_volume = sum(max(numeric_value(row.get("volume")), 0.0) for row in tradable)
    filtered_volume = sum(max(numeric_value(row.get("volume")), 0.0) for row in filtered)
    excluded_volume = max(raw_volume - filtered_volume, 0.0)
    return {
        "rows": filtered,
        "excludedRows": excluded,
        "rawOrderCount": len(tradable),
        "filteredOrderCount": len(filtered),
        "excludedOrderCount": len(excluded),
        "rawVolume": rounded(raw_volume, 4),
        "filteredVolume": rounded(filtered_volume, 4),
        "excludedVolume": rounded(excluded_volume, 4),
        "excludedVolumeRatio": rounded(excluded_volume / raw_volume * 100, 1) if raw_volume else 0,
        "medianVolume": rounded(median_volume, 4),
        "volumeCutoff": min(volume_cutoffs.values(), default=0.0),
        "volumeCutoffs": volume_cutoffs,
        "targetVolumeCoverage": 95.0,
    }


def toxic_dynamic_push_sessions(rows: list[dict]) -> dict:
    timed = sorted(
        (
            (
                parse_trade_time(row.get("open_time_msc") or row.get("open_time")),
                parse_trade_time(row.get("close_time_msc") or row.get("close_time")),
                row,
            )
            for row in rows
        ),
        key=lambda item: item[0] or datetime.min,
    )
    timed = [item for item in timed if item[0]]
    if not timed:
        return {
            "sessions": [], "sessionGapsHours": [], "sessionBreakMinutes": 0,
            "typicalOrderGapMinutes": 0, "quietGapScore": 0, "singleSessionOnly": False,
        }
    micro_batches: list[list[tuple[datetime, datetime | None, dict]]] = []
    for opened, closed, row in timed:
        if not micro_batches or (opened - micro_batches[-1][-1][0]).total_seconds() > 5:
            micro_batches.append([])
        micro_batches[-1].append((opened, closed, row))
    batch_gaps = [
        max((current[0][0] - previous[-1][0]).total_seconds(), 0.0)
        for previous, current in zip(micro_batches, micro_batches[1:])
    ]
    positive_gaps = sorted(gap for gap in batch_gaps if gap > 0)
    intra_batch_gaps = [
        max((current[0] - previous[0]).total_seconds(), 0.0)
        for batch in micro_batches for previous, current in zip(batch, batch[1:])
        if (current[0] - previous[0]).total_seconds() > 0
    ]
    holding_seconds = [
        max((closed - opened).total_seconds(), 0.0)
        for opened, closed, _ in timed if closed and closed >= opened
    ]
    fallback_gap = statistics.median(intra_batch_gaps) if intra_batch_gaps else (statistics.median(holding_seconds) if holding_seconds else 60.0)
    fallback_gap = max(fallback_gap, 1.0)
    if positive_gaps:
        lower_count = max(1, math.ceil(len(positive_gaps) * 0.6))
        typical_gap = min(statistics.median(positive_gaps[:lower_count]), fallback_gap) if len(positive_gaps) == 1 else statistics.median(positive_gaps[:lower_count])
        best_jump = 0.0
        jump_threshold = None
        for lower, upper in zip(positive_gaps, positive_gaps[1:]):
            ratio = upper / max(lower, 1.0)
            if upper >= 60 and ratio > best_jump:
                best_jump = ratio
                jump_threshold = math.sqrt(max(lower, 1.0) * upper)
        session_break = jump_threshold if best_jump >= 3 and jump_threshold is not None else typical_gap * 8
        session_break = max(60.0, min(session_break, max(typical_gap * 30, 60.0)))
    else:
        typical_gap = fallback_gap
        session_break = max(60.0, typical_gap * 8)
    sessions: list[list[tuple[datetime, dict]]] = []
    session_closes: list[list[datetime]] = []
    for batch in micro_batches:
        opened = batch[0][0]
        if not sessions or (opened - sessions[-1][-1][0]).total_seconds() > session_break:
            sessions.append([])
            session_closes.append([])
        for item_opened, item_closed, row in batch:
            sessions[-1].append((item_opened, row))
            if item_closed:
                session_closes[-1].append(item_closed)
    active_gaps = [gap for gap in batch_gaps if 0 < gap <= session_break]
    if active_gaps:
        typical_gap = statistics.median(active_gaps)
    session_gaps = []
    for index in range(1, len(sessions)):
        previous_end = max(session_closes[index - 1], default=sessions[index - 1][-1][0])
        session_gaps.append(max((sessions[index][0][0] - previous_end).total_seconds(), 0.0))
    if session_gaps:
        gap_scores = [toxic_ramp(gap / max(typical_gap, 1.0), 6, 30) for gap in session_gaps]
        quiet_score = statistics.mean(gap_scores)
    else:
        quiet_score = 0.0
    return {
        "sessions": sessions,
        "sessionGapsHours": [gap / 3600 for gap in session_gaps],
        "sessionBreakMinutes": rounded(session_break / 60, 1),
        "typicalOrderGapMinutes": rounded(typical_gap / 60, 2),
        "quietGapScore": rounded(quiet_score, 1),
        "singleSessionOnly": len(sessions) == 1 and len(timed) >= 8,
        "quietEvidenceAvailable": bool(session_gaps),
    }


def toxic_push_has_session_evidence(push_behavior: dict) -> bool:
    return mysql_int(push_behavior.get("concentratedSessions")) >= 2 or bool(push_behavior.get("singleSessionOnly"))


def toxic_push_position_campaigns(sessions: list[list[tuple[datetime, dict]]]) -> dict:
    """Measure staged position building inside dynamic trading campaigns.

    A push campaign can enter the same direction over tens of seconds or several
    minutes and still close the whole position together.  Treating every gap over
    five seconds as a new batch incorrectly turns that pattern into grid adding.
    Here the dynamic session is the campaign; cohesive exits protect staged
    entries, while non-cohesive campaigns are split into entry waves using the
    account's own typical entry gap.
    """
    staggered_addon_orders = 0
    staggered_addon_volume = 0.0
    cohesive_campaign_orders = 0
    cohesive_wave_orders = 0
    multi_order_campaign_orders = 0
    cohesive_campaigns = 0
    entry_wave_count = 0
    close_tolerances = []
    for items in sessions:
        campaign = []
        for opened, row in sorted(items, key=lambda item: item[0]):
            closed = parse_trade_time(row.get("close_time_msc") or row.get("close_time"))
            campaign.append((opened, closed, row))
        if not campaign:
            continue
        if len(campaign) >= 2:
            multi_order_campaign_orders += len(campaign)
        open_gaps = [
            max((campaign[index][0] - campaign[index - 1][0]).total_seconds(), 0.0)
            for index in range(1, len(campaign))
        ]
        positive_open_gaps = [gap for gap in open_gaps if gap > 0]
        typical_open_gap = statistics.median(positive_open_gaps) if positive_open_gaps else 0.0
        holding_seconds = [
            max((closed - opened).total_seconds(), 0.0)
            for opened, closed, _ in campaign
            if closed and closed >= opened
        ]
        median_holding = statistics.median(holding_seconds) if holding_seconds else 0.0
        close_tolerance = min(120.0, max(10.0, typical_open_gap * 1.5, median_holding * .05))
        close_tolerances.append(close_tolerance)
        closes = [closed for _, closed, _ in campaign if closed]
        close_span = (max(closes) - min(closes)).total_seconds() if len(closes) == len(campaign) and closes else math.inf
        directions = defaultdict(int)
        symbols = defaultdict(int)
        for _, _, row in campaign:
            directions[clean_trade_type(row.get("type", ""))] += 1
            symbols[toxic_base_symbol(row.get("symbol")) or "-"] += 1
        dominant_direction_ratio = max(directions.values(), default=0) / len(campaign) * 100
        dominant_symbol_ratio = max(symbols.values(), default=0) / len(campaign) * 100
        cohesive_campaign = bool(
            len(campaign) >= 2
            and len(closes) == len(campaign)
            and close_span <= close_tolerance
            and dominant_direction_ratio >= 70
            and dominant_symbol_ratio >= 70
        )
        if cohesive_campaign:
            cohesive_campaign_orders += len(campaign)
            cohesive_wave_orders += len(campaign)
            cohesive_campaigns += 1
            entry_wave_count += 1
            continue

        entry_wave_gap = min(60.0, max(2.0, typical_open_gap * 3 if typical_open_gap else 10.0))
        entry_waves: list[list[tuple[datetime, datetime | None, dict]]] = []
        for opened, closed, row in campaign:
            if not entry_waves or (opened - entry_waves[-1][-1][0]).total_seconds() > entry_wave_gap:
                entry_waves.append([])
            entry_waves[-1].append((opened, closed, row))
        entry_wave_count += len(entry_waves)
        previous_latest_close: datetime | None = None
        for wave in entry_waves:
            wave_start = wave[0][0]
            if previous_latest_close and wave_start < previous_latest_close:
                staggered_addon_orders += len(wave)
                staggered_addon_volume += sum(max(numeric_value(row.get("volume")), 0.0) for _, _, row in wave)
            wave_closes = [closed for _, closed, _ in wave if closed]
            wave_open_gaps = [
                max((wave[index][0] - wave[index - 1][0]).total_seconds(), 0.0)
                for index in range(1, len(wave))
            ]
            wave_positive_gaps = [gap for gap in wave_open_gaps if gap > 0]
            wave_typical_gap = statistics.median(wave_positive_gaps) if wave_positive_gaps else 0.0
            wave_holding = [
                max((closed - opened).total_seconds(), 0.0)
                for opened, closed, _ in wave
                if closed and closed >= opened
            ]
            wave_median_holding = statistics.median(wave_holding) if wave_holding else 0.0
            wave_close_tolerance = min(120.0, max(10.0, wave_typical_gap * 1.5, wave_median_holding * .05))
            wave_close_span = (max(wave_closes) - min(wave_closes)).total_seconds() if len(wave_closes) == len(wave) and wave_closes else math.inf
            wave_directions = defaultdict(int)
            wave_symbols = defaultdict(int)
            for _, _, row in wave:
                wave_directions[clean_trade_type(row.get("type", ""))] += 1
                wave_symbols[toxic_base_symbol(row.get("symbol")) or "-"] += 1
            if (
                len(wave) >= 2
                and len(wave_closes) == len(wave)
                and wave_close_span <= wave_close_tolerance
                and max(wave_directions.values(), default=0) / len(wave) * 100 >= 70
                and max(wave_symbols.values(), default=0) / len(wave) * 100 >= 70
            ):
                cohesive_wave_orders += len(wave)
            if wave_closes:
                wave_latest_close = max(wave_closes)
                if previous_latest_close is None or wave_latest_close > previous_latest_close:
                    previous_latest_close = wave_latest_close
    return {
        "campaignCount": len(sessions),
        "entryWaveCount": entry_wave_count,
        "multiOrderCampaignOrders": multi_order_campaign_orders,
        "cohesiveCampaignOrders": cohesive_campaign_orders,
        "cohesiveWaveOrders": cohesive_wave_orders,
        "cohesiveCampaigns": cohesive_campaigns,
        "staggeredAddOnOrders": staggered_addon_orders,
        "staggeredAddOnVolume": staggered_addon_volume,
        "closeToleranceSecondsMedian": rounded(statistics.median(close_tolerances), 1) if close_tolerances else 0,
    }


def toxic_push_behavior(rows: list[dict], order_filter: dict | None = None) -> dict:
    order_filter = order_filter or toxic_filter_push_orders(rows)
    ordered = sorted(order_filter["rows"], key=lambda row: parse_trade_time(row.get("open_time_msc") or row.get("open_time")) or datetime.min)
    total = len(ordered)
    if not total:
        return {"orders": 0, **{key: value for key, value in order_filter.items() if key not in {"rows", "excludedRows"}}}
    symbols: dict[str, int] = defaultdict(int)
    day_rows: dict[str, list[dict]] = defaultdict(list)
    nets = []
    total_volume = 0.0
    loss_volume = 0.0
    max_run = run = 0
    previous_direction = ""
    non_overlapping = 0
    latest_close: datetime | None = None
    for row in ordered:
        symbols[toxic_base_symbol(row.get("symbol")) or "-"] += 1
        opened = parse_trade_time(row.get("open_time_msc") or row.get("open_time"))
        closed = parse_trade_time(row.get("close_time_msc") or row.get("close_time"))
        day_rows[opened.strftime("%Y-%m-%d") if opened else "-"] .append(row)
        if opened and (latest_close is None or opened >= latest_close):
            non_overlapping += 1
        if closed and (latest_close is None or closed > latest_close):
            latest_close = closed
        direction = clean_trade_type(row.get("type", ""))
        run = run + 1 if direction == previous_direction else 1
        previous_direction = direction
        max_run = max(max_run, run)
        net = toxic_trade_net(row)
        volume = max(numeric_value(row.get("volume")), 0.0)
        nets.append(net)
        total_volume += volume
        if net < 0:
            loss_volume += volume
    daily_lot_matches = 0
    for items in day_rows.values():
        counts: dict[float, int] = defaultdict(int)
        for row in items:
            counts[rounded(numeric_value(row.get("volume")), 4)] += 1
        daily_lot_matches += max(counts.values(), default=0)
    wins = [value for value in nets if value > 0]
    losses = [value for value in nets if value < 0]
    profit_factor = sum(wins) / abs(sum(losses)) if losses else (99.0 if wins else 0.0)
    median_volume = numeric_value(order_filter.get("medianVolume"))
    camouflage_cutoff = numeric_value(order_filter.get("volumeCutoff"))
    core_rows = ordered
    camouflage_rows = order_filter["excludedRows"]
    core_volume = sum(max(numeric_value(row.get("volume")), 0.0) for row in core_rows)
    camouflage_volume = sum(max(numeric_value(row.get("volume")), 0.0) for row in camouflage_rows)
    session_profile = toxic_dynamic_push_sessions(core_rows)
    sessions = session_profile["sessions"]
    concentrated_sessions = [items for items in sessions if len(items) >= 2]
    concentrated_volume = sum(max(numeric_value(row.get("volume")), 0.0) for items in concentrated_sessions for _, row in items)
    short_core_volume = sum(max(numeric_value(row.get("volume")), 0.0) for row in core_rows if toxic_trade_holding(row) <= 3600)
    core_open_times = [opened for items in sessions for opened, _ in items]
    core_open_span_seconds = (
        max((opened for opened in core_open_times), default=datetime.min)
        - min((opened for opened in core_open_times), default=datetime.min)
    ).total_seconds() if core_open_times else 0.0
    position_campaigns = toxic_push_position_campaigns(sessions)
    staggered_addon_orders = mysql_int(position_campaigns.get("staggeredAddOnOrders"))
    staggered_addon_volume = numeric_value(position_campaigns.get("staggeredAddOnVolume"))
    cohesive_batch_orders = mysql_int(position_campaigns.get("cohesiveWaveOrders"))
    cohesive_campaign_orders = mysql_int(position_campaigns.get("cohesiveCampaignOrders"))
    multi_order_batch_orders = mysql_int(position_campaigns.get("multiOrderCampaignOrders"))
    quiet_gaps = session_profile["sessionGapsHours"]
    outside_camouflage = 0
    for row in camouflage_rows:
        opened = parse_trade_time(row.get("open_time_msc") or row.get("open_time"))
        if opened and core_open_times and all(abs((opened - core_opened).total_seconds()) > 90 * 60 for core_opened in core_open_times):
            outside_camouflage += 1
    total_positive_volume = core_volume + camouflage_volume
    return {
        "orders": total,
        "rawOrders": order_filter.get("rawOrderCount", total),
        "filteredOrders": order_filter.get("filteredOrderCount", total),
        "excludedOrders": order_filter.get("excludedOrderCount", 0),
        "excludedVolumeRatio": order_filter.get("excludedVolumeRatio", 0),
        "activeDays": len(day_rows),
        "ordersPerDay": rounded(total / max(len(day_rows), 1), 1),
        "symbolConcentration": rounded(max(symbols.values(), default=0) / total * 100, 1),
        "dailyLotConsistency": rounded(daily_lot_matches / total * 100, 1),
        "nonOverlapRatio": rounded(non_overlapping / total * 100, 1),
        "maxSameDirectionRun": max_run,
        "maxSameDirectionRunRatio": rounded(max_run / total * 100, 1),
        "winRate": rounded(len(wins) / total * 100, 1),
        "lossRate": rounded(len(losses) / total * 100, 1),
        "lossVolumeRatio": rounded(loss_volume / total_volume * 100, 1) if total_volume else 0,
        "profitFactor": rounded(profit_factor, 2),
        "netProfit": rounded(sum(nets)),
        "netPerLot": rounded(sum(nets) / total_volume, 2) if total_volume else 0,
        "medianVolume": rounded(median_volume, 4),
        "camouflageCutoff": rounded(camouflage_cutoff, 4),
        "coreOrders": len(core_rows),
        "camouflageOrders": len(camouflage_rows),
        "camouflageOrderRatio": rounded(len(camouflage_rows) / max(mysql_int(order_filter.get("rawOrderCount")), 1) * 100, 1),
        "camouflageVolumeRatio": rounded(camouflage_volume / total_positive_volume * 100, 1) if total_positive_volume else 0,
        "camouflageOutsideRatio": rounded(outside_camouflage / len(camouflage_rows) * 100, 1) if camouflage_rows else 0,
        "coreVolumeRatio": rounded(core_volume / total_positive_volume * 100, 1) if total_positive_volume else 0,
        "sessionCount": len(sessions),
        "concentratedSessions": len(concentrated_sessions),
        "concentratedCoreVolumeRatio": rounded(concentrated_volume / core_volume * 100, 1) if core_volume else 0,
        "coreShortHoldVolumeRatio": rounded(short_core_volume / core_volume * 100, 1) if core_volume else 0,
        "coreOpenSpanSeconds": rounded(core_open_span_seconds, 1),
        "entryBatchCount": position_campaigns.get("campaignCount", len(sessions)),
        "entryWaveCount": position_campaigns.get("entryWaveCount", 0),
        "cohesiveCampaigns": position_campaigns.get("cohesiveCampaigns", 0),
        "campaignCloseToleranceSecondsMedian": position_campaigns.get("closeToleranceSecondsMedian", 0),
        "multiOrderBatchRatio": rounded(multi_order_batch_orders / len(core_rows) * 100, 1) if core_rows else 0,
        "cohesiveBatchOrderRatio": rounded(cohesive_batch_orders / len(core_rows) * 100, 1) if core_rows else 0,
        "cohesiveCampaignOrderRatio": rounded(cohesive_campaign_orders / len(core_rows) * 100, 1) if core_rows else 0,
        "staggeredAddOnOrderRatio": rounded(staggered_addon_orders / len(core_rows) * 100, 1) if core_rows else 0,
        "staggeredAddOnVolumeRatio": rounded(staggered_addon_volume / core_volume * 100, 1) if core_volume else 0,
        "quietGapRatio": session_profile["quietGapScore"],
        "medianQuietHours": rounded(statistics.median(quiet_gaps), 1) if quiet_gaps else 0,
        "typicalOrderGapMinutes": session_profile["typicalOrderGapMinutes"],
        "sessionBreakMinutes": session_profile["sessionBreakMinutes"],
        "singleSessionOnly": session_profile["singleSessionOnly"],
        "quietEvidenceAvailable": session_profile["quietEvidenceAvailable"],
    }


def toxic_single_burst_chain(push_behavior: dict, tick: dict | None) -> bool:
    """Confirm a dense one-sided campaign only when post-entry Tick evidence supports it."""
    core_orders = mysql_int(push_behavior.get("coreOrders"))
    return bool(
        core_orders >= 12
        and push_behavior.get("singleSessionOnly")
        and numeric_value(push_behavior.get("coreOpenSpanSeconds")) <= 120
        and numeric_value(push_behavior.get("symbolConcentration")) >= 95
        and numeric_value(push_behavior.get("maxSameDirectionRunRatio")) >= 90
        and numeric_value(push_behavior.get("dailyLotConsistency")) >= 80
        and numeric_value(push_behavior.get("concentratedCoreVolumeRatio")) >= 95
        and numeric_value(push_behavior.get("coreShortHoldVolumeRatio")) >= 90
        and numeric_value(push_behavior.get("winRate")) >= 90
        and numeric_value(push_behavior.get("lossVolumeRatio")) <= 5
        and numeric_value(push_behavior.get("nonOverlapRatio")) <= 10
        and numeric_value(push_behavior.get("staggeredAddOnVolumeRatio")) <= 35
        and tick
        and tick.get("available")
        and mysql_int(tick.get("analyzedOrders")) >= 8
        and numeric_value(tick.get("eventImpact10VolumeRatio")) >= 50
        and (
            numeric_value(tick.get("eventPersistence60VolumeRatio")) >= 30
            or numeric_value(tick.get("positiveImpact20VolumeRatio")) >= 30
        )
    )


def toxic_coordinated_tick_chain(push_behavior: dict, sync: dict | None, tick: dict | None) -> bool:
    """Confirm repeated mirrored execution only with synchronized exits and Tick support."""
    return bool(
        mysql_int(push_behavior.get("coreOrders")) >= 8
        and numeric_value(push_behavior.get("concentratedCoreVolumeRatio")) >= 90
        and numeric_value(push_behavior.get("coreShortHoldVolumeRatio")) >= 80
        and numeric_value(push_behavior.get("staggeredAddOnVolumeRatio")) <= 35
        and sync
        and sync.get("available")
        and mysql_int(sync.get("sampledOrders")) >= 8
        and numeric_value(sync.get("coordinatedMatchedRatio")) >= 80
        and numeric_value(sync.get("coordinatedVolumeRatio")) >= 80
        and numeric_value(sync.get("coordinatedCloseRatio")) >= 60
        and mysql_int(sync.get("maxPeerMatches")) >= 5
        and numeric_value(sync.get("maxPeerRatio")) >= 60
        and tick
        and tick.get("available")
        and mysql_int(tick.get("analyzedOrders")) >= 8
        and numeric_value(tick.get("eventImpact10VolumeRatio")) >= 50
        and numeric_value(tick.get("positiveImpact20VolumeRatio")) >= 60
        and numeric_value(tick.get("favorableTickRatio50Median")) >= 65
    )


def toxic_sudden_exposure_chain(
    push_behavior: dict,
    tick: dict | None,
    finance: dict,
    ea_attention: bool,
) -> bool:
    """Confirm a tiny-sample exposure burst only with execution, funding, and Tick evidence."""
    core_orders = mysql_int(push_behavior.get("coreOrders"))
    deposit_total = numeric_value(finance.get("depositTotal")) if finance.get("available") else 0.0
    withdrawal_total = numeric_value(finance.get("withdrawalTotal")) if finance.get("available") else 0.0
    first_trade_hours = finance.get("firstDepositToTradeHours") if finance.get("available") else None
    withdrawal_hours = finance.get("lastTradeToWithdrawalHours") if finance.get("available") else None
    return bool(
        2 <= core_orders <= 7
        and numeric_value(push_behavior.get("coreOpenSpanSeconds")) <= 5
        and numeric_value(push_behavior.get("symbolConcentration")) >= 95
        and numeric_value(push_behavior.get("maxSameDirectionRunRatio")) >= 95
        and numeric_value(push_behavior.get("cohesiveCampaignOrderRatio")) >= 90
        and numeric_value(push_behavior.get("coreShortHoldVolumeRatio")) >= 90
        and numeric_value(push_behavior.get("winRate")) >= 90
        and numeric_value(push_behavior.get("lossVolumeRatio")) <= 5
        and ea_attention
        and tick
        and tick.get("available")
        and mysql_int(tick.get("analyzedOrders")) >= 2
        and numeric_value(tick.get("eventImpact10VolumeRatio")) >= 80
        and numeric_value(tick.get("eventPersistence60VolumeRatio")) >= 80
        and numeric_value(tick.get("favorableTickRatio50Median")) >= 80
        and finance.get("available")
        and deposit_total > 0
        and withdrawal_total / deposit_total >= .9
        and first_trade_hours is not None
        and 0 <= numeric_value(first_trade_hours) <= 24
        and withdrawal_hours is not None
        and 0 <= numeric_value(withdrawal_hours) <= 24
        and numeric_value(finance.get("highestHoldingVolume")) >= 1
    )


def toxic_push_mixed_episode_summary(rows: list[dict], sync: dict | None, tick: dict | None = None) -> dict:
    """Find repeated coordinated push episodes inside an otherwise mixed account.

    Global grid or averaging behavior must not erase separate directional episodes.
    To avoid flagging ordinary traders, an episode needs economic size, direction and
    short-hold consistency, repeated coordination with the same peer, synchronized
    exits, and supporting Tick behavior.  The account-level override only activates
    when that pattern repeats and covers a meaningful share of core volume.
    """
    if not (sync and sync.get("available")):
        return {"available": False, "reason": "跨账户同步不可用"}
    sampled_matches = sync.get("sampledOrderMatches") or []
    if not sampled_matches:
        return {"available": False, "reason": "当前同步结果没有订单级事件映射"}
    sampled_by_ticket = {normalize_text(item.get("ticket")): item for item in sampled_matches if normalize_text(item.get("ticket"))}
    session_profile = toxic_dynamic_push_sessions(rows)
    total_volume = sum(max(numeric_value(row.get("volume")), 0.0) for row in rows)
    ea_rows = [row for row in rows if is_ea_trade(row)]
    copy_rows = [row for row in rows if is_copy_trade(row)]
    ea_volume = sum(max(numeric_value(row.get("volume")), 0.0) for row in ea_rows)
    copy_volume = sum(max(numeric_value(row.get("volume")), 0.0) for row in copy_rows)
    ea_order_ratio = len(ea_rows) / len(rows) * 100 if rows else 0.0
    copy_order_ratio = len(copy_rows) / len(rows) * 100 if rows else 0.0
    ea_volume_ratio = ea_volume / total_volume * 100 if total_volume else 0.0
    copy_volume_ratio = copy_volume / total_volume * 100 if total_volume else 0.0
    ea_attention = bool(ea_volume_ratio >= 80 or ea_order_ratio >= 70)
    copy_attention = bool(copy_volume_ratio >= 20 or copy_order_ratio >= 20)
    execution_mode = "EA主导" if ea_attention else "跟单标记" if copy_attention else "人工/未知"
    candidates = []
    for session_index, items in enumerate(session_profile.get("sessions") or [], start=1):
        session_rows = [row for _, row in items]
        if len(session_rows) < 5:
            continue
        session_volume = sum(max(numeric_value(row.get("volume")), 0.0) for row in session_rows)
        if session_volume <= 0:
            continue
        direction_volumes: dict[str, float] = defaultdict(float)
        symbol_volumes: dict[str, float] = defaultdict(float)
        short_volume = 0.0
        for row in session_rows:
            volume = max(numeric_value(row.get("volume")), 0.0)
            direction_volumes[clean_trade_type(row.get("type", ""))] += volume
            symbol_volumes[toxic_base_symbol(row.get("symbol")) or "-"] += volume
            if toxic_trade_holding(row) <= 3600:
                short_volume += volume
        direction_ratio = max(direction_volumes.values(), default=0.0) / session_volume * 100
        symbol_ratio = max(symbol_volumes.values(), default=0.0) / session_volume * 100
        short_ratio = short_volume / session_volume * 100
        position = toxic_push_position_campaigns([items])
        staggered_volume_ratio = numeric_value(position.get("staggeredAddOnVolume")) / session_volume * 100
        samples = [sampled_by_ticket.get(normalize_text(row.get("ticket"))) for row in session_rows]
        samples = [item for item in samples if item is not None]
        if len(samples) < 3:
            continue
        coordinated_samples = [item for item in samples if item.get("peers")]
        close_samples = [item for item in samples if item.get("closePeers")]
        peer_counts: dict[str, int] = defaultdict(int)
        peer_close_counts: dict[str, int] = defaultdict(int)
        for item in samples:
            for peer in item.get("peers") or []:
                peer_counts[peer] += 1
            for peer in item.get("closePeers") or []:
                peer_close_counts[peer] += 1
        top_peer, top_peer_matches = max(peer_counts.items(), key=lambda item: (item[1], item[0]), default=("", 0))
        sync_ratio = len(coordinated_samples) / len(samples) * 100
        close_ratio = len(close_samples) / len(samples) * 100
        top_peer_ratio = top_peer_matches / len(samples) * 100
        top_peer_close_ratio = peer_close_counts.get(top_peer, 0) / len(samples) * 100 if top_peer else 0.0
        volume_ratio = session_volume / total_volume * 100 if total_volume else 0.0
        episode_ea_volume = sum(max(numeric_value(row.get("volume")), 0.0) for row in session_rows if is_ea_trade(row))
        episode_copy_volume = sum(max(numeric_value(row.get("volume")), 0.0) for row in session_rows if is_copy_trade(row))
        episode_score = (
            sync_ratio * .25
            + top_peer_ratio * .2
            + top_peer_close_ratio * .2
            + direction_ratio * .1
            + short_ratio * .1
            + toxic_ramp(volume_ratio, 1, 10) * .1
            + max(0.0, 100 - staggered_volume_ratio) * .05
        )
        qualifies = bool(
            volume_ratio >= 1
            and direction_ratio >= 70
            and symbol_ratio >= 70
            and short_ratio >= 70
            and staggered_volume_ratio <= 35
            and sync_ratio >= 60
            and top_peer_ratio >= 40
            and top_peer_close_ratio >= 30
        )
        if qualifies:
            opened = [parse_trade_time(row.get("open_time_msc") or row.get("open_time")) for row in session_rows]
            opened = [value for value in opened if value]
            candidates.append({
                "session": session_index,
                "start": trade_time_text(min(opened)) if opened else "",
                "orders": len(session_rows),
                "volume": rounded(session_volume, 4),
                "volumeRatio": rounded(volume_ratio, 1),
                "eaVolumeRatio": rounded(episode_ea_volume / session_volume * 100, 1) if session_volume else 0,
                "copyVolumeRatio": rounded(episode_copy_volume / session_volume * 100, 1) if session_volume else 0,
                "directionVolumeRatio": rounded(direction_ratio, 1),
                "shortHoldVolumeRatio": rounded(short_ratio, 1),
                "staggeredVolumeRatio": rounded(staggered_volume_ratio, 1),
                "sampledOrders": len(samples),
                "syncRatio": rounded(sync_ratio, 1),
                "closeSyncRatio": rounded(close_ratio, 1),
                "topPeer": top_peer,
                "topPeerRatio": rounded(top_peer_ratio, 1),
                "topPeerCloseRatio": rounded(top_peer_close_ratio, 1),
                "score": rounded(episode_score, 1),
            })
    peer_episode_counts: dict[str, int] = defaultdict(int)
    for episode in candidates:
        if episode.get("topPeer"):
            peer_episode_counts[episode["topPeer"]] += 1
    repeated_peers = {peer for peer, count in peer_episode_counts.items() if count >= 2}
    confirmed_episodes = [episode for episode in candidates if episode.get("topPeer") in repeated_peers]
    confirmed_orders = sum(mysql_int(episode.get("orders")) for episode in confirmed_episodes)
    confirmed_volume = sum(numeric_value(episode.get("volume")) for episode in confirmed_episodes)
    confirmed_volume_ratio = confirmed_volume / total_volume * 100 if total_volume else 0.0
    strongest_peer, strongest_peer_episodes = max(peer_episode_counts.items(), key=lambda item: (item[1], item[0]), default=("", 0))
    tick_support = bool(
        tick and tick.get("available")
        and (
            (numeric_value(tick.get("positiveImpact20VolumeRatio")) >= 45 and numeric_value(tick.get("favorableTickRatio50Median")) >= 50)
            or numeric_value(tick.get("win10VolumeRatio")) >= 45
        )
    )
    global_sync_threshold = 50 if ea_attention else 60
    global_sync_support = bool(
        numeric_value(sync.get("coordinatedMatchedRatio")) >= global_sync_threshold
        and numeric_value(sync.get("coordinatedCloseRatio")) >= 30
        and numeric_value(sync.get("maxPeerRatio")) >= 30
        and mysql_int(sync.get("recurringPeerAccounts")) >= 2
    )
    volume_threshold = 8 if ea_attention else 10
    confirmed = bool(
        len(confirmed_episodes) >= 2
        and confirmed_orders >= 10
        and confirmed_volume_ratio >= volume_threshold
        and global_sync_support
        and tick_support
    )
    score = max((numeric_value(item.get("score")) for item in confirmed_episodes), default=0.0)
    if confirmed:
        score = max(score, 85)
        if (
            numeric_value(sync.get("coordinatedMatchedRatio")) >= 80
            and numeric_value(sync.get("coordinatedCloseRatio")) >= 60
            and numeric_value(sync.get("maxPeerRatio")) >= 50
        ):
            score = max(score, 90)
    return {
        "available": True,
        "confirmed": confirmed,
        "score": rounded(min(score, 95), 1),
        "candidateEpisodes": len(candidates),
        "confirmedEpisodes": len(confirmed_episodes),
        "candidateSessionIds": [mysql_int(episode.get("session")) for episode in candidates],
        "confirmedSessionIds": [mysql_int(episode.get("session")) for episode in confirmed_episodes],
        "confirmedOrders": confirmed_orders,
        "confirmedVolume": rounded(confirmed_volume, 4),
        "confirmedVolumeRatio": rounded(confirmed_volume_ratio, 1),
        "strongestPeer": strongest_peer,
        "strongestPeerEpisodes": strongest_peer_episodes,
        "tickSupport": tick_support,
        "globalSyncSupport": global_sync_support,
        "globalSyncThreshold": global_sync_threshold,
        "volumeThreshold": volume_threshold,
        "eaAttention": ea_attention,
        "copyAttention": copy_attention,
        "executionMode": execution_mode,
        "eaOrderRatio": rounded(ea_order_ratio, 1),
        "eaVolumeRatio": rounded(ea_volume_ratio, 1),
        "copyOrderRatio": rounded(copy_order_ratio, 1),
        "copyVolumeRatio": rounded(copy_volume_ratio, 1),
        "topEpisodes": sorted(confirmed_episodes or candidates, key=lambda item: (-numeric_value(item.get("score")), -numeric_value(item.get("volumeRatio"))))[:8],
        "definition": "局部轮次需至少5单、占核心手数1%以上、同方向和短持仓手数均不低于70%，同一固定账户开仓同步不低于40%、平仓同步不低于30%；至少2个轮次由同一账户反复配合，合计覆盖核心手数10%以上并有Tick支持后才提升风险。若EA占核心手数80%以上，可将合计手数门槛降至8%、全局协调同步门槛降至50%，但仍必须满足固定同伙重复、同步平仓和Tick交叉证据；跟单标记只触发重点复核，不单独降低门槛。",
    }


def toxic_build_push_context(rows: list[dict]) -> dict:
    order_filter = toxic_filter_push_orders(rows)
    core_rows = order_filter["rows"]
    return {
        "rows": core_rows,
        "filter": order_filter,
        "behavior": toxic_push_behavior(core_rows, order_filter=order_filter),
    }


def toxic_sample_even(rows: list[dict], limit: int) -> list[dict]:
    if len(rows) <= limit:
        return list(rows)
    indexes = sorted({round(index * (len(rows) - 1) / (limit - 1)) for index in range(limit)})
    return [rows[index] for index in indexes]


def toxic_tick_candidate_key(row: dict) -> str:
    ticket = normalize_text(row.get("ticket"))
    identity = [
        normalize_text(row.get("platform")).upper(),
        normalize_text(row.get("server")),
        ticket,
    ]
    if not ticket:
        identity.extend([
            normalize_text(row.get("open_time_msc") or row.get("open_time")),
            normalize_text(row.get("close_time_msc") or row.get("close_time")),
            normalize_text(row.get("symbol")),
            clean_trade_type(row.get("type", "")),
            numeric_value(row.get("open_price")),
            numeric_value(row.get("volume")),
        ])
    return json.dumps(identity, ensure_ascii=False, separators=(",", ":"))


def toxic_tick_winning_sample_results(results: list[dict], winning_sample_keys: set[str]) -> list[dict]:
    return [
        item for item in results
        if item.get("_candidateKey") in winning_sample_keys and numeric_value(item.get("realizedNet")) > 0
    ]


def toxic_sync_candidates_for_source(
    source: dict,
    login: str,
    target_seconds: set[str],
    target_symbols: set[str] | None = None,
    target_directions: set[str] | None = None,
) -> list[dict]:
    candidates = []
    close_by_position: dict[int, datetime] = {}
    symbols = sorted({toxic_base_symbol(value) for value in (target_symbols or set()) if toxic_base_symbol(value)})
    actions = sorted({0 if value == "buy" else 1 for value in (target_directions or set()) if value in {"buy", "sell"}})
    with mysql_trade_connect(source) as conn:
        with conn.cursor() as cur:
            values = sorted(target_seconds)
            for offset in range(0, len(values), 250):
                batch = values[offset:offset + 250]
                placeholders = ",".join(["%s"] * len(batch))
                if source.get("kind") == "mt5_deals":
                    action_sql = f"Action in ({','.join(['%s'] * len(actions))})" if actions else "Action in (0,1)"
                    symbol_sql = (
                        f" and REPLACE(UPPER(SUBSTRING_INDEX(Symbol, '.', 1)), 'ROLL', '') in ({','.join(['%s'] * len(symbols))})"
                        if symbols else ""
                    )
                    cur.execute(
                        f"select Login, PositionID, Action, Symbol, Time, TimeMsc, Volume from `{source['schema']}`.`{source['table']}` where Entry=0 and {action_sql} and Login<>%s and Time in ({placeholders}){symbol_sql}",
                        [*actions, int(login), *batch, *symbols],
                    )
                else:
                    action_sql = f"CMD in ({','.join(['%s'] * len(actions))})" if actions else "CMD in (0,1)"
                    symbol_sql = (
                        f" and REPLACE(UPPER(SUBSTRING_INDEX(SYMBOL, '.', 1)), 'ROLL', '') in ({','.join(['%s'] * len(symbols))})"
                        if symbols else ""
                    )
                    cur.execute(
                        f"select LOGIN as Login, TICKET as PositionID, CMD as Action, SYMBOL as Symbol, OPEN_TIME as Time, OPEN_TIME as TimeMsc, VOLUME as Volume, CLOSE_TIME as CloseTime from `{source['schema']}`.`{source['table']}` where {action_sql} and LOGIN<>%s and OPEN_TIME in ({placeholders}){symbol_sql}",
                        [*actions, int(login), *batch, *symbols],
                    )
                candidates.extend(cur.fetchall())
            if source.get("kind") == "mt5_deals" and candidates:
                position_ids = list({mysql_int(item.get("PositionID")) for item in candidates if mysql_int(item.get("PositionID"))})
                for offset in range(0, len(position_ids), 500):
                    batch = position_ids[offset:offset + 500]
                    placeholders = ",".join(["%s"] * len(batch))
                    cur.execute(
                        f"select PositionID, min(TimeMsc) as CloseTime from `{source['schema']}`.`{source['table']}` where Entry=1 and Action in (0,1) and PositionID in ({placeholders}) group by PositionID",
                        batch,
                    )
                    for item in cur.fetchall():
                        close_by_position[mysql_int(item.get("PositionID"))] = parse_trade_time(item.get("CloseTime"))
            else:
                close_by_position = {mysql_int(item.get("PositionID")): parse_trade_time(item.get("CloseTime")) for item in candidates}
    server = normalize_text(source.get("server") or source.get("name"))
    platform = normalize_text(source.get("platform"))
    return [{
        "login": f'{server}/{normalize_text(item.get("Login"))}',
        "account": normalize_text(item.get("Login")),
        "platform": platform,
        "server": server,
        "position": mysql_int(item.get("PositionID")),
        "direction": "buy" if mysql_int(item.get("Action")) == 0 else "sell",
        "symbol": toxic_base_symbol(item.get("Symbol")),
        "volume": normalize_mt5_volume(item.get("Volume")) if source.get("kind") == "mt5_deals" else normalize_mt4_volume(item.get("Volume")),
        "opened": parse_trade_time(item.get("TimeMsc") or item.get("Time")),
        "closed": close_by_position.get(mysql_int(item.get("PositionID"))),
    } for item in candidates]


def toxic_sync_candidates_across_sources(
    sources: list[dict],
    login: str,
    target_seconds: set[str],
    target_symbols: set[str],
    target_directions: set[str],
) -> list[dict]:
    """Query independent trading sources concurrently and retain source order."""
    if not sources:
        return []

    def load(index: int, source: dict) -> tuple[int, dict]:
        started = time.monotonic()
        server = normalize_text(source.get("server") or source.get("name"))
        try:
            candidates = toxic_sync_candidates_for_source(
                source,
                login,
                target_seconds,
                target_symbols,
                target_directions,
            )
            return index, {
                "server": server,
                "candidates": candidates,
                "error": "",
                "seconds": time.monotonic() - started,
            }
        except Exception as exc:
            return index, {
                "server": server,
                "candidates": [],
                "error": str(exc),
                "seconds": time.monotonic() - started,
            }

    ordered: list[dict | None] = [None] * len(sources)
    with ThreadPoolExecutor(
        max_workers=min(4, len(sources)),
        thread_name_prefix="toxic-sync-source",
    ) as executor:
        futures = {
            executor.submit(load, index, source): index
            for index, source in enumerate(sources)
        }
        for future in as_completed(futures):
            index, result = future.result()
            ordered[index] = result
    return [result for result in ordered if result is not None]


def toxic_suspected_accounts(top_peers: list[tuple[str, int]], recurring_peers: set[str], peer_close_counts: dict[str, int], sampled_count: int) -> list[dict]:
    suspected = []
    for peer, count in top_peers:
        if peer not in recurring_peers:
            continue
        server, _, account = peer.partition("/")
        source = next((item for item in MYSQL_SOURCES if normalize_text(item.get("server") or item.get("name")) == server), None)
        close_matches = peer_close_counts.get(peer, 0)
        suspected.append({
            "platform": normalize_text((source or {}).get("platform")),
            "server": server,
            "account": account,
            "matches": count,
            "matchRatio": rounded(count / sampled_count * 100, 1) if sampled_count else 0,
            "closeMatches": close_matches,
            "closeMatchRatio": rounded(close_matches / sampled_count * 100, 1) if sampled_count else 0,
        })
    return suspected[:20]


def toxic_cross_account_sync(login: str, rows: list[dict], limit: int = 200) -> dict:
    sync_started = time.monotonic()
    mysql_rows = [row for row in rows if row.get("data_source") == "mysql" and parse_trade_time(row.get("open_time_msc") or row.get("open_time"))]
    sampled = toxic_sample_even(sorted(mysql_rows, key=lambda row: normalize_text(row.get("open_time_msc") or row.get("open_time"))), limit)
    if not sampled:
        return {
            "available": False,
            "reason": "当前订单没有可用于跨账户匹配的数据库毫秒时间",
            "performance": {"totalSeconds": rounded(time.monotonic() - sync_started, 3)},
        }
    peer_counts: dict[str, int] = defaultdict(int)
    peer_close_counts: dict[str, int] = defaultdict(int)
    order_peer_sets: list[tuple[str, float, set[str], set[str]]] = []
    order_match_details: list[dict] = []
    sampled_order_matches = {
        normalize_text(row.get("ticket")): {
            "ticket": normalize_text(row.get("ticket")),
            "volume": max(numeric_value(row.get("volume")), 0.0),
            "peers": set(),
            "closePeers": set(),
        }
        for row in sampled
        if normalize_text(row.get("ticket"))
    }
    matched_tickets: list[str] = []
    matched_orders = 0
    matched_volume = 0.0
    close_matched = 0
    sampled_volume = sum(max(numeric_value(row.get("volume")), 0.0) for row in sampled)
    errors = []
    query_seconds = 0.0
    match_seconds = 0.0
    candidate_count = 0
    source_timings: dict[str, float] = {}
    by_source: dict[tuple[str, str], list[dict]] = defaultdict(list)
    for row in sampled:
        by_source[(normalize_text(row.get("platform")), normalize_text(row.get("server")))].append(row)
    searched_servers: set[str] = set()
    for (platform, server), source_rows in by_source.items():
        target_source = next((item for item in MYSQL_SOURCES if item.get("platform") == platform and item.get("server") == server), None)
        if not target_source:
            errors.append(f"{platform}/{server} 没有跨账户查询配置")
            continue
        target_seconds = set()
        for row in source_rows:
            opened = parse_trade_time(row.get("open_time_msc") or row.get("open_time"))
            for delta in range(-2, 3):
                target_seconds.add(trade_time_text(opened + timedelta(seconds=delta)))
        target_symbols = {toxic_base_symbol(row.get("symbol")) for row in source_rows if toxic_base_symbol(row.get("symbol"))}
        target_directions = {clean_trade_type(row.get("type", "")) for row in source_rows}
        normalized_candidates = []
        candidate_sources = [item for item in MYSQL_SOURCES if item.get("platform") == platform]
        source_results = toxic_sync_candidates_across_sources(
            candidate_sources,
            login,
            target_seconds,
            target_symbols,
            target_directions,
        )
        for source_result in source_results:
            candidate_server = source_result["server"]
            source_candidates = source_result["candidates"]
            elapsed = numeric_value(source_result["seconds"])
            normalized_candidates.extend(source_candidates)
            candidate_count += len(source_candidates)
            query_seconds += elapsed
            source_timings[candidate_server] = source_timings.get(candidate_server, 0.0) + elapsed
            if source_result["error"]:
                errors.append(f"{candidate_server}: {source_result['error']}")
            else:
                searched_servers.add(candidate_server)
        match_started = time.monotonic()
        candidate_index: dict[tuple[str, str, datetime], list[dict]] = defaultdict(list)
        for item in normalized_candidates:
            if item["opened"]:
                candidate_index[(item["symbol"], item["direction"], item["opened"].replace(microsecond=0))].append(item)
        for row in source_rows:
            opened = parse_trade_time(row.get("open_time_msc") or row.get("open_time"))
            closed = parse_trade_time(row.get("close_time_msc") or row.get("close_time"))
            direction = clean_trade_type(row.get("type", ""))
            symbol = toxic_base_symbol(row.get("symbol"))
            peers = [
                item
                for delta in range(-2, 3)
                for item in candidate_index.get((symbol, direction, (opened + timedelta(seconds=delta)).replace(microsecond=0)), [])
                if abs((item["opened"] - opened).total_seconds()) <= 2
            ]
            if not peers:
                continue
            matched_orders += 1
            matched_volume += max(numeric_value(row.get("volume")), 0.0)
            matched_tickets.append(normalize_text(row.get("ticket")))
            peer_logins = {item["login"] for item in peers if item["login"]}
            close_peer_logins = {
                item["login"] for item in peers
                if item["login"] and closed and item["closed"]
                and abs((item["closed"] - closed).total_seconds()) <= 2
            }
            for peer in peer_logins:
                peer_counts[peer] += 1
            for peer in close_peer_logins:
                peer_close_counts[peer] += 1
            ticket = normalize_text(row.get("ticket"))
            order_peer_sets.append((ticket, max(numeric_value(row.get("volume")), 0.0), peer_logins, close_peer_logins))
            order_match_details.append({
                "targetTicket": ticket,
                "targetPlatform": normalize_text(row.get("platform")),
                "targetServer": normalize_text(row.get("server")),
                "targetSymbol": symbol,
                "targetDirection": direction,
                "targetVolume": max(numeric_value(row.get("volume")), 0.0),
                "targetOpened": opened,
                "targetClosed": closed,
                "candidates": peers,
            })
            if ticket in sampled_order_matches:
                sampled_order_matches[ticket]["peers"].update(peer_logins)
                sampled_order_matches[ticket]["closePeers"].update(close_peer_logins)
            if close_peer_logins:
                close_matched += 1
        match_seconds += time.monotonic() - match_started
    sampled_count = len(sampled)
    top_peers = sorted(peer_counts.items(), key=lambda item: (-item[1], item[0]))
    recurring_min_matches = max(2, math.ceil(sampled_count * 0.05))
    recurring_peers = {peer for peer, count in peer_counts.items() if count >= recurring_min_matches}
    coordinated_orders = [item for item in order_peer_sets if item[2] & recurring_peers]
    coordinated_close_orders = [item for item in order_peer_sets if item[3] & recurring_peers]
    coordinated_volume = sum(item[1] for item in coordinated_orders)
    coordinated_tickets = [item[0] for item in coordinated_orders]
    suspected_accounts = toxic_suspected_accounts(top_peers, recurring_peers, peer_close_counts, sampled_count)
    episode_peer_universe = {peer for peer, _ in top_peers[:20] if peer in recurring_peers}
    peer_rank = {peer: index for index, (peer, _) in enumerate(top_peers)}
    comparison_rows = []
    for detail in order_match_details:
        target_opened = detail["targetOpened"]
        target_closed = detail["targetClosed"]
        best_by_peer: dict[str, tuple[tuple, dict, float, float | None]] = {}
        for candidate in detail["candidates"]:
            peer = normalize_text(candidate.get("login"))
            peer_opened = candidate.get("opened")
            if peer not in recurring_peers or not peer_opened:
                continue
            open_delta = (peer_opened - target_opened).total_seconds()
            peer_closed = candidate.get("closed")
            close_delta = (
                (peer_closed - target_closed).total_seconds()
                if peer_closed and target_closed else None
            )
            choice = (
                abs(open_delta),
                0 if close_delta is not None and abs(close_delta) <= 2 else 1,
                abs(close_delta) if close_delta is not None else float("inf"),
                mysql_int(candidate.get("position")),
            )
            if peer not in best_by_peer or choice < best_by_peer[peer][0]:
                best_by_peer[peer] = (choice, candidate, open_delta, close_delta)
        for peer, (_, candidate, open_delta, close_delta) in best_by_peer.items():
            comparison_rows.append({
                "targetTicket": detail["targetTicket"],
                "targetPlatform": detail["targetPlatform"],
                "targetServer": detail["targetServer"],
                "targetAccount": normalize_text(login),
                "targetSymbol": detail["targetSymbol"],
                "targetDirection": detail["targetDirection"],
                "targetVolume": rounded(detail["targetVolume"], 4),
                "targetOpened": mysql_datetime_text(detail["targetOpened"]),
                "targetClosed": mysql_datetime_text(detail["targetClosed"]),
                "peerPlatform": normalize_text(candidate.get("platform")),
                "peerServer": normalize_text(candidate.get("server")),
                "peerAccount": normalize_text(candidate.get("account")),
                "peerTicket": normalize_text(candidate.get("position")),
                "peerVolume": rounded(candidate.get("volume"), 4),
                "peerOpened": mysql_datetime_text(candidate.get("opened")),
                "peerClosed": mysql_datetime_text(candidate.get("closed")),
                "openDeltaSeconds": rounded(open_delta, 3),
                "closeDeltaSeconds": rounded(close_delta, 3) if close_delta is not None else None,
                "closeSynchronized": bool(close_delta is not None and abs(close_delta) <= 2),
            })
    comparison_rows.sort(key=lambda item: (
        peer_rank.get(f'{item["peerServer"]}/{item["peerAccount"]}', len(peer_rank)),
        item["targetOpened"],
        item["targetTicket"],
    ))
    comparison_limit = 1000
    return {
        "available": True,
        "sampledOrders": sampled_count,
        "matchedOrders": matched_orders,
        "matchedRatio": rounded(matched_orders / sampled_count * 100, 1) if sampled_count else 0,
        "matchedVolumeRatio": rounded(matched_volume / sampled_volume * 100, 1) if sampled_volume else 0,
        "closeMatchedRatio": rounded(close_matched / sampled_count * 100, 1) if sampled_count else 0,
        "peerAccounts": len(peer_counts),
        "maxPeerMatches": top_peers[0][1] if top_peers else 0,
        "maxPeerRatio": rounded(top_peers[0][1] / sampled_count * 100, 1) if top_peers and sampled_count else 0,
        "recurringMinMatches": recurring_min_matches,
        "recurringPeerAccounts": len(recurring_peers),
        "coordinatedMatchedRatio": rounded(len(coordinated_orders) / sampled_count * 100, 1) if sampled_count else 0,
        "coordinatedVolumeRatio": rounded(coordinated_volume / sampled_volume * 100, 1) if sampled_volume else 0,
        "coordinatedCloseRatio": rounded(len(coordinated_close_orders) / sampled_count * 100, 1) if sampled_count else 0,
        "topPeers": [
            {"server": peer.split("/", 1)[0], "account": peer.split("/", 1)[1], "matches": count, "closeMatches": peer_close_counts.get(peer, 0)}
            for peer, count in top_peers[:10]
        ],
        "suspectedAccounts": suspected_accounts,
        "sampledOrderMatches": [
            {
                "ticket": item["ticket"],
                "volume": rounded(item["volume"], 4),
                "peers": sorted(item["peers"] & episode_peer_universe),
                "closePeers": sorted(item["closePeers"] & episode_peer_universe),
            }
            for item in sampled_order_matches.values()
        ],
        "comparisonTotal": len(comparison_rows),
        "comparisonLimit": comparison_limit,
        "comparisonTruncated": len(comparison_rows) > comparison_limit,
        "comparisonRows": comparison_rows[:comparison_limit],
        "searchedServers": sorted(searched_servers),
        "evidenceOrders": coordinated_tickets[:20] or matched_tickets[:20],
        "errors": errors[:5],
        "performance": {
            "candidateOrders": candidate_count,
            "querySeconds": rounded(query_seconds, 3),
            "matchSeconds": rounded(match_seconds, 3),
            "totalSeconds": rounded(time.monotonic() - sync_started, 3),
            "sourceSeconds": {server: rounded(seconds, 3) for server, seconds in sorted(source_timings.items())},
        },
        "definition": f"查询同平台全部已接入服务器；同品种、同方向、开仓相差不超过2秒，同一服务器/账号至少重复匹配{recurring_min_matches}单才计入协调同步，平仓同步单独统计。",
    }


def toxic_finance_summary(login: str, rows: list[dict], metrics: dict) -> dict:
    servers = {normalize_text(row.get("server")) for row in rows if normalize_text(row.get("server"))}
    platforms = {normalize_text(row.get("platform")) for row in rows if normalize_text(row.get("platform"))}
    if len(servers) != 1 or len(platforms) != 1:
        return {"available": False, "reason": "资金行为深查需要选择单一平台和服务器"}
    server = next(iter(servers))
    platform = next(iter(platforms))
    source = next((item for item in MYSQL_SOURCES if item.get("server") == server and item.get("platform") == platform), None)
    if not source:
        return {"available": False, "reason": "当前服务器没有资金快照配置"}
    try:
        if source.get("kind") == "mt5_deals":
            panel = query_mt5_finance_panel(source, login, rows, metrics)
        elif source.get("kind") == "mt4_trades":
            panel = query_mt4_finance_panel(source, login, rows, metrics)
        else:
            return {"available": False, "reason": "当前服务器没有可识别的资金流水配置"}
        return {"available": True, **panel}
    except Exception as exc:
        return {"available": False, "reason": str(exc)}


def toxic_weighted_median(values: list[tuple[float, float]]) -> float | None:
    values = sorted((value, max(weight, 0.0)) for value, weight in values if value is not None and weight > 0)
    total = sum(weight for _, weight in values)
    running = 0.0
    for value, weight in values:
        running += weight
        if running >= total / 2:
            return value
    return values[-1][0] if values else None


def toxic_cached_tick_mapping(login: str, report_symbol: str) -> dict | None:
    report_symbol = normalize_text(report_symbol).upper()
    paths = sorted(KLINE_OUT_DIR.glob(f"{safe_stem_text(login)}_*_mapping.json"), key=lambda path: path.stat().st_mtime, reverse=True)
    for path in paths[:20]:
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, ValueError, TypeError):
            continue
        for key, mapping in payload.items():
            if normalize_text(key).upper() == report_symbol and isinstance(mapping, dict):
                return {
                    "reportSymbol": report_symbol,
                    "mt5Symbol": normalize_text(mapping.get("mt5_symbol")) or report_symbol,
                    "timeMode": normalize_text(mapping.get("time_mode")) or "report_is_GMT",
                    "hourDelta": mysql_int(mapping.get("hour_delta")),
                    "source": "图表校准缓存",
                    "mappingFile": path.name,
                }
    return None


def toxic_mt5_symbol_for(mt5, report_symbol: str) -> str:
    raw = normalize_text(report_symbol).upper()
    base = raw.split(".")[0]
    base_no_roll = base.replace("ROLL", "")
    alias_map = {
        "CHINA50": "CN50Roll", "CN50": "CN50Roll", "HKG50": "HKG50Roll", "HK50": "HKG50Roll",
        "NAS100": "NAS100Roll", "US30": "US30Roll", "SPX500": "SPX500Roll", "UK100": "UK100Roll",
        "GER40": "GER40Roll", "JPN225": "JPN225Roll", "AUS200": "AUS200Roll",
        "UKOIL": "UKOILRoll", "USOIL": "USOILRoll", "NGAS": "NGASRoll",
    }
    aliases = [alias_map.get(base), alias_map.get(base_no_roll)]
    candidates = []
    for value in (raw, base, base_no_roll, f"{base_no_roll}Roll", f"{base_no_roll}Roll.ECN", f"{base_no_roll}Roll.PRO", *aliases):
        if value and value not in candidates:
            candidates.append(value)
    for symbol in candidates:
        if mt5.symbol_info(symbol):
            mt5.symbol_select(symbol, True)
            return symbol
    found = []
    for key in (raw, base, base_no_roll, *aliases):
        if not key:
            continue
        for info in mt5.symbols_get(f"*{key}*") or []:
            if info.name not in found:
                found.append(info.name)
    if found:
        symbol = sorted(found, key=lambda name: (0 if name.upper() == raw else 1, 0 if name.endswith(".ECN") else 1, len(name), name))[0]
        mt5.symbol_select(symbol, True)
        return symbol
    raise RuntimeError(f"Terminal 中找不到品种 {report_symbol}")


def toxic_live_tick_mapping(mt5, report_symbol: str, mt5_symbol: str, rows: list[dict]) -> dict | None:
    ordered = sorted(rows, key=lambda row: parse_trade_time(row.get("open_time_msc") or row.get("open_time")) or datetime.min)
    if len(ordered) > 5:
        indexes = sorted({round(index * (len(ordered) - 1) / 4) for index in range(5)})
        sample = [ordered[index] for index in indexes]
    else:
        sample = ordered
    alignments = []
    for label, hour_delta in (("report_is_GMT", 0), ("report_is_GMT+3", -3)):
        distances = []
        inside = 0
        for row in sample:
            opened = parse_trade_time(row.get("open_time_msc") or row.get("open_time"))
            if not opened:
                continue
            query_time = opened + timedelta(hours=hour_delta)
            start = (query_time - timedelta(minutes=2)).replace(tzinfo=timezone.utc)
            end = (query_time + timedelta(minutes=2)).replace(tzinfo=timezone.utc)
            rates = mt5.copy_rates_range(mt5_symbol, mt5.TIMEFRAME_M1, start, end)
            if rates is None or not len(rates):
                continue
            target = int(query_time.replace(second=0, microsecond=0, tzinfo=timezone.utc).timestamp())
            bar = min(rates, key=lambda item: abs(int(item["time"]) - target))
            if abs(int(bar["time"]) - target) > 60:
                continue
            price = numeric_value(row.get("open_price"))
            low, high = float(bar["low"]), float(bar["high"])
            distance = 0.0 if low <= price <= high else min(abs(price - low), abs(price - high))
            distances.append(distance)
            inside += int(distance == 0)
        if distances:
            alignments.append({
                "timeMode": label, "hourDelta": hour_delta, "matched": len(distances), "inside": inside,
                "medianDistance": rounded(statistics.median(distances), 8),
            })
    if not alignments:
        return None
    best = sorted(alignments, key=lambda item: (-item["inside"], item["medianDistance"], -item["matched"]))[0]
    return {"reportSymbol": report_symbol, "mt5Symbol": mt5_symbol, "source": "Terminal M1 实时校准", **best}


def toxic_copy_ticks_retry(mt5, symbol: str, start: datetime, end: datetime, attempts: int = 3):
    ticks = None
    for attempt in range(attempts):
        ticks = mt5.copy_ticks_range(symbol, start, end, mt5.COPY_TICKS_ALL)
        if ticks is not None and len(ticks):
            return ticks
        if attempt + 1 < attempts:
            time.sleep(0.25)
    return ticks


def toxic_prefetch_tick_windows(
    mt5,
    symbol: str,
    rows: list[dict],
    hour_delta: int,
) -> dict[str, list]:
    """Merge overlapping quote windows, then restore each order's exact slice."""
    windows = []
    for row in rows:
        opened = parse_trade_time(row.get("open_time_msc") or row.get("open_time"))
        if not opened:
            continue
        offset = timedelta(hours=hour_delta)
        fetch_start = (opened - timedelta(seconds=31) + offset).replace(tzinfo=timezone.utc)
        fetch_end = (opened + timedelta(seconds=181) + offset).replace(tzinfo=timezone.utc)
        windows.append({
            "key": toxic_tick_candidate_key(row),
            "start": fetch_start,
            "end": fetch_end,
            "startMs": int(fetch_start.timestamp() * 1000),
            "endMs": int(fetch_end.timestamp() * 1000),
        })
    windows.sort(key=lambda item: (item["start"], item["end"], item["key"]))
    groups: list[dict] = []
    for window in windows:
        can_merge = bool(
            groups
            and (
                window["start"] <= groups[-1]["end"]
                or window["end"] - groups[-1]["start"] <= timedelta(minutes=15)
            )
        )
        if not can_merge:
            groups.append({"start": window["start"], "end": window["end"], "windows": [window]})
            continue
        groups[-1]["end"] = max(groups[-1]["end"], window["end"])
        groups[-1]["windows"].append(window)

    prefetched: dict[str, list] = {}
    for group in groups:
        ticks = toxic_copy_ticks_retry(mt5, symbol, group["start"], group["end"])
        if (ticks is None or not len(ticks)) and len(group["windows"]) > 1:
            for window in group["windows"]:
                individual = toxic_copy_ticks_retry(mt5, symbol, window["start"], window["end"])
                prefetched[window["key"]] = list(individual) if individual is not None else []
            continue
        available = list(ticks) if ticks is not None else []
        for window in group["windows"]:
            sliced = [
                item
                for item in available
                if window["startMs"] <= int(item["time_msc"]) <= window["endMs"]
            ]
            if not sliced and len(group["windows"]) > 1:
                individual = toxic_copy_ticks_retry(mt5, symbol, window["start"], window["end"])
                sliced = list(individual) if individual is not None else []
            prefetched[window["key"]] = sliced
    return prefetched


def toxic_closeable_move(direction: str, entry_bid: float, entry_ask: float, bid: float, ask: float) -> float:
    if direction == "buy":
        return bid - entry_ask
    if direction == "sell":
        return entry_bid - ask
    return 0.0


def toxic_winning_ticks(login: str, rows: list[dict], limit: int = 30) -> dict:
    eligible_rows = [
        row for row in rows
        if normalize_text(row.get("platform")).upper() in {"MT4", "MT5"}
        and parse_trade_time(row.get("open_time_msc") or row.get("open_time"))
        and numeric_value(row.get("open_price"))
    ]
    eligible_count = len(eligible_rows)
    winning_candidates = [row for row in eligible_rows if toxic_trade_net(row) > 0]
    winning_eligible_count = len(winning_candidates)
    winning_candidates.sort(key=lambda row: parse_trade_time(row.get("open_time_msc") or row.get("open_time")) or datetime.min)
    winning_sample = toxic_sample_even(winning_candidates, limit)
    all_sample = toxic_sample_even(sorted(eligible_rows, key=lambda row: parse_trade_time(row.get("open_time_msc") or row.get("open_time")) or datetime.min), limit)
    winning_sample_keys = {toxic_tick_candidate_key(row) for row in winning_sample}
    candidate_map = {toxic_tick_candidate_key(row): row for row in [*winning_sample, *all_sample]}
    candidates = list(candidate_map.values())
    if not candidates:
        return {"available": False, "reason": "没有可对齐的 MT4/MT5 订单", "accountOrders": eligible_count, "candidateOrders": winning_eligible_count}
    try:
        import MetaTrader5 as mt5
    except ImportError:
        return {"available": False, "reason": "当前 Python 环境没有 MetaTrader5 Tick 读取模块"}
    results: list[dict] = []
    errors: list[str] = []
    sources: list[dict] = []
    mappings: list[dict] = []
    groups: dict[tuple[str, str, str], list[dict]] = defaultdict(list)
    for row in candidates:
        platform = normalize_text(row.get("platform")).upper()
        server = normalize_text(row.get("server")) or "未指定服务器"
        configured_terminal = normalize_text(TOXIC_MT5_TERMINALS.get(server))
        if platform == "MT4" and not configured_terminal:
            errors.append(f"{server} 未配置供 MT4 订单使用的 MT5 Terminal 行情源")
            continue
        quote_config = TOXIC_MT5_QUOTE_ACCOUNTS.get(server)
        if platform == "MT4" and (not isinstance(quote_config, dict) or not mysql_int(quote_config.get("login"))):
            errors.append(f"{server} 未配置供 MT4 订单使用的 MT5 行情账号")
            continue
        terminal_path = configured_terminal or str(TRADE_KLINE_TERMINAL)
        groups[(platform, server, terminal_path)].append(row)
    with TOXIC_TICK_LOCK:
        for (database_platform, database_server, terminal_path), server_rows in groups.items():
            if not Path(terminal_path).exists():
                errors.append(f"{database_server} 未配置可用 Terminal：{terminal_path}")
                continue
            if not mt5.initialize(path=terminal_path, timeout=10000):
                errors.append(f"{database_server} Terminal 连接失败：{mt5.last_error()}")
                continue
            original_account = mt5.account_info()
            original_login = mysql_int(getattr(original_account, "login", 0))
            original_server = normalize_text(getattr(original_account, "server", ""))
            quote_config = TOXIC_MT5_QUOTE_ACCOUNTS.get(database_server) or {}
            quote_login = mysql_int(quote_config.get("login")) if isinstance(quote_config, dict) else 0
            quote_server = normalize_text(quote_config.get("server")) if isinstance(quote_config, dict) else ""
            switched_login = False
            if quote_login and (original_login != quote_login or (quote_server and original_server != quote_server)):
                mt5.login(quote_login, server=quote_server or None, timeout=10000)
                active_quote = mt5.account_info()
                active_login = mysql_int(getattr(active_quote, "login", 0))
                active_server = normalize_text(getattr(active_quote, "server", ""))
                if active_login != quote_login or (quote_server and active_server != quote_server):
                    errors.append(f"{database_server} 行情账号登录失败：需要 {quote_login} / {quote_server or '-'}，当前 {active_login or '-'} / {active_server or '-'}")
                    if original_login:
                        mt5.login(original_login, server=original_server or None, timeout=10000)
                    mt5.shutdown()
                    continue
                switched_login = True
            account_info = mt5.account_info()
            terminal_info = mt5.terminal_info()
            terminal_server = normalize_text(getattr(account_info, "server", ""))
            source_info = {
                "databasePlatform": database_platform,
                "databaseServer": database_server,
                "terminalPath": terminal_path,
                "terminalServer": terminal_server,
                "terminalLogin": normalize_text(getattr(account_info, "login", "")),
                "connected": bool(getattr(terminal_info, "connected", False)),
                "configuredForServer": database_server in TOXIC_MT5_TERMINALS,
                "quoteAccountConfigured": bool(quote_login),
                "temporaryQuoteLogin": switched_login,
            }
            sources.append(source_info)
            try:
                by_symbol: dict[str, list[dict]] = defaultdict(list)
                for row in server_rows:
                    by_symbol[normalize_text(row.get("symbol"))].append(row)
                for report_symbol, symbol_rows in by_symbol.items():
                    try:
                        mt5_symbol = toxic_mt5_symbol_for(mt5, report_symbol)
                    except Exception as exc:
                        errors.append(f"{database_server} / {report_symbol}：{exc}")
                        continue
                    mapping = None if database_platform == "MT4" else toxic_cached_tick_mapping(login, report_symbol)
                    mapping = mapping or toxic_live_tick_mapping(mt5, report_symbol, mt5_symbol, symbol_rows)
                    if not mapping:
                        mapping = {"reportSymbol": report_symbol, "mt5Symbol": mt5_symbol, "timeMode": "未校准，按GMT尝试", "hourDelta": 0, "source": "默认"}
                    mapping = {**mapping, "databaseServer": database_server, "terminalServer": terminal_server}
                    mappings.append(mapping)
                    hour_delta = mysql_int(mapping.get("hourDelta"))
                    prefetched_ticks = toxic_prefetch_tick_windows(
                        mt5,
                        mt5_symbol,
                        symbol_rows,
                        hour_delta,
                    )
                    for row in symbol_rows:
                        symbol = normalize_text(row.get("symbol"))
                        opened = parse_trade_time(row.get("open_time_msc") or row.get("open_time"))
                        closed = parse_trade_time(row.get("close_time_msc") or row.get("close_time")) or (opened + timedelta(seconds=60))
                        query_opened = opened + timedelta(hours=hour_delta)
                        legacy_end = min(closed, opened + timedelta(seconds=60))
                        query_window_seconds = max((legacy_end - opened).total_seconds(), 1.0)
                        ticks = prefetched_ticks.get(toxic_tick_candidate_key(row), [])
                        if ticks is None or not len(ticks):
                            terminal_label = terminal_server or Path(terminal_path).parent.name
                            errors.append(f"{database_server} / {terminal_label} / {mt5_symbol} / {trade_time_text(opened)} ({mapping.get('timeMode')}) 返回0 Tick")
                            continue
                        open_ms = int(query_opened.replace(tzinfo=timezone.utc).timestamp() * 1000)
                        legacy_end_ms = int((legacy_end + timedelta(hours=hour_delta)).replace(tzinfo=timezone.utc).timestamp() * 1000)
                        analysis_start_ms = open_ms - 30_000
                        analysis_end_ms = open_ms + 180_000
                        previous = None
                        event_previous = None
                        event_points: list[tuple[int, float, float]] = []
                        price_win = net_win = None
                        effective_index = 0
                        raw_tick_index = 0
                        direction = clean_trade_type(row.get("type", ""))
                        baseline_mark = baseline_spread = None
                        entry_bid = entry_ask = None
                        favorable_ticks_50 = observed_ticks_50 = 0
                        impact_at_20 = None
                        max_favorable_50 = 0.0
                        open_price = numeric_value(row.get("open_price"))
                        volume = max(numeric_value(row.get("volume")), 0.0)
                        tick_value = numeric_value(row.get("tick_value"))
                        tick_size = numeric_value(row.get("tick_size"))
                        realized_move = abs(numeric_value(row.get("close_price")) - open_price)
                        money_per_price = abs(numeric_value(row.get("profit"))) / realized_move if realized_move > 0 else 0.0
                        costs = sum(numeric_value(row.get(key)) for key in ("commission", "fee", "swap", "taxes"))
                        for tick_item in ticks:
                            tick_ms = int(tick_item["time_msc"])
                            bid, ask = float(tick_item["bid"]), float(tick_item["ask"])
                            if analysis_start_ms <= tick_ms <= analysis_end_ms and event_previous != (bid, ask):
                                event_points.append((tick_ms, bid, ask))
                                event_previous = (bid, ask)
                            if tick_ms <= open_ms:
                                entry_bid, entry_ask = bid, ask
                                baseline_mark = bid if direction == "buy" else ask
                                baseline_spread = max(ask - bid, 0.0)
                                previous = (bid, ask)
                                continue
                            if entry_bid is None or entry_ask is None:
                                entry_bid, entry_ask = bid, ask
                                baseline_mark = bid if direction == "buy" else ask
                                baseline_spread = max(ask - bid, 0.0)
                                previous = (bid, ask)
                                continue
                            if tick_ms > legacy_end_ms:
                                continue
                            raw_tick_index += 1
                            closeable_move = toxic_closeable_move(direction, entry_bid, entry_ask, bid, ask)
                            if price_win is None and closeable_move > 0:
                                price_win = raw_tick_index
                            if net_win is None and ((tick_size > 0 and tick_value > 0 and volume > 0) or money_per_price > 0):
                                floating = (closeable_move / tick_size * tick_value * volume if tick_size > 0 and tick_value > 0 else closeable_move * money_per_price) + costs
                                if floating > 0:
                                    net_win = raw_tick_index
                            if previous == (bid, ask):
                                continue
                            previous = (bid, ask)
                            effective_index += 1
                            mark = bid if direction == "buy" else ask
                            gross_move = mark - open_price if direction == "buy" else open_price - mark
                            if baseline_mark is not None and effective_index <= 50:
                                quote_impact = mark - baseline_mark if direction == "buy" else baseline_mark - mark
                                observed_ticks_50 += 1
                                favorable_ticks_50 += int(quote_impact > 0)
                                max_favorable_50 = max(max_favorable_50, quote_impact)
                                if effective_index == 20:
                                    impact_at_20 = quote_impact
                        pre_points = [item for item in event_points if item[0] <= open_ms]
                        post_points_60 = [item for item in event_points if open_ms < item[0] <= open_ms + 60_000]
                        post_points_180 = [item for item in event_points if open_ms < item[0] <= analysis_end_ms]
                        baseline_point = pre_points[-1] if pre_points else (post_points_180[0] if post_points_180 else None)

                        def event_mark(item):
                            return item[1] if direction == "buy" else item[2]

                        def directional_move(start_mark, end_mark):
                            if start_mark is None or end_mark is None:
                                return None
                            return end_mark - start_mark if direction == "buy" else start_mark - end_mark

                        def mark_at(seconds):
                            values = [item for item in post_points_180 if item[0] <= open_ms + seconds * 1000]
                            return event_mark(values[-1]) if values else (event_mark(baseline_point) if baseline_point else None)

                        event_baseline_mark = event_mark(baseline_point) if baseline_point else None
                        pre_start_mark = event_mark(pre_points[0]) if pre_points else None
                        pre_move_30 = directional_move(pre_start_mark, event_baseline_mark)
                        post_move_10 = directional_move(event_baseline_mark, mark_at(10))
                        post_move_60 = directional_move(event_baseline_mark, mark_at(60))
                        post_move_180 = directional_move(event_baseline_mark, mark_at(180))
                        acceleration_10 = None
                        if post_move_10 is not None:
                            acceleration_10 = post_move_10 - (pre_move_30 or 0.0) / 3
                        persistence_60 = None
                        if post_move_10 is not None and post_move_10 > 0 and post_move_60 is not None:
                            persistence_60 = max(-200.0, min(300.0, post_move_60 / post_move_10 * 100))
                        reversal_180 = post_move_180 - post_move_60 if post_move_180 is not None and post_move_60 is not None else None
                        pre_spreads = [max(item[2] - item[1], 0.0) for item in pre_points]
                        post_spreads = [max(item[2] - item[1], 0.0) for item in post_points_60]
                        pre_spread_median = statistics.median(pre_spreads) if pre_spreads else baseline_spread
                        post_spread_median = statistics.median(post_spreads) if post_spreads else None
                        spread_expansion = post_spread_median / pre_spread_median if pre_spread_median and post_spread_median is not None else None
                        post_gaps = [
                            (post_points_180[index][0] - post_points_180[index - 1][0]) / 1000
                            for index in range(1, len(post_points_180))
                        ]
                        pre_tick_rate = len(pre_points) * 2.0
                        post_tick_rate = float(len(post_points_60))
                        results.append({
                            "_candidateKey": toxic_tick_candidate_key(row),
                            "ticket": normalize_text(row.get("ticket")), "symbol": symbol, "mt5Symbol": mt5_symbol,
                            "databaseServer": database_server, "terminalServer": terminal_server,
                            "timeMode": mapping.get("timeMode"), "hourDelta": hour_delta, "volume": volume,
                            "priceWinTick": price_win, "netWinTick": net_win, "rawTicks": raw_tick_index, "effectiveTicks": effective_index,
                            "tickRatePerMinute": rounded(effective_index / query_window_seconds * 60, 1),
                            "preTickRatePerMinute": rounded(pre_tick_rate, 1),
                            "postTickRatePerMinute": rounded(post_tick_rate, 1),
                            "maxPostQuoteGapSeconds": rounded(max(post_gaps), 3) if post_gaps else None,
                            "spreadExpansionRatio": rounded(spread_expansion, 3) if spread_expansion is not None else None,
                            "preMove30": rounded(pre_move_30, 8) if pre_move_30 is not None else None,
                            "postMove10": rounded(post_move_10, 8) if post_move_10 is not None else None,
                            "postMove60": rounded(post_move_60, 8) if post_move_60 is not None else None,
                            "postMove180": rounded(post_move_180, 8) if post_move_180 is not None else None,
                            "acceleration10": rounded(acceleration_10, 8) if acceleration_10 is not None else None,
                            "persistence60": rounded(persistence_60, 1) if persistence_60 is not None else None,
                            "reversal180": rounded(reversal_180, 8) if reversal_180 is not None else None,
                            "lowLiquidityTimeWindow": query_opened.hour >= 21 or query_opened.hour < 1,
                            "realizedNet": rounded(toxic_trade_net(row)),
                            "favorableTickRatio50": rounded(favorable_ticks_50 / observed_ticks_50 * 100, 1) if observed_ticks_50 else None,
                            "impactAt20": rounded(impact_at_20, 8) if impact_at_20 is not None else None,
                            "impactSpreadMultiple": rounded(max_favorable_50 / baseline_spread, 3) if baseline_spread and baseline_spread > 0 else None,
                        })
            finally:
                if switched_login and original_login:
                    mt5.login(original_login, server=original_server or None, timeout=10000)
                    restored = mt5.account_info()
                    if mysql_int(getattr(restored, "login", 0)) != original_login:
                        errors.append(f"Terminal 原登录恢复失败：需要 {original_login} / {original_server or '-'}")
                mt5.shutdown()
    if not results:
        reason = errors[0] if errors else "候选时间段没有可用 Tick"
        if sources and any("demo" in normalize_text(item.get("terminalServer")).lower() for item in sources):
            reason += "；当前 Terminal 连接的是 Demo 服务器，请在对应实盘行情 Terminal 登录后重试"
        return {"available": False, "reason": reason, "accountOrders": eligible_count, "candidateOrders": winning_eligible_count, "sampledOrders": len(winning_sample), "sources": sources, "mappings": mappings, "errors": errors[:12]}
    winning_results = toxic_tick_winning_sample_results(results, winning_sample_keys)
    weighted = [(item["priceWinTick"], item["volume"]) for item in winning_results if item["priceWinTick"] is not None]
    net_weighted = [(item["netWinTick"], item["volume"]) for item in winning_results if item["netWinTick"] is not None]
    favorable_weighted = [(item["favorableTickRatio50"], item["volume"]) for item in winning_results if item.get("favorableTickRatio50") is not None]
    spread_impact_weighted = [(item["impactSpreadMultiple"], item["volume"]) for item in winning_results if item.get("impactSpreadMultiple") is not None]
    pre_rate_weighted = [(item["preTickRatePerMinute"], item["volume"]) for item in results if item.get("preTickRatePerMinute") is not None]
    post_rate_weighted = [(item["postTickRatePerMinute"], item["volume"]) for item in results if item.get("postTickRatePerMinute") is not None]
    spread_expansion_weighted = [(item["spreadExpansionRatio"], item["volume"]) for item in results if item.get("spreadExpansionRatio") is not None]
    max_gap_weighted = [(item["maxPostQuoteGapSeconds"], item["volume"]) for item in results if item.get("maxPostQuoteGapSeconds") is not None]
    total_volume = sum(item["volume"] for item in winning_results)
    all_sampled_volume = sum(item["volume"] for item in results)
    win1_volume = sum(item["volume"] for item in winning_results if item["priceWinTick"] == 1)
    win3_volume = sum(item["volume"] for item in winning_results if item["priceWinTick"] is not None and item["priceWinTick"] <= 3)
    win10_volume = sum(item["volume"] for item in winning_results if item["priceWinTick"] is not None and item["priceWinTick"] <= 10)
    win20_volume = sum(item["volume"] for item in winning_results if item["priceWinTick"] is not None and item["priceWinTick"] <= 20)
    win50_volume = sum(item["volume"] for item in winning_results if item["priceWinTick"] is not None and item["priceWinTick"] <= 50)
    net_win1_volume = sum(item["volume"] for item in winning_results if item["netWinTick"] == 1)
    net_win3_volume = sum(item["volume"] for item in winning_results if item["netWinTick"] is not None and item["netWinTick"] <= 3)
    positive_impact_20_volume = sum(item["volume"] for item in winning_results if item.get("impactAt20") is not None and item["impactAt20"] > 0)
    win1_orders = sum(1 for item in winning_results if item["priceWinTick"] == 1)
    win3_orders = sum(1 for item in winning_results if item["priceWinTick"] is not None and item["priceWinTick"] <= 3)
    win10_orders = sum(1 for item in winning_results if item["priceWinTick"] is not None and item["priceWinTick"] <= 10)
    all_volume = sum(item["volume"] for item in results)
    low_liquidity_results = [item for item in results if numeric_value(item.get("tickRatePerMinute")) < 60]
    low_liquidity_volume = sum(item["volume"] for item in low_liquidity_results)
    rollover_results = [item for item in results if item.get("lowLiquidityTimeWindow")]
    rollover_volume = sum(item["volume"] for item in rollover_results)
    event_impact_10_volume = sum(item["volume"] for item in winning_results if numeric_value(item.get("postMove10")) > 0)
    event_acceleration_10_volume = sum(item["volume"] for item in winning_results if numeric_value(item.get("acceleration10")) > 0)
    event_persistence_60_volume = sum(item["volume"] for item in winning_results if numeric_value(item.get("persistence60")) >= 50)
    preexisting_trend_volume = sum(item["volume"] for item in winning_results if numeric_value(item.get("preMove30")) > 0)
    reversal_180_volume = sum(item["volume"] for item in winning_results if numeric_value(item.get("reversal180")) < 0 and numeric_value(item.get("postMove60")) > 0)
    liquidity_drop_results = [
        item for item in results
        if numeric_value(item.get("preTickRatePerMinute")) > 0
        and numeric_value(item.get("postTickRatePerMinute")) < max(30, numeric_value(item.get("preTickRatePerMinute")) * 0.5)
    ]
    liquidity_drop_volume = sum(item["volume"] for item in liquidity_drop_results)
    return {
        "available": True,
        "accountOrders": eligible_count,
        "candidateOrders": winning_eligible_count,
        "sampledOrders": len(winning_sample),
        "analyzedOrders": len(winning_results),
        "allSampledOrders": len(all_sample),
        "allAnalyzedOrders": len(results),
        "coverageRatio": rounded(len(winning_results) / len(winning_sample) * 100, 1) if winning_sample else 0,
        "accountCoverageRatio": rounded(len(winning_results) / winning_eligible_count * 100, 1) if winning_eligible_count else 0,
        "priceWinTickMedian": rounded(toxic_weighted_median(weighted), 1) if weighted else None,
        "netWinTickMedian": rounded(toxic_weighted_median(net_weighted), 1) if net_weighted else None,
        "win1VolumeRatio": rounded(win1_volume / total_volume * 100, 1) if total_volume else 0,
        "win3VolumeRatio": rounded(win3_volume / total_volume * 100, 1) if total_volume else 0,
        "win10VolumeRatio": rounded(win10_volume / total_volume * 100, 1) if total_volume else 0,
        "win1OrderRatio": rounded(win1_orders / len(winning_results) * 100, 1) if winning_results else 0,
        "win3OrderRatio": rounded(win3_orders / len(winning_results) * 100, 1) if winning_results else 0,
        "win10OrderRatio": rounded(win10_orders / len(winning_results) * 100, 1) if winning_results else 0,
        "win20VolumeRatio": rounded(win20_volume / total_volume * 100, 1) if total_volume else 0,
        "win50VolumeRatio": rounded(win50_volume / total_volume * 100, 1) if total_volume else 0,
        "netWin1VolumeRatio": rounded(net_win1_volume / total_volume * 100, 1) if total_volume else 0,
        "netWin3VolumeRatio": rounded(net_win3_volume / total_volume * 100, 1) if total_volume else 0,
        "positiveImpact20VolumeRatio": rounded(positive_impact_20_volume / total_volume * 100, 1) if total_volume else 0,
        "favorableTickRatio50Median": rounded(toxic_weighted_median(favorable_weighted), 1) if favorable_weighted else None,
        "impactSpreadMultipleMedian": rounded(toxic_weighted_median(spread_impact_weighted), 3) if spread_impact_weighted else None,
        "preTickRatePerMinuteMedian": rounded(toxic_weighted_median(pre_rate_weighted), 1) if pre_rate_weighted else None,
        "postTickRatePerMinuteMedian": rounded(toxic_weighted_median(post_rate_weighted), 1) if post_rate_weighted else None,
        "spreadExpansionRatioMedian": rounded(toxic_weighted_median(spread_expansion_weighted), 3) if spread_expansion_weighted else None,
        "maxPostQuoteGapSecondsMedian": rounded(toxic_weighted_median(max_gap_weighted), 3) if max_gap_weighted else None,
        "eventImpact10VolumeRatio": rounded(event_impact_10_volume / total_volume * 100, 1) if total_volume else 0,
        "eventAcceleration10VolumeRatio": rounded(event_acceleration_10_volume / total_volume * 100, 1) if total_volume else 0,
        "eventPersistence60VolumeRatio": rounded(event_persistence_60_volume / total_volume * 100, 1) if total_volume else 0,
        "preexistingTrendVolumeRatio": rounded(preexisting_trend_volume / total_volume * 100, 1) if total_volume else 0,
        "reversal180VolumeRatio": rounded(reversal_180_volume / total_volume * 100, 1) if total_volume else 0,
        "liquidityDropOrderRatio": rounded(len(liquidity_drop_results) / len(results) * 100, 1) if results else 0,
        "liquidityDropVolumeRatio": rounded(liquidity_drop_volume / all_sampled_volume * 100, 1) if all_sampled_volume else 0,
        "lowLiquidityOrderRatio": rounded(len(low_liquidity_results) / len(results) * 100, 1),
        "lowLiquidityVolumeRatio": rounded(low_liquidity_volume / all_volume * 100, 1) if all_volume else 0,
        "rolloverWindowOrderRatio": rounded(len(rollover_results) / len(results) * 100, 1),
        "rolloverWindowVolumeRatio": rounded(rollover_volume / all_volume * 100, 1) if all_volume else 0,
        "tickRatePerMinuteMedian": rounded(statistics.median(item["tickRatePerMinute"] for item in results), 1),
        "unresolvedOrders": sum(1 for item in winning_results if item["priceWinTick"] is None),
        "netUnresolvedOrders": sum(1 for item in winning_results if item["netWinTick"] is None),
        "orders": [{key: value for key, value in item.items() if key != "_candidateKey"} for item in results[:20]],
        "sources": sources,
        "mappings": mappings,
        "errors": errors[:12],
        "definition": "赢点只统计最终盈利单，并按开仓后的原始Tick顺序计算：买单首次满足Bid>入场Ask，卖单首次满足Ask<入场Bid；动态流动性使用开仓前30秒和开仓后180秒的去重Bid/Ask，比较Tick频率、点差、入场前趋势、入场后加速、持续和回撤。",
    }


def toxic_result(type_id: str, score: float, stage: str, summary: str, metrics: list[dict], triggers: list[str], evidence: list[str], limitations: list[str], confidence: int, analysis: list[dict] | None = None) -> dict:
    item = TOXIC_CHECK_TYPE_MAP[type_id]
    score = rounded(max(0.0, min(100.0, score)), 1)
    return {
        "type": type_id, "label": item["label"], "score": score, "level": toxic_level(score),
        "stage": stage, "confidence": max(0, min(100, int(confidence))), "summary": summary,
        "metrics": metrics, "triggeredRules": triggers, "evidenceOrders": [value for value in evidence if value],
        "limitations": limitations, "requiresTick": item["requiresTick"], "analysis": analysis or [],
    }


def toxic_push_analysis(score: float, sync_score: float, serial_score: float, burst_score: float, push_behavior: dict, sync: dict | None, tick: dict | None, finance: dict, mixed_episodes: dict | None = None) -> list[dict]:
    if score >= 90:
        conclusion = "高度疑似存在推盘或有组织的协同交易，建议优先人工复核。"
    elif score >= 75:
        conclusion = "推盘嫌疑较高，已经形成可重复的异常交易模式。"
    elif score >= 60:
        conclusion = "发现了一些可疑交易习惯，但现有证据还不足以直接认定推盘。"
    else:
        conclusion = "目前没有形成足够完整的推盘证据链，暂时只建议继续观察。"

    dominant = max((sync_score, "sync"), (serial_score, "serial"), (burst_score, "burst"))[1]
    if dominant == "sync" and sync and sync.get("available"):
        sampled = mysql_int(sync.get("sampledOrders"))
        ratio = numeric_value(sync.get("coordinatedMatchedRatio", sync.get("matchedRatio")))
        peers = mysql_int(sync.get("recurringPeerAccounts", sync.get("peerAccounts")))
        strongest = numeric_value(sync.get("maxPeerRatio"))
        close_ratio = numeric_value(sync.get("coordinatedCloseRatio", sync.get("closeMatchedRatio")))
        episode_ready = (
            mysql_int(push_behavior.get("coreOrders")) >= 8
            and toxic_push_has_session_evidence(push_behavior)
            and numeric_value(push_behavior.get("concentratedCoreVolumeRatio")) >= 70
            and numeric_value(push_behavior.get("coreShortHoldVolumeRatio")) >= 70
            and numeric_value(push_behavior.get("quietGapRatio")) >= 60
        )
        if episode_ready:
            detail = f"抽查的 {sampled} 笔订单中，约 {ratio}% 会和同一批账户在前后 2 秒内交易同一品种、同一方向；这种关系反复出现在 {peers} 个账户上，最紧密的单一账户覆盖约 {strongest}% 的订单。"
            if close_ratio >= 20:
                detail += f"其中约 {close_ratio}% 连平仓时间也接近，因此不像活跃行情中偶然同向下单。"
        else:
            detail = f"原始时间匹配虽然覆盖约 {ratio}% 的订单并涉及 {peers} 个账户，但本账户只有约 {push_behavior.get('concentratedCoreVolumeRatio', 0)}% 的主要仓位落在成组交易时段，没有形成集中交易后停手的节奏；这些同步先按活跃行情中的碰撞处理，不作为主要推盘证据。"
    elif dominant == "serial":
        if numeric_value(push_behavior.get("concentratedCoreVolumeRatio")) >= 70:
            detail = (
                f"主要仓位集中出现在 {push_behavior.get('concentratedSessions', 0)} 段交易时段，这些时段覆盖约 {push_behavior.get('concentratedCoreVolumeRatio', 0)}% 的主要仓位；"
                f"约 {push_behavior.get('coreShortHoldVolumeRatio', 0)}% 的主要仓位在一小时内结束。相对该账户典型做单间隔，动态停手强度约为 {push_behavior.get('quietGapRatio', 0)}%。"
                "这个“集中进场、短时持仓、随后长时间安静”的组合，是本次判断的核心。"
            )
        else:
            detail = (
                f"目前只有约 {push_behavior.get('concentratedCoreVolumeRatio', 0)}% 的主要仓位落在成组交易时段，尚不能说明订单明显集中；"
                f"约 {push_behavior.get('coreShortHoldVolumeRatio', 0)}% 的主要仓位在一小时内结束，相对自身做单节奏的动态停手强度约为 {push_behavior.get('quietGapRatio', 0)}%。"
                "这些特征只满足了一部分，彼此没有组成完整的推盘节奏。"
            )
        if numeric_value(push_behavior.get("camouflageOrderRatio")) >= 5 and numeric_value(push_behavior.get("camouflageVolumeRatio")) <= 10:
            if numeric_value(push_behavior.get("concentratedCoreVolumeRatio")) >= 70:
                detail += f"另外有 {push_behavior.get('camouflageOrders', 0)} 笔明显偏小的订单，只占约 {push_behavior.get('camouflageVolumeRatio', 0)}% 的手数，符合小手数穿插交易的特征，但它只作为辅助证据。"
            else:
                detail += f"另外有 {push_behavior.get('camouflageOrders', 0)} 笔明显偏小的订单，只占约 {push_behavior.get('camouflageVolumeRatio', 0)}% 的手数；由于主要仓位本身并不集中，这些小单也可能只是普通试单，不能解释成伪装。"
    else:
        detail = "主要异常来自短时间内连续堆入同方向仓位。单次集中下单不能单独证明推盘，但如果反复出现，就需要结合下单后的报价变化复核。"
    if mixed_episodes and mixed_episodes.get("confirmed"):
        detail = (
            f"账户整体同时存在普通、网格或摊平交易，但另外识别到 {mixed_episodes.get('confirmedEpisodes', 0)} 个独立协同轮次，"
            f"合计覆盖 {mixed_episodes.get('confirmedOrders', 0)} 笔核心订单和约 {mixed_episodes.get('confirmedVolumeRatio', 0)}% 的核心手数；"
            f"固定关联账户 {mixed_episodes.get('strongestPeer') or '-'} 在 {mixed_episodes.get('strongestPeerEpisodes', 0)} 个轮次中反复出现。"
            "这些轮次分别满足方向集中、短持仓、固定同伙同步开平仓和Tick支持，因此不被账户其他时段的网格交易抵消。"
        )

    loss_rate = numeric_value(push_behavior.get("lossRate"))
    loss_volume_ratio = numeric_value(push_behavior.get("lossVolumeRatio"))
    if loss_rate:
        if score >= 60:
            loss_text = f"账户有约 {loss_rate}% 的亏损单，涉及约 {loss_volume_ratio}% 的交易手数。亏损并没有抹掉上面的同步或重复行为；它们可能是失败尝试、方向做反或用于打乱画像的离群单。只有当亏损交易占据主要手数和主要交易时段时，才会明显削弱判断。"
        else:
            loss_text = f"账户有约 {loss_rate}% 的亏损单，涉及约 {loss_volume_ratio}% 的交易手数。亏损本身既不能证明、也不能排除推盘；本次没有判为高风险，是因为同步、重复执行和报价冲击都还不够强，而不是因为账户出现了亏损。"
    else:
        loss_text = "没有明显亏损单。盈亏只作为辅助信息，真正决定判断的是下单行为能否稳定重复。"

    staggered_ratio = numeric_value(push_behavior.get("staggeredAddOnOrderRatio"))
    cohesive_ratio = numeric_value(push_behavior.get("cohesiveBatchOrderRatio"))
    cohesive_campaign_ratio = numeric_value(push_behavior.get("cohesiveCampaignOrderRatio"))
    non_overlap_ratio = numeric_value(push_behavior.get("nonOverlapRatio"))
    coordinated_close_ratio = numeric_value((sync or {}).get("coordinatedCloseRatio", (sync or {}).get("closeMatchedRatio")))
    recurring_peer_accounts = mysql_int((sync or {}).get("recurringPeerAccounts", (sync or {}).get("peerAccounts")))
    synchronized_cohesive_override = bool(
        sync and sync.get("available")
        and cohesive_campaign_ratio >= 50
        and coordinated_close_ratio >= 30
        and recurring_peer_accounts >= 2
    )
    mixed_episode_confirmed = bool(mixed_episodes and mixed_episodes.get("confirmed"))
    if staggered_ratio >= 35 and mixed_episode_confirmed:
        position_text = f"账户整体动态轮次叠仓约占 {staggered_ratio}%，说明确实混有网格或摊平交易；但已独立确认 {mixed_episodes.get('confirmedEpisodes', 0)} 个协同轮次、覆盖约 {mixed_episodes.get('confirmedVolumeRatio', 0)}% 核心手数。网格只作为账户策略标签，不否决其他时间段的协同打盘证据。"
    elif staggered_ratio >= 35 and not synchronized_cohesive_override:
        position_text = f"约 {staggered_ratio}% 的主要订单是在动态做单轮次内压着旧仓继续分波加入，只有约 {non_overlap_ratio}% 的订单等上一轮结束后再开；这些订单没有形成足够集中的整轮平仓，更像网格加仓、摊平成本或扛单，因此作为强反证。"
    elif staggered_ratio >= 35 and synchronized_cohesive_override:
        position_text = f"虽然动态轮次叠仓约占 {staggered_ratio}%，但约 {cohesive_campaign_ratio}% 的主要订单属于同品种同方向分段进场后集中平仓的完整轮次，跨账户平仓同步约为 {coordinated_close_ratio}%。这种结构更像协同建仓，不按普通网格一票否决。"
    elif cohesive_ratio >= 50:
        position_text = f"约 {cohesive_ratio}% 的主要订单属于按账户自身节奏形成并集中平仓的建仓波次，其中完整动态轮次占 {cohesive_campaign_ratio}%；跨轮次继续叠仓约占 {staggered_ratio}%。仓位结构与成堆推盘较一致。"
    else:
        position_text = f"动态轮次叠仓约占 {staggered_ratio}%，集中平仓轮次约占 {cohesive_ratio}%，单仓或上一轮结束后再开的比例约为 {non_overlap_ratio}%。目前仓位结构没有提供强烈的推盘支持或反证。"

    if tick and tick.get("available"):
        analyzed = mysql_int(tick.get("analyzedOrders"))
        candidates = mysql_int(tick.get("candidateOrders"))
        win1 = numeric_value(tick.get("win1OrderRatio", tick.get("win1VolumeRatio")))
        win3 = numeric_value(tick.get("win3OrderRatio", tick.get("win3VolumeRatio")))
        win1_volume = numeric_value(tick.get("win1VolumeRatio"))
        win3_volume = numeric_value(tick.get("win3VolumeRatio"))
        win_median = tick.get("priceWinTickMedian")
        impact = numeric_value(tick.get("positiveImpact20VolumeRatio"))
        favorable = numeric_value(tick.get("favorableTickRatio50Median"))
        tick_text = f"赢点只统计最终盈利单。本次检查了 {candidates} 笔可分析盈利单中的 {analyzed} 笔：约 {win1}% 的盈利单从第 1 个原始 Tick 就开始盈利，约 {win3}% 在前 3 个原始 Tick 内开始盈利；按手数加权后分别为 {win1_volume}% 和 {win3_volume}%。典型盈利单从第 {win_median} 个原始 Tick 开始盈利。赢点越低，说明下单后越快获得报价优势，嫌疑越高。另有约 {impact}% 的样本手数在第 20 个有效报价变化后仍处于有利方向，前 50 个有效报价变化的有利比例中位数为 {favorable}%。赢点使用原始 Tick，流动性和持续冲击使用去重后的有效报价变化，两者不能混用，也不能单独定性。"
        liquidity_text = (
            f"全体订单的 Tick 抽样中，约 {tick.get('lowLiquidityOrderRatio', 0)}% 的订单开仓后每分钟有效报价变化少于 60 次，"
            f"按手数计算为 {tick.get('lowLiquidityVolumeRatio', 0)}%；开仓前30秒与开仓后60秒的有效报价频率中位数分别约为 "
            f"{tick.get('preTickRatePerMinuteMedian', 0)} 和 {tick.get('postTickRatePerMinuteMedian', 0)} 次/分钟，"
            f"约 {tick.get('liquidityDropVolumeRatio', 0)}% 的样本手数出现开仓后流动性明显下降。"
            f"点差扩张中位倍数为 {tick.get('spreadExpansionRatioMedian', 0)}，开仓后最大无报价间隔中位数为 {tick.get('maxPostQuoteGapSecondsMedian', 0)} 秒。"
            f"约 {tick.get('eventAcceleration10VolumeRatio', 0)}% 的盈利样本手数在入场后10秒出现超过入场前节奏的同向加速，"
            f"其中约 {tick.get('eventPersistence60VolumeRatio', 0)}% 的手数把有利移动保持到60秒；"
            f"约 {tick.get('preexistingTrendVolumeRatio', 0)}% 的手数在入场前已经处于同向趋势。"
            f"另有约 {tick.get('rolloverWindowOrderRatio', 0)}% 位于校准后的 UTC 21:00–01:00 rollover 时段。"
            "这些指标用于区分订单后的异常报价变化与追随既有行情，但没有外部基准时仍不能单独证明因果。"
        )
    else:
        tick_text = "本次没有取得可用 Tick，因此只能确认交易行为异常，不能判断下单后报价是否受到推动。"
        liquidity_text = "本次没有取得可用 Tick，无法验证下单时的实际报价密度；不能仅凭服务器时间推断低流动性。"

    if finance.get("available") and mysql_int(finance.get("depositCount")):
        deposit_total = numeric_value(finance.get("depositTotal"))
        withdrawal_total = numeric_value(finance.get("withdrawalTotal"))
        withdrawal_hours = finance.get("lastTradeToWithdrawalHours")
        withdrawal_ratio = withdrawal_total / deposit_total * 100 if deposit_total else 0
        funding_text = f"累计入金约 {rounded(deposit_total)}，累计出金约 {rounded(withdrawal_total)}，出金约占入金 {rounded(withdrawal_ratio,1)}%。"
        if withdrawal_hours is not None:
            funding_text += f"最后一笔交易结束后约 {withdrawal_hours} 小时出现后续出金。"
        else:
            funding_text += "没有找到发生在最后一笔交易之后的出金记录。"
        funding_text += "小额入金和快速出金只能加强已有推盘证据，不能单独定性。"
    elif finance.get("available"):
        funding_text = "资金流水中没有识别到真实入金，无法判断入金规模和交易后出金速度。"
    else:
        funding_text = f"资金行为未验证：{finance.get('reason', '当前数据源没有可用的出入金明细')}。"

    if sync and sync.get("available"):
        sync_ratio = numeric_value(sync.get("coordinatedMatchedRatio", sync.get("matchedRatio")))
        sync_peers = mysql_int(sync.get("recurringPeerAccounts", sync.get("peerAccounts")))
        episode_confirmed_for_text = (
            mysql_int(push_behavior.get("coreOrders")) >= 8
            and toxic_push_has_session_evidence(push_behavior)
            and numeric_value(push_behavior.get("concentratedCoreVolumeRatio")) >= 70
            and numeric_value(push_behavior.get("coreShortHoldVolumeRatio")) >= 70
            and numeric_value(push_behavior.get("quietGapRatio")) >= 60
        )
        if sync_peers >= 2 and sync_ratio >= 30 and episode_confirmed_for_text:
            coordination_text = f"全平台找到 {sync_peers} 个反复关联账户，约 {sync_ratio}% 的订单与这些账户在前后 2 秒内交易同一品种和方向。多人协同条件成立。"
        elif sync_peers >= 2 and sync_ratio >= 30:
            coordination_text = f"找到 {sync_peers} 个反复关联账户，涉及约 {sync_ratio}% 的订单。这些账户作为疑似同伙线索展示并提供加分，但同步本身不能单独认定推盘。"
        else:
            coordination_text = f"当前规则未找到达到反复匹配门槛的疑似同伙，协调订单约占 {sync_ratio}%。这只表示暂未定位协同账户，不构成排除或降低本账户推盘嫌疑的依据。"
    else:
        coordination_text = f"全平台协同未验证：{(sync or {}).get('reason', '跨账户数据不可用')}。协同只用于定位疑似同伙和增加证据，不影响对本账户自身行为与报价推动的判断。"

    if mixed_episodes and mixed_episodes.get("available"):
        if mixed_episodes.get("confirmed"):
            mixed_text = f"确认 {mixed_episodes.get('confirmedEpisodes', 0)} 个重复协同轮次，覆盖约 {mixed_episodes.get('confirmedVolumeRatio', 0)}% 核心手数；最常出现的固定同伙为 {mixed_episodes.get('strongestPeer') or '-'}，反复出现在 {mixed_episodes.get('strongestPeerEpisodes', 0)} 个轮次。"
        else:
            mixed_text = f"找到 {mixed_episodes.get('candidateEpisodes', 0)} 个局部候选轮次，但只有 {mixed_episodes.get('confirmedEpisodes', 0)} 个能由同一固定账户反复配合；合计手数、重复次数、同步平仓或Tick交叉证据尚未同时达到门槛。"
    else:
        mixed_text = f"局部事件簇未验证：{(mixed_episodes or {}).get('reason', '当前没有订单级同步映射')}。"
    if mixed_episodes and mixed_episodes.get("eaAttention"):
        execution_text = f"EA订单约占 {mixed_episodes.get('eaOrderRatio', 0)}%，EA手数约占 {mixed_episodes.get('eaVolumeRatio', 0)}%。EA说明执行高度自动化、重复模式更值得复核，但EA本身不代表违规；只有固定同伙重复、同步平仓、经济手数和Tick证据同时成立时才提升风险。"
    elif mixed_episodes and mixed_episodes.get("copyAttention"):
        execution_text = f"带跟单标记的订单约占 {mixed_episodes.get('copyOrderRatio', 0)}%，手数约占 {mixed_episodes.get('copyVolumeRatio', 0)}%。跟单可能是正常复制交易，只作为重点复核线索，不单独定性或加分。"
    else:
        execution_text = "当前核心订单没有达到EA主导或显著跟单标记门槛。"
    boundary = "当前最强证据仍是平台内的协同或重复交易。要最终认定推盘，还需要把本平台报价与同期外部基准报价对比：只有本平台在下单后出现异常移动，而外部市场没有同步变化，证据才完整。"
    return [
        {"title": "结论", "text": conclusion},
        {"title": "主要依据", "text": detail},
        {"title": "亏损怎么理解", "text": loss_text},
        {"title": "仓位结构", "text": position_text},
        {"title": "混合型事件簇", "text": mixed_text},
        {"title": "EA/跟单", "text": execution_text},
        {"title": "报价证据", "text": tick_text},
        {"title": "低流动性", "text": liquidity_text},
        {"title": "资金行为", "text": funding_text},
        {"title": "全平台协同", "text": coordination_text},
        {"title": "还不能确定什么", "text": boundary},
    ]


def toxic_push_evidence_chain(
    score: float,
    confidence: int,
    push_behavior: dict,
    tick: dict | None,
    sync: dict | None,
    finance: dict,
    mixed_episodes: dict | None,
    structure_score: float,
    tick_group_score: float,
    coordination_score: float,
    context_score: float,
    consistency_bonus: float,
    counterevidence_deduction: float,
    single_burst_confirmed: bool = False,
    coordinated_tick_confirmed: bool = False,
    sudden_exposure_confirmed: bool = False,
) -> dict:
    """Turn the numeric result into an auditable reasoning chain for operators."""
    core_orders = mysql_int(push_behavior.get("coreOrders"))
    raw_orders = mysql_int(push_behavior.get("rawOrders"))
    concentrated = numeric_value(push_behavior.get("concentratedCoreVolumeRatio"))
    short_hold = numeric_value(push_behavior.get("coreShortHoldVolumeRatio"))
    quiet = numeric_value(push_behavior.get("quietGapRatio"))
    open_span = numeric_value(push_behavior.get("coreOpenSpanSeconds"))
    same_direction = numeric_value(push_behavior.get("maxSameDirectionRunRatio"))
    win_rate = numeric_value(push_behavior.get("winRate"))
    loss_volume = numeric_value(push_behavior.get("lossVolumeRatio"))
    grid_ratio = max(
        numeric_value(push_behavior.get("staggeredAddOnOrderRatio")),
        numeric_value(push_behavior.get("staggeredAddOnVolumeRatio")),
    )
    facts: list[dict] = []
    observed_risks: list[str] = []
    counterpoints: list[str] = []
    uncertainties: list[str] = []
    next_checks: list[str] = []

    if score >= 90:
        status = "confirmed"
        headline = "多组独立证据已经形成闭环，进入优先人工复核。"
    elif score >= 75:
        status = "high"
        headline = "推盘嫌疑较高；关键证据已经互相支持，但仍需外部报价复核完成因果确认。"
    elif score >= 60:
        status = "warning"
        headline = "存在明确风险信号，当前适合预警和持续复核，尚未达到完整推盘证据链。"
    else:
        status = "unconfirmed"
        headline = "当前未形成完整推盘证据链；这不是风险为零，而是关键证据仍不够闭合。"

    if sudden_exposure_confirmed:
        order_text = (
            f"账户只有 {core_orders} 笔核心订单，却在约 {open_span} 秒内形成同品种、同方向的集中仓位；"
            f"主要仓位全部在一小时内结束，胜率 {win_rate}%，同方向连单比例 {same_direction}%。"
        )
    elif single_burst_confirmed:
        order_text = (
            f"{core_orders} 笔核心订单集中在约 {open_span} 秒内，主要仓位覆盖 {concentrated}%、"
            f"一小时内结束比例 {short_hold}%，同方向连单比例 {same_direction}%，胜率 {win_rate}%。"
        )
    elif mixed_episodes and mixed_episodes.get("confirmed"):
        order_text = (
            f"账户整体可能混有普通或网格交易，但另外确认了 {mixed_episodes.get('confirmedEpisodes', 0)} 个独立协同轮次，"
            f"覆盖约 {mixed_episodes.get('confirmedVolumeRatio', 0)}% 核心手数。"
        )
    else:
        order_text = (
            f"共读取 {raw_orders} 笔订单，过滤后分析 {core_orders} 笔；主要仓位集中度 {concentrated}%，"
            f"一小时内结束比例 {short_hold}%，动态停手强度 {quiet}%，同方向连单比例 {same_direction}%。"
        )
    facts.append({"title": "1. 订单行为", "text": order_text, "strength": "strong" if structure_score >= 30 else "partial"})

    if tick and tick.get("available"):
        analyzed = mysql_int(tick.get("analyzedOrders"))
        impact = numeric_value(tick.get("eventImpact10VolumeRatio"))
        acceleration = numeric_value(tick.get("eventAcceleration10VolumeRatio"))
        persistence = numeric_value(tick.get("eventPersistence60VolumeRatio"))
        positive = numeric_value(tick.get("positiveImpact20VolumeRatio"))
        favorable = numeric_value(tick.get("favorableTickRatio50Median"))
        pretrend = numeric_value(tick.get("preexistingTrendVolumeRatio"))
        reversal = numeric_value(tick.get("reversal180VolumeRatio"))
        tick_text = (
            f"取得 {analyzed} 笔可用 Tick；入场后10秒同向加速 {acceleration}%，"
            f"20个有效报价后仍有利 {positive}%，60秒持续 {persistence}%，"
            f"前50个报价有利比例中位数 {favorable}%。"
        )
        if pretrend >= 60:
            tick_text += f"但其中 {pretrend}% 的手数在入场前已经处于同向趋势，存在正常趋势跟随的替代解释。"
        if reversal >= 60:
            tick_text += f"另外 {reversal}% 的手数在180秒内出现回撤，报价冲击并非全程持续。"
        facts.append({"title": "2. Tick反应", "text": tick_text, "strength": "strong" if tick_group_score >= 18 else "partial"})
        uncertainties.append("当前只验证了本平台 Tick，尚未与同期外部基准报价对照，因此不能最终证明价格移动由该账户造成。")
    else:
        facts.append({"title": "2. Tick反应", "text": "没有取得可用 Tick，因此只能确认订单形态，不能确认下单后报价是否出现异常反应。", "strength": "missing"})
        uncertainties.append("缺少可用 Tick，无法验证入场后报价冲击，也不能仅凭服务器时间推断因果。")

    if sync and sync.get("available"):
        sync_ratio = numeric_value(sync.get("coordinatedMatchedRatio", sync.get("matchedRatio")))
        close_ratio = numeric_value(sync.get("coordinatedCloseRatio", sync.get("closeMatchedRatio")))
        peers = mysql_int(sync.get("recurringPeerAccounts", sync.get("peerAccounts")))
        max_peer = numeric_value(sync.get("maxPeerRatio"))
        if coordinated_tick_confirmed:
            sync_text = f"固定关联账户在开仓同步 {sync_ratio}%、平仓同步 {close_ratio}% 的订单中反复出现，最强单一账户覆盖 {max_peer}%，且 Tick 同时支持。"
        elif peers or sync_ratio:
            sync_text = f"找到 {peers} 个反复关联账户，开仓同步约 {sync_ratio}%，平仓同步约 {close_ratio}%；这提供协同线索，但同步本身不能证明推盘。"
        else:
            sync_text = "本次没有找到达到反复匹配门槛的固定账户；这不会降低本账户自身的推盘嫌疑。"
        facts.append({"title": "3. 跨账户协同", "text": sync_text, "strength": "strong" if coordination_score >= 10 else "partial"})
        if sync_ratio >= 60 and close_ratio < 30:
            uncertainties.append("开仓同步明显高于平仓同步，可能是信号跟随或市场共振，尚不能确认共同操盘。")
        if not peers:
            uncertainties.append("暂未定位固定同伙；这只是协同证据缺失，不是对主体行为的排除。")
        elif sync_ratio < 30:
            uncertainties.append("当前固定账户同步比例较低，尚未提供额外协同确认；这一点不作为主体账户的免责依据。")
    else:
        facts.append({"title": "3. 跨账户协同", "text": "跨账户协同未验证；主体账户内的订单和 Tick 证据仍单独有效。", "strength": "missing"})
        uncertainties.append("跨账户数据不可用，无法判断是否存在固定协同账户。")

    ea_ratio = numeric_value((mixed_episodes or {}).get("eaOrderRatio"))
    copy_ratio = numeric_value((mixed_episodes or {}).get("copyOrderRatio"))
    if finance.get("available") and mysql_int(finance.get("depositCount")):
        deposit = numeric_value(finance.get("depositTotal"))
        withdrawal = numeric_value(finance.get("withdrawalTotal"))
        hours = finance.get("lastTradeToWithdrawalHours")
        funding_text = f"资金流水显示累计入金 {rounded(deposit)}、出金 {rounded(withdrawal)}"
        if hours is not None:
            funding_text += f"，最后交易后约 {hours} 小时出金"
        if sudden_exposure_confirmed:
            funding_text += "；这与突然高暴露链共同增强风险。"
        else:
            funding_text += "；资金行为只能作为辅助，不能单独定性。"
    else:
        funding_text = "没有可用的完整资金流水，无法判断入金规模和交易后出金速度。"
    if ea_ratio >= 80:
        funding_text += f" EA订单占比约 {ea_ratio}%，说明执行自动化程度高，但EA本身不是违规证明。"
    elif copy_ratio >= 30:
        funding_text += f" 跟单标记订单约占 {copy_ratio}%，可能是正常复制执行，需要结合平仓和Tick证据。"
    facts.append({"title": "4. 执行与资金", "text": funding_text, "strength": "supporting" if context_score else "partial"})

    if sudden_exposure_confirmed:
        observed_risks.append(f"少量订单在 {open_span} 秒内突然形成高暴露同向仓位，并伴随EA、强Tick反应和快速资金回转。")
    elif single_burst_confirmed:
        observed_risks.append(f"{core_orders} 笔核心订单在 {open_span} 秒内形成单向密集仓位，短持仓和高胜率同时成立。")
    elif concentrated >= 70 and short_hold >= 70:
        observed_risks.append(f"主要仓位集中度 {concentrated}%，一小时内结束比例 {short_hold}%，订单节奏明显偏离分散交易。")
    elif concentrated >= 50 or same_direction >= 70:
        observed_risks.append(f"订单已经出现局部集中或连续同方向执行，集中度 {concentrated}%，同方向连单比例 {same_direction}%。")
    if tick and tick.get("available") and (
        numeric_value(tick.get("eventImpact10VolumeRatio")) >= 50
        or numeric_value(tick.get("eventPersistence60VolumeRatio")) >= 50
    ):
        observed_risks.append(
            f"订单后报价反应明显：10秒冲击手数 {numeric_value(tick.get('eventImpact10VolumeRatio'))}%，"
            f"60秒持续手数 {numeric_value(tick.get('eventPersistence60VolumeRatio'))}%。"
        )
    if sync and sync.get("available") and numeric_value(sync.get("coordinatedMatchedRatio", sync.get("matchedRatio"))) >= 30:
        observed_risks.append(
            f"跨账户协调开仓约 {numeric_value(sync.get('coordinatedMatchedRatio', sync.get('matchedRatio')))}%，"
            f"协调平仓约 {numeric_value(sync.get('coordinatedCloseRatio', sync.get('closeMatchedRatio')))}%。"
        )
    if ea_ratio >= 80:
        observed_risks.append(f"EA订单占比约 {ea_ratio}%，自动化执行使异常模式更容易重复，需要重点复核。")
    if not observed_risks:
        observed_risks.append("当前只发现零散异常指标，尚未形成稳定、重复且可交叉验证的高风险行为。")

    if grid_ratio >= 35:
        counterpoints.append(f"动态轮次中约 {grid_ratio}% 的订单或手数带有叠仓/摊平特征，普通网格是当前最主要的替代解释。")
    if pretrend >= 60 if tick and tick.get("available") else False:
        counterpoints.append(f"入场前已有同向趋势的手数约 {pretrend}%，部分结果可能是趋势跟随而不是订单推动。")
    if reversal >= 60 if tick and tick.get("available") else False:
        counterpoints.append(f"入场后180秒回撤手数约 {reversal}%，报价优势并非全程稳定。")
    if loss_volume >= 30:
        counterpoints.append(f"亏损手数约 {loss_volume}%，说明账户并非每次都成功；亏损不能自动排除推盘，但会降低单次因果解释。")
    if not counterpoints:
        counterpoints.append("目前没有发现足以推翻主体异常形态的强反证。")

    if counterevidence_deduction > 0:
        uncertainties.append(f"网格/叠仓反证扣分 {rounded(counterevidence_deduction, 1)} 分，模型已将其从总分中扣除。")
    if core_orders < 8:
        uncertainties.append(f"核心订单只有 {core_orders} 笔，重复性有限；当前分数不设少单上限，但置信度会降低。")
    if tick and tick.get("available") and tick_group_score < 18:
        uncertainties.append("Tick只提供部分支持，尚未形成稳定的入场后冲击、持续或赢点链。")
    if not tick or not tick.get("available"):
        next_checks.append("补充对应服务器和品种的 Tick，复核入场前30秒与入场后180秒的报价变化。")
    else:
        next_checks.append("用同期外部基准报价复核：只有本平台单独出现订单后移动，才能完成因果确认。")
    if sync and sync.get("available") and numeric_value(sync.get("coordinatedMatchedRatio", sync.get("matchedRatio"))) >= 30:
        next_checks.append("抽查最强关联账户的开仓、平仓时间和手数，区分固定协同与普通信号跟随。")
    if score < 75:
        missing = []
        if structure_score < 30:
            missing.append("主体订单结构还不够集中或重复")
        if tick_group_score < 18:
            missing.append("Tick冲击/持续证据不足")
        if counterevidence_deduction >= 5:
            missing.append("网格/叠仓反证较强")
        if not missing:
            missing.append("多组证据尚未同时达到高危阈值")
        observed_summary = "；".join(item.rstrip("。") for item in observed_risks[:2])
        reasoning = f"当前分数 {rounded(score, 1)} 只是证据汇总，不代表账户安全。已观察到：{observed_summary}。尚未达到高危的主要原因是：" + "、".join(missing) + "。"
    else:
        reasoning = (
            f"当前分数 {rounded(score, 1)} 的原因不是单项指标，而是订单行为、Tick反应"
            + ("、固定账户同步开平仓" if coordination_score >= 8 else "")
            + ("、EA/资金辅助" if context_score else "")
            + "形成了相互支持的证据链；网格、趋势跟随和普通复制交易等替代解释不足以解释全部现象。"
        )
    return {
        "status": status,
        "headline": headline,
        "reasoning": reasoning,
        "facts": facts,
        "observedRisks": list(dict.fromkeys(observed_risks)),
        "riskPoints": list(dict.fromkeys(observed_risks)),
        "counterpoints": list(dict.fromkeys(counterpoints)),
        "uncertainties": list(dict.fromkeys(uncertainties)),
        "nextChecks": list(dict.fromkeys(next_checks)),
        "scoreBasis": {
            "structure": rounded(structure_score, 1),
            "tick": rounded(tick_group_score, 1),
            "coordination": rounded(coordination_score, 1),
            "context": rounded(context_score, 1),
            "consistencyBonus": rounded(consistency_bonus, 1),
            "counterevidence": rounded(counterevidence_deduction, 1),
            "confidence": confidence,
        },
    }


def calculate_toxic_results(login: str, rows: list[dict], type_ids: list[str], stage: str, finance: dict, tick: dict | None = None, sync: dict | None = None, push_context: dict | None = None) -> list[dict]:
    metrics = trade_metrics(rows)
    push_context = push_context or toxic_build_push_context(rows)
    push_rows = push_context["rows"]
    push_behavior = push_context["behavior"]
    mixed_push_episodes = toxic_push_mixed_episode_summary(push_rows, sync, tick) if "market_pushing" in type_ids else {"available": False}
    push_burst1, push_burst5 = (toxic_burst(push_rows, value) for value in (1, 5))
    burst1, burst5, burst60, burst120, burst300, burst600 = (toxic_burst(rows, value) for value in (1, 5, 60, 120, 300, 600))
    pair = toxic_best_opposite_pair(rows)
    short10 = toxic_volume_ratio(rows, lambda row: toxic_trade_holding(row) <= 10)
    short60 = toxic_volume_ratio(rows, lambda row: toxic_trade_holding(row) <= 60)
    short300 = toxic_volume_ratio(rows, lambda row: toxic_trade_holding(row) <= 300)
    short_rows = [row for row in rows if toxic_trade_holding(row) <= 60]
    short_win_rate = sum(1 for row in short_rows if toxic_trade_net(row) > 0) / len(short_rows) if short_rows else 0.0
    weekend_rows = []
    for row in rows:
        opened = parse_trade_time(row.get("open_time"))
        closed = parse_trade_time(row.get("close_time"))
        if opened and closed and opened.weekday() == 4 and (closed - opened).total_seconds() >= 30 * 3600:
            weekend_rows.append(row)
    equity = numeric_value(finance.get("equity")) if finance.get("available") else 0.0
    credit = numeric_value(finance.get("credit")) if finance.get("available") else 0.0
    cash_equity = max(equity - max(credit, 0.0), 50.0)
    margin = numeric_value(finance.get("margin")) if finance.get("available") else 0.0
    margin_level = numeric_value(finance.get("marginLevel")) if finance.get("available") else 0.0
    margin_pressure = min(margin / max(equity, 1.0), 2.0) / 2.0 if equity else 0.0
    rebate = numeric_value(finance.get("rebate")) if finance.get("available") else 0.0
    net_deposit = max(abs(numeric_value(finance.get("netDeposit"))), 50.0) if finance.get("available") else 50.0
    credit_ratio = credit / cash_equity if credit > 0 else 0.0
    tick_score = 0.0
    tick_context_metrics: list[dict] = []
    if tick and tick.get("available"):
        median_tick = tick.get("priceWinTickMedian")
        median_score = toxic_ramp(median_tick, 1, 6, reverse=True) if median_tick is not None else 0.0
        tick_score = median_score * 0.45 + numeric_value(tick.get("win1VolumeRatio")) * 0.3 + numeric_value(tick.get("win3VolumeRatio")) * 0.25
        terminal_servers = sorted({normalize_text(item.get("terminalServer")) for item in tick.get("sources", []) if normalize_text(item.get("terminalServer"))})
        time_modes = sorted({normalize_text(item.get("timeMode")) for item in tick.get("mappings", []) if normalize_text(item.get("timeMode"))})
        tick_context_metrics = [
            {"label": "盈利单Tick有效样本", "value": f'{tick.get("analyzedOrders", 0)}/{tick.get("sampledOrders", tick.get("candidateOrders", 0))} 单 ({tick.get("coverageRatio", 0)}%)'},
            {"label": "盈利单Tick抽样覆盖", "value": f'{tick.get("analyzedOrders", 0)}/{tick.get("candidateOrders", 0)} 单 ({tick.get("accountCoverageRatio", tick.get("coverageRatio", 0))}%)'},
            {"label": "Terminal行情源", "value": "/".join(terminal_servers) or "未识别"},
            {"label": "时间校准", "value": "/".join(time_modes) or "未识别"},
        ]
    results = []
    for type_id in type_ids:
        triggers: list[str] = []
        limitations: list[str] = []
        evidence: list[str] = []
        item_metrics: list[dict] = []
        confidence = 82 if stage == "deep" else 58
        score = 0.0
        summary = "未见明显形态"
        analysis: list[dict] = []
        evidence_chain: dict | None = None
        if type_id == "market_pushing":
            concentrated_ratio = numeric_value(push_behavior.get("concentratedCoreVolumeRatio"))
            short_core_ratio = numeric_value(push_behavior.get("coreShortHoldVolumeRatio"))
            quiet_gap_ratio = numeric_value(push_behavior.get("quietGapRatio"))
            core_volume_ratio = numeric_value(push_behavior.get("coreVolumeRatio"))
            staggered_addon_ratio = numeric_value(push_behavior.get("staggeredAddOnOrderRatio"))
            staggered_addon_volume_ratio = numeric_value(push_behavior.get("staggeredAddOnVolumeRatio"))
            cohesive_batch_ratio = numeric_value(push_behavior.get("cohesiveBatchOrderRatio"))
            cohesive_campaign_ratio = numeric_value(push_behavior.get("cohesiveCampaignOrderRatio"))
            non_overlap_ratio = numeric_value(push_behavior.get("nonOverlapRatio"))
            coordinated_ratio = numeric_value((sync or {}).get("coordinatedMatchedRatio", (sync or {}).get("matchedRatio")))
            coordinated_volume_ratio = numeric_value((sync or {}).get("coordinatedVolumeRatio", (sync or {}).get("matchedVolumeRatio")))
            coordinated_close_ratio = numeric_value((sync or {}).get("coordinatedCloseRatio", (sync or {}).get("closeMatchedRatio")))
            recurring_peer_accounts = mysql_int((sync or {}).get("recurringPeerAccounts", (sync or {}).get("peerAccounts")))
            synchronized_cohesive_override = bool(
                sync and sync.get("available")
                and cohesive_campaign_ratio >= 50
                and coordinated_close_ratio >= 30
                and recurring_peer_accounts >= 2
            )
            position_structure_ok = bool(
                (staggered_addon_ratio <= 20 and (non_overlap_ratio >= 70 or cohesive_batch_ratio >= 50))
                or synchronized_cohesive_override
            )
            camouflage_score = min(
                toxic_ramp(push_behavior.get("camouflageOrderRatio"), 5, 25),
                toxic_ramp(push_behavior.get("camouflageVolumeRatio"), 1, 15, reverse=True),
            )
            episode_score = concentrated_ratio * .3 + short_core_ratio * .3 + quiet_gap_ratio * .2 + core_volume_ratio * .1 + camouflage_score * .1
            episode_score -= toxic_ramp(staggered_addon_ratio, 20, 60) * .35
            episode_score = max(0.0, episode_score)
            episode_structure = (
                mysql_int(push_behavior.get("coreOrders")) >= 8
                and toxic_push_has_session_evidence(push_behavior)
                and concentrated_ratio >= 70
                and short_core_ratio >= 70
                and position_structure_ok
            )
            episode_confirmed = episode_structure and quiet_gap_ratio >= 60
            raw_burst_score = toxic_ramp(push_burst5["count"], 3, 12) * .35 + toxic_ramp(push_burst5["volume"], 1, 20) * .3 + toxic_ramp(push_burst1["count"], 2, 8) * .2 + margin_pressure * 100 * .15
            burst_score = raw_burst_score * .65 + short_core_ratio * .2 + quiet_gap_ratio * .15
            burst_confirmed = push_burst5["count"] >= 3 and (push_burst5["volume"] >= 1 or margin_pressure >= .5) and short_core_ratio >= 70 and quiet_gap_ratio >= 60 and position_structure_ok
            sync_score = 0.0
            if sync and sync.get("available"):
                raw_sync_score = coordinated_ratio * .35 + coordinated_volume_ratio * .15 + coordinated_close_ratio * .15 + numeric_value(sync.get("maxPeerRatio")) * .2 + toxic_ramp(recurring_peer_accounts, 1, 5) * .15
                sync_score = raw_sync_score * .75 + episode_score * .25
            serial_tick_score = 0.0
            if tick and tick.get("available"):
                win1_signal = (numeric_value(tick.get("win1OrderRatio", tick.get("win1VolumeRatio"))) + numeric_value(tick.get("win1VolumeRatio"))) / 2
                win3_signal = (numeric_value(tick.get("win3OrderRatio", tick.get("win3VolumeRatio"))) + numeric_value(tick.get("win3VolumeRatio"))) / 2
                win10_signal = (numeric_value(tick.get("win10OrderRatio", tick.get("win10VolumeRatio"))) + numeric_value(tick.get("win10VolumeRatio"))) / 2
                winning_point_score = (
                    win1_signal * .55
                    + win3_signal * .3
                    + win10_signal * .15
                )
                liquidity_signal = (
                    numeric_value(tick.get("lowLiquidityOrderRatio"))
                    + numeric_value(tick.get("lowLiquidityVolumeRatio"))
                ) / 2
                causal_liquidity_signal = (
                    numeric_value(tick.get("eventImpact10VolumeRatio")) * .3
                    + numeric_value(tick.get("eventAcceleration10VolumeRatio")) * .3
                    + numeric_value(tick.get("eventPersistence60VolumeRatio")) * .2
                    + numeric_value(tick.get("reversal180VolumeRatio")) * .1
                    + toxic_ramp(tick.get("spreadExpansionRatioMedian"), 1.05, 1.5) * .1
                )
                serial_tick_score = (
                    winning_point_score * .45
                    + numeric_value(tick.get("positiveImpact20VolumeRatio")) * .15
                    + numeric_value(tick.get("favorableTickRatio50Median")) * .1
                    + toxic_ramp(tick.get("impactSpreadMultipleMedian"), .25, 2) * .05
                    + liquidity_signal * .05
                    + causal_liquidity_signal * .2
                )
                causal_liquidity_chain = bool(
                    mysql_int(tick.get("analyzedOrders")) >= 2
                    and numeric_value(tick.get("eventImpact10VolumeRatio")) >= 65
                    and numeric_value(tick.get("eventAcceleration10VolumeRatio")) >= 60
                    and numeric_value(tick.get("eventPersistence60VolumeRatio")) >= 60
                    and numeric_value(tick.get("preexistingTrendVolumeRatio")) <= 60
                    and numeric_value(tick.get("spreadExpansionRatioMedian")) <= 1.5
                )
            else:
                causal_liquidity_chain = False
            core_order_count = mysql_int(push_behavior.get("coreOrders"))
            small_sample_sync_threshold = 100 if core_order_count <= 2 else 50
            small_sample_warning = bool(
                2 <= core_order_count < 8
                and sync and sync.get("available")
                and mysql_int(sync.get("maxPeerMatches")) >= 2
                and numeric_value(sync.get("maxPeerRatio")) >= small_sample_sync_threshold
                and coordinated_ratio >= small_sample_sync_threshold
                and coordinated_close_ratio >= small_sample_sync_threshold
                and recurring_peer_accounts >= 1
                and concentrated_ratio >= 50
                and short_core_ratio >= 80
                and quiet_gap_ratio >= 40
                and position_structure_ok
                and tick and tick.get("available")
                and mysql_int(tick.get("analyzedOrders")) >= 2
                and numeric_value(tick.get("eventImpact10VolumeRatio")) >= 70
                and numeric_value(tick.get("positiveImpact20VolumeRatio")) >= 70
                and numeric_value(tick.get("favorableTickRatio50Median")) >= 50
                and numeric_value(tick.get("eventPersistence60VolumeRatio")) >= 70
                and (
                    bool(mixed_push_episodes.get("eaAttention"))
                    or numeric_value(tick.get("eventAcceleration10VolumeRatio")) >= 60
                )
            )
            serial_score = episode_score * (.85 if tick and tick.get("available") else 1.0) + serial_tick_score * (.15 if tick and tick.get("available") else 0.0)
            score = max(burst_score, sync_score, serial_score)
            evidence = list(dict.fromkeys([*(sync or {}).get("evidenceOrders", []), *push_burst5["tickets"]]))[:20]
            item_metrics = [
                {"label": "订单过滤", "value": f'{push_behavior.get("filteredOrders", 0)}/{push_behavior.get("rawOrders", 0)} 笔核心订单，过滤 {push_behavior.get("excludedOrders", 0)} 笔'},
                {"label": "过滤订单手数占比", "value": f'{push_behavior.get("excludedVolumeRatio", 0)}%'},
                {"label": "EA订单 / 手数", "value": f'{mixed_push_episodes.get("eaOrderRatio", 0)}% / {mixed_push_episodes.get("eaVolumeRatio", 0)}%'},
                {"label": "跟单标记订单 / 手数", "value": f'{mixed_push_episodes.get("copyOrderRatio", 0)}% / {mixed_push_episodes.get("copyVolumeRatio", 0)}%'},
                {"label": "执行方式重点复核", "value": mixed_push_episodes.get("executionMode", "人工/未知")},
                {"label": "同步型子分", "value": rounded(sync_score, 1)},
                {"label": "时段集中型子分", "value": rounded(serial_score, 1)},
                {"label": "集中堆单型子分", "value": rounded(burst_score, 1)},
                {"label": "反复协调同步", "value": f'{(sync or {}).get("coordinatedMatchedRatio", (sync or {}).get("matchedRatio", 0))}% / {(sync or {}).get("recurringPeerAccounts", (sync or {}).get("peerAccounts", 0))} 个反复关联账号'},
                {"label": "最强单一关联", "value": f'{(sync or {}).get("maxPeerRatio", 0)}%'},
                {"label": "主要仓位集中时段", "value": f'{push_behavior.get("concentratedCoreVolumeRatio", 0)}% / {push_behavior.get("concentratedSessions", 0)} 段'},
                {"label": "主要仓位一小时内结束", "value": f'{push_behavior.get("coreShortHoldVolumeRatio", 0)}%'},
                {"label": "相对自身节奏的动态停手强度", "value": f'{push_behavior.get("quietGapRatio", 0)}%'},
                {"label": "典型订单间隔 / 动态分段线", "value": f'{push_behavior.get("typicalOrderGapMinutes", 0)} / {push_behavior.get("sessionBreakMinutes", 0)} 分钟'},
                {"label": "核心建仓跨度 / 同方向连单", "value": f'{push_behavior.get("coreOpenSpanSeconds", 0)} 秒 / {push_behavior.get("maxSameDirectionRunRatio", 0)}%'},
                {"label": "小手数穿插单", "value": f'{push_behavior.get("camouflageOrders", 0)} 单 / {push_behavior.get("camouflageVolumeRatio", 0)}% 手数'},
                {"label": "动态轮次叠仓", "value": f'{push_behavior.get("staggeredAddOnOrderRatio", 0)}% 订单 / {push_behavior.get("staggeredAddOnVolumeRatio", 0)}% 手数'},
                {"label": "动态轮次集中平仓", "value": f'{push_behavior.get("cohesiveCampaignOrderRatio", 0)}% 订单'},
                {"label": "动态建仓波次集中平仓", "value": f'{push_behavior.get("cohesiveBatchOrderRatio", 0)}% 订单'},
                {"label": "独立协同疑似轮次", "value": f'{mixed_push_episodes.get("confirmedEpisodes", 0)}/{mixed_push_episodes.get("candidateEpisodes", 0)} 段，{mixed_push_episodes.get("confirmedVolumeRatio", 0)}% 核心手数'},
                {"label": "重复轮次同伙", "value": f'{mixed_push_episodes.get("strongestPeer") or "未识别"} / {mixed_push_episodes.get("strongestPeerEpisodes", 0)} 段'},
                {"label": "亏损单 / 亏损手数", "value": f'{push_behavior.get("lossRate", 0)}% / {push_behavior.get("lossVolumeRatio", 0)}%'},
            ]
            if stage == "deep":
                if tick and tick.get("available"):
                    item_metrics += [
                        {"label": "价格赢点中位数", "value": tick.get("priceWinTickMedian")},
                        {"label": "扣费净赢点中位数", "value": tick.get("netWinTickMedian")},
                        {"label": "盈利单第1 Tick开始盈利", "value": f'{tick.get("win1VolumeRatio")}% 手数'},
                        {"label": "盈利单前3 Tick开始盈利", "value": f'{tick.get("win3VolumeRatio")}% 手数'},
                        {"label": "盈利单前10 Tick开始盈利", "value": f'{tick.get("win10VolumeRatio")}% 手数'},
                        {"label": "盈利单第1 Tick盈利概率", "value": f'{tick.get("win1OrderRatio", tick.get("win1VolumeRatio"))}% 订单'},
                        {"label": "盈利单前3 Tick盈利概率", "value": f'{tick.get("win3OrderRatio", tick.get("win3VolumeRatio"))}% 订单'},
                        {"label": "20 Tick后有利手数", "value": f'{tick.get("positiveImpact20VolumeRatio")}%'},
                        {"label": "前50 Tick有利报价中位数", "value": f'{tick.get("favorableTickRatio50Median")}%'},
                        {"label": "50 Tick内开始盈利手数", "value": f'{tick.get("win50VolumeRatio")}%'},
                        {"label": "低Tick密度订单", "value": f'{tick.get("lowLiquidityOrderRatio", 0)}% 订单 / {tick.get("lowLiquidityVolumeRatio", 0)}% 手数'},
                        {"label": "Rollover时段订单", "value": f'{tick.get("rolloverWindowOrderRatio", 0)}% 订单'},
                        {"label": "每分钟有效Tick中位数", "value": tick.get("tickRatePerMinuteMedian")},
                        {"label": "开仓前 / 后有效Tick中位数", "value": f'{tick.get("preTickRatePerMinuteMedian", 0)} / {tick.get("postTickRatePerMinuteMedian", 0)} 每分钟'},
                        {"label": "点差扩张中位倍数", "value": tick.get("spreadExpansionRatioMedian")},
                        {"label": "入场后10秒有利移动手数", "value": f'{tick.get("eventImpact10VolumeRatio", 0)}%'},
                        {"label": "超过入场前节奏的10秒加速手数", "value": f'{tick.get("eventAcceleration10VolumeRatio", 0)}%'},
                        {"label": "有利移动持续60秒手数", "value": f'{tick.get("eventPersistence60VolumeRatio", 0)}%'},
                        {"label": "入场前已有同向趋势手数", "value": f'{tick.get("preexistingTrendVolumeRatio", 0)}%'},
                        {"label": "180秒内回撤手数", "value": f'{tick.get("reversal180VolumeRatio", 0)}%'},
                        {"label": "开仓后流动性下降", "value": f'{tick.get("liquidityDropOrderRatio", 0)}% 订单 / {tick.get("liquidityDropVolumeRatio", 0)}% 手数'},
                        {"label": "开仓后最大无报价间隔中位数", "value": f'{tick.get("maxPostQuoteGapSecondsMedian", 0)} 秒'},
                    ] + tick_context_metrics
                    if tick.get("errors"):
                        limitations.append(f'部分候选未取到Tick：{tick.get("errors")[0]}')
                else:
                    limitations.append((tick or {}).get("reason", "Tick 未加载，无法计算报价冲击")); confidence = 55
            else:
                limitations.append("初筛未读取跨账户同步和Tick报价冲击；只展示账户内重复执行候选")
            if burst_confirmed:
                triggers.append("5秒内集中建仓且主要仓位持仓较短"); score = max(score, 60)
            sync_matched_ratio = numeric_value((sync or {}).get("coordinatedMatchedRatio", (sync or {}).get("matchedRatio")))
            sync_peer_accounts = mysql_int((sync or {}).get("recurringPeerAccounts", (sync or {}).get("peerAccounts")))
            if sync and sync.get("available") and sync_matched_ratio >= 60 and sync_peer_accounts >= 2 and short_core_ratio >= 70 and concentrated_ratio >= 60 and quiet_gap_ratio >= 40 and position_structure_ok:
                triggers.append("多数订单与多个账户在2秒内同向同步"); score = max(score, 75)
            if sync and sync.get("available") and sync_matched_ratio >= 80 and sync_peer_accounts >= 5 and numeric_value(sync.get("maxPeerRatio")) >= 20 and short_core_ratio >= 70 and concentrated_ratio >= 60 and quiet_gap_ratio >= 40 and position_structure_ok:
                triggers.append("高比例多账户协调开平仓"); score = max(score, 90)
            if tick and tick.get("available") and episode_score >= 60 and numeric_value(tick.get("positiveImpact20VolumeRatio")) >= 45 and numeric_value(tick.get("favorableTickRatio50Median")) >= 50:
                triggers.append("主要仓位集中交易，并在开仓后持续出现有利报价"); score = max(score, 75)
            elif episode_confirmed:
                triggers.append("主要仓位集中在短时段，结束后长时间停手"); score = max(score, 75)
                if numeric_value(push_behavior.get("camouflageOrderRatio")) >= 5 and numeric_value(push_behavior.get("camouflageVolumeRatio")) <= 10:
                    triggers.append("非集中时段穿插了少量小手数订单")
                if not (tick and tick.get("available")):
                    limitations.append("集中交易形态已成立；仍需Tick确认下单后是否出现报价冲击")
            if causal_liquidity_chain and (
                sync_matched_ratio >= 40
                or concentrated_ratio >= 60
            ):
                triggers.append("订单后出现超过入场前节奏的持续报价冲击，建议核对外部基准")
                score = max(score, 75)
                limitations.append("动态流动性指标只证明平台报价在订单后加速或持续，尚未证明价格移动由该账户造成；需外部基准报价复核")
            if small_sample_warning:
                triggers.append("订单虽少，但固定账户同步开平仓、短持仓和Tick持续性同时明显")
                score = max(score, 75)
                limitations.append("当前订单样本较少，风险分不设上限，但结果置信度相应降低；仍需更多重复交易或外部基准报价确认")
            if tick and tick.get("available") and episode_score >= 60 and (
                win1_signal >= 35
                or win3_signal >= 65
                or (numeric_value(tick.get("priceWinTickMedian")) <= 3 and win3_signal >= 50)
            ):
                triggers.append("多数盈利单在开仓后极少数原始Tick内开始盈利")
                score = max(score, 75)
            if sync and sync.get("errors"):
                limitations.append(f'跨账户查询部分失败：{sync.get("errors")[0]}')
            deposit_total = numeric_value(finance.get("depositTotal")) if finance.get("available") else 0
            withdrawal_total = numeric_value(finance.get("withdrawalTotal")) if finance.get("available") else 0
            withdrawal_hours = finance.get("lastTradeToWithdrawalHours") if finance.get("available") else None
            withdrawal_ratio = withdrawal_total / deposit_total * 100 if deposit_total else 0
            funding_evidence = bool(
                deposit_total > 0
                and deposit_total <= 20000
                and withdrawal_total > 0
                and withdrawal_ratio >= 30
                and withdrawal_hours is not None
                and 0 <= numeric_value(withdrawal_hours) <= 72
            )
            item_metrics += [
                {"label": "累计入金 / 出金", "value": f'{rounded(deposit_total)} / {rounded(withdrawal_total)}'},
                {"label": "交易结束至后续出金", "value": f'{withdrawal_hours} 小时' if withdrawal_hours is not None else "未发现后续出金"},
            ]
            coordinated_ratio = numeric_value((sync or {}).get("coordinatedMatchedRatio", (sync or {}).get("matchedRatio")))
            recurring_peers = mysql_int((sync or {}).get("recurringPeerAccounts", (sync or {}).get("peerAccounts")))
            platform_coordinated = bool(sync and sync.get("available") and recurring_peers >= 2 and coordinated_ratio >= 30)
            if platform_coordinated and episode_confirmed:
                triggers.append("全平台存在多个反复同步账户"); score = max(score, 75)
            if platform_coordinated and episode_structure and quiet_gap_ratio >= 40 and position_structure_ok and tick and tick.get("available") and liquidity_signal >= 50:
                triggers.append("多人集中交易发生在低Tick密度时段"); score = min(100, score + 7)
            if platform_coordinated and episode_structure and quiet_gap_ratio >= 40 and position_structure_ok and funding_evidence:
                triggers.append("小额入金后在交易结束不久快速出金"); score = min(100, score + 7)
            if stage == "deep" and not platform_coordinated:
                if sync and sync.get("available"):
                    limitations.append("当前未找到达到反复匹配门槛的疑似同伙；该结果只表示暂未定位协同账户，不降低主体推盘嫌疑")
                else:
                    limitations.append("全平台协同尚未验证；协同仅作为同伙线索和加分项，不影响主体推盘嫌疑")
            mixed_episode_confirmed = bool(mixed_push_episodes.get("confirmed"))
            if mixed_episode_confirmed:
                triggers.append("正常或网格交易中反复出现独立协同打盘轮次")
                score = max(score, numeric_value(mixed_push_episodes.get("score")))
                if mixed_push_episodes.get("eaAttention"):
                    triggers.append("EA主导执行中重复出现固定同伙协同轮次")
                if mixed_push_episodes.get("copyAttention"):
                    limitations.append("检测到跟单标记；跟单本身可能是正常复制执行，已要求同步平仓、重复轮次和Tick证据共同成立")
            elif mixed_push_episodes.get("eaAttention"):
                limitations.append("核心订单以EA执行为主；EA只触发重点复核，不单独增加推盘分数或直接定性")
            elif mixed_push_episodes.get("copyAttention"):
                limitations.append("检测到跟单标记；跟单只作为重点复核线索，不单独增加推盘分数或直接定性")
            if (staggered_addon_ratio >= 35 or staggered_addon_volume_ratio >= 35):
                if mixed_episode_confirmed:
                    limitations.append("账户整体包含网格或叠仓交易，但独立轮次已通过固定同伙、同步平仓、手数占比和Tick的交叉验证；网格作为反证扣分，不否决独立证据")
                elif causal_liquidity_chain:
                    limitations.append("账户存在网格或叠仓结构，但订单后的动态报价加速、持续性和入场前基线同时达到重点复核门槛；网格作为反证扣分，不直接封顶")
                else:
                    limitations.append("动态轮次内存在叠仓或摊平成本，已作为反证扣分；不再将账户总分硬封顶")
            # Evidence groups replace the old all-or-nothing caps. Order count lowers
            # confidence only; it no longer suppresses a strong, cross-validated signal.
            position_signal = max(non_overlap_ratio, cohesive_batch_ratio, cohesive_campaign_ratio)
            structure_score = min(40.0, (
                concentrated_ratio * .10
                + short_core_ratio * .10
                + quiet_gap_ratio * .08
                + position_signal * .07
                + core_volume_ratio * .05
            ))
            tick_available = bool(tick and tick.get("available"))
            tick_win_signal = 0.0
            tick_group_score = 0.0
            if tick_available:
                tick_win_signal = (
                    numeric_value(tick.get("win1OrderRatio", tick.get("win1VolumeRatio"))) * .55
                    + numeric_value(tick.get("win3OrderRatio", tick.get("win3VolumeRatio"))) * .30
                    + numeric_value(tick.get("win10OrderRatio", tick.get("win10VolumeRatio"))) * .15
                )
                tick_group_score = min(35.0, (
                    numeric_value(tick.get("eventImpact10VolumeRatio")) * .07
                    + numeric_value(tick.get("eventAcceleration10VolumeRatio")) * .07
                    + numeric_value(tick.get("eventPersistence60VolumeRatio")) * .06
                    + numeric_value(tick.get("positiveImpact20VolumeRatio")) * .05
                    + numeric_value(tick.get("favorableTickRatio50Median")) * .04
                    + tick_win_signal * .04
                    + toxic_ramp(tick.get("impactSpreadMultipleMedian"), .25, 2) * .02
                ))
            coordination_score = 0.0
            if sync and sync.get("available"):
                coordination_score = min(15.0, (
                    coordinated_ratio * .04
                    + coordinated_volume_ratio * .02
                    + coordinated_close_ratio * .05
                    + numeric_value(sync.get("maxPeerRatio")) * .05
                    + min(mysql_int(sync.get("maxPeerMatches")) / 2, 1) * 5
                ))
            execution_score = 6.0 if mixed_push_episodes.get("eaAttention") else 4.0 if mixed_push_episodes.get("copyAttention") else 0.0
            context_score = min(10.0, execution_score + (4.0 if funding_evidence else 0.0))
            grid_strength = max(staggered_addon_ratio, staggered_addon_volume_ratio)
            counterevidence_deduction = toxic_ramp(grid_strength, 35, 80) * .15
            if mixed_episode_confirmed:
                counterevidence_deduction *= .25
            elif synchronized_cohesive_override or causal_liquidity_chain:
                counterevidence_deduction *= .5
            single_burst_confirmed = toxic_single_burst_chain(push_behavior, tick)
            if single_burst_confirmed:
                structure_score = max(structure_score, 40.0)
                tick_group_score = max(tick_group_score, 25.0)
                triggers.append("单轮次内密集同向建仓，短持仓高胜率并伴随订单后报价冲击")
                limitations.append("单轮次爆发链由密集同向建仓、持仓重叠和订单后Tick反应交叉确认；仍需外部基准排除正常趋势跟随或批量执行")
            coordinated_tick_confirmed = toxic_coordinated_tick_chain(push_behavior, sync, tick)
            if coordinated_tick_confirmed:
                triggers.append("固定账户高比例同步开平仓，并伴随订单后报价冲击")
                limitations.append("强协同链要求固定账户重复同步开仓、同步平仓和Tick反应共同成立；账户内少量反向单不作为免责证据")
            sudden_exposure_confirmed = toxic_sudden_exposure_chain(
                push_behavior, tick, finance, bool(mixed_push_episodes.get("eaAttention"))
            )
            if sudden_exposure_confirmed:
                structure_score = max(structure_score, 35.0)
                tick_group_score = max(tick_group_score, 28.0)
                context_score = max(context_score, 10.0)
                triggers.append("少量订单突然形成高暴露同向仓位，并在强Tick反应后快速完成资金回转")
                limitations.append("突然高暴露链同时要求EA执行、秒级集中建平仓、强Tick反应和短周期资金回转；缺少任一证据均不提升为高危")
            consistency_bonus = market_pushing_consistency_bonus(
                structure_score, tick_group_score, counterevidence_deduction
            )
            if coordinated_tick_confirmed:
                consistency_bonus = max(consistency_bonus, 10.0)
            evidence_score = max(
                0.0,
                structure_score + tick_group_score + coordination_score + context_score
                + consistency_bonus - counterevidence_deduction,
            )
            minimum_score = 0.0
            if not tick_available and episode_confirmed:
                minimum_score = max(minimum_score, 75.0)
            if not tick_available and sync and sync.get("available") and coordinated_ratio >= 60 and recurring_peer_accounts >= 2 and episode_structure and quiet_gap_ratio >= 40 and position_structure_ok:
                minimum_score = max(minimum_score, 75.0)
            if not tick_available and sync and sync.get("available") and coordinated_ratio >= 80 and recurring_peer_accounts >= 5 and numeric_value(sync.get("maxPeerRatio")) >= 20 and episode_structure and quiet_gap_ratio >= 40 and position_structure_ok:
                minimum_score = max(minimum_score, 90.0)
            if mixed_episode_confirmed:
                structure_score = max(structure_score, 40.0)
                tick_group_score = max(tick_group_score, 35.0)
                coordination_score = max(coordination_score, 15.0)
                minimum_score = max(minimum_score, 90.0)
            score = max(evidence_score, minimum_score)
            core_count_for_confidence = mysql_int(push_behavior.get("coreOrders"))
            if core_count_for_confidence < 12:
                confidence = min(confidence, 60 + min(core_count_for_confidence, 11) * 2)
            item_metrics += [
                {"label": "主体结构证据 / 40", "value": rounded(structure_score, 1)},
                {"label": "Tick市场反应 / 35", "value": rounded(tick_group_score, 1)},
                {"label": "跨账户协同 / 15", "value": rounded(coordination_score, 1)},
                {"label": "EA与资金辅助 / 10", "value": rounded(context_score, 1)},
                {"label": "核心证据一致性加分", "value": rounded(consistency_bonus, 1)},
                {"label": "网格/叠仓反证扣分", "value": rounded(counterevidence_deduction, 1)},
                {"label": "单轮次集中爆发链", "value": "成立" if single_burst_confirmed else "未成立"},
            ]
            if coordinated_tick_confirmed:
                item_metrics.append({"label": "固定同伙开平仓+Tick交叉链", "value": "成立"})
            if sudden_exposure_confirmed:
                item_metrics.append({"label": "突然高暴露+资金回转链", "value": "成立"})
            limitations.append("跨账户同步只能证明协调执行；推盘定性仍需结合本平台与外部基准报价差异")
            summary = f'证据组得分：结构 {rounded(structure_score,1)}/40 / Tick {rounded(tick_group_score,1)}/35 / 协同 {rounded(coordination_score,1)}/15 / 辅助 {rounded(context_score,1)}/10；一致性加分 {rounded(consistency_bonus,1)}；反证扣分 {rounded(counterevidence_deduction,1)}'
            analysis = toxic_push_analysis(score, sync_score, serial_score, burst_score, push_behavior, sync, tick, finance, mixed_push_episodes)
            evidence_chain = toxic_push_evidence_chain(
                score, confidence, push_behavior, tick, sync, finance, mixed_push_episodes,
                structure_score, tick_group_score, coordination_score, context_score,
                consistency_bonus, counterevidence_deduction,
                single_burst_confirmed, coordinated_tick_confirmed, sudden_exposure_confirmed,
            )
        elif type_id == "quote_latency_arbitrage":
            score = short10 * 35 + short60 * 20 + short_win_rate * 20 + toxic_ramp(len(short_rows), 3, 20) * .15 + toxic_ramp(burst60["count"], 3, 20) * .1
            item_metrics = [{"label": "10秒内手数占比", "value": f'{rounded(short10*100,1)}%'}, {"label": "60秒内手数占比", "value": f'{rounded(short60*100,1)}%'}, {"label": "短单胜率", "value": f'{rounded(short_win_rate*100,1)}%'}]
            evidence = [normalize_text(row.get("ticket")) for row in sorted(short_rows, key=toxic_trade_holding)[:12]]
            if stage == "deep":
                if tick and tick.get("available"):
                    score = score * .7 + tick_score * .3
                    item_metrics += [
                        {"label": "价格赢点中位数", "value": tick.get("priceWinTickMedian")},
                        {"label": "扣费净赢点中位数", "value": tick.get("netWinTickMedian")},
                        {"label": "价格赢点≤3手数占比", "value": f'{tick.get("win3VolumeRatio")}%'},
                        {"label": "净赢点≤3手数占比", "value": f'{tick.get("netWin3VolumeRatio")}%'},
                    ] + tick_context_metrics
                    if tick.get("errors"):
                        limitations.append(f'部分候选未取到Tick：{tick.get("errors")[0]}')
                else:
                    limitations.append((tick or {}).get("reason", "Tick 未加载，无法计算赢点")); confidence = 45
            else:
                limitations.append("需深度读取 Tick 后才能判断赢点和报价优势")
            if len(short_rows) >= 3 and short60 >= .85 and short_win_rate >= .6:
                triggers.append("短时间内存在稳定的秒级平仓簇"); score = max(score, 60)
            if stage == "deep" and tick and tick.get("available") and numeric_value(tick.get("win3VolumeRatio")) >= 70 and len(short_rows) >= 3:
                triggers.append("多数可疑手数在3个有效Tick内盈利"); score = max(score, 75)
            summary = f'{len(short_rows)} 笔在60秒内平仓，手数占比 {rounded(short60*100,1)}%'
        elif type_id == "cross_platform_spread_arbitrage":
            balance_score = max(0.0, min(100.0, (100 - burst300["direction"]) * 2))
            score = toxic_ramp(burst300["count"], 10, 100) * .25 + toxic_ramp(burst300["volume"], 2, 50) * .2 + short60 * 25 + balance_score * .15 + numeric_value(pair.get("match")) * .15
            item_metrics = [{"label": "5分钟订单峰值", "value": burst300["count"]}, {"label": "5分钟手数峰值", "value": burst300["volume"]}, {"label": "60秒内手数", "value": f'{rounded(short60*100,1)}%'}]
            evidence = burst300["tickets"]
            if stage == "deep" and tick and tick.get("available"):
                score = score * .85 + tick_score * .15
                item_metrics += [{"label": "价格赢点≤3手数占比", "value": f'{tick.get("win3VolumeRatio")}%'}, {"label": "净赢点≤3手数占比", "value": f'{tick.get("netWin3VolumeRatio")}%'}, *tick_context_metrics]
            elif stage == "deep":
                limitations.append((tick or {}).get("reason", "没有精确 Tick 源")); confidence = 48
            limitations.append("当前只能确认平台内形态；没有外部平台另一条腿时不能定性")
            if burst300["count"] >= 20 and burst300["volume"] >= 2 and short60 >= .85:
                triggers.append("5分钟内出现高周转短持仓簇"); score = max(score, 75)
            summary = f'5分钟峰值 {burst300["count"]} 单 / {burst300["volume"]} 手'
        elif type_id == "rebate_churning":
            score = numeric_value(pair.get("match")) * .25 + toxic_ramp(pair.get("volume"), 5, 20) * .25 + short10 * 15 + toxic_ramp(rebate / net_deposit, .1, 1) * .25 + (100 if metrics.get("netProfit", 0) < 0 else 0) * .1
            item_metrics = [{"label": "反向匹配度", "value": f'{pair.get("match",0)}%'}, {"label": "配对手数", "value": pair.get("volume",0)}, {"label": "返佣 / 净入金", "value": f'{rounded(rebate/net_deposit*100,1)}%'}]
            evidence = pair.get("tickets", [])
            if not finance.get("available"):
                limitations.append(finance.get("reason", "返佣资金数据不可用")); confidence -= 20
            if pair.get("found") and numeric_value(pair.get("match")) >= 95 and numeric_value(pair.get("volume")) >= 5:
                triggers.append("出现秒级大手数反向匹配"); score = max(score, 75)
            if rebate / net_deposit >= 1:
                triggers.append("累计返佣已达到净入金规模"); score = max(score, 90)
            summary = f'返佣 {rounded(rebate)}，最强反向配对 {pair.get("volume",0)} 手'
        elif type_id == "bonus_arbitrage":
            score = toxic_ramp(credit_ratio, .5, 1) * .3 + margin_pressure * 100 * .3 + toxic_ramp(burst600["direction"], 70, 100) * .15 + numeric_value(pair.get("match")) * .15 + (100 if any(not normalize_text(row.get("sl")) for row in rows) else 0) * .1
            item_metrics = [{"label": "Credit / 现金权益", "value": f'{rounded(credit_ratio*100,1)}%'}, {"label": "保证金 / 总权益", "value": f'{rounded(margin/max(equity,1)*100,1) if equity else 0}%'}, {"label": "最强方向集中", "value": f'{burst600["direction"]}%'}]
            evidence = burst600["tickets"]
            if not finance.get("available"):
                limitations.append(finance.get("reason", "Credit和保证金数据不可用")); confidence -= 25
            limitations.append("外部对锁不可见；只能输出牺牲账户嫌疑")
            if credit_ratio >= .5 and (margin_pressure >= .25 or burst600["direction"] >= 90):
                triggers.append("赠金占比较高且仓位风险集中"); score = max(score, 75)
            if credit_ratio >= 1 and margin_level and margin_level <= 80:
                triggers.append("Credit不低于现金且保证金水平极低"); score = max(score, 90)
            summary = f'Credit占现金 {rounded(credit_ratio*100,1)}%，保证金水平 {rounded(margin_level,1) if margin_level else "-"}%'
        elif type_id == "short_close_trading":
            score = short60 * 30 + toxic_ramp(burst120["count"], 3, 20) * .2 + toxic_ramp(burst120["volume"], 1, 10) * .25 + toxic_ramp(burst120["direction"], 70, 100) * .15 + short_win_rate * 10
            item_metrics = [{"label": "2分钟订单峰值", "value": burst120["count"]}, {"label": "2分钟手数峰值", "value": burst120["volume"]}, {"label": "60秒内手数", "value": f'{rounded(short60*100,1)}%'}]
            evidence = burst120["tickets"]
            if burst120["count"] >= 3 and burst120["volume"] >= 1 and short60 >= .5:
                triggers.append("两分钟内形成有经济规模的短平簇"); score = max(score, 60)
            if burst120["count"] >= 5 and burst120["volume"] >= 5 and short60 >= .8:
                triggers.append("大手数短平集中爆发"); score = max(score, 75)
            summary = f'两分钟峰值 {burst120["count"]} 单 / {burst120["volume"]} 手'
        elif type_id in {"internal_lock_arbitrage", "high_leverage_lock_arbitrage"}:
            pair_score = numeric_value(pair.get("match")) * .35 + toxic_ramp(pair.get("seconds") if pair.get("seconds") is not None else 60, 0, 60, reverse=True) * .2 + toxic_ramp(pair.get("volume"), .2, 5) * .2 + numeric_value(pair.get("overlap")) * .15
            score = pair_score + margin_pressure * 100 * .1
            if type_id == "high_leverage_lock_arbitrage":
                score = pair_score * .55 + toxic_ramp(margin_pressure, .35, 1) * .25 + toxic_ramp(margin_level or 300, 80, 180, reverse=True) * .2
            item_metrics = [{"label": "反向匹配度", "value": f'{pair.get("match",0)}%'}, {"label": "开仓时间差", "value": f'{pair.get("seconds") if pair.get("seconds") is not None else "-"} 秒'}, {"label": "匹配手数", "value": pair.get("volume",0)}, {"label": "保证金水平", "value": f'{rounded(margin_level,1) if margin_level else "-"}%'}]
            evidence = pair.get("tickets", [])
            limitations.append("当前结果先检查本账号内部反向腿；跨全平台账号匹配将在后续数据索引中增强")
            if pair.get("found") and numeric_value(pair.get("match")) >= 85 and numeric_value(pair.get("seconds")) <= 60:
                triggers.append("同品种反向仓高度同步"); score = max(score, 75)
            if type_id == "high_leverage_lock_arbitrage" and pair.get("found") and margin_level and margin_level <= 80:
                triggers.append("对锁同时保证金水平低于80%"); score = max(score, 90)
            summary = f'最强反向匹配 {pair.get("match",0)}%，时间差 {pair.get("seconds") if pair.get("seconds") is not None else "-"} 秒'
        elif type_id == "weekend_gap_trading":
            weekend_volume = sum(numeric_value(row.get("volume")) for row in weekend_rows)
            score = toxic_ramp(len(weekend_rows), 1, 5) * .3 + toxic_ramp(weekend_volume, .2, 5) * .25 + margin_pressure * 100 * .25 + numeric_value(pair.get("match")) * .1 + (100 if weekend_rows and any(not normalize_text(row.get("sl")) for row in weekend_rows) else 0) * .1
            item_metrics = [{"label": "跨周末订单", "value": len(weekend_rows)}, {"label": "跨周末手数", "value": rounded(weekend_volume,4)}, {"label": "保证金水平", "value": f'{rounded(margin_level,1) if margin_level else "-"}%'}]
            evidence = [normalize_text(row.get("ticket")) for row in weekend_rows[:12]]
            limitations.append("需配置每个品种的真实收盘时间和历史Gap后，才能计算预计穿仓金额")
            if weekend_rows and (weekend_volume >= .2 or margin_pressure >= .25):
                triggers.append("存在有经济规模的跨周末仓位"); score = max(score, 60)
            summary = f'{len(weekend_rows)} 笔跨周末订单，共 {rounded(weekend_volume,4)} 手'
        elif type_id == "open_betting":
            score = toxic_ramp(burst600["count"], 3, 15) * .25 + toxic_ramp(burst600["volume"], .5, 10) * .25 + toxic_ramp(burst600["direction"], 70, 100) * .2 + margin_pressure * 100 * .3
            item_metrics = [{"label": "10分钟订单峰值", "value": burst600["count"]}, {"label": "10分钟手数峰值", "value": burst600["volume"]}, {"label": "同向度", "value": f'{burst600["direction"]}%'}]
            evidence = burst600["tickets"]
            limitations.append("需要品种交易时段表确认候选簇是否位于真实开盘窗口")
            if stage == "deep" and tick and tick.get("available"):
                item_metrics += [{"label": "价格赢点≤3手数占比", "value": f'{tick.get("win3VolumeRatio")}%'}, {"label": "净赢点≤3手数占比", "value": f'{tick.get("netWin3VolumeRatio")}%'}, *tick_context_metrics]
            if burst600["count"] >= 3 and burst600["direction"] >= 90 and (burst600["volume"] >= .5 or margin_pressure >= .25):
                triggers.append("10分钟内出现高集中单向下注簇"); score = max(score, 60)
            summary = f'10分钟峰值 {burst600["count"]} 单 / {burst600["volume"]} 手'
        elif type_id == "news_event_betting":
            score = toxic_ramp(burst600["count"], 3, 15) * .2 + toxic_ramp(burst600["volume"], .5, 10) * .25 + toxic_ramp(burst600["direction"], 70, 100) * .15 + margin_pressure * 100 * .4
            item_metrics = [{"label": "10分钟订单峰值", "value": burst600["count"]}, {"label": "10分钟手数峰值", "value": burst600["volume"]}, {"label": "保证金水平", "value": f'{rounded(margin_level,1) if margin_level else "-"}%'}]
            evidence = burst600["tickets"]
            limitations.append("经济日历尚未接入，当前只能筛选高波动下注候选，不能确认非农/CPI等事件")
            if stage == "deep" and tick and tick.get("available"):
                item_metrics += [{"label": "价格赢点≤3手数占比", "value": f'{tick.get("win3VolumeRatio")}%'}, {"label": "净赢点≤3手数占比", "value": f'{tick.get("netWin3VolumeRatio")}%'}, *tick_context_metrics]
            else:
                confidence = min(confidence, 45)
            summary = "等待经济日历与事件窗口确认"
        result_item = toxic_result(type_id, score, stage, summary, item_metrics, triggers, evidence, limitations, confidence, analysis)
        if type_id == "market_pushing":
            result_item["suspectedAccomplices"] = list((sync or {}).get("suspectedAccounts") or [])
            result_item["mixedEpisodes"] = mixed_push_episodes
            result_item["evidenceChain"] = evidence_chain or {}
        results.append(result_item)
    results.sort(key=lambda item: (-numeric_value(item.get("score")), item.get("label", "")))
    return results


def update_toxic_job(job_id: str, **updates) -> None:
    with TOXIC_JOBS_LOCK:
        job = TOXIC_JOBS.setdefault(job_id, {"id": job_id, "status": "queued"})
        job.update(updates)


def get_toxic_job(job_id: str) -> dict:
    with TOXIC_JOBS_LOCK:
        return dict(TOXIC_JOBS.get(job_id, {"id": job_id, "status": "missing", "message": "任务不存在"}))


def update_push_discovery_job(job_id: str, **updates) -> None:
    with PUSH_DISCOVERY_JOBS_LOCK:
        job = PUSH_DISCOVERY_JOBS.setdefault(job_id, {"id": job_id, "status": "queued"})
        job.update(updates)


def get_push_discovery_job(job_id: str) -> dict:
    with PUSH_DISCOVERY_JOBS_LOCK:
        return dict(PUSH_DISCOVERY_JOBS.get(job_id, {"id": job_id, "status": "missing", "message": "任务不存在"}))


def push_discovery_result_rows(output_dir: Path) -> list[dict]:
    result_path = output_dir / "deep_results.json"
    if not result_path.exists():
        return []
    rows = json.loads(result_path.read_text(encoding="utf-8"))
    return [{
        "deepRank": mysql_int(item.get("deepRank")),
        "platform": normalize_text(item.get("platform")),
        "server": normalize_text(item.get("server")),
        "account": normalize_text(item.get("login")),
        "orders": mysql_int(item.get("exactOrders") or item.get("closedOrders")),
        "periodNetRaw": rounded(item.get("periodNetRaw")),
        "initialScore": rounded(item.get("initialScore"), 1),
        "deepScore": rounded(item.get("deepScore"), 1),
        "level": normalize_text(item.get("deepLevel")),
        "tickAvailable": bool(item.get("tickAvailable")),
        "coordinatedMatchedRatio": rounded(item.get("coordinatedMatchedRatio"), 1),
        "headline": normalize_text(item.get("headline")),
        "routeReasons": list(item.get("routeReasons") or []),
        "suspectedAccomplices": list(item.get("suspectedAccomplices") or []),
    } for item in rows[:200]]


def run_push_discovery_job(job_id: str, options: dict) -> None:
    started = time.time()
    if not PUSH_DISCOVERY_RUN_LOCK.acquire(blocking=False):
        update_push_discovery_job(
            job_id, status="failed", percent=100,
            message="已有全平台推盘扫描正在运行，请等待完成后重试",
            elapsedSeconds=rounded(time.time() - started, 1),
        )
        return
    try:
        update_push_discovery_job(job_id, status="running", percent=2, message="正在启动全平台推盘扫描", startedAt=now_text())
        cmd = [
            str(K_DESK_PYTHON), str(PUSH_DISCOVERY_SCRIPT),
            "--days", str(options["days"]),
            "--max-orders", str(options["maxOrders"]),
            "--small-order-priority", str(options["smallOrderPriority"]),
            "--deep-limit", str(options["deepLimit"]),
            "--workers", str(options["workers"]),
        ]
        env = os.environ.copy()
        env["PYTHONIOENCODING"] = "utf-8"
        env.setdefault("K_DESK_ROOT", str(ROOT))
        if MYSQL_PASSWORD:
            env["ACCOUNT_TRADE_MYSQL_PASSWORD"] = MYSQL_PASSWORD
        proc = subprocess.Popen(
            cmd, cwd=str(ROOT), stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
            text=True, encoding="utf-8", errors="replace", env=env,
        )
        output_dir = None
        assert proc.stdout is not None
        for line in proc.stdout:
            text_line = line.strip()
            if text_line.startswith("PROGRESS "):
                try:
                    progress = json.loads(text_line[len("PROGRESS "):])
                    updates = {
                        "stage": normalize_text(progress.get("stage")),
                        "message": normalize_text(progress.get("message")),
                        "percent": min(max(mysql_int(progress.get("percent")), 0), 100),
                    }
                    if progress.get("summary"):
                        updates["summary"] = progress["summary"]
                        output_dir = Path(progress["summary"]["outputDir"])
                    update_push_discovery_job(job_id, **updates)
                except (ValueError, TypeError, json.JSONDecodeError):
                    continue
            elif text_line.startswith("RESULT "):
                summary_path = Path(text_line[len("RESULT "):].strip())
                if summary_path.exists():
                    output_dir = summary_path.parent
        return_code = proc.wait()
        if return_code != 0 or not output_dir:
            raise RuntimeError(f"全平台扫描进程异常结束（代码 {return_code}）")
        summary = json.loads((output_dir / "summary.json").read_text(encoding="utf-8"))
        update_push_discovery_job(
            job_id,
            status="done",
            percent=100,
            message="全平台推盘扫描完成",
            summary=summary,
            results=push_discovery_result_rows(output_dir),
            outputDir=str(output_dir),
            elapsedSeconds=rounded(time.time() - started, 1),
        )
    except Exception as exc:
        update_push_discovery_job(
            job_id, status="failed", percent=100, message=str(exc), error=str(exc),
            elapsedSeconds=rounded(time.time() - started, 1),
        )
    finally:
        PUSH_DISCOVERY_RUN_LOCK.release()


def query_toxic_scope_trades(login: str, filters: dict) -> list[dict]:
    platform = normalize_text(filters.get("platform"))
    server = normalize_text(filters.get("server"))
    symbol = normalize_text(filters.get("symbol"))
    start = normalize_text(filters.get("start"))
    end = normalize_text(filters.get("end"))
    start_dt = parse_trade_time(start)
    end_dt = parse_trade_time(end)
    sources = [source for source in MYSQL_SOURCES if source_allowed(source, platform=platform, server=server)]
    if start_dt and end_dt and len(sources) == 1 and sources[0].get("kind") == "mt5_deals":
        compatibility_root = Path(__file__).resolve().parents[2]
        if str(compatibility_root) not in sys.path:
            sys.path.insert(0, str(compatibility_root))
        from scripts import run_ac_mt5_push_validation as mt5_runner

        rows = mt5_runner.load_recent_closed_trades(sources[0], login, start_dt, end_dt)["rows"]
        return [row for row in rows if not symbol or normalize_text(row.get("symbol")) == symbol]

    rows = query_db_trades(
        login,
        platform=platform,
        server=server,
        symbol=symbol,
        start=start,
        end=end,
        limit=50000,
    )
    if start_dt:
        rows = [
            row for row in rows
            if (closed := parse_trade_time(row.get("close_time_msc") or row.get("close_time"))) and closed >= start_dt
        ]
    if end_dt:
        rows = [
            row for row in rows
            if (closed := parse_trade_time(row.get("close_time_msc") or row.get("close_time"))) and closed < end_dt
        ]
    return rows


def run_toxic_job(job_id: str, login: str, mode: str, type_ids: list[str], filters: dict) -> None:
    started = time.time()
    stage_started = time.monotonic()
    stage_timings: dict[str, float] = {}
    update_toxic_job(job_id, status="running", message="正在读取订单进行专项检测", percent=8, startedAt=now_text())
    try:
        rows = query_toxic_scope_trades(login, filters)
        stage_timings["orderQuerySeconds"] = rounded(time.monotonic() - stage_started, 3)
        if not rows:
            raise RuntimeError("当前筛选范围没有可检测订单")
        selected = list(TOXIC_CHECK_TYPE_MAP) if mode == "screen" else type_ids
        uses_tick = mode == "selected" and any(TOXIC_CHECK_TYPE_MAP[type_id]["requiresTick"] for type_id in selected)
        fingerprint = "|".join([
            login, mode, ",".join(sorted(selected)), str(len(rows)),
            max((normalize_text(row.get("close_time_msc") or row.get("close_time") or row.get("open_time")) for row in rows), default=""),
            json.dumps(filters, ensure_ascii=False, sort_keys=True),
        ])
        cache_key = hashlib.sha256(fingerprint.encode("utf-8")).hexdigest()
        cached = None if uses_tick else TOXIC_RESULT_CACHE.get(cache_key)
        if cached:
            update_toxic_job(
                job_id,
                status="done",
                message="已使用相同订单版本的缓存结果",
                percent=100,
                cached=True,
                result=cached,
                stageTimings=stage_timings,
                elapsedSeconds=rounded(time.time()-started, 1),
            )
            return
        stage_started = time.monotonic()
        push_context = toxic_build_push_context(rows) if "market_pushing" in selected else None
        push_rows = push_context["rows"] if push_context else rows
        push_filter = push_context["filter"] if push_context else None
        filter_summary = ({key: value for key, value in push_filter.items() if key not in {"rows", "excludedRows"}} if push_filter else None)
        filter_message = f"，过滤后 {len(push_rows)} 条核心订单" if push_context and len(push_rows) != len(rows) else ""
        update_toxic_job(job_id, message=f"已读取 {len(rows)} 条订单{filter_message}，正在计算资金与事件簇", percent=35)
        metrics = trade_metrics(rows)
        finance = toxic_finance_summary(login, rows, metrics)
        push_metrics = trade_metrics(push_rows) if push_context else metrics
        push_finance = toxic_finance_summary(login, push_rows, push_metrics) if push_context else finance
        stage_timings["prepareSeconds"] = rounded(time.monotonic() - stage_started, 3)
        push_tick = None
        raw_tick = None
        sync = None
        if mode == "selected" and "market_pushing" in selected:
            update_toxic_job(job_id, message="正在查询同品种同方向的跨账户同步订单", percent=48)
            stage_started = time.monotonic()
            sync = toxic_cross_account_sync(login, push_rows)
            stage_timings["crossAccountSyncSeconds"] = rounded(time.monotonic() - stage_started, 3)
            update_toxic_job(job_id, stageTimings=dict(stage_timings))
        if mode == "selected" and "market_pushing" in selected:
            sync_elapsed = stage_timings.get("crossAccountSyncSeconds", 0)
            update_toxic_job(job_id, message=f"跨账户同步完成（{sync_elapsed}s），正在读取候选时间段 Tick", percent=62)
            stage_started = time.monotonic()
            push_tick = toxic_winning_ticks(login, push_rows)
            stage_timings["pushTickSeconds"] = rounded(time.monotonic() - stage_started, 3)
        other_types = [type_id for type_id in selected if type_id != "market_pushing"]
        if mode == "selected" and any(TOXIC_CHECK_TYPE_MAP[type_id]["requiresTick"] for type_id in other_types):
            update_toxic_job(job_id, message="正在读取其他专项所需的完整订单 Tick", percent=68)
            stage_started = time.monotonic()
            raw_tick = toxic_winning_ticks(login, rows)
            stage_timings["otherTickSeconds"] = rounded(time.monotonic() - stage_started, 3)
        update_toxic_job(job_id, message="正在汇总各类型分数和证据", percent=82)
        stage_started = time.monotonic()
        stage = "initial" if mode == "screen" else "deep"
        results = []
        if "market_pushing" in selected:
            results.extend(calculate_toxic_results(login, push_rows, ["market_pushing"], stage, push_finance, push_tick, sync, push_context=push_context))
        if other_types:
            results.extend(calculate_toxic_results(login, rows, other_types, stage, finance, raw_tick, None))
        results.sort(key=lambda item: (-numeric_value(item.get("score")), item.get("label", "")))
        stage_timings["scoringSeconds"] = rounded(time.monotonic() - stage_started, 3)
        tick = push_tick if "market_pushing" in selected else raw_tick
        stage_timings["totalSeconds"] = rounded(time.time() - started, 3)
        result = {
            "account": login, "mode": mode, "orderCount": len(rows), "analysisOrderCount": len(push_rows) if push_context else len(rows), "results": results,
            "financeAvailable": bool(finance.get("available")), "tick": tick, "pushSync": sync,
            "rawTick": raw_tick if raw_tick is not tick else None, "pushOrderFilter": filter_summary, "filters": filters,
            "performance": stage_timings, "completedAt": now_text(),
        }
        if not uses_tick:
            TOXIC_RESULT_CACHE[cache_key] = result
        sync_text = f"，跨账户同步 {stage_timings['crossAccountSyncSeconds']}s" if "crossAccountSyncSeconds" in stage_timings else ""
        tick_seconds = stage_timings.get("pushTickSeconds", 0) + stage_timings.get("otherTickSeconds", 0)
        tick_text = f"，Tick {rounded(tick_seconds, 3)}s" if tick_seconds else ""
        update_toxic_job(
            job_id,
            status="done",
            message=f"Toxic 检测完成（订单读取 {stage_timings['orderQuerySeconds']}s{sync_text}{tick_text}）",
            percent=100,
            cached=False,
            result=result,
            stageTimings=stage_timings,
            elapsedSeconds=rounded(time.time()-started, 1),
        )
    except Exception as exc:
        stage_timings["totalSeconds"] = rounded(time.time() - started, 3)
        update_toxic_job(
            job_id,
            status="failed",
            message=str(exc),
            percent=100,
            error=str(exc),
            stageTimings=stage_timings,
            elapsedSeconds=rounded(time.time()-started, 1),
        )


def safe_stem_text(value: str) -> str:
    text = re.sub(r"[^0-9A-Za-z_.-]+", "_", normalize_text(value)).strip("_")
    return text or "account"


def canonical_trade_row(row: dict) -> dict[str, object]:
    open_time = parse_trade_time(row.get("open_time"))
    close_time = parse_trade_time(row.get("close_time")) or open_time
    holding_seconds = row.get("holding_seconds")
    if holding_seconds in (None, "") and open_time and close_time:
        holding_seconds = (close_time - open_time).total_seconds()
    return {
        "Ticket": normalize_text(row.get("ticket")),
        "Open Time": trade_time_text(open_time),
        "Close Time": trade_time_text(close_time),
        "Type": clean_trade_type(row.get("type", "")),
        "Volume": row.get("volume") or 0,
        "Item": normalize_text(row.get("symbol")),
        "Open Price": row.get("open_price") or 0,
        "Close Price": row.get("close_price") or 0,
        "Commission": row.get("commission") or 0,
        "Taxes": row.get("taxes") or 0,
        "Swap": row.get("swap") or 0,
        "Profit": row.get("profit") or 0,
        "S/L": "",
        "T/P": "",
        "Reason": normalize_text(row.get("reason")),
        "Comment": normalize_text(row.get("comment")),
        "Expert ID": normalize_text(row.get("expert_id")),
        "Holding Seconds": holding_seconds or 0,
        "Account Currency": normalize_text(row.get("account_currency")),
        "Display Currency": normalize_text(row.get("display_currency")),
        "Money Scale": numeric_value(row.get("money_scale")) or 1.0,
        "Is Cent Account": bool(row.get("is_cent_account")),
        "Platform": normalize_text(row.get("platform")),
        "Server": normalize_text(row.get("server")),
    }


def write_trades_csv_from_db(account: str, rows: list[dict]) -> tuple[Path, str]:
    chartable = [row for row in rows if is_chartable_trade(row)]
    if not chartable:
        raise RuntimeError("筛选后没有可画图的 buy/sell 平仓订单")
    canonical = [canonical_trade_row(row) for row in chartable]
    start = min(parse_trade_time(row["Open Time"]) for row in canonical if row["Open Time"])
    end = max(parse_trade_time(row["Close Time"]) for row in canonical if row["Close Time"])
    stem = f"{safe_stem_text(account)}_{start:%Y%m%d_%H%M%S}_{end:%Y%m%d_%H%M%S}"
    KLINE_OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = KLINE_OUT_DIR / f"{stem}_trades.csv"
    fields = ["Ticket", "Open Time", "Close Time", "Type", "Volume", "Item", "Open Price", "Close Price", "Commission", "Taxes", "Swap", "Profit", "S/L", "T/P", "Reason", "Comment", "Expert ID", "Holding Seconds", "Account Currency", "Display Currency", "Money Scale", "Is Cent Account", "Platform", "Server"]
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        writer.writerows(canonical)
    return path, stem


def update_kline_job(job_id: str, **updates) -> None:
    with KLINE_JOBS_LOCK:
        job = KLINE_JOBS.setdefault(job_id, {"id": job_id, "status": "queued", "logs": ""})
        job.update(updates)


def append_kline_job_log(job_id: str, text: str) -> None:
    with KLINE_JOBS_LOCK:
        job = KLINE_JOBS.setdefault(job_id, {"id": job_id, "status": "queued", "logs": ""})
        job["logs"] = (job.get("logs", "") + text)[-20000:]
        if text.strip():
            job["message"] = text.strip()[-240:]


def get_kline_job(job_id: str) -> dict:
    with KLINE_JOBS_LOCK:
        return dict(KLINE_JOBS.get(job_id, {"id": job_id, "status": "missing", "message": "任务不存在"}))


def parse_generated_html_path(output: str) -> Path | None:
    matches = re.findall(r"[A-Za-z]:\\[^\r\n]+?_trade_kline\.html", output)
    for match in reversed(matches):
        path = Path(match.strip())
        if path.exists():
            return path
    return None


def unavailable_kline_timeline(reason: str) -> dict:
    return {
        "version": 1,
        "available": False,
        "reason": reason,
        "events": [],
        "curve": [],
        "liquidationPoints": [],
        "summary": {"currency": "USD", "eventCount": 0, "allEventCount": 0},
        "openingState": {"timestamp": "", "balance": None, "credit": None, "known": False},
    }


def build_db_kline_timeline(account: str, filters: dict) -> dict:
    """Slice a cached full-account Balance/Credit replay for a K-line chart."""
    route_filters = {
        "platform": normalize_text(filters.get("platform")),
        "server": normalize_text(filters.get("server")),
    }
    try:
        from kdesk.application.kline_timeline_cache import KlineTimelineCache
        from kdesk.domain.historical_funds import build_historical_funds
        from kdesk.domain.kline_timeline import build_kline_timeline

        def build_full_replay() -> dict:
            raw = account_historical_funds_source_payload(account, route_filters)
            if not raw.get("available"):
                return {"available": False, "reason": normalize_text(raw.get("reason")) or "账户资金时间线不可用"}
            return {
                "available": True,
                "replay": build_historical_funds(
                    platform=raw.get("platform", ""),
                    currency=raw.get("currency", "USD"),
                    money_scale=numeric_value(raw.get("moneyScale")) or 1.0,
                    events=raw.get("events", []),
                    anchors=raw.get("anchors", []),
                    current_anchor=raw.get("currentAnchor"),
                ),
                "source": {
                    "platform": raw.get("platform", ""),
                    "server": raw.get("server", ""),
                    "source": raw.get("source", ""),
                    "coverage": raw.get("coverage", {}),
                },
            }

        cached, cache_status = KlineTimelineCache(KLINE_TIMELINE_CACHE_DIR).get_or_build(
            account,
            route_filters["platform"],
            route_filters["server"],
            build_full_replay,
            refresh=bool(filters.get("refreshTimelineCache")),
        )
        if not cached.get("available"):
            return unavailable_kline_timeline(normalize_text(cached.get("reason")) or "账户资金时间线不可用")
        source = cached.get("source") if isinstance(cached.get("source"), dict) else {}
        timeline = build_kline_timeline(
            cached["replay"],
            start=normalize_text(filters.get("start")),
            end=normalize_text(filters.get("end")),
        )
        timeline.update({
            "available": True,
            "account": normalize_text(account),
            "platform": source.get("platform", route_filters["platform"]),
            "server": source.get("server", route_filters["server"]),
            "source": source.get("source", ""),
            "coverage": source.get("coverage", {}),
            "cacheStatus": cache_status,
        })
        return timeline
    except Exception as exc:
        logger.warning("K-line funds timeline unavailable for %s: %s", account, exc)
        return unavailable_kline_timeline("账户资金时间线读取失败")


def run_db_kline_job(job_id: str, account: str, filters: dict) -> None:
    started = time.time()
    update_kline_job(job_id, status="running", startedAt=now_text(), message="正在读取数据库订单", percent=8)
    try:
        rows = query_db_trades(
            account,
            platform=normalize_text(filters.get("platform")),
            server=normalize_text(filters.get("server")),
            symbol=normalize_text(filters.get("symbol")),
            start=normalize_text(filters.get("start")),
            end=normalize_text(filters.get("end")),
            limit=50000,
        )
        recent_orders = max(0, min(1000, int(filters.get("recentOrders") or 0)))
        if recent_orders:
            rows = recent_chartable_kline_trades(rows, recent_orders)
        update_kline_job(job_id, message=f"已读取 {len(rows)} 条订单，正在写入标准 CSV", percent=20)
        trades_path, stem = write_trades_csv_from_db(account, rows)
        include_timeline = bool(filters.get("includeTimeline"))
        timeline = None
        timeline_path = None
        if include_timeline:
            update_kline_job(job_id, message="正在读取或复用全量资金与 Credit 缓存", percent=27)
            timeline = build_db_kline_timeline(account, filters)
            timeline_path = KLINE_OUT_DIR / f"{stem}_timeline.json"
            timeline_path.write_text(json.dumps(timeline, ensure_ascii=False), encoding="utf-8")
        update_kline_job(job_id, stem=stem, tradesCsv=str(trades_path), message="正在调用 K 线生成脚本", percent=35)
        cmd = [
            str(K_DESK_PYTHON),
            str(TRADE_KLINE_GENERATOR),
            "--trades-csv",
            str(trades_path),
            "--account",
            account,
            "--out-dir",
            str(KLINE_OUT_DIR),
            "--terminal",
            str(TRADE_KLINE_TERMINAL),
            "--platform",
            normalize_text(filters.get("platform")),
            "--server",
            normalize_text(filters.get("server")),
        ]
        if timeline_path is not None:
            cmd.extend(["--timeline-json", str(timeline_path)])
        env = os.environ.copy()
        env["PYTHONIOENCODING"] = "utf-8"
        env.setdefault("K_DESK_ROOT", str(ROOT))
        env.setdefault("TRADE_KLINE_OUT_DIR", str(KLINE_OUT_DIR))
        env.setdefault("TRADE_KLINE_PYDEPS", str(TRADE_KLINE_PYDEPS))
        proc = subprocess.Popen(
            cmd,
            cwd=str(ROOT),
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            encoding="utf-8",
            errors="replace",
            env=env,
        )
        output_parts: list[str] = []
        assert proc.stdout is not None
        for line in proc.stdout:
            output_parts.append(line)
            append_kline_job_log(job_id, line)
            if "cache saved" in line or "cache hit" in line:
                update_kline_job(job_id, percent=70)
            elif "outputs" in line:
                update_kline_job(job_id, percent=90)
        code = proc.wait()
        output = "".join(output_parts)
        generation = {}
        for line in reversed(output.splitlines()):
            if line.startswith("KLINE_RESULT "):
                try:
                    generation = json.loads(line[len("KLINE_RESULT ") :])
                except (TypeError, ValueError, json.JSONDecodeError):
                    generation = {}
                break
        html_path = parse_generated_html_path(output) or (KLINE_OUT_DIR / f"{stem}_trade_kline.html")
        if code == 0 and html_path.exists() and generation.get("symbols"):
            chart = parse_chart_file(html_path)
            update_kline_job(
                job_id,
                status="done",
                message=generation.get("message") or "生成完成",
                percent=100,
                finishedAt=now_text(),
                elapsedSeconds=round(time.time() - started, 1),
                stem=chart["stem"],
                chart=chart,
                htmlPath=str(html_path),
                partial=bool(generation.get("partial")),
                symbols=generation.get("symbols") or [],
                failures=generation.get("failures") or [],
                quoteSources=generation.get("quoteSources") or [],
                timeline=(
                    {
                        "available": bool(timeline.get("available")),
                        "eventCount": timeline.get("summary", {}).get("eventCount", 0),
                        "allEventCount": timeline.get("summary", {}).get("allEventCount", 0),
                        "liquidationCount": timeline.get("summary", {}).get("liquidationCount", 0),
                        "reason": timeline.get("reason", ""),
                        "cacheStatus": timeline.get("cacheStatus", ""),
                    }
                    if timeline
                    else None
                ),
            )
        else:
            failures = generation.get("failures") or []
            update_kline_job(
                job_id,
                status="failed",
                message=(failures[0].get("reason") if failures else generation.get("message")) or "生成失败，请查看日志",
                percent=100,
                finishedAt=now_text(),
                elapsedSeconds=round(time.time() - started, 1),
                returnCode=code,
                partial=False,
                symbols=[],
                failures=failures,
                quoteSources=generation.get("quoteSources") or [],
            )
    except Exception as exc:
        append_kline_job_log(job_id, f"\nERROR: {exc}\n")
        update_kline_job(job_id, status="failed", message=str(exc), percent=100, finishedAt=now_text(), elapsedSeconds=round(time.time() - started, 1))


def public_chart_url(path: Path) -> str:
    return TRADE_KLINE_WEB_URL.rstrip("/") + "/output/" + quote(path.name)


def ai_result_path_for_stem(stem: str) -> Path:
    return KLINE_OUT_DIR / f"{Path(stem).name}_ai_analysis.json"


def read_ai_result_for_stem(stem: str) -> dict | None:
    path = ai_result_path_for_stem(stem)
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None


def parse_chart_file(path: Path) -> dict[str, str | int | bool]:
    name = path.name
    stem = name[: -len("_trade_kline.html")] if name.endswith("_trade_kline.html") else path.stem
    parts = stem.split("_")
    account = parts[0] if parts else ""
    start = ""
    end = ""
    if len(parts) >= 5:
        start = f"{parts[1]}_{parts[2]}"
        end = f"{parts[3]}_{parts[4]}"
    stat = path.stat()
    item = {
        "account": account,
        "name": name,
        "stem": stem,
        "path": str(path),
        "url": public_chart_url(path),
        "size": stat.st_size,
        "sizeText": f"{stat.st_size / 1024 / 1024:.1f} MB",
        "mtime": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
        "start": start,
        "end": end,
    }
    ai = read_ai_result_for_stem(stem)
    if ai:
        item["aiRiskLevel"] = normalize_text(ai.get("risk_level", ""))
        item["aiConclusion"] = normalize_text(ai.get("conclusion", ""))
        item["aiSuggestedNote"] = normalize_text(ai.get("suggested_ledger_note", ""))
        item["aiAnalysisTime"] = normalize_text(ai.get("created_at", ""))
        item["aiResultPath"] = str(ai_result_path_for_stem(stem))
        item["hasAiAnalysis"] = True
    else:
        item["hasAiAnalysis"] = False
    return item


def scan_chart_files(records: list[dict[str, str]] | None = None) -> list[dict]:
    records = records or load_records()
    known_accounts = {record["账号"] for record in records if record["账号"]}
    charts = []
    for path in sorted(KLINE_OUT_DIR.glob("*_trade_kline.html"), key=lambda p: p.stat().st_mtime, reverse=True):
        item = parse_chart_file(path)
        item["inRegistry"] = item["account"] in known_accounts
        item["recordId"] = ""
        for record in records:
            if record["账号"] == item["account"]:
                item["recordId"] = record["记录ID"]
                item["status"] = record["状态"]
                item["action"] = record["建议动作"]
                break
        charts.append(item)
    return charts


def is_today_text(value: str, today: str | None = None) -> bool:
    today = today or datetime.now().strftime("%Y-%m-%d")
    return normalize_text(value).startswith(today)


def daily_report_records() -> list[dict[str, str]]:
    today = datetime.now().strftime("%Y-%m-%d")
    records = load_records()
    history_rows = read_history_rows()
    today_ids = {
        record["记录ID"]
        for record in records
        if is_today_text(record["加入时间"], today) or is_today_text(record["修改时间"], today)
    }
    today_ids.update(
        row["记录ID"]
        for row in history_rows
        if is_today_text(row["修改时间"], today) and row["操作"] != "删除"
    )
    selected = [record for record in records if record["记录ID"] in today_ids]
    return sorted(selected, key=lambda record: (record["加入时间"], record["账号"] or record["关联账号/主体"]))


def daily_report_docx_bytes() -> tuple[str, bytes]:
    today = datetime.now()
    records = daily_report_records()
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"问题账户日报 {today.strftime('%m%d')}")
    run.bold = True
    run.font.size = Pt(16)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["账号", "建议", "备注"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for record in records:
        row = table.add_row().cells
        row[0].text = record["账号"] or record["关联账号/主体"] or record["记录ID"]
        row[1].text = record["建议动作"] or ""
        row[2].text = record["风险/问题备注"] or ""

    output_path = OUT_DIR / f"journal_{today.strftime('%m%d')}.docx"
    document.save(output_path)
    return output_path.name, output_path.read_bytes()


def normalize_quick_action(value: object) -> str:
    action = normalize_text(value).strip()
    if not action:
        raise ValueError("快捷标记不能为空")
    if len(action) > 40:
        raise ValueError("快捷标记最多 40 个字符")
    if any(ord(char) < 32 for char in action):
        raise ValueError("快捷标记包含无效字符")
    return action


def normalize_quick_actions(values: object) -> list[str]:
    actions = []
    if isinstance(values, list):
        for value in values:
            try:
                action = normalize_quick_action(value)
            except ValueError:
                continue
            if action not in actions:
                actions.append(action)
    if "自定义" not in actions:
        actions.append("自定义")
    return actions[:60]


def load_quick_actions() -> list[str]:
    with QUICK_ACTIONS_LOCK:
        try:
            payload = json.loads(QUICK_ACTIONS_PATH.read_text(encoding="utf-8"))
            values = payload.get("actions") if isinstance(payload, dict) else payload
            actions = normalize_quick_actions(values)
            return actions or list(DEFAULT_QUICK_ACTIONS)
        except (FileNotFoundError, OSError, ValueError, json.JSONDecodeError):
            return list(DEFAULT_QUICK_ACTIONS)


def save_quick_actions(actions: list[str]) -> list[str]:
    normalized = normalize_quick_actions(actions)
    with QUICK_ACTIONS_LOCK:
        QUICK_ACTIONS_PATH.parent.mkdir(parents=True, exist_ok=True)
        temp_path = QUICK_ACTIONS_PATH.with_suffix(QUICK_ACTIONS_PATH.suffix + ".tmp")
        temp_path.write_text(json.dumps({"actions": normalized}, ensure_ascii=False, indent=2), encoding="utf-8")
        temp_path.replace(QUICK_ACTIONS_PATH)
    return normalized


def add_quick_action(value: object) -> list[str]:
    action = normalize_quick_action(value)
    actions = load_quick_actions()
    if action not in actions:
        if len(actions) >= 60:
            raise ValueError("快捷标记最多保存 60 个")
        custom_index = actions.index("自定义") if "自定义" in actions else len(actions)
        actions.insert(custom_index, action)
    return save_quick_actions(actions)


def delete_quick_action(value: object) -> list[str]:
    action = normalize_quick_action(value)
    if action in PROTECTED_QUICK_ACTIONS:
        raise ValueError(f"{action} 是系统保留标记，不能删除")
    actions = [item for item in load_quick_actions() if item != action]
    return save_quick_actions(actions)


def action_choices_for(records: list[dict[str, str]]) -> list[str]:
    choices = list(ACTION_CHOICES)
    for record in records:
        action = normalize_text(record.get("建议动作"))
        if action and action not in choices:
            insert_at = max(len(choices) - 2, 0)
            choices.insert(insert_at, action)
    return choices


def json_response(handler: BaseHTTPRequestHandler, payload: object, status: int = 200) -> None:
    data = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Cache-Control", "no-store")
    handler.send_header("Content-Length", str(len(data)))
    handler.end_headers()
    handler.wfile.write(data)


def html_response(handler: BaseHTTPRequestHandler, body: str, status: int = 200) -> None:
    data = body.encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "text/html; charset=utf-8")
    handler.send_header("Cache-Control", "no-store")
    handler.send_header("Content-Length", str(len(data)))
    handler.end_headers()
    handler.wfile.write(data)


def error_response(handler: BaseHTTPRequestHandler, message: str, status: int = 400) -> None:
    json_response(handler, {"ok": False, "error": message}, status)


def parse_body(handler: BaseHTTPRequestHandler) -> dict:
    length = int(handler.headers.get("Content-Length", "0") or "0")
    raw = handler.rfile.read(length) if length else b"{}"
    return json.loads(raw.decode("utf-8") or "{}")


def normalize_payload(payload: dict, existing: dict[str, str] | None = None) -> dict[str, str]:
    existing = existing or {}
    rec = {header: normalize_text(existing.get(header, "")) for header in HEADERS}
    for field in EDITABLE_FIELDS:
        if field in payload:
            rec[field] = normalize_text(payload[field])
    if not rec["记录类型"]:
        rec["记录类型"] = "账户" if rec["账号"] else "其他"
    if not rec["建议动作"]:
        rec["建议动作"] = derive_action(rec["风险/问题备注"])
    if not rec["当前分组"]:
        rec["当前分组"] = derive_group(rec["建议动作"])
    if not rec["风险标签"]:
        rec["风险标签"] = derive_tags(rec["风险/问题备注"], rec["记录类型"])
    if not rec["状态"]:
        rec["状态"] = "待复核"
    if not rec["加入时间"]:
        rec["加入时间"] = now_text()
    else:
        rec["加入时间"] = normalize_datetime(rec["加入时间"])
    rec["修改时间"] = now_text()
    rec["记录ID"] = existing.get("记录ID") or make_record_id(
        rec["账号"], f"{rec['关联账号/主体']}:{rec['风险/问题备注']}:{uuid.uuid4().hex}"
    )
    return rec


def normalized_mark_payload(payload: dict, login: str) -> dict[str, str]:
    field_map = {
        "recordType": "记录类型",
        "related": "关联账号/主体",
        "action": "建议动作",
        "group": "当前分组",
        "tags": "风险标签",
        "note": "风险/问题备注",
        "rawNote": "原始记录",
        "status": "状态",
        "owner": "处理人/来源",
    }
    normalized = {"账号": login}
    for source, target in field_map.items():
        if source in payload:
            normalized[target] = payload[source]
    for field in EDITABLE_FIELDS:
        if field in payload:
            normalized[field] = payload[field]
    return normalized


def mark_account_records(payload: dict, accounts: list[object]) -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    logins = []
    for value in accounts:
        login = normalize_text(value)
        if not login:
            continue
        if len(login) > 64 or not re.fullmatch(r"[0-9A-Za-z_.-]+", login):
            raise ValueError(f"账号格式无效: {login}")
        if login not in logins:
            logins.append(login)
    if not logins:
        raise ValueError("请输入账号")
    if len(logins) > 50:
        raise ValueError("单次最多批量保存 50 个账号")

    records = load_records()
    history = []
    saved_records = []
    for login in logins:
        normalized = normalized_mark_payload(payload, login)
        existing = ledger_record_for_login(login, records)
        if existing:
            saved = normalize_payload(normalized, existing)
            for index, record in enumerate(records):
                if record["记录ID"] == existing["记录ID"]:
                    records[index] = saved
                    break
            changed = history_changed_fields(existing, saved)
            if changed:
                history.append(make_history_row(existing, saved, "修改", changed))
        else:
            saved = normalize_payload(normalized)
            records.append(saved)
            history.append(make_history_row({}, saved, "加入", "新增台账记录"))
        saved_records.append(saved)
    records.sort(key=record_sort_key)
    save_records(records, history)
    return saved_records, records


def mark_account_record(payload: dict) -> tuple[dict[str, str], list[dict[str, str]]]:
    login = normalize_text(payload.get("account") or payload.get("账号"))
    saved, records = mark_account_records(payload, [login])
    return saved[0], records


def account_logs_payload(account: str, start: str, end: str) -> dict:
    """Read-only MySQL trade-export query, filtered by account and exact time range."""
    return query_account_logs(account, start, end, MYSQL_SOURCES, mysql_trade_connect)


REBATE_CHURNING_SERVICE = rebate_churning.RebateChurningService(
    MYSQL_SOURCES,
    mysql_trade_connect,
    classify_mt5_cashflows=classify_mt5_cashflows,
    classify_mt4_cashflows=classify_mt4_cashflows,
    now_text=now_text,
)


def rebate_churning_account_audit_payload(
    account: str,
    start: str = "",
    end: str = "",
    environment: str = "",
    server_code: str = "",
) -> dict:
    """Return the bounded, read-only account-to-IB rebate audit."""
    return REBATE_CHURNING_SERVICE.target_account_audit(
        account,
        start,
        end,
        environment,
        server_code,
    )


def rebate_churning_ib_payload(
    environment: str,
    ib_id: int,
    start: str = "",
    end: str = "",
) -> dict:
    """Return one recipient IB's read-only audit tree for full-platform drill-down."""
    return REBATE_CHURNING_SERVICE.ib_detail(environment, ib_id, start, end)


class Handler(BaseHTTPRequestHandler):
    server_version = "AccountRegistry/1.0"

    def log_message(self, fmt: str, *args) -> None:
        sys.stderr.write("[%s] %s\n" % (now_text(), fmt % args))

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        path = parsed.path
        if path == "/":
            html_response(self, WORKBENCH_HTML.replace("http://127.0.0.1:8766", TRADE_KLINE_WEB_URL))
            return
        account_page_match = re.fullmatch(r"/account/([^/]+)", path)
        if account_page_match:
            login = unquote(account_page_match.group(1))
            if len(login) > 64 or not re.fullmatch(r"[0-9A-Za-z_.-]+", login):
                html_response(self, "<h1>账号格式无效</h1>", 400)
                return
            page = ACCOUNT_DETAIL_HTML.replace("__ACCOUNT_LOGIN_JSON__", json.dumps(login, ensure_ascii=False))
            html_response(self, page.replace("http://127.0.0.1:8766", TRADE_KLINE_WEB_URL))
            return
        if path == "/api/account-lookup":
            params = parse_qs(parsed.query)
            login = normalize_text((params.get("account") or [""])[0])
            if not login:
                error_response(self, "请输入账号")
                return
            records = load_records()
            record = ledger_record_for_login(login, records)
            databases = account_lookup_databases(login)
            database = databases[0] if databases else {
                "exists": False,
                "account": login,
                "orderCount": 0,
                "chartableOrderCount": 0,
                "symbols": [],
                "latestSource": {},
                "accountMeta": account_money_meta(),
                "refreshedAt": now_text(),
            }
            json_response(self, {
                "ok": True,
                "account": login,
                "marked": bool(record),
                "record": public_ledger_record(record),
                "database": database,
                "databases": databases,
            })
            return
        if path == "/api/account-lookup-finance":
            params = parse_qs(parsed.query)
            login = normalize_text((params.get("account") or [""])[0])
            platform = normalize_text((params.get("platform") or [""])[0])
            server = normalize_text((params.get("server") or [""])[0])
            try:
                json_response(self, account_lookup_finance_payload(login, platform, server))
            except ValueError as exc:
                error_response(self, str(exc))
            return
        if path == "/api/account-logs":
            params = parse_qs(parsed.query)
            account = normalize_text((params.get("account") or [""])[0])
            start = normalize_text((params.get("start") or [""])[0])
            end = normalize_text((params.get("end") or [""])[0])
            try:
                json_response(self, account_logs_payload(account, start, end))
            except ValueError as exc:
                error_response(self, str(exc))
            except RuntimeError as exc:
                error_response(self, str(exc), 503)
            except Exception as exc:
                self.log_message("account log query failed for %s: %s", account, exc)
                error_response(self, "日志查询失败，请稍后重试", 500)
            return
        if path == "/api/hierarchy-products":
            try:
                json_response(self, hierarchy_products_payload())
            except Exception as exc:
                self.log_message("hierarchy product query failed: %s", exc)
                error_response(self, "产品列表读取失败，请稍后重试", 500)
            return
        if path == "/api/hierarchy-net-deposit":
            params = parse_qs(parsed.query)
            target = normalize_text((params.get("target") or [""])[0])
            start = normalize_text((params.get("start") or [""])[0])
            end = normalize_text((params.get("end") or [""])[0])
            product = normalize_text((params.get("product") or [""])[0])
            activity_rules = normalize_text((params.get("activityRules") or [""])[0])
            try:
                json_response(self, hierarchy_net_deposit_payload(target, start, end, product, activity_rules))
            except hierarchy_net_deposit.AmbiguousTargetError as exc:
                json_response(self, {"ok": False, "error": str(exc), "candidates": exc.candidates}, 409)
            except ValueError as exc:
                error_response(self, str(exc))
            except Exception as exc:
                self.log_message("hierarchy net deposit query failed: %s", exc)
                error_response(self, "下线净入金查询失败，请检查数据库连接后重试", 500)
            return
        detail_match = re.fullmatch(r"/api/accounts/by-login/([^/]+)/detail", path)
        if detail_match:
            login = unquote(detail_match.group(1))
            filters = {key: values[0] for key, values in parse_qs(parsed.query).items() if values}
            json_response(self, account_detail_payload(login, filters))
            return
        risk_panels_match = re.fullmatch(r"/api/accounts/by-login/([^/]+)/risk-panels", path)
        if risk_panels_match:
            login = unquote(risk_panels_match.group(1))
            filters = {key: values[0] for key, values in parse_qs(parsed.query).items() if values}
            json_response(self, account_risk_panels_payload(login, filters))
            return
        copy_origins_match = re.fullmatch(r"/api/accounts/by-login/([^/]+)/copy-origins", path)
        if copy_origins_match:
            login = unquote(copy_origins_match.group(1))
            filters = {key: values[0] for key, values in parse_qs(parsed.query).items() if values}
            try:
                json_response(self, account_copy_origins_payload(login, filters))
            except ValueError as exc:
                error_response(self, str(exc))
            return
        copy_group_profit_match = re.fullmatch(r"/api/accounts/by-login/([^/]+)/copy-group-profit", path)
        if copy_group_profit_match:
            login = unquote(copy_group_profit_match.group(1))
            filters = {key: values[0] for key, values in parse_qs(parsed.query).items() if values}
            try:
                json_response(self, account_copy_group_profit_payload(login, filters))
            except ValueError as exc:
                error_response(self, str(exc))
            return
        ea_comment_profit_match = re.fullmatch(r"/api/accounts/by-login/([^/]+)/ea-comment-profit", path)
        if ea_comment_profit_match:
            login = unquote(ea_comment_profit_match.group(1))
            filters = {key: values[0] for key, values in parse_qs(parsed.query).items() if values}
            try:
                json_response(self, account_ea_comment_profit_payload(login, filters))
            except ValueError as exc:
                error_response(self, str(exc))
            return
        automation_match = re.fullmatch(r"/api/accounts/by-login/([^/]+)/automation-analysis", path)
        if automation_match:
            login = unquote(automation_match.group(1))
            filters = {key: values[0] for key, values in parse_qs(parsed.query).items() if values}
            try:
                json_response(self, account_automation_payload(login, filters))
            except ValueError as exc:
                error_response(self, str(exc))
            except Exception as exc:
                self.log_message("automation analysis failed for %s: %s", login, exc)
                error_response(self, "跟单 / EA 分析失败，请稍后重试", 500)
            return
        ledger_match = re.fullmatch(r"/api/accounts/by-login/([^/]+)/ledger", path)
        if ledger_match:
            json_response(self, account_ledger_payload(unquote(ledger_match.group(1))))
            return
        login_ips_match = re.fullmatch(r"/api/accounts/by-login/([^/]+)/login-ips", path)
        if login_ips_match:
            try:
                json_response(self, account_login_ips_payload(unquote(login_ips_match.group(1))))
            except ValueError as exc:
                error_response(self, str(exc))
            return
        orders_match = re.fullmatch(r"/api/accounts/by-login/([^/]+)/orders", path)
        if orders_match:
            login = unquote(orders_match.group(1))
            params = parse_qs(parsed.query)
            try:
                page = max(mysql_int((params.get("page") or [1])[0], 1), 1)
                page_size = min(max(mysql_int((params.get("pageSize") or [100])[0], 100), 20), 200)
                platform = normalize_text((params.get("platform") or [""])[0])
                server = normalize_text((params.get("server") or [""])[0])
                json_response(self, account_orders_payload(login, page, page_size, platform, server))
            except Exception as exc:
                error_response(self, str(exc))
            return
        if path == "/api/accounts":
            records = load_records()
            history_counts: dict[str, int] = {}
            for row in read_history_rows():
                if row["操作"] == "修改":
                    history_counts[row["记录ID"]] = history_counts.get(row["记录ID"], 0) + 1
            json_response(
                self,
                {
                    "ok": True,
                    "summary": summarize(records),
                    "records": records,
                    "historyCounts": history_counts,
                    "statuses": STATUS_CHOICES,
                },
            )
            return
        if path == "/api/quick-actions":
            json_response(self, {
                "ok": True,
                "actions": load_quick_actions(),
                "protected": sorted(PROTECTED_QUICK_ACTIONS),
            })
            return
        if path == "/api/charts":
            records = load_records()
            charts = scan_chart_files(records)
            json_response(
                self,
                {
                    "ok": True,
                    "charts": charts,
                    "summary": {
                        "total": len(charts),
                        "linked": sum(1 for chart in charts if chart.get("inRegistry")),
                        "unlinked": sum(1 for chart in charts if not chart.get("inRegistry")),
                        "uploadUrl": TRADE_KLINE_WEB_URL,
                    },
                },
            )
            return
        chart_match = re.match(r"^/api/accounts/(.+)/charts$", path)
        if chart_match:
            record_id = unquote(chart_match.group(1))
            records = load_records()
            record = next((row for row in records if row["记录ID"] == record_id), None)
            if not record:
                error_response(self, "记录不存在", 404)
                return
            charts = [
                chart
                for chart in scan_chart_files(records)
                if chart["account"] == record["账号"]
            ]
            json_response(self, {"ok": True, "record": record, "charts": charts})
            return
        trades_match = re.match(r"^/api/accounts/(.+)/trades/summary$", path)
        if trades_match:
            record_id = unquote(trades_match.group(1))
            records = load_records()
            record = next((row for row in records if row["记录ID"] == record_id), None)
            if not record:
                error_response(self, "记录不存在", 404)
                return
            filters = {key: values[0] for key, values in parse_qs(parsed.query).items() if values}
            summary = trade_summary_for_account(record["账号"], filters)
            json_response(self, {"ok": True, "record": record, "summary": summary})
            return
        if path == "/api/trades/summary":
            filters = {key: values[0] for key, values in parse_qs(parsed.query).items() if values}
            account = normalize_text(filters.pop("account", ""))
            summary = trade_summary_for_account(account, filters)
            json_response(self, {"ok": True, "summary": summary})
            return
        job_match = re.match(r"^/api/kline/jobs/(.+)$", path)
        if job_match:
            json_response(self, {"ok": True, "job": get_kline_job(unquote(job_match.group(1)))})
            return
        toxic_job_match = re.match(r"^/api/toxic/jobs/(.+)$", path)
        if toxic_job_match:
            json_response(self, {"ok": True, "job": get_toxic_job(unquote(toxic_job_match.group(1)))})
            return
        push_discovery_job_match = re.match(r"^/api/push-discovery/jobs/(.+)$", path)
        if push_discovery_job_match:
            json_response(self, {"ok": True, "job": get_push_discovery_job(unquote(push_discovery_job_match.group(1)))})
            return
        history_match = re.match(r"^/api/accounts/(.+)/history$", path)
        if history_match:
            record_id = unquote(history_match.group(1))
            rows = [row for row in read_history_rows() if row["记录ID"] == record_id]
            rows.sort(key=lambda row: row["修改时间"], reverse=True)
            json_response(self, {"ok": True, "history": rows})
            return
        if path == "/api/meta":
            records = load_records()
            json_response(
                self,
                {
                    "ok": True,
                    "summary": summarize(records),
                    "actions": action_choices_for(records),
                    "statuses": STATUS_CHOICES,
                    "types": TYPE_CHOICES,
                },
            )
            return
        if path == "/download/problematic_accounts.xlsx":
            init_workbook()
            data = WORKBOOK_PATH.read_bytes()
            self.send_response(200)
            self.send_header(
                "Content-Type",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            self.send_header(
                "Content-Disposition",
                f"attachment; filename*=UTF-8''{quote(WORKBOOK_PATH.name)}",
            )
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        if path == "/download/daily-report":
            filename, data = daily_report_docx_bytes()
            self.send_response(200)
            self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header(
                "Content-Disposition",
                f"attachment; filename*=UTF-8''{quote(filename)}",
            )
            self.send_header("Cache-Control", "no-store")
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        file_match = re.match(r"^/chart-file/(.+)$", path)
        if file_match:
            name = Path(unquote(file_match.group(1))).name
            file_path = (KLINE_OUT_DIR / name).resolve()
            if file_path.parent != KLINE_OUT_DIR.resolve() or not file_path.exists() or not name.endswith("_trade_kline.html"):
                error_response(self, "图表文件不存在", 404)
                return
            data = file_path.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Cache-Control", "no-store")
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        error_response(self, "Not found", 404)

    def do_POST(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/api/quick-actions":
            try:
                payload = parse_body(self)
                actions = add_quick_action(payload.get("action") or payload.get("name"))
            except ValueError as exc:
                error_response(self, str(exc))
                return
            json_response(self, {"ok": True, "actions": actions, "protected": sorted(PROTECTED_QUICK_ACTIONS)})
            return
        if parsed.path == "/api/accounts/mark":
            try:
                saved, records = mark_account_record(parse_body(self))
            except ValueError as exc:
                error_response(self, str(exc))
                return
            json_response(self, {"ok": True, "record": public_ledger_record(saved), "summary": summarize(records)})
            return
        if parsed.path == "/api/accounts/mark-batch":
            try:
                payload = parse_body(self)
                accounts = payload.get("accounts")
                if not isinstance(accounts, list):
                    raise ValueError("批量账号格式无效")
                saved, records = mark_account_records(payload, accounts)
            except ValueError as exc:
                error_response(self, str(exc))
                return
            json_response(self, {
                "ok": True,
                "records": [public_ledger_record(record) for record in saved],
                "savedCount": len(saved),
                "summary": summarize(records),
            })
            return
        if parsed.path == "/api/kline/generate-from-db":
            payload = parse_body(self)
            account = normalize_text(payload.get("account"))
            if not account:
                error_response(self, "请输入要查询的账号")
                return
            job_id = f"DBK-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
            filters = {
                "platform": normalize_text(payload.get("platform")),
                "server": normalize_text(payload.get("server")),
                "symbol": normalize_text(payload.get("symbol")),
                "start": normalize_text(payload.get("start")),
                "end": normalize_text(payload.get("end")),
                "includeTimeline": payload_bool(payload.get("includeTimeline")),
                "refreshTimelineCache": payload_bool(payload.get("refreshTimelineCache")),
            }
            if not filters["includeTimeline"]:
                filters["refreshTimelineCache"] = False
            update_kline_job(job_id, status="queued", message="已提交，等待生成", percent=0, account=account, filters=filters, createdAt=now_text())
            threading.Thread(target=run_db_kline_job, args=(job_id, account, filters), daemon=True).start()
            json_response(self, {"ok": True, "job": get_kline_job(job_id)})
            return
        if parsed.path == "/api/push-discovery/start":
            try:
                payload = parse_body(self)
                days = max(mysql_int(payload.get("days"), 7), 1)
                max_orders = mysql_int(payload.get("maxOrders"), 200)
                if days > 30:
                    raise ValueError("扫描天数不能超过30天")
                if not 20 <= max_orders <= 1000:
                    raise ValueError("最大订单数必须在20到1000之间")
                with PUSH_DISCOVERY_JOBS_LOCK:
                    active = next(
                        (dict(job) for job in PUSH_DISCOVERY_JOBS.values() if job.get("status") in {"queued", "running"}),
                        None,
                    )
                if active:
                    json_response(self, {"ok": True, "alreadyRunning": True, "job": active})
                    return
            except ValueError as exc:
                error_response(self, str(exc))
                return
            job_id = f"PDS-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
            options = {
                "days": days,
                "maxOrders": max_orders,
                "smallOrderPriority": min(max_orders, 200),
                "deepLimit": 50,
                "workers": 4,
            }
            update_push_discovery_job(
                job_id, status="queued", percent=0, message="已提交，等待扫描",
                options=options, createdAt=now_text(), results=[],
            )
            threading.Thread(target=run_push_discovery_job, args=(job_id, options), daemon=True).start()
            json_response(self, {"ok": True, "job": get_push_discovery_job(job_id)})
            return
        toxic_match = re.fullmatch(r"/api/accounts/by-login/([^/]+)/toxic-checks", parsed.path)
        if toxic_match:
            login = normalize_text(unquote(toxic_match.group(1)))
            try:
                if not login or len(login) > 64 or not re.fullmatch(r"[0-9A-Za-z_.-]+", login):
                    raise ValueError("账号格式无效")
                payload = parse_body(self)
                mode = normalize_text(payload.get("mode") or "selected")
                if mode not in {"screen", "selected"}:
                    raise ValueError("Toxic 检测模式无效")
                raw_types = payload.get("types") or []
                if not isinstance(raw_types, list):
                    raise ValueError("检测项目格式无效")
                type_ids = [normalize_text(value) for value in raw_types if normalize_text(value)]
                invalid = [value for value in type_ids if value not in TOXIC_CHECK_TYPE_MAP]
                if invalid:
                    raise ValueError(f"未知检测项目：{invalid[0]}")
                if mode == "selected" and not type_ids:
                    raise ValueError("请至少选择一个检测项目")
                filters = {
                    "platform": normalize_text(payload.get("platform")),
                    "server": normalize_text(payload.get("server")),
                    "symbol": normalize_text(payload.get("symbol")),
                    "start": normalize_text(payload.get("start")),
                    "end": normalize_text(payload.get("end")),
                }
            except ValueError as exc:
                error_response(self, str(exc))
                return
            job_id = f"TOX-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
            update_toxic_job(job_id, status="queued", message="已提交，等待检测", percent=0, account=login, mode=mode, types=type_ids, createdAt=now_text())
            threading.Thread(target=run_toxic_job, args=(job_id, login, mode, type_ids, filters), daemon=True).start()
            json_response(self, {"ok": True, "job": get_toxic_job(job_id), "types": TOXIC_CHECK_TYPES})
            return
        gen_match = re.match(r"^/api/accounts/(.+)/kline/generate-from-db$", parsed.path)
        if gen_match:
            record_id = unquote(gen_match.group(1))
            payload = parse_body(self)
            records = load_records()
            record = next((row for row in records if row["记录ID"] == record_id), None)
            if not record:
                error_response(self, "记录不存在", 404)
                return
            account = normalize_text(record["账号"])
            if not account:
                error_response(self, "当前记录没有账号，无法从数据库生成图表")
                return
            job_id = f"DBK-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
            filters = {
                "platform": normalize_text(payload.get("platform")),
                "server": normalize_text(payload.get("server")),
                "symbol": normalize_text(payload.get("symbol")),
                "start": normalize_text(payload.get("start")),
                "end": normalize_text(payload.get("end")),
                "includeTimeline": payload_bool(payload.get("includeTimeline")),
                "refreshTimelineCache": payload_bool(payload.get("refreshTimelineCache")),
            }
            if not filters["includeTimeline"]:
                filters["refreshTimelineCache"] = False
            update_kline_job(job_id, status="queued", message="已提交，等待生成", percent=0, account=account, filters=filters, createdAt=now_text())
            threading.Thread(target=run_db_kline_job, args=(job_id, account, filters), daemon=True).start()
            json_response(self, {"ok": True, "job": get_kline_job(job_id)})
            return
        ai_match = re.match(r"^/api/accounts/(.+)/ai-note/confirm$", parsed.path)
        if ai_match:
            record_id = unquote(ai_match.group(1))
            payload = parse_body(self)
            stem = normalize_text(payload.get("stem") or payload.get("chartStem") or "")
            chart_name = normalize_text(payload.get("chart") or payload.get("chartName") or "")
            if not stem and chart_name.endswith("_trade_kline.html"):
                stem = chart_name[: -len("_trade_kline.html")]
            if not stem:
                error_response(self, "missing AI analysis stem")
                return
            ai = read_ai_result_for_stem(stem)
            if not ai:
                error_response(self, "AI analysis result not found", 404)
                return
            records = load_records()
            for idx, record in enumerate(records):
                if record["记录ID"] == record_id:
                    updated = dict(record)
                    updated["AI风险等级"] = normalize_text(ai.get("risk_level", ""))
                    updated["AI备注"] = normalize_text(ai.get("suggested_ledger_note") or ai.get("conclusion") or "")
                    updated["AI分析时间"] = normalize_text(ai.get("created_at", now_text()))
                    updated["AI证据图表"] = f"{stem}_trade_kline.html"
                    updated["修改时间"] = now_text()
                    records[idx] = updated
                    changed = history_changed_fields(record, updated)
                    history = [make_history_row(record, updated, "修改", changed)] if changed else []
                    save_records(records, history)
                    json_response(self, {"ok": True, "record": updated, "summary": summarize(records)})
                    return
            error_response(self, "record not found", 404)
            return
        if parsed.path == "/api/accounts":
            payload = parse_body(self)
            records = load_records()
            rec = normalize_payload(payload)
            if rec["账号"] and any(r["账号"] == rec["账号"] for r in records):
                error_response(self, f"账号 {rec['账号']} 已存在")
                return
            records.append(rec)
            records.sort(key=record_sort_key)
            save_records(records)
            json_response(self, {"ok": True, "record": rec, "summary": summarize(records)})
            return
        if parsed.path == "/api/rebuild":
            init_workbook(force=True)
            records = load_records()
            json_response(self, {"ok": True, "summary": summarize(records), "records": records})
            return
        error_response(self, "Not found", 404)

    def do_PUT(self) -> None:
        parsed = urlparse(self.path)
        match = re.match(r"^/api/accounts/(.+)$", parsed.path)
        if not match:
            error_response(self, "Not found", 404)
            return
        record_id = unquote(match.group(1))
        payload = parse_body(self)
        records = load_records()
        for idx, record in enumerate(records):
            if record["记录ID"] == record_id:
                updated = normalize_payload(payload, record)
                if updated["账号"] and any(
                    r["账号"] == updated["账号"] and r["记录ID"] != record_id for r in records
                ):
                    error_response(self, f"账号 {updated['账号']} 已存在")
                    return
                records[idx] = updated
                history = []
                changed = history_changed_fields(record, updated)
                if changed:
                    history.append(make_history_row(record, updated, "修改", changed))
                save_records(records, history)
                json_response(self, {"ok": True, "record": updated, "summary": summarize(records)})
                return
        error_response(self, "记录不存在", 404)

    def do_DELETE(self) -> None:
        parsed = urlparse(self.path)
        quick_action_match = re.fullmatch(r"/api/quick-actions/([^/]+)", parsed.path)
        if quick_action_match:
            try:
                actions = delete_quick_action(unquote(quick_action_match.group(1)))
            except ValueError as exc:
                error_response(self, str(exc))
                return
            json_response(self, {"ok": True, "actions": actions, "protected": sorted(PROTECTED_QUICK_ACTIONS)})
            return
        match = re.match(r"^/api/accounts/(.+)$", parsed.path)
        if not match:
            error_response(self, "Not found", 404)
            return
        record_id = unquote(match.group(1))
        records = load_records()
        kept = [record for record in records if record["记录ID"] != record_id]
        if len(kept) == len(records):
            error_response(self, "记录不存在", 404)
            return
        removed = [record for record in records if record["记录ID"] == record_id]
        history = [make_history_row(removed[0], None, "删除", "删除记录")] if removed else []
        save_records(kept, history)
        json_response(self, {"ok": True, "summary": summarize(kept)})


INDEX_HTML = r"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>问题账户台账</title>
  <style>
    :root {
      --bg: #f3f6f8;
      --panel: #ffffff;
      --panel-soft: #f8fbfc;
      --ink: #1f2937;
      --muted: #64737f;
      --line: #d7e0e6;
      --line-soft: #edf2f5;
      --teal: #00796b;
      --teal-soft: #e3f3f0;
      --red: #c62828;
      --red-soft: #ffebee;
      --amber: #b77900;
      --amber-soft: #fff7df;
      --green: #2e7d32;
      --green-soft: #e8f5e9;
      --blue: #2563eb;
      --blue-soft: #eaf1ff;
      --gray: #455a64;
      --focus: #00838f;
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      background: var(--bg);
      color: var(--ink);
      font-family: "Microsoft YaHei", "Segoe UI", Arial, sans-serif;
      font-size: 14px;
    }
    header {
      min-height: 58px;
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      padding: 10px 24px;
      background: #ffffff;
      border-bottom: 1px solid var(--line);
    }
    h1 { margin: 0; font-size: 20px; font-weight: 800; letter-spacing: 0; }
    main { padding: 14px 20px 22px; }
    .summary {
      display: grid;
      grid-template-columns: repeat(4, minmax(150px, 1fr));
      gap: 10px;
      margin-bottom: 10px;
    }
    .metric {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 10px 12px;
      min-height: 62px;
    }
    .metric b { display: block; font-size: 22px; margin-top: 3px; }
    .metric span { color: var(--muted); font-size: 13px; }
    .toolbar {
      display: grid;
      grid-template-columns: minmax(260px, 1fr) 142px 142px 164px auto auto auto auto;
      gap: 8px;
      align-items: center;
      margin-bottom: 10px;
    }
    input, select, textarea {
      width: 100%;
      border: 1px solid var(--line);
      border-radius: 7px;
      background: #fff;
      color: var(--ink);
      font: inherit;
      padding: 8px 10px;
      outline: none;
      min-height: 36px;
    }
    input:focus, select:focus, textarea:focus {
      border-color: var(--focus);
      box-shadow: 0 0 0 3px rgba(0, 151, 167, .14);
    }
    button, .button-link {
      border: 1px solid transparent;
      border-radius: 7px;
      padding: 8px 11px;
      font: inherit;
      cursor: pointer;
      background: var(--gray);
      color: #fff;
      text-decoration: none;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      gap: 6px;
      min-height: 36px;
      white-space: nowrap;
    }
    button:disabled { opacity: .55; cursor: not-allowed; }
    button.primary { background: var(--teal); }
    button.primary:hover { background: #00695c; }
    button.danger { background: var(--red); }
    button.light, .button-link.light {
      background: #fff;
      color: var(--ink);
      border-color: var(--line);
    }
    button.light:hover, .button-link.light:hover { border-color: #aebcc5; background: #f8fbfc; }
    .layout {
      display: grid;
      grid-template-columns: minmax(660px, 1fr) minmax(430px, 480px);
      gap: 12px;
      align-items: start;
    }
    .table-wrap, .editor, .chart-panel {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      overflow: hidden;
    }
    table {
      width: 100%;
      border-collapse: collapse;
      table-layout: fixed;
    }
    th, td {
      border-bottom: 1px solid var(--line);
      padding: 9px 10px;
      vertical-align: top;
      font-size: 13px;
    }
    th {
      position: sticky;
      top: 0;
      z-index: 1;
      background: #263238;
      color: #fff;
      text-align: left;
      font-weight: 700;
      height: 38px;
    }
    tr:hover td { background: #f4fbfa; }
    tr.selected td { background: #edf7f5; box-shadow: inset 3px 0 0 var(--teal); }
    .col-account { width: 132px; }
    .col-action { width: 96px; }
    .col-tags { width: 220px; }
    .col-time { width: 156px; }
    .col-status { width: 100px; }
    .col-tools { width: 118px; text-align: right; }
    .row-actions { display: flex; gap: 6px; justify-content: flex-end; flex-wrap: wrap; }
    .row-actions button { min-height: 30px; padding: 5px 8px; font-size: 12px; }
    .account-main { font-size: 14px; font-weight: 800; word-break: break-word; }
    .account-sub { color: var(--muted); font-size: 12px; margin-top: 2px; display: flex; gap: 5px; flex-wrap: wrap; align-items: center; }
    .row-markers { display: flex; gap: 5px; flex-wrap: wrap; margin-top: 6px; }
    .mini-badge {
      display: inline-flex;
      align-items: center;
      min-height: 19px;
      padding: 1px 6px;
      border-radius: 999px;
      border: 1px solid var(--line);
      color: var(--muted);
      background: #fff;
      font-size: 11px;
      line-height: 16px;
    }
    .note {
      display: -webkit-box;
      -webkit-line-clamp: 2;
      -webkit-box-orient: vertical;
      overflow: hidden;
      line-height: 1.45;
      color: #334155;
    }
    .tags {
      display: flex;
      flex-wrap: wrap;
      gap: 5px;
    }
    .tag {
      display: inline-flex;
      align-items: center;
      max-width: 100%;
      padding: 2px 7px;
      border-radius: 999px;
      background: #eef6f5;
      color: #0f766e;
      font-size: 12px;
      line-height: 18px;
      word-break: break-word;
    }
    .pill {
      display: inline-flex;
      align-items: center;
      min-height: 24px;
      padding: 2px 8px;
      border-radius: 999px;
      background: #eef2f6;
      color: #334155;
      font-size: 12px;
      font-weight: 700;
    }
    .pill.b { background: var(--green-soft); color: #166534; }
    .pill.m { background: var(--blue-soft); color: #1d4ed8; }
    .pill.p { background: var(--amber-soft); color: #8a5a00; }
    .pill.t { background: var(--red-soft); color: #b71c1c; }
    .pill.a { background: #f3e8ff; color: #6b21a8; }
    .pill.custom { background: #eef2f6; color: #334155; }
    .status-pill {
      display: inline-flex;
      align-items: center;
      min-height: 24px;
      padding: 2px 8px;
      border-radius: 999px;
      background: #eef2f6;
      color: #334155;
      font-size: 12px;
      font-weight: 700;
      max-width: 100%;
      word-break: break-word;
    }
    .editor { padding: 12px; position: sticky; top: 72px; max-height: calc(100vh - 86px); overflow: auto; }
    .editor h2 { margin: 0; font-size: 18px; }
    .detail-header {
      display: grid;
      gap: 5px;
      padding-bottom: 10px;
      border-bottom: 1px solid var(--line);
      margin-bottom: 12px;
    }
    .detail-subtitle { color: var(--muted); font-size: 12px; line-height: 1.45; }
    .detail-section {
      border-top: 1px solid var(--line);
      padding-top: 12px;
      margin-top: 12px;
    }
    .form-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 9px; }
    .form-grid label { display: block; color: var(--muted); font-size: 12px; margin-bottom: 5px; }
    .span-2 { grid-column: 1 / -1; }
    textarea { min-height: 78px; resize: vertical; line-height: 1.45; }
    #raw { min-height: 58px; }
    .editor-actions { display: flex; gap: 8px; justify-content: flex-end; margin-top: 12px; }
    .statusbar { min-height: 22px; margin-top: 10px; color: var(--muted); font-size: 13px; }
    .empty { padding: 36px; color: var(--muted); text-align: center; }
    .time-text { color: var(--muted); font-size: 12px; line-height: 1.45; }
    .list-time { font-size: 13px; line-height: 1.5; white-space: normal; word-break: keep-all; }
    .list-time b { color: var(--ink); font-weight: 700; }
    dialog {
      width: min(980px, calc(100vw - 32px));
      max-height: 86vh;
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 0;
      overflow: hidden;
    }
    dialog::backdrop { background: rgba(31, 41, 51, .32); }
    .history-head {
      display: flex;
      justify-content: space-between;
      align-items: center;
      gap: 10px;
      padding: 14px 16px;
      border-bottom: 1px solid var(--line);
      background: #fff;
    }
    .history-body {
      max-height: 72vh;
      overflow: auto;
      padding: 14px 16px 18px 28px;
      background: #f8fafb;
      position: relative;
    }
    .history-item {
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 0;
      margin-bottom: 12px;
      background: #fff;
      overflow: hidden;
      position: relative;
    }
    .history-item::before {
      content: "";
      position: absolute;
      left: -16px;
      top: 17px;
      width: 9px;
      height: 9px;
      border-radius: 999px;
      background: var(--teal);
      box-shadow: 0 0 0 4px #dff3ef;
    }
    .version-head {
      display: flex;
      justify-content: space-between;
      gap: 10px;
      align-items: flex-start;
      padding: 12px 14px;
      border-bottom: 1px solid var(--line);
      background: #fbfdfe;
    }
    .version-title { display: flex; flex-wrap: wrap; gap: 8px; align-items: center; }
    .version-badge {
      display: inline-flex;
      align-items: center;
      min-height: 24px;
      padding: 2px 9px;
      border-radius: 999px;
      background: var(--teal-soft);
      color: #00695c;
      font-weight: 700;
      font-size: 12px;
    }
    .version-badge.add { background: var(--green-soft); color: #166534; }
    .version-badge.delete { background: #ffebee; color: #b71c1c; }
    .version-time { color: var(--muted); font-size: 12px; line-height: 1.5; text-align: right; }
    .version-body { display: grid; grid-template-columns: 1fr; gap: 12px; padding: 12px 14px 14px; }
    .version-panel {
      border: 1px solid var(--line);
      border-radius: 8px;
      overflow: hidden;
      background: #fff;
    }
    .version-panel h4 {
      margin: 0;
      padding: 8px 10px;
      background: #eceff1;
      font-size: 13px;
    }
    .version-panel.after h4 { background: #e0f2ef; color: #00695c; }
    .version-panel.before h4 { background: #fff8e1; color: #7a5600; }
    .version-grid { display: grid; grid-template-columns: 92px minmax(0, 1fr); }
    .version-label, .version-value {
      padding: 8px 10px;
      border-top: 1px solid #edf1f3;
      font-size: 13px;
      line-height: 1.45;
      min-width: 0;
    }
    .version-label { color: var(--muted); background: #fafafa; }
    .version-value { white-space: pre-wrap; word-break: break-word; }
    .version-value.changed {
      background: #fffde7;
      box-shadow: inset 3px 0 0 var(--amber);
    }
    .pagination {
      display: flex;
      justify-content: space-between;
      align-items: center;
      gap: 12px;
      padding: 12px 14px;
      background: #fff;
      border-top: 1px solid var(--line);
      flex-wrap: wrap;
    }
    .pager-buttons { display: flex; gap: 8px; align-items: center; }
    .pager-buttons button { min-height: 34px; padding: 6px 10px; }
    .pager-info { color: var(--muted); font-size: 13px; }
    @media (max-width: 820px) {
      .version-time { text-align: left; }
    }
    .chart-section { margin-top: 14px; border-top: 1px solid var(--line); padding-top: 12px; }
    .section-title { display: flex; justify-content: space-between; align-items: center; gap: 8px; margin-bottom: 8px; }
    .section-title b { font-size: 15px; }
    .chart-list { display: grid; gap: 8px; max-height: 240px; overflow: auto; padding-right: 2px; }
    .chart-item {
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 9px;
      background: #fff;
      display: grid;
      gap: 7px;
    }
    .chart-item.active { border-color: var(--teal); box-shadow: 0 0 0 3px rgba(0, 121, 107, .12); }
    .chart-name { font-weight: 700; word-break: break-all; font-size: 13px; line-height: 1.35; }
    .chart-meta { color: var(--muted); font-size: 12px; line-height: 1.45; }
    .chart-actions { display: flex; flex-wrap: wrap; gap: 6px; }
    .chart-actions a, .chart-actions button { min-height: 32px; padding: 6px 9px; font-size: 13px; }
    .chart-frame {
      width: 100%;
      height: 360px;
      border: 1px solid var(--line);
      border-radius: 8px;
      margin-top: 10px;
      background: #fff;
    }
    .db-trade-card {
      margin-top: 14px;
      border-top: 1px solid var(--line);
      padding-top: 12px;
      display: grid;
      gap: 10px;
    }
    .db-trade-status {
      border: 1px solid var(--line);
      border-radius: 8px;
      background: #fbfdfe;
      padding: 10px;
      font-size: 13px;
      line-height: 1.5;
    }
    .db-trade-status.loading { border-color: #bfdbfe; background: #eff6ff; color: #1d4ed8; }
    .db-trade-status.empty { border-color: var(--line); background: #f8fafc; color: var(--muted); }
    .db-trade-status.error { border-color: #fecaca; background: #fff1f2; color: #991b1b; }
    .db-trade-status.ready { border-color: #bbf7d0; background: #f0fdf4; }
    .db-trade-status b { font-size: 16px; }
    .db-summary-grid {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 8px;
      margin-top: 8px;
    }
    .db-summary-item {
      border: 1px solid rgba(15, 118, 110, .16);
      border-radius: 7px;
      background: #fff;
      padding: 7px 8px;
      min-width: 0;
    }
    .db-summary-item span { display: block; color: var(--muted); font-size: 11px; margin-bottom: 2px; }
    .db-summary-item strong { display: block; font-size: 13px; color: var(--ink); word-break: break-word; }
    .db-trade-controls { display: grid; grid-template-columns: 1fr 1fr; gap: 8px; }
    .db-trade-controls label { color: var(--muted); font-size: 12px; display: grid; gap: 4px; }
    .db-range-buttons { display: flex; gap: 6px; flex-wrap: wrap; }
    .db-range-buttons button { min-height: 30px; padding: 5px 8px; font-size: 12px; }
    .db-range-buttons button.active { background: var(--teal); color: #fff; border-color: var(--teal); }
    .db-job {
      display: none;
      border: 1px solid #cbd5e1;
      border-radius: 8px;
      padding: 9px;
      background: #f8fafc;
      font-size: 12px;
      line-height: 1.5;
    }
    .db-job.active { display: block; }
    .db-progress { height: 6px; background: #e5e7eb; border-radius: 999px; overflow: hidden; margin-top: 6px; }
    .db-progress span { display: block; height: 100%; width: 0%; background: var(--teal); transition: width .2s ease; }
    .db-log { max-height: 90px; overflow: auto; white-space: pre-wrap; color: var(--muted); margin-top: 6px; }
    .generated-actions {
      display: none;
      gap: 6px;
      flex-wrap: wrap;
      margin-top: 8px;
    }
    .generated-actions.active { display: flex; }
    .muted { color: var(--muted); }
    @media (max-width: 1250px) {
      .summary { grid-template-columns: repeat(2, minmax(140px, 1fr)); }
      .layout { grid-template-columns: 1fr; }
      .editor { position: static; }
      .toolbar { grid-template-columns: 1fr 1fr 1fr; }
    }
    @media (max-width: 680px) {
      header { height: auto; gap: 10px; padding: 14px; flex-wrap: wrap; }
      main { padding: 14px; }
      .summary { grid-template-columns: 1fr; }
      .toolbar { grid-template-columns: 1fr; }
      .form-grid { grid-template-columns: 1fr; }
      .db-summary-grid { grid-template-columns: 1fr; }
      .span-2 { grid-column: auto; }
      .table-wrap { overflow-x: auto; }
      table { min-width: 860px; }
    }
  </style>
</head>
<body>
  <header>
    <h1>问题账户台账</h1>
    <a class="button-link light" href="/download/problematic_accounts.xlsx">下载 Excel</a>
  </header>
  <main>
    <section class="summary">
      <div class="metric"><span>记录数</span><b id="total">0</b></div>
      <div class="metric"><span>账号记录</span><b id="accountRecords">0</b></div>
      <div class="metric"><span>IB/组记录</span><b id="groupRecords">0</b></div>
      <div class="metric"><span>更新时间</span><b id="updatedAt" style="font-size:15px">-</b></div>
    </section>
    <section class="toolbar">
      <input id="search" placeholder="搜索账号、标签、备注、IB" />
      <select id="actionFilter"></select>
      <select id="statusFilter"></select>
      <select id="sortBy">
        <option value="joined_desc">加入时间 新到旧</option>
        <option value="account">按账号排序</option>
        <option value="joined_asc">加入时间 旧到新</option>
        <option value="updated_desc">更新时间 新到旧</option>
        <option value="updated_asc">更新时间 旧到新</option>
      </select>
      <button class="primary" id="newBtn">＋ 新增</button>
      <button class="light" id="chartLibraryBtn">图表库</button>
      <a class="button-link light" href="http://127.0.0.1:8766" target="_blank" rel="noopener">生成图表</a>
      <a class="button-link light" href="/download/daily-report">导出日报</a>
      <button class="light" id="reloadBtn">刷新</button>
    </section>
    <section class="layout">
      <div class="table-wrap">
        <table>
          <thead>
            <tr>
              <th class="col-account">账号</th>
              <th class="col-action">建议</th>
              <th class="col-tags">标签</th>
              <th>备注</th>
              <th class="col-time">时间</th>
              <th class="col-status">状态</th>
              <th class="col-tools"></th>
            </tr>
          </thead>
          <tbody id="rows"></tbody>
        </table>
        <div class="empty" id="empty" hidden>没有匹配记录</div>
        <div class="pagination">
          <div class="pager-info" id="pageInfo">第 1 页</div>
          <div class="pager-buttons">
            <button class="light" id="prevPageBtn">上一页</button>
            <button class="light" id="nextPageBtn">下一页</button>
          </div>
        </div>
      </div>
      <aside class="editor">
        <div class="detail-header">
          <h2 id="formTitle">新增记录</h2>
          <div class="detail-subtitle" id="detailSubtitle">未选择记录，可以直接新增或输入账号查询数据库订单。</div>
        </div>
        <div class="section-title">
          <b>基础信息</b>
          <span class="muted">保存后自动同步时间</span>
        </div>
        <div class="form-grid">
          <div>
            <label>账号</label>
            <input id="account" />
          </div>
          <div>
            <label>记录类型</label>
            <select id="type"></select>
          </div>
          <div>
            <label>建议动作</label>
      <select id="action"></select>
            <input id="customAction" placeholder="输入自定义建议动作" style="margin-top:8px;display:none" />
          </div>
          <div>
            <label>当前分组</label>
            <input id="group" />
          </div>
          <div class="span-2">
            <label>关联账号/主体</label>
            <input id="related" />
          </div>
          <div class="span-2">
            <label>风险标签</label>
            <input id="tags" />
          </div>
          <div class="span-2">
            <label>风险/问题备注</label>
            <textarea id="note"></textarea>
          </div>
          <div class="span-2">
            <label>原始记录</label>
            <textarea id="raw"></textarea>
          </div>
          <div>
            <label>状态</label>
            <select id="status"></select>
          </div>
          <div>
            <label>处理人/来源</label>
            <input id="source" />
          </div>
          <div>
            <label>加入时间</label>
            <input id="joinedAt" readonly />
          </div>
          <div>
            <label>修改时间</label>
            <input id="updatedAtField" readonly />
          </div>
        </div>
        <div class="editor-actions">
          <button class="light" id="clearBtn">清空</button>
          <button class="light" id="historyBtn" hidden>历史</button>
          <button class="danger" id="deleteBtn" hidden>删除</button>
          <button class="primary" id="saveBtn">保存</button>
        </div>
        <div class="statusbar" id="statusbar"></div>
        <div class="db-trade-card">
          <div class="section-title">
            <b>数据库查询</b>
            <button class="light" id="dbRefreshBtn" type="button">刷新订单</button>
          </div>
          <div class="db-trade-status" id="dbTradeStatus">选择账号后查询数据库订单</div>
          <div class="db-trade-controls">
            <label class="span-2">查询账号
              <input id="dbAccount" placeholder="输入账号查询数据库订单" />
            </label>
            <label>平台
              <select id="dbPlatform"></select>
            </label>
            <label>服务器
              <select id="dbServer"></select>
            </label>
            <label class="span-2">品种
              <select id="dbSymbol"></select>
            </label>
            <div class="span-2">
              <div class="db-range-buttons" id="dbRangeButtons">
                <button class="light active" type="button" data-range-days="30">30天</button>
                <button class="light" type="button" data-range-days="7">7天</button>
                <button class="light" type="button" data-range-days="90">3个月</button>
                <button class="light" type="button" data-range-days="180">半年</button>
                <button class="light" type="button" data-range-days="365">一年</button>
                <button class="light" type="button" data-range-days="all">全部</button>
              </div>
            </div>
            <label>开始
              <input id="dbStart" placeholder="YYYY-MM-DD HH:MM:SS" />
            </label>
            <label>结束
              <input id="dbEnd" placeholder="YYYY-MM-DD HH:MM:SS" />
            </label>
          </div>
          <button class="primary" id="dbGenerateBtn" type="button">从数据库生成K线图</button>
          <div class="generated-actions" id="generatedActions">
            <a class="button-link light" id="generatedOpenLink" target="_blank" rel="noopener">打开 AI 图表</a>
            <button class="light" id="generatedPreviewBtn" type="button">预览图表</button>
            <button class="light" id="generatedCopyBtn" type="button">复制链接</button>
          </div>
          <div class="db-job" id="dbJobBox">
            <div id="dbJobText">等待生成</div>
            <div class="db-progress"><span id="dbJobProgress"></span></div>
            <div class="db-log" id="dbJobLog"></div>
          </div>
        </div>
        <div class="chart-section">
          <div class="section-title">
            <b>匹配图表</b>
            <span class="muted" id="chartHint">选择账号后显示</span>
          </div>
          <div class="chart-list" id="accountCharts"></div>
          <iframe class="chart-frame" id="chartPreview" title="买卖点图预览" hidden></iframe>
        </div>
      </aside>
    </section>
    <dialog id="historyDialog">
      <div class="history-head">
        <b id="historyTitle">修改历史</b>
        <button class="light" id="closeHistoryBtn">关闭</button>
      </div>
      <div class="history-body" id="historyBody"></div>
    </dialog>
    <dialog id="chartLibraryDialog">
      <div class="history-head">
        <b id="chartLibraryTitle">图表库</b>
        <button class="light" id="closeChartLibraryBtn">关闭</button>
      </div>
      <div class="history-body">
        <div class="toolbar" style="grid-template-columns: minmax(180px, 1fr) 140px auto; padding:0; margin-bottom:10px">
          <input id="chartSearch" placeholder="搜索账号或图表文件" />
          <select id="chartLinkedFilter">
            <option value="">全部图表</option>
            <option value="linked">已入台账</option>
            <option value="unlinked">未入台账</option>
          </select>
          <a class="button-link light" href="http://127.0.0.1:8766" target="_blank" rel="noopener">上传生成</a>
        </div>
        <div class="chart-list" id="chartLibraryList" style="max-height:60vh"></div>
      </div>
    </dialog>
  </main>
  <script>
    const state = { records: [], meta: {}, editingId: null, historyCounts: {}, accountCharts: [], allCharts: [], tradeSummary: null, klineJobTimer: null, page: 1, pageSize: 20, lastFilteredCount: 0 };
    const $ = (id) => document.getElementById(id);
    const fields = {
      "账号": "account",
      "记录类型": "type",
      "关联账号/主体": "related",
      "建议动作": "action",
      "当前分组": "group",
      "风险标签": "tags",
      "风险/问题备注": "note",
      "原始记录": "raw",
      "状态": "status",
      "处理人/来源": "source",
    };

    function setStatus(text, good = true) {
      $("statusbar").textContent = text || "";
      $("statusbar").style.color = good ? "var(--muted)" : "var(--red)";
    }
    function fillSelect(select, values, first) {
      select.innerHTML = "";
      if (first !== undefined) select.append(new Option(first, ""));
      values.forEach((value) => select.append(new Option(value || "未填", value)));
    }
    function ensureSelectOption(select, value) {
      const text = String(value || "").trim();
      if (!text) return;
      if (!Array.from(select.options).some((opt) => opt.value === text)) {
        const customIndex = Array.from(select.options).findIndex((opt) => opt.value === "自定义");
        const option = new Option(text, text);
        if (customIndex >= 0) select.add(option, select.options[customIndex]);
        else select.append(option);
      }
    }
    function actionClass(action) {
      const text = String(action || "").trim().toUpperCase();
      if (!text) return "custom";
      if (text === "A" || text.includes("P->A")) return "a";
      if (text.includes("T")) return "t";
      if (text.includes("P")) return "p";
      if (text.includes("M")) return "m";
      if (text.includes("B")) return "b";
      if (text.includes("自定义")) return "custom";
      return "";
    }
    function statusClass(status) {
      const text = String(status || "");
      if (text.includes("已") || text.includes("完成")) return "status-pill";
      if (text.includes("待") || text.includes("观察")) return "status-pill";
      return "status-pill";
    }
    function escapeText(text) {
      return String(text || "").replace(/[&<>"']/g, (c) => ({
        "&": "&amp;", "<": "&lt;", ">": "&gt;", '"': "&quot;", "'": "&#39;"
      }[c]));
    }
    function splitTags(text) {
      return String(text || "").split(/[；;,，\n]+/).map((x) => x.trim()).filter(Boolean);
    }
    function visibleTime(record) {
      const hasHistory = (state.historyCounts[record["记录ID"]] || 0) > 0;
      return {
        label: hasHistory ? "修改" : "加入",
        value: hasHistory ? (record["修改时间"] || record["加入时间"] || "-") : (record["加入时间"] || "-"),
      };
    }
    function effectiveTimeValue(record) {
      return visibleTime(record).value;
    }
    function renderSummary(summary) {
      $("total").textContent = summary.total || 0;
      $("accountRecords").textContent = summary.accountRecords || 0;
      $("groupRecords").textContent = summary.groupRecords || 0;
      $("updatedAt").textContent = summary.updatedAt || "-";
    }
    function filteredRecords() {
      const q = $("search").value.trim().toLowerCase();
      const action = $("actionFilter").value;
      const status = $("statusFilter").value;
      const rows = state.records.filter((record) => {
        const hay = Object.values(record).join(" ").toLowerCase();
        return (!q || hay.includes(q)) &&
          (!action || record["建议动作"] === action) &&
          (!status || record["状态"] === status);
      });
      const asTime = (value) => {
        const text = String(value || "").trim();
        if (!text) return 0;
        const normalized = text.length <= 10 ? `${text} 00:00:00` : text;
        const parsed = Date.parse(normalized.replace(/-/g, "/"));
        return Number.isNaN(parsed) ? 0 : parsed;
      };
      const asAccount = (value) => {
        const parsed = Number.parseInt(String(value || "").replace(/\D/g, ""), 10);
        return Number.isNaN(parsed) ? Number.MAX_SAFE_INTEGER : parsed;
      };
      const sortBy = $("sortBy").value;
      rows.sort((a, b) => {
        if (sortBy === "joined_desc") return asTime(effectiveTimeValue(b)) - asTime(effectiveTimeValue(a));
        if (sortBy === "joined_asc") return asTime(a["加入时间"]) - asTime(b["加入时间"]);
        if (sortBy === "updated_desc") return asTime(b["修改时间"]) - asTime(a["修改时间"]);
        if (sortBy === "updated_asc") return asTime(a["修改时间"]) - asTime(b["修改时间"]);
        return asAccount(a["账号"]) - asAccount(b["账号"]);
      });
      return rows;
    }
    function updatePagination(total) {
      const pages = Math.max(1, Math.ceil(total / state.pageSize));
      if (state.page > pages) state.page = pages;
      const start = total ? (state.page - 1) * state.pageSize + 1 : 0;
      const end = Math.min(total, state.page * state.pageSize);
      $("pageInfo").textContent = `第 ${state.page} / ${pages} 页 · ${start}-${end} / ${total}`;
      $("prevPageBtn").disabled = state.page <= 1;
      $("nextPageBtn").disabled = state.page >= pages;
    }
    function renderRows() {
      const tbody = $("rows");
      const rows = filteredRecords();
      state.lastFilteredCount = rows.length;
      updatePagination(rows.length);
      const pageRows = rows.slice((state.page - 1) * state.pageSize, state.page * state.pageSize);
      tbody.innerHTML = pageRows.map((record) => {
        const tags = splitTags(record["风险标签"]).map((tag) => `<span class="tag">${escapeText(tag)}</span>`).join("");
        const account = record["账号"] || record["关联账号/主体"] || "未填";
        const time = visibleTime(record);
        const historyCount = state.historyCounts[record["记录ID"]] || 0;
        const markers = [
          historyCount ? `<span class="mini-badge">历史 ${historyCount}</span>` : "",
          record["账号"] ? `<span class="mini-badge">账号</span>` : `<span class="mini-badge">主体</span>`,
        ].filter(Boolean).join("");
        const selected = record["记录ID"] === state.editingId ? ' class="selected"' : "";
        return `<tr${selected}>
          <td class="col-account">
            <div class="account-main">${escapeText(account)}</div>
            <div class="account-sub">${escapeText(record["记录类型"])}</div>
            <div class="row-markers">${markers}</div>
          </td>
          <td class="col-action"><span class="pill ${actionClass(record["建议动作"])}">${escapeText(record["建议动作"] || "待定")}</span></td>
          <td class="col-tags"><div class="tags">${tags || '<span style="color:var(--muted)">-</span>'}</div></td>
          <td><div class="note">${escapeText(record["风险/问题备注"])}</div></td>
          <td class="col-time"><div class="list-time"><b>${escapeText(time.label)}</b><br>${escapeText(time.value)}</div></td>
          <td class="col-status"><span class="${statusClass(record["状态"])}">${escapeText(record["状态"] || "-")}</span></td>
          <td class="col-tools"><div class="row-actions"><button class="light" data-history="${escapeText(record["记录ID"])}">历史</button><button class="light" data-edit="${escapeText(record["记录ID"])}">编辑</button></div></td>
        </tr>`;
      }).join("");
      $("empty").hidden = rows.length > 0;
      tbody.querySelectorAll("[data-edit]").forEach((btn) => {
        btn.addEventListener("click", () => editRecord(btn.dataset.edit));
      });
      tbody.querySelectorAll("[data-history]").forEach((btn) => {
        btn.addEventListener("click", () => showHistory(btn.dataset.history));
      });
    }
    function chartItemHtml(chart, activeName = "") {
      const linked = chart.inRegistry ? `已入台账 · ${escapeText(chart.status || "-")} · ${escapeText(chart.action || "-")}` : "未入台账";
      const active = chart.name === activeName ? " active" : "";
      const aiLine = chart.hasAiAnalysis ? `<br><b>AI ${escapeText(chart.aiRiskLevel || "-")}</b> · ${escapeText(chart.aiConclusion || "")}` : "";
      const aiButton = chart.hasAiAnalysis && chart.recordId
        ? `<button class="light" data-confirm-ai="${escapeText(chart.recordId)}" data-ai-stem="${escapeText(chart.stem || "")}">确认AI备注</button>`
        : "";
      return `<div class="chart-item${active}" data-chart-name="${escapeText(chart.name)}">
        <div class="chart-name">${escapeText(chart.name)}</div>
        <div class="chart-meta">账号 ${escapeText(chart.account)} · ${linked}<br>${escapeText(chart.mtime)} · ${escapeText(chart.sizeText || "")}${aiLine}</div>
        <div class="chart-actions">
          <button class="light" data-preview-chart="${escapeText(chart.name)}">预览</button>
          <a class="button-link light" href="${escapeText(chart.url)}" target="_blank" rel="noopener">打开 AI 图表</a>
          <button class="light" data-copy-chart="${escapeText(chart.url)}">复制链接</button>
          ${aiButton}
        </div>
      </div>`;
    }
    function showGeneratedChart(chart) {
      if (!chart || !chart.url) return;
      const normalized = {
        ...chart,
        account: chart.account || $("dbAccount").value.trim(),
        mtime: chart.mtime || "刚刚生成",
        sizeText: chart.sizeText || "",
        inRegistry: Boolean(chart.inRegistry),
      };
      $("chartHint").textContent = "刚生成";
      $("accountCharts").innerHTML = chartItemHtml(normalized, normalized.name);
      bindChartButtons($("accountCharts"), [normalized]);
      $("chartPreview").src = normalized.url;
      $("chartPreview").hidden = false;
      $("generatedActions").classList.add("active");
      $("generatedOpenLink").href = normalized.url;
      $("generatedPreviewBtn").dataset.previewUrl = normalized.url;
      $("generatedCopyBtn").dataset.copyUrl = normalized.url;
    }
    function fillDbSelect(select, items, firstLabel) {
      const current = select.value;
      select.innerHTML = "";
      select.append(new Option(firstLabel, ""));
      (items || []).forEach((item) => {
        if (typeof item === "string") select.append(new Option(item, item));
        else select.append(new Option(item.label || item.value || "未指定", item.value || ""));
      });
      if (Array.from(select.options).some((opt) => opt.value === current)) select.value = current;
    }
    function parseDateTime(text) {
      const value = String(text || "").trim();
      if (!value) return null;
      const parsed = Date.parse(value.replace(/-/g, "/"));
      return Number.isNaN(parsed) ? null : new Date(parsed);
    }
    function formatDateTime(date) {
      const pad = (n) => String(n).padStart(2, "0");
      return `${date.getFullYear()}-${pad(date.getMonth() + 1)}-${pad(date.getDate())} ${pad(date.getHours())}:${pad(date.getMinutes())}:${pad(date.getSeconds())}`;
    }
    function activeDbRange() {
      const active = $("dbRangeButtons").querySelector("button.active");
      return active ? active.dataset.rangeDays : "30";
    }
    function applyDbRange(days) {
      $("dbRangeButtons").querySelectorAll("button").forEach((btn) => btn.classList.toggle("active", btn.dataset.rangeDays === String(days)));
      if (days === "all") {
        $("dbStart").value = "";
        $("dbEnd").value = "";
        return;
      }
      const last = parseDateTime(state.tradeSummary?.lastTime) || new Date();
      const start = new Date(last.getTime() - Number(days || 30) * 24 * 60 * 60 * 1000);
      $("dbStart").value = formatDateTime(start);
      $("dbEnd").value = formatDateTime(last);
    }
    function dbTradeFilters() {
      return {
        account: $("dbAccount").value.trim(),
        platform: $("dbPlatform").value,
        server: $("dbServer").value,
        symbol: $("dbSymbol").value,
        start: $("dbStart").value.trim(),
        end: $("dbEnd").value.trim(),
      };
    }
    function queryStringFromObject(obj) {
      const params = new URLSearchParams();
      Object.entries(obj).forEach(([key, value]) => {
        if (value) params.set(key, value);
      });
      const text = params.toString();
      return text ? `?${text}` : "";
    }
    function renderTradeSummary(summary) {
      state.tradeSummary = summary;
      $("dbTradeStatus").className = "db-trade-status";
      if (!summary || summary.error) {
        $("dbTradeStatus").classList.add("error");
        $("dbTradeStatus").innerHTML = `<span class="muted">${escapeText(summary?.error || "暂无数据库信息")}</span>`;
        $("dbGenerateBtn").disabled = true;
        return;
      }
      if (!summary.exists) {
        $("dbTradeStatus").classList.add("empty");
        $("dbTradeStatus").innerHTML = `<b>${escapeText(summary.account || "")}</b><br><span class="muted">账户暂未做单</span>`;
        $("dbGenerateBtn").disabled = true;
        return;
      }
      $("dbTradeStatus").classList.add("ready");
      const source = summary.latestSource?.label ? `<br>最新来源：${escapeText(summary.latestSource.label)}` : "";
      $("dbTradeStatus").innerHTML = `
        <b>${escapeText(summary.account || $("dbAccount").value.trim())}</b>
        <div class="db-summary-grid">
          <div class="db-summary-item"><span>可画图订单</span><strong>${escapeText(summary.filteredChartableRows)} / ${escapeText(summary.chartableRows)}</strong></div>
          <div class="db-summary-item"><span>当前筛选范围</span><strong>${escapeText(summary.filteredFirstTime || "-")} → ${escapeText(summary.filteredLastTime || "-")}</strong></div>
          <div class="db-summary-item"><span>完整范围</span><strong>${escapeText(summary.firstTime || "-")} → ${escapeText(summary.lastTime || "-")}</strong></div>
          <div class="db-summary-item"><span>最新来源</span><strong>${escapeText(summary.latestSource?.label || "-")}</strong></div>
        </div>
        ${source ? "" : ""}
      `;
      $("dbGenerateBtn").disabled = !(summary.filteredChartableRows > 0 && $("dbAccount").value.trim());
    }
    async function refreshDbTradeSummary() {
      if (!$("dbAccount").value.trim()) {
        renderTradeSummary({ error: "请输入账号后查询数据库订单" });
        return;
      }
      $("dbTradeStatus").className = "db-trade-status loading";
      $("dbTradeStatus").innerHTML = '<span>正在查询数据库订单...</span>';
      try {
        const data = await requestJson(`/api/trades/summary${queryStringFromObject(dbTradeFilters())}`);
        renderTradeSummary(data.summary);
      } catch (err) {
        renderTradeSummary({ error: err.message });
      }
    }
    async function loadTradeSummary(recordId) {
      $("dbJobBox").classList.remove("active");
      $("dbJobLog").textContent = "";
      $("dbJobProgress").style.width = "0%";
      $("generatedActions").classList.remove("active");
      $("dbTradeStatus").className = "db-trade-status loading";
      $("dbTradeStatus").innerHTML = '<span>正在查询数据库订单...</span>';
      $("dbGenerateBtn").disabled = true;
      const record = state.records.find((item) => item["记录ID"] === recordId);
      $("dbAccount").value = record ? (record["账号"] || "") : "";
      if (!$("dbAccount").value.trim()) {
        renderTradeSummary({ error: "当前记录没有账号，也可以手动输入账号查询" });
        return;
      }
      try {
        const data = await requestJson(`/api/trades/summary${queryStringFromObject({ account: $("dbAccount").value.trim() })}`);
        const summary = data.summary || {};
        state.tradeSummary = summary;
        fillDbSelect($("dbPlatform"), summary.platforms || [], "全部平台");
        fillDbSelect($("dbServer"), summary.servers || [], "全部服务器");
        fillDbSelect($("dbSymbol"), summary.symbols || [], "全部品种");
        if (summary.latestSource?.platform) $("dbPlatform").value = summary.latestSource.platform;
        if (summary.latestSource?.server) $("dbServer").value = summary.latestSource.server;
        applyDbRange(activeDbRange());
        await refreshDbTradeSummary();
      } catch (err) {
        renderTradeSummary({ error: err.message });
      }
    }
    async function generateDbKline() {
      if (!$("dbAccount").value.trim()) {
        renderTradeSummary({ error: "请输入账号后生成K线图" });
        return;
      }
      $("dbGenerateBtn").disabled = true;
      $("generatedActions").classList.remove("active");
      $("dbJobBox").classList.add("active");
      $("dbJobText").textContent = "正在提交生成任务...";
      $("dbJobLog").textContent = "";
      $("dbJobProgress").style.width = "3%";
      try {
        const data = await requestJson(`/api/kline/generate-from-db`, {
          method: "POST",
          headers: { "Content-Type": "application/json; charset=utf-8" },
          body: JSON.stringify(dbTradeFilters()),
        });
        pollDbKlineJob(data.job.id);
      } catch (err) {
        $("dbJobText").textContent = err.message || String(err);
        $("dbGenerateBtn").disabled = false;
      }
    }
    async function pollDbKlineJob(jobId) {
      if (state.klineJobTimer) clearTimeout(state.klineJobTimer);
      try {
        const data = await requestJson(`/api/kline/jobs/${encodeURIComponent(jobId)}`);
        const job = data.job || {};
        const percent = Math.max(0, Math.min(100, Number(job.percent || 0)));
        $("dbJobBox").classList.add("active");
        $("dbJobText").textContent = `${job.status || "-"} · ${job.message || ""}${job.elapsedSeconds ? ` · ${job.elapsedSeconds}s` : ""}`;
        $("dbJobProgress").style.width = `${percent}%`;
        $("dbJobLog").textContent = job.logs || "";
        if (job.status === "done") {
          $("dbGenerateBtn").disabled = false;
          if (job.chart) showGeneratedChart(job.chart);
          if (state.editingId) await loadAccountCharts(state.editingId);
          return;
        }
        if (job.status === "failed" || job.status === "missing") {
          $("dbGenerateBtn").disabled = false;
          return;
        }
        state.klineJobTimer = setTimeout(() => pollDbKlineJob(jobId), 1000);
      } catch (err) {
        $("dbJobText").textContent = err.message || String(err);
        $("dbGenerateBtn").disabled = false;
      }
    }
    async function loadAccountCharts(recordId) {
      $("accountCharts").innerHTML = '<div class="muted">加载中...</div>';
      $("chartPreview").hidden = true;
      $("chartPreview").removeAttribute("src");
      try {
        const data = await requestJson(`/api/accounts/${encodeURIComponent(recordId)}/charts`);
        state.accountCharts = data.charts || [];
        $("chartHint").textContent = state.accountCharts.length ? `${state.accountCharts.length} 个图表` : "暂无图表";
        if (!state.accountCharts.length) {
          $("accountCharts").innerHTML = '<div class="muted">暂无匹配图表，可点击顶部“生成图表”上传 statement。</div>';
          return;
        }
        $("accountCharts").innerHTML = state.accountCharts.map((chart) => chartItemHtml(chart)).join("");
        bindChartButtons($("accountCharts"), state.accountCharts);
      } catch (err) {
        $("accountCharts").innerHTML = `<div class="muted">${escapeText(err.message)}</div>`;
      }
    }
    function bindChartButtons(container, charts) {
      container.querySelectorAll("[data-preview-chart]").forEach((btn) => {
        btn.addEventListener("click", () => {
          const chart = charts.find((item) => item.name === btn.dataset.previewChart);
          if (!chart) return;
          $("chartPreview").src = chart.url;
          $("chartPreview").hidden = false;
          container.querySelectorAll(".chart-item").forEach((item) => item.classList.toggle("active", item.dataset.chartName === chart.name));
        });
      });
      container.querySelectorAll("[data-confirm-ai]").forEach((btn) => {
        btn.addEventListener("click", async () => {
          const recordId = btn.dataset.confirmAi;
          const stem = btn.dataset.aiStem;
          if (!recordId || !stem) return;
          btn.disabled = true;
          try {
            await requestJson(`/api/accounts/${encodeURIComponent(recordId)}/ai-note/confirm`, {
              method: "POST",
              body: JSON.stringify({ stem }),
            });
            await loadData();
            if (state.editingId === recordId) await loadAccountCharts(recordId);
          } catch (err) {
            alert(err.message || String(err));
          } finally {
            btn.disabled = false;
          }
        });
      });
      container.querySelectorAll("[data-copy-chart]").forEach((btn) => {
        btn.addEventListener("click", async () => {
          const url = btn.dataset.copyChart || "";
          try {
            await navigator.clipboard.writeText(url);
            btn.textContent = "已复制";
            setTimeout(() => btn.textContent = "复制链接", 1200);
          } catch {
            prompt("复制图表链接", url);
          }
        });
      });
    }
    function clearForm() {
      state.editingId = null;
      $("formTitle").textContent = "新增记录";
      $("detailSubtitle").textContent = "未选择记录，可以直接新增或输入账号查询数据库订单。";
      Object.values(fields).forEach((id) => $(id).value = "");
      $("type").value = "账户";
      $("status").value = "待复核";
      $("customAction").value = "";
      $("joinedAt").value = "";
      $("updatedAtField").value = "";
      syncCustomAction();
      $("deleteBtn").hidden = true;
      $("historyBtn").hidden = true;
      $("chartHint").textContent = "选择账号后显示";
      $("accountCharts").innerHTML = "";
      $("chartPreview").hidden = true;
      $("chartPreview").removeAttribute("src");
      state.tradeSummary = null;
      if (state.klineJobTimer) clearTimeout(state.klineJobTimer);
      $("dbTradeStatus").textContent = "选择账号后查询数据库订单";
      $("dbTradeStatus").className = "db-trade-status";
      $("dbGenerateBtn").disabled = true;
      $("generatedActions").classList.remove("active");
      $("generatedOpenLink").removeAttribute("href");
      $("generatedPreviewBtn").removeAttribute("data-preview-url");
      $("generatedCopyBtn").removeAttribute("data-copy-url");
      $("dbJobBox").classList.remove("active");
      $("dbJobLog").textContent = "";
      $("dbJobProgress").style.width = "0%";
      fillDbSelect($("dbPlatform"), [], "全部平台");
      fillDbSelect($("dbServer"), [], "全部服务器");
      fillDbSelect($("dbSymbol"), [], "全部品种");
      $("dbAccount").value = "";
      $("dbStart").value = "";
      $("dbEnd").value = "";
      renderRows();
      setStatus("");
    }
    function syncCustomAction() {
      const show = $("action").value === "自定义";
      $("customAction").style.display = show ? "block" : "none";
    }
    function editRecord(id) {
      const record = state.records.find((item) => item["记录ID"] === id);
      if (!record) return;
      state.editingId = id;
      $("formTitle").textContent = record["账号"] ? `编辑 ${record["账号"]}` : "编辑记录";
      $("detailSubtitle").textContent = [
        record["记录类型"] || "记录",
        record["建议动作"] ? `建议 ${record["建议动作"]}` : "",
        record["状态"] || "",
        record["修改时间"] ? `修改 ${record["修改时间"]}` : `加入 ${record["加入时间"] || "-"}`
      ].filter(Boolean).join(" · ");
      for (const [name, idName] of Object.entries(fields)) {
        if (name === "建议动作") ensureSelectOption($(idName), record[name]);
        $(idName).value = record[name] || "";
      }
      $("joinedAt").value = record["加入时间"] || "";
      $("updatedAtField").value = record["修改时间"] || "";
      $("customAction").value = "";
      syncCustomAction();
      $("deleteBtn").hidden = false;
      $("historyBtn").hidden = false;
      renderRows();
      loadAccountCharts(id);
      loadTradeSummary(id);
      setStatus("");
    }
    function formPayload() {
      const payload = {};
      for (const [name, idName] of Object.entries(fields)) payload[name] = $(idName).value.trim();
      if ($("action").value === "自定义") payload["建议动作"] = $("customAction").value.trim() || "自定义";
      return payload;
    }
    async function requestJson(url, options) {
      const res = await fetch(url, options);
      const data = await res.json();
      if (!res.ok || !data.ok) throw new Error(data.error || `HTTP ${res.status}`);
      return data;
    }
    async function loadData() {
      const data = await requestJson("/api/accounts");
      state.records = data.records || [];
      state.meta = data.summary || {};
      state.historyCounts = data.historyCounts || {};
      renderSummary(state.meta);
      renderRows();
    }
    async function loadMeta() {
      const data = await requestJson("/api/meta");
      fillSelect($("actionFilter"), data.actions || [], "全部建议");
      fillSelect($("statusFilter"), data.statuses || [], "全部状态");
      fillSelect($("action"), data.actions || []);
      fillSelect($("status"), data.statuses || []);
      fillSelect($("type"), data.types || []);
    }
    async function saveCurrent() {
      try {
        const payload = formPayload();
        const url = state.editingId ? `/api/accounts/${encodeURIComponent(state.editingId)}` : "/api/accounts";
        const method = state.editingId ? "PUT" : "POST";
        await requestJson(url, {
          method,
          headers: { "Content-Type": "application/json; charset=utf-8" },
          body: JSON.stringify(payload),
        });
        await loadData();
        clearForm();
        setStatus("已保存");
      } catch (err) {
        setStatus(err.message, false);
      }
    }
    async function deleteCurrent() {
      if (!state.editingId) return;
      const record = state.records.find((item) => item["记录ID"] === state.editingId);
      const label = record ? (record["账号"] || record["关联账号/主体"] || record["记录ID"]) : state.editingId;
      if (!confirm(`删除 ${label}？`)) return;
      try {
        await requestJson(`/api/accounts/${encodeURIComponent(state.editingId)}`, { method: "DELETE" });
        await loadData();
        clearForm();
        setStatus("已删除");
      } catch (err) {
        setStatus(err.message, false);
      }
    }
    function renderChartLibrary() {
      const q = $("chartSearch").value.trim().toLowerCase();
      const filter = $("chartLinkedFilter").value;
      const charts = state.allCharts.filter((chart) => {
        const hay = `${chart.account} ${chart.name}`.toLowerCase();
        return (!q || hay.includes(q)) &&
          (!filter || (filter === "linked" ? chart.inRegistry : !chart.inRegistry));
      });
      $("chartLibraryTitle").textContent = `图表库 (${charts.length}/${state.allCharts.length})`;
      $("chartLibraryList").innerHTML = charts.length
        ? charts.map((chart) => chartItemHtml(chart)).join("")
        : '<div class="empty">没有匹配图表</div>';
      bindChartButtons($("chartLibraryList"), charts);
    }
    async function showChartLibrary() {
      $("chartLibraryList").innerHTML = '<div class="muted">加载中...</div>';
      $("chartLibraryDialog").showModal();
      try {
        const data = await requestJson("/api/charts");
        state.allCharts = data.charts || [];
        renderChartLibrary();
      } catch (err) {
        $("chartLibraryList").innerHTML = `<div class="muted">${escapeText(err.message)}</div>`;
      }
    }
    function parseHistoryJson(text) {
      if (!text) return null;
      try {
        return JSON.parse(text);
      } catch {
        return null;
      }
    }
    function valueOrDash(value) {
      const text = String(value || "").trim();
      return text || "-";
    }
    function versionPanel(title, record, changedFields, tone) {
      if (!record) {
        return `<div class="version-panel ${tone}"><h4>${escapeText(title)}</h4><div class="version-grid"><div class="version-label">状态</div><div class="version-value">无</div></div></div>`;
      }
      const fields = [
        ["账号", record["账号"] || record["关联账号/主体"] || record["记录ID"]],
        ["建议", record["建议动作"]],
        ["状态", record["状态"]],
        ["分组", record["当前分组"]],
        ["标签", record["风险标签"]],
        ["加入时间", record["加入时间"]],
        ["修改时间", record["修改时间"]],
        ["备注", record["风险/问题备注"]],
      ];
      return `<div class="version-panel ${tone}">
        <h4>${escapeText(title)}</h4>
        <div class="version-grid">
          ${fields.map(([label, value]) => {
            const changed = changedFields.has(label) || changedFields.has(label === "建议" ? "建议动作" : label === "分组" ? "当前分组" : label === "标签" ? "风险标签" : label === "备注" ? "风险/问题备注" : label);
            return `<div class="version-label">${escapeText(label)}</div><div class="version-value ${changed ? "changed" : ""}">${escapeText(valueOrDash(value))}</div>`;
          }).join("")}
        </div>
      </div>`;
    }
    function historyItemHtml(item, index) {
      const before = parseHistoryJson(item["修改前JSON"]);
      const after = parseHistoryJson(item["修改后JSON"]);
      const changedFields = new Set(String(item["修改字段"] || "").split(/[；;,，]+/).map((x) => x.trim()).filter(Boolean));
      const operation = item["操作"] || "修改";
      const badgeClass = operation === "删除" ? "version-badge delete" : (operation === "新增" || operation === "加入" ? "version-badge add" : "version-badge");
      const account = item["账号"] || before?.["账号"] || after?.["账号"] || before?.["关联账号/主体"] || after?.["关联账号/主体"] || item["记录ID"];
      return `<div class="history-item">
        <div class="version-head">
          <div>
            <div class="version-title">
              <span class="${badgeClass}">${escapeText(operation)}</span>
              <b>${escapeText(account || `版本 ${index + 1}`)}</b>
            </div>
            <div class="time-text">修改字段：${escapeText(item["修改字段"] || "-")}</div>
          </div>
          <div class="version-time">操作时间<br>${escapeText(item["修改时间"] || "-")}</div>
        </div>
        <div class="version-body">
          ${versionPanel(operation === "删除" ? "删除前版本" : "修改前版本", before, changedFields, "before")}
          ${versionPanel(operation === "删除" ? "删除后版本" : "修改后版本", after, changedFields, "after")}
        </div>
      </div>`;
    }
    async function showHistory(recordId) {
      try {
        const record = state.records.find((item) => item["记录ID"] === recordId);
        const data = await requestJson(`/api/accounts/${encodeURIComponent(recordId)}/history`);
        $("historyTitle").textContent = `修改历史 ${record ? (record["账号"] || record["关联账号/主体"] || "") : ""}`;
        const rows = (data.history || []).slice().sort((a, b) => {
          const at = Date.parse(String(a["修改时间"] || "").replace(/-/g, "/")) || 0;
          const bt = Date.parse(String(b["修改时间"] || "").replace(/-/g, "/")) || 0;
          return bt - at;
        });
        if (!rows.length) {
          $("historyBody").innerHTML = '<div class="empty">暂无修改历史</div>';
        } else {
          $("historyBody").innerHTML = rows.map((item, index) => historyItemHtml(item, index)).join("");
        }
        $("historyDialog").showModal();
      } catch (err) {
        setStatus(err.message, false);
      }
    }
    function resetPageAndRender() {
      state.page = 1;
      renderRows();
    }
    $("search").addEventListener("input", resetPageAndRender);
    $("actionFilter").addEventListener("change", resetPageAndRender);
    $("statusFilter").addEventListener("change", resetPageAndRender);
    $("sortBy").addEventListener("change", resetPageAndRender);
    $("prevPageBtn").addEventListener("click", () => {
      if (state.page > 1) {
        state.page -= 1;
        renderRows();
      }
    });
    $("nextPageBtn").addEventListener("click", () => {
      const pages = Math.max(1, Math.ceil(state.lastFilteredCount / state.pageSize));
      if (state.page < pages) {
        state.page += 1;
        renderRows();
      }
    });
    $("action").addEventListener("change", syncCustomAction);
    $("newBtn").addEventListener("click", clearForm);
    $("reloadBtn").addEventListener("click", loadData);
    $("clearBtn").addEventListener("click", clearForm);
    $("saveBtn").addEventListener("click", saveCurrent);
    $("deleteBtn").addEventListener("click", deleteCurrent);
    $("historyBtn").addEventListener("click", () => state.editingId && showHistory(state.editingId));
    $("closeHistoryBtn").addEventListener("click", () => $("historyDialog").close());
    $("chartLibraryBtn").addEventListener("click", showChartLibrary);
    $("closeChartLibraryBtn").addEventListener("click", () => $("chartLibraryDialog").close());
    $("chartSearch").addEventListener("input", renderChartLibrary);
    $("chartLinkedFilter").addEventListener("change", renderChartLibrary);
    $("dbRefreshBtn").addEventListener("click", refreshDbTradeSummary);
    $("dbGenerateBtn").addEventListener("click", generateDbKline);
    $("generatedPreviewBtn").addEventListener("click", () => {
      const url = $("generatedPreviewBtn").dataset.previewUrl;
      if (!url) return;
      $("chartPreview").src = url;
      $("chartPreview").hidden = false;
    });
    $("generatedCopyBtn").addEventListener("click", async () => {
      const url = $("generatedCopyBtn").dataset.copyUrl || "";
      if (!url) return;
      try {
        await navigator.clipboard.writeText(url);
        $("generatedCopyBtn").textContent = "已复制";
        setTimeout(() => $("generatedCopyBtn").textContent = "复制链接", 1200);
      } catch {
        prompt("复制图表链接", url);
      }
    });
    $("dbAccount").addEventListener("change", refreshDbTradeSummary);
    $("dbAccount").addEventListener("keydown", (ev) => {
      if (ev.key === "Enter") refreshDbTradeSummary();
    });
    ["dbPlatform", "dbServer", "dbSymbol"].forEach((id) => $(id).addEventListener("change", refreshDbTradeSummary));
    ["dbStart", "dbEnd"].forEach((id) => $(id).addEventListener("change", refreshDbTradeSummary));
    $("dbRangeButtons").querySelectorAll("button").forEach((btn) => {
      btn.addEventListener("click", () => {
        applyDbRange(btn.dataset.rangeDays);
        refreshDbTradeSummary();
      });
    });
    (async function init() {
      try {
        await loadMeta();
        await loadData();
        clearForm();
      } catch (err) {
        setStatus(err.message, false);
      }
    })();
  </script>
</body>
</html>
"""


WORKBENCH_HTML = r"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>账号风控台账</title>
  <style>
    :root { --ink:#182126; --muted:#667278; --line:#dce2e4; --soft:#f4f6f6; --paper:#fff; --accent:#087f78; --accent-dark:#075e59; --warn:#a15c00; --danger:#b42318; --good:#177245; }
    * { box-sizing:border-box; }
    body { margin:0; color:var(--ink); background:#eef1f1; font-family:"Microsoft YaHei","Segoe UI",sans-serif; letter-spacing:0; }
    button,input,select { font:inherit; letter-spacing:0; }
    button,a { -webkit-tap-highlight-color:transparent; }
    button,a,select,.result { transition:border-color .15s,background-color .15s,color .15s,box-shadow .15s; }
    .topbar { height:58px; background:#172226; color:#fff; display:flex; align-items:center; justify-content:space-between; padding:0 28px; }
    .brand { display:flex; align-items:baseline; gap:12px; min-width:0; }
    .brand strong { font-size:18px; white-space:nowrap; }
    .brand span { color:#afbdc1; font-size:12px; white-space:nowrap; }
    .top-actions { display:flex; gap:8px; }
    .top-actions a { color:#e9f0f1; text-decoration:none; border:1px solid #46565b; padding:7px 11px; border-radius:5px; font-size:13px; }
    .top-actions a:hover,.top-actions a:focus-visible { border-color:#9ec7c3; background:#26383d; color:#fff; }
    main { width:min(1420px,calc(100% - 36px)); margin:20px auto 44px; }
    .lookup-band { background:var(--paper); border:1px solid var(--line); border-top:3px solid var(--accent); padding:22px 24px; }
    .lookup-title { display:flex; justify-content:space-between; gap:20px; align-items:end; margin-bottom:12px; }
    h1 { font-size:22px; margin:0; }
    .muted { color:var(--muted); font-size:13px; }
    .lookup-row { display:grid; grid-template-columns:minmax(220px,540px) 110px; gap:8px; }
    input,select { width:100%; min-height:38px; border:1px solid #bbc5c8; border-radius:5px; background:#fff; padding:8px 10px; color:var(--ink); }
    input:focus,select:focus { outline:2px solid #a8d8d4; border-color:var(--accent); }
    button { border:1px solid #aebbbf; border-radius:5px; min-height:38px; padding:8px 14px; cursor:pointer; background:#fff; color:var(--ink); }
    button:hover:not(:disabled),button:focus-visible { border-color:var(--accent); background:#edf8f7; box-shadow:0 0 0 2px #cce7e4; }
    button.primary { color:#fff; background:var(--accent); }
    button.primary:hover { background:var(--accent-dark); }
    button:disabled { cursor:not-allowed; opacity:.55; }
    .lookup-status { min-height:20px; margin-top:10px; color:var(--muted); font-size:13px; }
    .result { margin-top:14px; border:1px solid var(--line); border-left:4px solid var(--accent); display:grid; grid-template-columns:minmax(160px,1fr) minmax(340px,2.4fr) auto; align-items:center; gap:18px; padding:16px 18px; background:#fbfcfc; cursor:pointer; }
    .result:hover { border-color:#8cbab6; background:#f5fbfa; }
    .result-account { font-size:24px; font-weight:700; }
    .badges { display:flex; flex-wrap:wrap; gap:6px; margin-top:7px; }
    .badge { border-radius:999px; padding:3px 8px; background:#e8eeee; color:#465459; font-size:12px; }
    .badge.marked { color:#075e59; background:#d9efec; }
    .badge.empty { color:#8b4d00; background:#fff1d6; }
    .result-facts { display:grid; grid-template-columns:repeat(3,minmax(120px,1fr)); gap:12px; }
    .fact b { display:block; margin-top:3px; font-size:14px; overflow-wrap:anywhere; }
    .fact span { color:var(--muted); font-size:11px; }
    .fact b.positive { color:#177245; }
    .fact b.negative { color:#b42318; }
    .enter { white-space:nowrap; color:var(--accent-dark); font-weight:700; }
    .summary { display:grid; grid-template-columns:repeat(4,1fr); border:1px solid var(--line); border-top:0; background:#fff; }
    .summary-item { padding:14px 18px; border-right:1px solid var(--line); }
    .summary-item:last-child { border-right:0; }
    .summary-item b { display:block; font-size:21px; }
    .summary-item span { font-size:12px; color:var(--muted); }
    .push-discovery-panel { margin-top:18px; background:#fff; border:1px solid var(--line); }
    .push-discovery-head { display:flex; justify-content:space-between; align-items:center; gap:16px; padding:14px 18px; border-bottom:1px solid var(--line); }
    .push-discovery-head h2 { margin:0; font-size:17px; }
    .push-discovery-form { display:grid; grid-template-columns:180px 220px 150px 1fr; gap:9px; align-items:end; padding:14px 18px; border-bottom:1px solid var(--line); }
    .push-discovery-form label { display:grid; gap:6px; color:var(--muted); font-size:12px; }
    .push-discovery-status { min-height:20px; padding:10px 18px; color:var(--muted); font-size:13px; }
    .push-discovery-table { min-width:1180px; }
    .push-discovery-note { padding:0 18px 12px; color:var(--muted); font-size:12px; }
    .ledger { margin-top:18px; background:#fff; border:1px solid var(--line); }
    .ledger-head { padding:16px 18px; border-bottom:1px solid var(--line); display:grid; grid-template-columns:1fr minmax(260px,520px) 150px 150px; align-items:center; gap:8px; }
    .ledger-head h2 { font-size:17px; margin:0; }
    .table-wrap { overflow:auto; }
    table { width:100%; border-collapse:collapse; min-width:940px; }
    th { background:#f3f5f5; color:#526066; font-size:12px; text-align:left; font-weight:600; padding:10px 12px; border-bottom:1px solid var(--line); }
    td { padding:11px 12px; border-bottom:1px solid #e8ecec; font-size:13px; vertical-align:top; }
    tbody tr[data-login] { cursor:pointer; }
    tbody tr[data-login] td:first-child { border-left:3px solid transparent; }
    tbody tr[data-login]:hover { background:#edf8f7; }
    tbody tr[data-login]:hover td:first-child { border-left-color:var(--accent); }
    .account-link { display:inline-block; color:var(--accent-dark); background:#eef8f7; border:1px solid #9bc9c5; border-radius:4px; padding:4px 7px; text-decoration:none; font-weight:700; font-size:14px; }
    .account-link:hover,.account-link:focus-visible { color:#fff; background:var(--accent); border-color:var(--accent); }
    .inline-status { min-width:112px; min-height:32px; padding:5px 28px 5px 8px; border:1px solid #69aaa5; background:#f1fbfa; color:var(--accent-dark); font-weight:650; cursor:pointer; }
    .inline-status:hover,.inline-status:focus { border-color:var(--accent); background:#e1f4f2; box-shadow:0 0 0 2px #cce7e4; }
    .inline-status.saving { opacity:.6; }
    .note { max-width:420px; white-space:nowrap; overflow:hidden; text-overflow:ellipsis; }
    .action { display:inline-block; min-width:34px; text-align:center; border:1px solid #c8d1d3; border-radius:4px; padding:3px 7px; font-weight:700; background:#fff; }
    .empty-state { padding:40px; text-align:center; color:var(--muted); }
    .pager { display:flex; justify-content:space-between; align-items:center; padding:12px 16px; }
    .pager button { border-color:#c8d1d3; min-height:34px; }
    .pager-controls { display:flex; align-items:center; gap:8px; }
    .hierarchy-panel { margin-top:18px; background:#fff; border:1px solid var(--line); }
    .hierarchy-head { min-height:58px; display:flex; align-items:center; justify-content:space-between; gap:16px; padding:14px 18px; border-bottom:1px solid var(--line); }
    .hierarchy-head h2 { margin:0; font-size:17px; }
    .hierarchy-form { display:grid; grid-template-columns:minmax(220px,1.35fr) 190px 190px minmax(140px,.7fr) 104px; gap:9px; align-items:end; padding:16px 18px; border-bottom:1px solid var(--line); }
    .hierarchy-form label { display:grid; gap:6px; color:var(--muted); font-size:12px; min-width:0; }
    .hierarchy-rule-toggle { grid-column:1/-1; display:flex!important; align-items:flex-start; gap:9px!important; padding:10px 12px; border:1px solid #28527b; border-radius:5px; background:#071a32; cursor:pointer; }
    .hierarchy-rule-toggle input { width:18px; min-height:18px; margin:1px 0 0; flex:0 0 auto; accent-color:var(--accent); }
    .hierarchy-rule-toggle b { display:block; color:var(--ink); font-size:13px; }
    .hierarchy-rule-toggle small { display:block; margin-top:3px; color:var(--muted); line-height:1.45; }
    .hierarchy-status { min-height:20px; padding:10px 18px; color:var(--muted); font-size:13px; }
    .hierarchy-candidates { display:grid; grid-template-columns:minmax(260px,1fr) auto; gap:8px; align-items:end; padding:12px 18px; border-top:1px solid var(--line); border-bottom:1px solid var(--line); background:#0a203b; }
    .hierarchy-candidates label { display:grid; gap:6px; color:var(--muted); font-size:12px; }
    .hierarchy-context { display:flex; flex-wrap:wrap; gap:8px 18px; padding:13px 18px; border-top:1px solid var(--line); border-bottom:1px solid var(--line); font-size:13px; }
    .hierarchy-context b { color:var(--ink); }
    .hierarchy-metrics { display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); border-bottom:1px solid var(--line); }
    .hierarchy-metric { min-height:82px; padding:14px 16px; border-right:1px solid var(--line); border-bottom:1px solid var(--line); }
    .hierarchy-metric:nth-child(4n) { border-right:0; }
    .hierarchy-metric:nth-last-child(-n+4) { border-bottom:0; }
    .hierarchy-metric span { display:block; color:var(--muted); font-size:11px; margin-bottom:5px; }
    .hierarchy-metric b { display:block; font-size:18px; overflow-wrap:anywhere; }
    .promotion-audit { border-bottom:1px solid var(--line); }
    .promotion-audit-head { display:flex; justify-content:space-between; align-items:flex-start; gap:16px; padding:14px 18px; background:#071a32; }
    .promotion-audit-head b { color:var(--ink); }
    .promotion-audit-head span { color:var(--muted); font-size:12px; line-height:1.55; }
    .promotion-result { flex:0 0 auto; padding:4px 8px; border:1px solid #805e20; border-radius:4px; color:#ffd38c!important; background:#31250f; font-weight:700; }
    .promotion-result.qualified { border-color:#287862; color:#79e6c4!important; background:#0b302d; }
    .promotion-table { min-width:980px; }
    .activity-status { display:inline-block; max-width:210px; padding:3px 7px; border:1px solid #385a7c; border-radius:4px; color:#9ab6d3; background:#0a213b; font-size:11px; line-height:1.4; }
    .activity-status.included { color:#79e6c4; border-color:#287862; background:#0b302d; }
    .positive { color:var(--good)!important; }
    .negative { color:var(--danger)!important; }
    .hierarchy-tools { display:grid; grid-template-columns:minmax(220px,1fr) 170px; gap:8px; padding:12px 18px; border-bottom:1px solid var(--line); }
    .hierarchy-table { min-width:1320px; }
    .depth-label { color:var(--muted); font-variant-numeric:tabular-nums; }
    .role-label { display:inline-block; min-width:68px; padding:3px 7px; border:1px solid #2f668f; border-radius:4px; color:#74c8ff; background:#0a2947; text-align:center; font-size:12px; }
    .role-label.customer { color:#9bd7c6; border-color:#327663; background:#0b302d; }
    .account-type { display:flex; align-items:center; gap:6px; min-width:150px; }
    .cent-flag { display:inline-block; padding:2px 5px; border:1px solid #9b6a20; border-radius:3px; color:#ffd38c; background:#3a2a0e; font-size:10px; }
    .hierarchy-panel [hidden] { display:none!important; }
    /* Risk control dashboard theme */
    :root { --ink:#dbeafe; --muted:#7895b8; --line:#173e6b; --soft:#0a1d38; --paper:#081a32; --accent:#168cff; --accent-dark:#0d6fd8; --warn:#f2ab35; --danger:#ff5f6d; --good:#34d399; }
    html { color-scheme:dark; }
    body { min-height:100vh; background:radial-gradient(circle at 65% -10%,#0b3970 0,transparent 34%),linear-gradient(145deg,#020b18 0%,#06152b 52%,#020a16 100%); color:var(--ink); }
    body::before { content:""; position:fixed; inset:0; pointer-events:none; opacity:.22; background-image:linear-gradient(#148cff12 1px,transparent 1px),linear-gradient(90deg,#148cff12 1px,transparent 1px); background-size:38px 38px; mask-image:linear-gradient(to bottom,#000,transparent 88%); }
    .topbar { height:64px; position:relative; z-index:1; background:#030d1eeb; border-bottom:1px solid #176fbd; box-shadow:0 4px 24px #006ed933; }
    .brand::before { content:""; width:30px; height:30px; border-radius:50%; border:4px solid #168cff; border-right-color:#22d3ee; box-shadow:0 0 15px #168cff99; }
    .brand strong { color:#f2f8ff; text-shadow:0 0 12px #168cff77; }
    .brand span { color:#668bb7; letter-spacing:1px; }
    .top-actions a { color:#b9d8fa; border-color:#245586; background:#081b36; }
    .top-actions a:hover,.top-actions a:focus-visible { border-color:#168cff; background:#0b3767; box-shadow:0 0 12px #168cff55; }
    main { position:relative; z-index:1; width:min(1540px,calc(100% - 32px)); }
    .lookup-band,.ledger,.hierarchy-panel { background:linear-gradient(145deg,#081b35ee,#06162cee); border-color:#194a7c; box-shadow:0 14px 34px #0008,inset 0 1px 0 #2a73ad33; border-radius:10px; overflow:hidden; }
    .lookup-band { border-top:1px solid #1f7fd0; padding:20px 22px; }
    .lookup-title h1,.ledger-head h2,.hierarchy-head h2 { color:#e9f5ff; text-shadow:0 0 14px #168cff55; }
    .muted,.lookup-status { color:var(--muted); }
    input,select { color:#dcecff; background:#06162b; border-color:#28527b; }
    input::placeholder { color:#537294; }
    input:focus,select:focus { outline:2px solid #168cff55; border-color:#168cff; box-shadow:0 0 16px #168cff33; }
    option { background:#07172c; color:#dbeafe; }
    button { color:#cce4ff; background:#0a2342; border-color:#2b5c8c; }
    button:hover:not(:disabled),button:focus-visible { color:#fff; border-color:#168cff; background:#0d3764; box-shadow:0 0 14px #168cff55; }
    button.primary { background:linear-gradient(135deg,#0875df,#159bff); border-color:#29a6ff; box-shadow:0 0 16px #168cff44; }
    button.primary:hover { background:linear-gradient(135deg,#168cff,#22b1ff); }
    .result { color:#dcecff; background:linear-gradient(135deg,#0a2443,#07192f); border-color:#225688; border-left-color:#20a4ff; box-shadow:inset 0 1px 0 #5ebcff22; }
    .result:hover,.result:focus-visible { background:#0b2d55; border-color:#2da9ff; box-shadow:0 0 20px #168cff33; outline:none; }
    .badge { color:#9bb7d5; background:#102942; border:1px solid #285276; }
    .badge.marked { color:#7ff3d0; background:#0c3a36; border-color:#1f7e68; }
    .badge.empty { color:#ffd38c; background:#3a2a0e; border-color:#8d6520; }
    .fact span { color:#6f91b6; } .enter { color:#42b5ff; }
    .fact b.positive { color:#34d399; text-shadow:0 0 10px #34d39933; }
    .fact b.negative { color:#ff6472; text-shadow:0 0 10px #ff647233; }
    .summary { margin-top:14px; grid-template-columns:repeat(4,1fr); gap:12px; border:0; background:transparent; }
    .summary-item { position:relative; overflow:hidden; border:1px solid #174674!important; border-radius:9px; background:linear-gradient(145deg,#0a213f,#07172d); box-shadow:0 10px 24px #0005,inset 0 1px 0 #3c8dcc22; }
    .summary-item::after { content:""; position:absolute; right:-18px; top:-24px; width:80px; height:80px; border-radius:50%; background:#168cff12; border:1px solid #168cff22; }
    .summary-item b { color:#f0f8ff; text-shadow:0 0 12px #168cff55; }
    .summary-item span { color:#7091b5; }
    .ledger,.push-discovery-panel { border-radius:10px; background:linear-gradient(145deg,#081b35ee,#06162cee); border-color:#194a7c; box-shadow:0 14px 34px #0008,inset 0 1px 0 #2a73ad33; }
    .push-discovery-head { border-color:#163d65; background:#07182eaa; }
    .push-discovery-head h2 { color:#e9f5ff; text-shadow:0 0 14px #168cff55; }
    .push-discovery-status,.push-discovery-note { color:#7895b8; }
    .ledger-head { border-color:#163d65; background:#07182eaa; }
    .table-wrap { scrollbar-color:#1d5d96 #061426; }
    table { color:#bed4ec; }
    th { color:#7598be; background:#0a2240; border-color:#17446f; text-transform:none; }
    td { border-color:#102f51; }
    tbody tr[data-login]:hover { background:#0b294c; }
    .account-link { color:#55c3ff; background:#0b2a4b; border-color:#236ca7; }
    .account-link:hover,.account-link:focus-visible { color:#fff; background:#147cda; border-color:#3cb7ff; box-shadow:0 0 12px #168cff55; }
    .action { color:#a9d8ff; background:#0b2848; border-color:#285f8e; }
    .inline-status { color:#68c8ff; background:#082643; border-color:#277cb6; }
    .inline-status:hover,.inline-status:focus { background:#0b355d; border-color:#2faeff; box-shadow:0 0 12px #168cff44; }
    .pager { border-top:1px solid #163b62; background:#06172c; }
    .log-panel { margin-top:18px; background:linear-gradient(145deg,#081b35ee,#06162cee); border:1px solid #194a7c; border-radius:10px; overflow:hidden; box-shadow:0 14px 34px #0008,inset 0 1px 0 #2a73ad33; }
    .log-head { display:flex; align-items:center; justify-content:space-between; gap:14px; padding:14px 18px; border-bottom:1px solid #163d65; }
    .log-head h2 { margin:0; font-size:17px; color:#e9f5ff; }
    .log-form { display:grid; grid-template-columns:minmax(180px,1fr) 220px 220px 100px; gap:9px; align-items:end; padding:14px 18px; border-bottom:1px solid #163d65; }
    .log-form label { display:grid; gap:6px; color:var(--muted); font-size:12px; }
    .log-status { min-height:20px; padding:10px 18px; color:var(--muted); font-size:13px; }
    .log-status.error { color:#ff8b96; }
    .log-table-wrap { max-height:460px; overflow:auto; }
    .log-table { min-width:840px; }
    .log-table td { vertical-align:top; }
    .log-message { margin:0; max-width:720px; color:#c9def5; white-space:pre-wrap; overflow-wrap:anywhere; font:12px/1.5 Consolas,"Microsoft YaHei",monospace; }
    @media (max-width:850px) {
      .topbar { padding:0 14px; } .brand span,.top-actions a:first-child { display:none; }
      main { width:calc(100% - 20px); margin-top:10px; }
      .lookup-band { padding:17px 14px; } .lookup-title { align-items:start; flex-direction:column; }
      .lookup-row { grid-template-columns:1fr 88px; }
      .result { grid-template-columns:1fr; gap:12px; } .enter { display:none; }
      .result-facts { grid-template-columns:repeat(2,1fr); }
      .summary { grid-template-columns:repeat(2,1fr); } .summary-item:nth-child(2) { border-right:0; } .summary-item:nth-child(-n+2) { border-bottom:1px solid var(--line); }
      .push-discovery-form { grid-template-columns:1fr 1fr; padding:12px 14px; }
      .push-discovery-form button { grid-column:1/-1; }
      .ledger-head { grid-template-columns:1fr 1fr; } .ledger-head h2 { grid-column:1/-1; }
      .hierarchy-head { align-items:flex-start; flex-direction:column; }
      .hierarchy-form { grid-template-columns:1fr 1fr; padding:14px; }
      .hierarchy-form label:first-child,.hierarchy-form button { grid-column:1/-1; }
      .hierarchy-rule-toggle { grid-column:1/-1!important; }
      .promotion-audit-head { flex-direction:column; }
      .hierarchy-metrics { grid-template-columns:repeat(2,minmax(0,1fr)); }
      .hierarchy-metric:nth-child(4n) { border-right:1px solid var(--line); }
      .hierarchy-metric:nth-child(2n) { border-right:0; }
      .hierarchy-metric:nth-last-child(-n+4) { border-bottom:1px solid var(--line); }
      .hierarchy-metric:nth-last-child(-n+2) { border-bottom:0; }
      .hierarchy-tools { grid-template-columns:1fr; padding:12px 14px; }
      .hierarchy-candidates { grid-template-columns:1fr; padding:12px 14px; }
      .log-form { grid-template-columns:1fr 1fr; padding:12px 14px; }
      .log-form label:first-child,.log-form button { grid-column:1/-1; }
    }
  </style>
</head>
<body>
  <header class="topbar">
    <div class="brand"><strong>账号风控台账</strong><span>账号查询工作台</span></div>
    <nav class="top-actions"><a href="/download/daily-report">导出日报</a><a href="/download/problematic_accounts.xlsx">导出台账</a></nav>
  </header>
  <main>
    <section class="lookup-band">
      <div class="lookup-title"><div><h1>账号查询</h1><div class="muted">数据库订单与台账记录</div></div><div class="muted" id="updatedAt"></div></div>
      <form class="lookup-row" id="lookupForm"><input id="accountLookup" inputmode="numeric" autocomplete="off" placeholder="输入交易账号" aria-label="账号查询" /><button class="primary" id="lookupBtn">查询</button></form>
      <div class="lookup-status" id="lookupStatus"></div>
      <div id="lookupResult"></div>
    </section>
    <section class="log-panel">
      <div class="log-head"><h2>账号日志查询</h2><span class="muted">只读 · MySQL 交易库</span></div>
      <form class="log-form" id="logForm">
        <label>账号<input id="logAccount" inputmode="numeric" autocomplete="off" placeholder="输入交易账号" aria-label="日志查询账号" /></label>
        <label>开始时间<input id="logStart" type="datetime-local" aria-label="日志开始时间" /></label>
        <label>结束时间<input id="logEnd" type="datetime-local" aria-label="日志结束时间" /></label>
        <button class="primary" id="logQueryBtn" type="submit">查询日志</button>
      </form>
      <div class="log-status" id="logStatus">请输入账号和时间范围</div>
      <div class="log-table-wrap"><table class="log-table"><thead><tr><th>时间</th><th>数据源 / 类型</th><th>订单 / 成交号</th><th>数据库原始记录</th></tr></thead><tbody id="logRows"><tr><td colspan="4"><div class="empty-state">暂无查询结果</div></td></tr></tbody></table></div>
    </section>
    <section class="summary" id="summary"></section>
    <section class="push-discovery-panel">
      <div class="push-discovery-head"><h2>全平台推盘发现</h2><span class="muted">只读 · 近期开仓结构初筛 → Tick/协同深检</span></div>
      <form class="push-discovery-form" id="pushDiscoveryForm">
        <label>盈利窗口（天）<input id="pushDiscoveryDays" type="number" min="1" max="30" step="1" value="7" /></label>
        <label>近期开平仓订单上限<input id="pushDiscoveryMaxOrders" type="number" min="20" max="1000" step="10" value="200" /></label>
        <button class="primary" id="pushDiscoveryBtn" type="submit">开始全平台检测</button>
        <span class="muted">自动排除本地建议动作 T、TA、A、A/TA；订单上限由你每次调整。</span>
      </form>
      <div class="push-discovery-status" id="pushDiscoveryStatus">尚未运行</div>
      <div class="push-discovery-note" id="pushDiscoverySummary"></div>
      <div class="table-wrap" id="pushDiscoveryResults" hidden>
        <table class="push-discovery-table"><thead><tr><th>排名</th><th>账号</th><th>平台 / 服务器</th><th>近期开平仓单</th><th>初筛分</th><th>深检分</th><th>等级</th><th>Tick</th><th>协同开仓</th><th>结论</th></tr></thead><tbody id="pushDiscoveryRows"></tbody></table>
      </div>
    </section>
    <section class="hierarchy-panel" id="hierarchyPanel">
      <div class="hierarchy-head">
        <div><h2>下线净入金统计</h2></div>
        <div class="muted" id="hierarchyUpdatedAt"></div>
      </div>
      <form class="hierarchy-form" id="hierarchyForm">
        <label>IB / 客户<input id="hierarchyTarget" autocomplete="off" placeholder="交易账号、gb:CRM ID / cn:CRM ID 或精确姓名" /></label>
        <label>开始时间<input id="hierarchyStart" type="datetime-local" step="60" /></label>
        <label>结束时间<input id="hierarchyEnd" type="datetime-local" step="60" /></label>
        <label>产品<select id="hierarchyProduct"><option value="">全部产品</option><option value="@PROMOTION">本次活动产品（外汇 + 贵金属）</option></select></label>
        <button class="primary" id="hierarchyBtn" type="submit">统计</button>
        <label class="hierarchy-rule-toggle"><input id="hierarchyActivityRules" type="checkbox" /><span><b>按本次活动归属规则计算</b><small>60,000 USD + 600手；排除Cent；普通客户自行达标、下级IB达标后不再向上归集</small></span></label>
      </form>
      <div class="hierarchy-status" id="hierarchyStatus" role="status"></div>
      <div class="hierarchy-candidates" id="hierarchyCandidates" hidden>
        <label>选择具体IB / 客户<select id="hierarchyCandidateSelect"></select></label>
        <button id="hierarchyCandidateBtn" type="button">使用并查询</button>
      </div>
      <div id="hierarchyResult" hidden>
        <div class="hierarchy-context" id="hierarchyContext"></div>
        <div class="hierarchy-metrics" id="hierarchyMetrics"></div>
        <section class="promotion-audit" id="hierarchyPromotionAudit" hidden>
          <div class="promotion-audit-head"><div><b>活动归属逐级核验</b><br><span id="hierarchyPromotionNotes"></span></div><span class="promotion-result" id="hierarchyPromotionResult"></span></div>
          <div class="table-wrap"><table class="promotion-table"><thead><tr><th>层级</th><th>客户 / IB</th><th>角色</th><th>判断净入金</th><th>判断手数</th><th>判断结果</th><th>向上归集</th></tr></thead><tbody id="hierarchyPromotionRows"></tbody></table></div>
        </section>
        <div class="hierarchy-tools">
          <input id="hierarchyFilter" placeholder="过滤账号、姓名、CRM ID、服务器" aria-label="过滤下线账户" />
          <select id="hierarchyScope" aria-label="下线账户范围">
            <option value="all">全部账户</option>
            <option value="cashflow">有资金流水</option>
            <option value="trades">有产品交易</option>
            <option value="referral">Referral账户</option>
            <option value="customer">普通客户账户</option>
            <option value="cent">Cent账户</option>
            <option value="activity">计入当前活动业绩</option>
            <option value="activityExcluded">未计入当前活动业绩</option>
          </select>
        </div>
        <div class="table-wrap">
          <table class="hierarchy-table">
            <thead><tr><th>层级</th><th>客户 / IB</th><th>角色</th><th>服务器</th><th>交易账号</th><th>账户类型</th><th>活动归属</th><th>入金</th><th>出金</th><th>净入金</th><th>产品订单</th><th>产品手数</th><th>交易盈亏</th></tr></thead>
            <tbody id="hierarchyRows"></tbody>
          </table>
        </div>
        <div class="pager">
          <span class="muted" id="hierarchyPageInfo"></span>
          <div class="pager-controls"><button id="hierarchyPrevBtn" type="button">上一页</button><button id="hierarchyNextBtn" type="button">下一页</button></div>
        </div>
      </div>
    </section>
    <section class="ledger">
      <div class="ledger-head">
        <h2 id="ledgerTitle">已标记账号</h2>
        <input id="listFilter" placeholder="列表过滤：账号、标签、备注、IB" aria-label="列表过滤" />
        <select id="actionFilter" aria-label="建议动作"><option value="">全部动作</option></select>
        <select id="statusFilter" aria-label="状态"><option value="">全部状态</option></select>
      </div>
      <div class="table-wrap"><table><thead><tr><th>账号</th><th>建议</th><th>分组</th><th>风险标签</th><th>备注</th><th>状态</th><th>更新时间</th></tr></thead><tbody id="rows"></tbody></table></div>
      <div class="pager"><span class="muted" id="pageInfo"></span><div class="pager-controls"><button id="prevBtn" type="button">上一页</button><button id="nextBtn" type="button">下一页</button></div></div>
    </section>
  </main>
  <script>
    const $ = (id) => document.getElementById(id);
    const state = { records:[], summary:{}, statuses:[], page:1, pageSize:25, hierarchy:null, hierarchyPage:1, hierarchyPageSize:100, pushDiscoveryTimer:null };
    const esc = (value) => String(value ?? "").replace(/[&<>"']/g, c => ({"&":"&amp;","<":"&lt;",">":"&gt;",'"':"&quot;","'":"&#39;"}[c]));
    const fmt = (value) => new Intl.NumberFormat("zh-CN", { maximumFractionDigits:2 }).format(Number(value || 0));
    const money = (value) => new Intl.NumberFormat("zh-CN", { minimumFractionDigits:2, maximumFractionDigits:2 }).format(Number(value || 0));
    const valueClass = (value) => Number(value)>0?'positive':Number(value)<0?'negative':'';
    async function json(url, options) { const res=await fetch(url,options); const data=await res.json(); if(!res.ok||!data.ok){const error=new Error(data.error||`HTTP ${res.status}`);error.data=data;throw error;} return data; }
    function go(login) { location.href=`/account/${encodeURIComponent(login)}`; }
    function renderSummary() {
      const s=state.summary;
      $("summary").innerHTML=[['台账记录',s.total],['账号记录',s.accountRecords],['今日更新',state.records.filter(r=>String(r['修改时间']||r['加入时间']).startsWith(new Date().toISOString().slice(0,10))).length],['动作类型',Object.keys(s.actions||{}).length]].map(([k,v])=>`<div class="summary-item"><b>${fmt(v)}</b><span>${k}</span></div>`).join('');
      $("updatedAt").textContent=s.updatedAt?`台账更新 ${s.updatedAt}`:'';
    }
    function fillFilters() {
      const actions=[...new Set(state.records.map(r=>r['建议动作']).filter(Boolean))].sort();
      const statuses=[...new Set([...(state.statuses||[]),...state.records.map(r=>r['状态']).filter(Boolean)])];
      $("actionFilter").innerHTML='<option value="">全部动作</option>'+actions.map(v=>`<option>${esc(v)}</option>`).join('');
      $("statusFilter").innerHTML='<option value="">全部状态</option>'+statuses.map(v=>`<option>${esc(v)}</option>`).join('');
    }
    function filtered() {
      const q=$("listFilter").value.trim().toLowerCase(), action=$("actionFilter").value, status=$("statusFilter").value;
      return state.records.filter(r=>{
        const hay=[r['账号'],r['关联账号/主体'],r['风险标签'],r['风险/问题备注'],r['处理人/来源']].join(' ').toLowerCase();
        return (!q||hay.includes(q))&&(!action||r['建议动作']===action)&&(!status||r['状态']===status);
      }).sort((a,b)=>String(b['修改时间']||b['加入时间']).localeCompare(String(a['修改时间']||a['加入时间'])));
    }
    function renderRows() {
      const all=filtered(), pages=Math.max(1,Math.ceil(all.length/state.pageSize)); state.page=Math.min(state.page,pages);
      const shown=all.slice((state.page-1)*state.pageSize,state.page*state.pageSize);
      const statusOptions=current=>(state.statuses||[]).map(value=>`<option value="${esc(value)}" ${value===current?'selected':''}>${esc(value)}</option>`).join('');
      $("ledgerTitle").textContent=`已标记账号 (${all.length})`;
      $("rows").innerHTML=shown.length?shown.map(r=>`<tr data-login="${esc(r['账号'])}"><td><a class="account-link" href="/account/${encodeURIComponent(r['账号'])}">${esc(r['账号']||r['关联账号/主体']||'-')}</a></td><td><span class="action">${esc(r['建议动作']||'-')}</span></td><td>${esc(r['当前分组']||'-')}</td><td>${esc(r['风险标签']||'-')}</td><td><div class="note" title="${esc(r['风险/问题备注'])}">${esc(r['风险/问题备注']||'-')}</div></td><td><select class="inline-status" data-record-id="${esc(r['记录ID'])}" aria-label="修改 ${esc(r['账号'])} 的状态">${statusOptions(r['状态'])}</select></td><td>${esc(r['修改时间']||r['加入时间']||'-')}</td></tr>`).join(''):'<tr><td colspan="7"><div class="empty-state">没有匹配的台账记录</div></td></tr>';
      $("pageInfo").textContent=`第 ${state.page} / ${pages} 页 · ${all.length} 条`;
      $("prevBtn").disabled=state.page<=1; $("nextBtn").disabled=state.page>=pages;
      document.querySelectorAll('tr[data-login]').forEach(row=>row.addEventListener('click',event=>{if(!event.target.closest('a,button,select'))go(row.dataset.login);}));
      document.querySelectorAll('.inline-status').forEach(select=>{select.addEventListener('click',event=>event.stopPropagation());select.addEventListener('change',()=>updateInlineStatus(select));});
    }
    async function updateInlineStatus(select){
      const record=state.records.find(item=>item['记录ID']===select.dataset.recordId);if(!record)return;
      const previous=record['状态']||'';select.disabled=true;select.classList.add('saving');
      try{const data=await json(`/api/accounts/${encodeURIComponent(record['记录ID'])}`,{method:'PUT',headers:{'Content-Type':'application/json; charset=utf-8'},body:JSON.stringify({'状态':select.value})});record['状态']=data.record?.['状态']||select.value;record['修改时间']=data.record?.['修改时间']||record['修改时间'];state.summary=data.summary||state.summary;renderSummary();renderRows();}
      catch(err){record['状态']=previous;select.value=previous;select.title=err.message;}
      finally{select.disabled=false;select.classList.remove('saving');}
    }
    function lookupCard(data,db,index) {
      db=db||{}; const source=db.latestSource||{}, meta=db.accountMeta||{};
      const sourceText=[source.platform,source.server].filter(Boolean).join(' / ')||'-';
      const symbols=(db.symbols||[]).slice(0,5).join('、')||'-';
      const currencyText=meta.isCentAccount?'USC 美分账户 · 金额按 USD':(meta.currency?`${meta.currency} 账户`:'币种未识别');
      const href=`/account/${encodeURIComponent(data.account)}?platform=${encodeURIComponent(source.platform||'')}&server=${encodeURIComponent(source.server||'')}`;
      return `<article class="result" id="accountResult-${index}" data-href="${esc(href)}" tabindex="0"><div><div class="result-account">${esc(data.account)}</div><div class="badges"><span class="badge ${data.marked?'marked':''}">${data.marked?'已标记':'未标记'}</span><span class="badge ${db.exists?'':'empty'}">${db.exists?'数据库有订单':'账户暂未做单'}</span><span class="badge">${esc(currencyText)}</span>${data.record?.['建议动作']?`<span class="action">${esc(data.record['建议动作'])}</span>`:''}</div><div class="muted" style="margin-top:8px">刷新 ${esc(db.refreshedAt||'-')}</div></div><div class="result-facts"><div class="fact"><span>订单 / 可画图</span><b>${fmt(db.orderCount)} / ${fmt(db.chartableOrderCount)}</b></div><div class="fact"><span>平台 / 服务器</span><b>${esc(sourceText)}</b></div><div class="fact"><span>数据库状态 / 本地标记</span><b id="lookupStatus-${index}">加载中...</b></div><div class="fact"><span>综合盈利</span><b id="lookupProfit-${index}">加载中...</b></div><div class="fact"><span>交易时间</span><b>${esc(db.firstTime||'-')}<br>${esc(db.lastTime||'-')}</b></div><div class="fact"><span>品种</span><b>${esc(symbols)}</b></div></div><div class="enter">进入详情 →</div></article>`;
    }
    async function loadLookupFinance(data,matches){
      const results=await Promise.allSettled(matches.map(async(db,index)=>{const source=db.latestSource||{},q=new URLSearchParams({account:data.account,platform:source.platform||'',server:source.server||''}),finance=await json(`/api/account-lookup-finance?${q}`),statusEl=$(`lookupStatus-${index}`),profitEl=$(`lookupProfit-${index}`);if(statusEl){statusEl.textContent=`${finance.databaseStatus||'-'} / ${finance.localStatus||'-'}`;statusEl.title=`数据库状态 ${finance.databaseStatus||'-'}；本地标记 ${finance.localStatus||'-'}${finance.workflowStatus?`；流程状态 ${finance.workflowStatus}`:''}`;}if(profitEl){const value=Number(finance.comprehensiveProfit||0);profitEl.textContent=`${value>0?'+':''}${new Intl.NumberFormat('zh-CN',{minimumFractionDigits:2,maximumFractionDigits:2}).format(value)}${finance.currency?` ${finance.currency}`:''}`;profitEl.className=value>0?'positive':value<0?'negative':'';}}));
      results.forEach((result,index)=>{if(result.status==='rejected'){const statusEl=$(`lookupStatus-${index}`),profitEl=$(`lookupProfit-${index}`);if(statusEl)statusEl.textContent='暂不可用';if(profitEl)profitEl.textContent='暂不可用';}});
      return results;
    }
    async function lookup(event) {
      event?.preventDefault(); const account=$("accountLookup").value.trim(); if(!account){$("lookupStatus").textContent='请输入账号';return;}
      $("lookupBtn").disabled=true; $("lookupStatus").textContent='正在查询订单数据库...'; $("lookupResult").innerHTML='';
      try { const data=await json(`/api/account-lookup?account=${encodeURIComponent(account)}`), matches=data.databases||[]; $("lookupStatus").textContent=matches.length?`查询完成 · 找到 ${matches.length} 个平台/服务器账户 · 正在补充状态与综合盈利`:'未找到该账号'; $("lookupResult").innerHTML=matches.length?matches.map((db,index)=>lookupCard(data,db,index)).join(''):lookupCard(data,data.database||{},0); document.querySelectorAll('#lookupResult .result').forEach(card=>{const open=()=>location.href=card.dataset.href||`/account/${encodeURIComponent(data.account)}`;card.addEventListener('click',open);card.addEventListener('keydown',e=>{if(e.key==='Enter')open();});});if(matches.length){loadLookupFinance(data,matches).then(()=>{if($("accountLookup").value.trim()===account)$("lookupStatus").textContent=`查询完成 · 找到 ${matches.length} 个平台/服务器账户 · 状态与综合盈利已更新`;});} }
      catch(err){$("lookupStatus").textContent=err.message;} finally{$("lookupBtn").disabled=false;}
    }
    function localDateTimeValue(value){const shifted=new Date(value.getTime()-value.getTimezoneOffset()*60000);return shifted.toISOString().slice(0,16);}
    function initLogDates(){const end=new Date(),start=new Date(end.getTime()-24*60*60*1000);$("logStart").value=localDateTimeValue(start);$("logEnd").value=localDateTimeValue(end);}
    function logMessageText(value){if(typeof value==='string')return value;try{return JSON.stringify(value,null,2)||'';}catch{return String(value??'');}}
    function renderAccountLogs(data){
      const rows=data.rows||[];
      $("logStatus").className='log-status';
      const warnings=(data.warnings||[]).length?` · ${data.warnings.length} 个数据源不可用`:'';
      $("logStatus").textContent=`查询完成 · ${data.account} · ${data.start} 至 ${data.end} · 找到 ${data.matchedCount||0} 条${data.truncated?' · 结果已截断':''}${warnings}`;
      $("logRows").innerHTML=rows.length?rows.map(row=>`<tr><td>${esc(row.eventTime||'-')}</td><td>${esc(row.source||'-')}<br>${esc(row.platform||'')} ${esc(row.eventType||'')}</td><td>${esc(row.ticket||'-')}</td><td><pre class="log-message">${esc(logMessageText(row.details))}</pre></td></tr>`).join(''):'<tr><td colspan="4"><div class="empty-state">该账号在此时间范围内没有数据库记录</div></td></tr>';
    }
    async function queryAccountLogs(event){
      event?.preventDefault();
      const account=$("logAccount").value.trim(),start=$("logStart").value,end=$("logEnd").value;
      if(!account||!start||!end){$("logStatus").className='log-status error';$("logStatus").textContent='请输入账号、开始时间和结束时间';return;}
      $("logQueryBtn").disabled=true;$("logStatus").className='log-status';$("logStatus").textContent='正在查询日志...';$("logRows").innerHTML='<tr><td colspan="4"><div class="empty-state">读取中...</div></td></tr>';
      try{renderAccountLogs(await json(`/api/account-logs?${new URLSearchParams({account,start,end})}`));}
      catch(err){$("logStatus").className='log-status error';$("logStatus").textContent=err.message;$("logRows").innerHTML='<tr><td colspan="4"><div class="empty-state">数据库查询失败</div></td></tr>';}
      finally{$("logQueryBtn").disabled=false;}
    }
    function initHierarchyDates(){const end=new Date(),start=new Date(end);start.setDate(start.getDate()-7);$("hierarchyStart").value=localDateTimeValue(start);$("hierarchyEnd").value=localDateTimeValue(end);}
    async function loadHierarchyProducts(){
      const select=$("hierarchyProduct"),current=select.value;
      try{const data=await json('/api/hierarchy-products'),options=value=>`<option value="${esc(value)}">${esc(value)}</option>`,eligible=data.promotionProducts||[],other=data.otherProducts||[];select.innerHTML=`<option value="">全部产品</option><option value="${esc(data.promotionValue||'@PROMOTION')}">${esc(data.promotionLabel||'本次活动产品（外汇 + 贵金属）')}</option>${eligible.length?`<optgroup label="活动有效单品">${eligible.map(options).join('')}</optgroup>`:''}${other.length?`<optgroup label="其他产品">${other.map(options).join('')}</optgroup>`:''}`;if([...select.options].some(option=>option.value===current))select.value=current;}
      catch(err){select.title=err.message;}
    }
    function renderHierarchyCandidates(candidates){
      const rows=candidates||[],box=$("hierarchyCandidates"),select=$("hierarchyCandidateSelect");
      if(!rows.length){box.hidden=true;select.innerHTML='';return;}
      select.innerHTML=rows.map(row=>`<option value="${esc(row.target)}">${esc(row.name||'-')} · ${esc(row.roleLabel||'-')} · ${esc(row.environment||'-')} · CRM ${esc(row.userId)}</option>`).join('');box.hidden=false;
    }
    function hierarchyMetric(label,value,className=''){return `<div class="hierarchy-metric"><span>${esc(label)}</span><b class="${esc(className)}">${esc(value)}</b></div>`;}
    function hierarchyFilteredAccounts(){
      const data=state.hierarchy;if(!data)return[];
      const q=$("hierarchyFilter").value.trim().toLowerCase(),scope=$("hierarchyScope").value;
      return (data.accounts||[]).filter(row=>{
        const hay=[row.account,row.userId,row.name,row.server,row.platform,row.typeName,row.roleLabel].join(' ').toLowerCase();
         const scopeMatch=scope==='all'||(scope==='cashflow'&&(row.depositCount||row.withdrawalCount))||(scope==='trades'&&row.orders)||(scope==='referral'&&row.role==='referral')||(scope==='customer'&&row.role==='customer')||(scope==='cent'&&row.isCent)||(scope==='activity'&&row.activityIncluded)||(scope==='activityExcluded'&&row.activityIncluded===false);
        return scopeMatch&&(!q||hay.includes(q));
      });
    }
    function renderHierarchyRows(){
      const all=hierarchyFilteredAccounts(),pages=Math.max(1,Math.ceil(all.length/state.hierarchyPageSize));state.hierarchyPage=Math.min(state.hierarchyPage,pages);
      const shown=all.slice((state.hierarchyPage-1)*state.hierarchyPageSize,state.hierarchyPage*state.hierarchyPageSize);
       $("hierarchyRows").innerHTML=shown.length?shown.map(row=>{const href=`/account/${encodeURIComponent(row.account)}?platform=${encodeURIComponent(row.platform||'')}&server=${encodeURIComponent(row.server||'')}`,typeText=row.typeName||'-',activityText=row.activityReason||'-',activityClass=row.activityIncluded?'included':'';return `<tr><td><span class="depth-label">L${fmt(row.depth)}</span></td><td><b>${esc(row.name||'-')}</b><br><span class="muted">CRM ${esc(row.userId)}</span></td><td><span class="role-label ${row.role==='customer'?'customer':''}">${esc(row.roleLabel)}</span></td><td>${esc([row.platform,row.server].filter(Boolean).join(' · ')||'-')}</td><td><a class="account-link" href="${esc(href)}">${esc(row.account)}</a></td><td><span class="account-type">${esc(typeText)}${row.isCent?'<span class="cent-flag">CENT</span>':''}</span></td><td><span class="activity-status ${activityClass}" title="${esc(activityText)}">${esc(row.activityIncluded===true?'计入':row.activityIncluded===false?'排除':'-')}</span></td><td>${money(row.deposit)}</td><td>${money(row.withdrawal)}</td><td class="${valueClass(row.netDeposit)}"><b>${money(row.netDeposit)}</b></td><td>${fmt(row.orders)}</td><td>${fmt(row.lots)}</td><td class="${valueClass(row.tradingProfit)}">${money(row.tradingProfit)}</td></tr>`;}).join(''):'<tr><td colspan="13"><div class="empty-state">没有匹配的账户</div></td></tr>';
      $("hierarchyPageInfo").textContent=`第 ${state.hierarchyPage} / ${pages} 页 · ${all.length} 个账户`;
      $("hierarchyPrevBtn").disabled=state.hierarchyPage<=1;$("hierarchyNextBtn").disabled=state.hierarchyPage>=pages;
    }
     function renderPromotionRules(data){
       const rules=data.promotionRules,box=$("hierarchyPromotionAudit");
       if(!rules){box.hidden=true;return;}
       box.hidden=false;
       const qualified=Boolean(rules.subjectQualified),result=$("hierarchyPromotionResult");
       result.textContent=qualified?'当前目标达标':'当前目标未达标';result.className=`promotion-result ${qualified?'qualified':''}`;
       $("hierarchyPromotionNotes").textContent=(rules.notes||[]).join('；');
       $("hierarchyPromotionRows").innerHTML=(rules.decisions||[]).map(row=>{
         const roll=row.rollsToUserId?`CRM ${esc(row.rollsToUserId)}`:'-';
         return `<tr><td>L${fmt(row.depth)}</td><td><b>${esc(row.name||'-')}</b><br><span class="muted">CRM ${esc(row.userId)}</span></td><td><span class="role-label ${row.role==='customer'?'customer':''}">${esc(row.roleLabel||'-')}</span></td><td>${money(row.netDeposit)} USD</td><td>${fmt(row.lots)} 手</td><td>${esc(row.statusLabel||'-')}</td><td>${roll}</td></tr>`;
       }).join('')||'<tr><td colspan="7"><div class="empty-state">没有可核验的下级记录</div></td></tr>';
     }
     function renderHierarchy(data){
       state.hierarchy=data;state.hierarchyPage=1;const s=data.summary||{},subject=data.subject||{},query=data.query||{},product=query.productLabel||query.product||'全部产品',activity=Boolean(query.activityRules);
       renderHierarchyCandidates([]);
       $("hierarchyResult").hidden=false;$("hierarchyUpdatedAt").textContent=data.refreshedAt?`数据刷新 ${data.refreshedAt}`:'';
       $("hierarchyContext").innerHTML=`<span><b>${esc(subject.name||'-')}</b></span><span>${esc(subject.roleLabel||'-')} · CRM ${esc(subject.userId||'-')}</span><span>${esc(subject.environment||subject.schema||'-')}</span><span>${esc(query.start||'-')} 至 ${esc(query.end||'-')}</span><span>产品：${esc(product)}</span><span>${activity?'活动归属：标准账户客户业绩 · Cent排除':'资金口径：全部产品 · USD'}</span>`;
       $("hierarchyMetrics").innerHTML=[
         hierarchyMetric(activity?'活动归属净入金':'净入金',`${money(s.netDeposit)} USD`,valueClass(s.netDeposit)),
         hierarchyMetric('入金 / 出金',`${money(s.deposit)} / ${money(s.withdrawal)} USD`),
         hierarchyMetric('标准 / Cent净入金',`${money(s.standardNetDeposit)} / ${money(s.centNetDeposit)} USD`),
         hierarchyMetric(activity?'计入客户 / 账户':'用户 / 账户',`${fmt(s.users)} / ${fmt(s.accounts)}`),
         hierarchyMetric('有资金流水账户',fmt(s.accountsWithCashflow)),
         hierarchyMetric(`${product}订单数`,fmt(s.orders)),
         hierarchyMetric(`${product}总手数`,`${fmt(s.lots)} 手`),
         hierarchyMetric(`${product}交易盈亏`,`${money(s.tradingProfit)} USD`,valueClass(s.tradingProfit)),
       ].join('');
       $("hierarchyStatus").textContent=activity?`活动归属查询完成 · ${data.promotionRules?.subjectQualified?'当前目标达标':'当前目标未达标'} · ${fmt(s.users)} 个计入客户 · ${fmt(s.accounts)} 个账户`:`查询完成 · ${fmt(s.users)} 个用户 · ${fmt(s.accounts)} 个账户`;
       renderPromotionRules(data);
       renderHierarchyRows();
     }
     async function queryHierarchy(event){
       event?.preventDefault();const target=$("hierarchyTarget").value.trim(),start=$("hierarchyStart").value,end=$("hierarchyEnd").value,product=$("hierarchyProduct").value.trim(),activityRules=$("hierarchyActivityRules").checked;
       if(!target){$("hierarchyStatus").textContent='请输入IB或客户';return;}if(!start||!end){$("hierarchyStatus").textContent='请选择开始和结束时间';return;}
       $("hierarchyBtn").disabled=true;$("hierarchyStatus").textContent='正在读取CRM层级、账户和资金流水...';$("hierarchyResult").hidden=true;renderHierarchyCandidates([]);
       try{const q=new URLSearchParams({target,start,end});if(product)q.set('product',product);if(activityRules)q.set('activityRules','1');renderHierarchy(await json(`/api/hierarchy-net-deposit?${q}`));}
      catch(err){$("hierarchyStatus").textContent=err.message;renderHierarchyCandidates(err.data?.candidates||[]);}
      finally{$("hierarchyBtn").disabled=false;}
    }
    function renderPushDiscovery(job){
      const summary=job.summary||{},rows=job.results||[];
      $("pushDiscoveryStatus").textContent=`${job.message||job.status||'-'} · ${fmt(job.percent||0)}%${job.elapsedSeconds?` · ${fmt(job.elapsedSeconds)}s`:''}`;
      $("pushDiscoverySummary").textContent=summary.screenedAccounts!==undefined?`候选 ${fmt(summary.aggregateCandidates)} · 排除已标记 ${fmt(summary.excludedKnownAccounts)} · 初筛 ${fmt(summary.screenedAccounts)} · 深检队列 ${fmt(summary.deepEligible)} · 本轮完成 ${fmt(summary.deepCompleted)}${summary.deepPending?` · 仍待深检 ${fmt(summary.deepPending)}`:''}`:'';
      $("pushDiscoveryResults").hidden=!rows.length;
      $("pushDiscoveryRows").innerHTML=rows.map(row=>{const href=`/account/${encodeURIComponent(row.account)}?platform=${encodeURIComponent(row.platform||'')}&server=${encodeURIComponent(row.server||'')}`;return `<tr><td>${fmt(row.deepRank)}</td><td><a class="account-link" href="${esc(href)}">${esc(row.account)}</a></td><td>${esc([row.platform,row.server].filter(Boolean).join(' / '))}</td><td>${fmt(row.orders)}</td><td>${fmt(row.initialScore)}</td><td><b>${fmt(row.deepScore)}</b></td><td>${esc(row.level||'-')}</td><td>${row.tickAvailable?'有':'无'}</td><td>${fmt(row.coordinatedMatchedRatio)}%</td><td><div class="note" title="${esc(row.headline||'')}">${esc(row.headline||'-')}</div></td></tr>`;}).join('');
    }
    async function pollPushDiscovery(id){
      clearTimeout(state.pushDiscoveryTimer);
      try{const data=await json(`/api/push-discovery/jobs/${encodeURIComponent(id)}`),job=data.job||{};renderPushDiscovery(job);if(job.status==='done'||job.status==='failed'||job.status==='missing'){$("pushDiscoveryBtn").disabled=false;return;}state.pushDiscoveryTimer=setTimeout(()=>pollPushDiscovery(id),1200);}
      catch(err){$("pushDiscoveryStatus").textContent=err.message;$("pushDiscoveryBtn").disabled=false;}
    }
    async function startPushDiscovery(event){
      event?.preventDefault();const days=Number($("pushDiscoveryDays").value||7),maxOrders=Number($("pushDiscoveryMaxOrders").value||200);
      if(days<1||days>30){$("pushDiscoveryStatus").textContent='盈利窗口必须在1到30天之间';return;}
      if(maxOrders<20||maxOrders>1000){$("pushDiscoveryStatus").textContent='订单上限必须在20到1000之间';return;}
      $("pushDiscoveryBtn").disabled=true;$("pushDiscoveryStatus").textContent='正在提交全平台扫描任务...';$("pushDiscoverySummary").textContent='';$("pushDiscoveryResults").hidden=true;
      try{const data=await json('/api/push-discovery/start',{method:'POST',body:JSON.stringify({days,maxOrders})});renderPushDiscovery(data.job||{});pollPushDiscovery(data.job.id);}
      catch(err){$("pushDiscoveryStatus").textContent=err.message;$("pushDiscoveryBtn").disabled=false;}
    }
    async function load() { const data=await json('/api/accounts'); state.records=(data.records||[]).filter(r=>r['账号']); state.summary=data.summary||{}; state.statuses=data.statuses||[]; renderSummary(); fillFilters(); renderRows(); }
    $("lookupForm").addEventListener('submit',lookup);
     $("logForm").addEventListener('submit',queryAccountLogs);
     $("hierarchyForm").addEventListener('submit',queryHierarchy);
     $("pushDiscoveryForm").addEventListener('submit',startPushDiscovery);
     $("hierarchyActivityRules").addEventListener('change',()=>{$("hierarchyProduct").value=$("hierarchyActivityRules").checked?'@PROMOTION':$("hierarchyProduct").value;$("hierarchyProduct").disabled=$("hierarchyActivityRules").checked;});
    $("hierarchyCandidateBtn").addEventListener('click',()=>{const target=$("hierarchyCandidateSelect").value;if(!target)return;$("hierarchyTarget").value=target;queryHierarchy();});
    $("hierarchyFilter").addEventListener('input',()=>{state.hierarchyPage=1;renderHierarchyRows();});
    $("hierarchyScope").addEventListener('change',()=>{state.hierarchyPage=1;renderHierarchyRows();});
    $("hierarchyPrevBtn").addEventListener('click',()=>{state.hierarchyPage--;renderHierarchyRows();});
    $("hierarchyNextBtn").addEventListener('click',()=>{state.hierarchyPage++;renderHierarchyRows();});
    ["listFilter","actionFilter","statusFilter"].forEach(id=>$(id).addEventListener(id==='listFilter'?'input':'change',()=>{state.page=1;renderRows();}));
    $("prevBtn").addEventListener('click',()=>{state.page--;renderRows();}); $("nextBtn").addEventListener('click',()=>{state.page++;renderRows();});
    initHierarchyDates();
    initLogDates();
    loadHierarchyProducts();
    load().catch(err=>{$("rows").innerHTML=`<tr><td colspan="7"><div class="empty-state">${esc(err.message)}</div></td></tr>`});
  </script>
</body>
</html>"""


ACCOUNT_DETAIL_HTML = r"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>账号详情</title>
  <style>
    :root { --ink:#182126; --muted:#68747a; --line:#dce2e4; --soft:#f3f6f6; --paper:#fff; --accent:#087f78; --accent-dark:#075e59; --warn:#a15c00; --danger:#b42318; --good:#177245; }
    * { box-sizing:border-box; }
    body { margin:0; background:#eef1f1; color:var(--ink); font-family:"Microsoft YaHei","Segoe UI",sans-serif; letter-spacing:0; }
    button,input,select,textarea { font:inherit; letter-spacing:0; }
    button { cursor:pointer; }
    a { color:inherit; }
    button,a,select,summary { transition:border-color .15s,background-color .15s,color .15s,box-shadow .15s; }
    .topbar { height:58px; background:#172226; color:#fff; display:flex; align-items:center; justify-content:space-between; padding:0 28px; }
    .back { color:#dfe7e9; text-decoration:none; font-size:14px; border:1px solid #536268; border-radius:4px; padding:6px 9px; }
    .back:hover,.back:focus-visible { color:#fff; border-color:#9ec7c3; background:#26383a; }
    .brand { font-weight:700; }
    main { width:min(1420px,calc(100% - 36px)); margin:20px auto 50px; }
    .account-head { background:#fff; border:1px solid var(--line); border-top:3px solid var(--accent); padding:20px 24px; display:flex; align-items:center; justify-content:space-between; gap:20px; }
    .account-id { font-size:30px; font-weight:750; line-height:1.2; }
    .badges { display:flex; flex-wrap:wrap; gap:7px; margin-top:9px; }
    .badge { border-radius:999px; padding:4px 9px; background:#e8eeee; color:#465459; font-size:12px; }
    .badge.marked { color:var(--accent-dark); background:#d9efec; }
    .badge.empty { color:#8b4d00; background:#fff1d6; }
    .badge.ea { color:#075c78; background:#dff6ff; border:1px solid #66b9d8; font-weight:700; }
    .badge.action { min-width:34px; text-align:center; border-radius:4px; padding:3px 7px; font-weight:700; }
    .head-meta { text-align:right; color:var(--muted); font-size:12px; line-height:1.8; }
    .toolbar { border:1px solid var(--line); border-top:0; background:#f8fafa; padding:12px 16px; display:grid; grid-template-columns:1fr 1fr 1fr 120px 150px 130px; gap:8px; align-items:end; }
    label { display:block; color:#536067; font-size:12px; }
    input,select,textarea { width:100%; border:1px solid #bcc6c9; border-radius:5px; background:#fff; padding:8px 10px; color:var(--ink); }
    input,select { min-height:38px; }
    textarea { min-height:104px; resize:vertical; line-height:1.55; }
    input:focus,select:focus,textarea:focus { outline:2px solid #a8d8d4; border-color:var(--accent); }
    button { border:1px solid #c5ced1; border-radius:5px; min-height:38px; background:#fff; padding:8px 13px; color:var(--ink); }
    button.primary { color:#fff; background:var(--accent); border-color:var(--accent); }
    button.primary:hover { background:var(--accent-dark); }
    button:hover:not(:disabled),button:focus-visible { border-color:var(--accent); box-shadow:0 0 0 2px #cce7e4; }
    button:disabled { opacity:.55; cursor:not-allowed; }
    .section { min-width:0; margin-top:18px; background:#fff; border:1px solid var(--line); }
    .section-head { min-height:52px; padding:12px 18px; border-bottom:1px solid var(--line); display:flex; align-items:center; justify-content:space-between; gap:14px; }
    .section-head h2 { margin:0; font-size:17px; }
    .section-sub { color:var(--muted); font-size:12px; text-align:right; }
    .metric-groups { display:grid; grid-template-columns:repeat(2,1fr); }
    .metric-group { padding:17px 18px 18px; border-right:1px solid var(--line); border-bottom:1px solid var(--line); }
    .metric-group:nth-child(2n) { border-right:0; }
    .metric-group:nth-last-child(-n+2) { border-bottom:0; }
    .metric-group h3 { margin:0 0 13px; font-size:13px; color:#47545a; }
    .metrics { display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); gap:14px; }
    .metric span { display:block; color:var(--muted); font-size:11px; margin-bottom:4px; }
    .metric b { font-size:17px; overflow-wrap:anywhere; }
    .metric b.positive { color:var(--good); } .metric b.negative { color:var(--danger); }
    .split { display:grid; grid-template-columns:minmax(0,1.25fr) minmax(380px,.75fr); gap:18px; }
    .split > * { min-width:0; }
    .table-wrap { overflow:auto; }
    table { width:100%; border-collapse:collapse; min-width:560px; }
    th { text-align:left; color:#526066; background:#f3f5f5; font-size:12px; padding:9px 12px; border-bottom:1px solid var(--line); }
    td { padding:10px 12px; border-bottom:1px solid #e8ecec; font-size:13px; }
    .source-list { padding:10px 18px; }
    .source-row { display:grid; grid-template-columns:1fr 1.7fr 80px 110px; gap:10px; padding:9px 0; border-bottom:1px solid #e8ecec; font-size:13px; }
    .source-row:last-child { border-bottom:0; }
    .risk-summary { display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); }
    .risk-stat { min-height:82px; padding:16px 18px; border-right:1px solid var(--line); border-bottom:1px solid var(--line); }
    .risk-stat:nth-child(4n) { border-right:0; }
    .risk-stat:nth-last-child(-n+4) { border-bottom:0; }
    .risk-stat span { display:block; color:var(--muted); font-size:12px; margin-bottom:7px; }
    .risk-stat b { font-size:18px; overflow-wrap:anywhere; }
    .risk-note { color:var(--muted); font-size:12px; padding:10px 18px; border-top:1px solid var(--line); }
    .risk-table { min-width:1000px; }
    .risk-table th,.risk-table td { text-align:right; white-space:nowrap; }
    .risk-table th:first-child,.risk-table td:first-child,.risk-table th:nth-child(2),.risk-table td:nth-child(2) { text-align:left; }
    .risk-table tr.total td { background:#f5f8f8; font-weight:700; }
    .risk-table tr.current-account td { background:#e3f4f2; border-top:1px solid #72b5af; border-bottom:1px solid #72b5af; font-weight:650; }
    .risk-table tr.current-account td:first-child { border-left:4px solid var(--accent); }
    .account-link { display:inline-block; color:#075e59; background:#eef8f7; border:1px solid #9bc9c5; border-radius:4px; padding:3px 6px; text-decoration:none; font-weight:700; }
    .account-link:hover,.account-link:focus-visible { color:#fff; background:var(--accent); border-color:var(--accent); }
    .order-details > summary { min-height:52px; padding:12px 18px; display:flex; align-items:center; justify-content:space-between; gap:14px; cursor:pointer; font-size:17px; font-weight:700; background:#f8fbfb; border:1px solid transparent; }
    .order-details > summary:hover,.order-details > summary:focus-visible { color:var(--accent-dark); background:#edf8f7; border-color:#86bbb6; }
    .order-details[open] > summary { border-bottom:1px solid var(--line); }
    .order-summary-meta { color:var(--muted); font-size:12px; font-weight:400; text-align:right; }
    .order-table { min-width:1480px; }
    .order-table th,.order-table td { white-space:nowrap; }
    .order-type { font-weight:700; }
    .order-type.buy { color:#177245; } .order-type.sell { color:#b42318; }
    .order-pager { min-height:54px; padding:9px 18px; display:flex; align-items:center; justify-content:flex-end; gap:10px; border-top:1px solid var(--line); }
    .order-pager button { min-height:32px; padding:5px 10px; }
    .inline-kline-frame { display:block; width:100%; height:680px; border:0; background:#061322; }
    .inline-kline-note { padding:9px 18px; color:var(--muted); font-size:12px; border-top:1px solid var(--line); }
    .quick-mark { border:1px solid var(--line); border-top:0; background:#fff; padding:13px 16px 14px; }
    .quick-mark-main { display:grid; grid-template-columns:minmax(0,1fr) 190px 120px; gap:12px; align-items:end; }
    .quick-label { color:#536067; font-size:12px; margin-bottom:6px; }
    .quick-action-head { display:flex; align-items:center; justify-content:space-between; gap:10px; margin-bottom:6px; }
    .quick-action-head .quick-label { margin:0; }
    .manage-actions-btn { min-height:28px; padding:4px 9px; font-size:11px; }
    .action-grid { display:flex; flex-wrap:wrap; gap:7px; margin:0; }
    .action-btn { min-width:52px; min-height:34px; padding:6px 10px; }
    .action-btn:hover:not(:disabled) { background:#e8f6f4; }
    .action-btn.active { border-color:var(--accent); background:#dff2ef; color:var(--accent-dark); font-weight:700; }
    .quick-status select { border:2px solid var(--accent); background:#effaf8; color:var(--accent-dark); font-weight:700; }
    .quick-fields { display:grid; grid-template-columns:150px 1fr 1fr minmax(220px,2fr); gap:10px; align-items:end; margin-top:10px; }
    .quick-fields input { min-height:34px; padding:6px 9px; }
    .batch-toggle { min-height:34px; display:flex; align-items:center; gap:8px; border:1px solid #9bc9c5; border-radius:5px; background:#f1faf9; padding:6px 9px; color:var(--accent-dark); cursor:pointer; }
    .batch-toggle:hover { border-color:var(--accent); background:#e4f5f3; }
    .batch-toggle input { width:auto; min-height:auto; margin:0; accent-color:var(--accent); }
    .batch-toggle.disabled { opacity:.55; cursor:not-allowed; }
    .action-manager { margin-top:10px; padding:10px; border:1px solid #9bc9c5; border-radius:6px; background:#f1faf9; }
    .action-manager[hidden] { display:none; }
    .action-add-row { display:grid; grid-template-columns:minmax(160px,320px) auto 1fr; gap:8px; align-items:center; }
    .action-add-row input { min-height:32px; }
    .action-add-row button { min-height:32px; padding:5px 10px; }
    .action-manage-list { display:flex; flex-wrap:wrap; gap:7px; margin-top:9px; }
    .manage-action-chip { min-height:30px; display:inline-flex; align-items:center; overflow:hidden; border:1px solid #9bc9c5; border-radius:5px; background:#fff; color:#075e59; font-size:12px; }
    .manage-action-chip span { padding:5px 8px; }
    .manage-action-chip button { min-width:28px; min-height:28px; padding:2px 7px; border:0; border-left:1px solid #9bc9c5; border-radius:0; color:#b42318; background:#fff; font-size:16px; line-height:1; }
    .manage-action-chip.protected { color:#68747a; background:#eef2f2; }
    .manage-action-chip.protected::after { content:"保留"; padding-right:7px; color:#7a878c; font-size:10px; }
    .quick-footer { min-height:20px; display:flex; align-items:center; justify-content:space-between; gap:12px; margin-top:8px; }
    .status { color:var(--muted); font-size:13px; }
    .chart-tools { padding:16px 18px; display:grid; grid-template-columns:minmax(0,1fr) auto; gap:14px; border-bottom:1px solid var(--line); }
    .chart-product-grid { display:grid; grid-template-columns:repeat(3,minmax(0,1fr)); gap:8px; max-height:258px; overflow:auto; padding:12px 18px; border-bottom:1px solid var(--line); }
    .chart-product { min-width:0; min-height:82px; display:grid; grid-template-columns:minmax(0,1fr) auto; grid-template-areas:"name profit" "meta meta" "time time"; gap:5px 10px; padding:10px 11px; color:#b9d8f4; text-align:left; border:1px solid #24527f; border-radius:6px; background:#071b32; }
    .chart-product:hover,.chart-product:focus-visible { border-color:#2da9ff; background:#0a294a; box-shadow:0 0 12px #168cff33; }
    .chart-product.active { border-color:#3bc1ff; background:#0b3b68; box-shadow:inset 0 0 15px #168cff33,0 0 12px #168cff33; }
    .chart-product-name { grid-area:name; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; color:#e7f5ff; font-weight:800; }
    .chart-product-profit { grid-area:profit; font-weight:800; font-variant-numeric:tabular-nums; }
    .chart-product-meta { grid-area:meta; color:#89a9c8; font-size:11px; }
    .chart-product-time { grid-area:time; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; color:#6386aa; font-size:10px; }
    .chart-filter-grid { display:grid; grid-template-columns:minmax(150px,.8fr) minmax(190px,1fr) minmax(190px,1fr); gap:10px; margin-bottom:12px; }
    .chart-filter-grid label { min-width:0; }
    .progress { height:7px; background:#e6ebec; margin-top:10px; overflow:hidden; }
    .progress > div { height:100%; width:0; background:var(--accent); transition:width .25s; }
    .job-text { color:var(--muted); font-size:12px; margin-top:7px; min-height:18px; }
    .chart-list { padding:4px 18px 12px; }
    .chart-row { display:grid; grid-template-columns:minmax(180px,1fr) auto; gap:14px; align-items:center; padding:12px 0; border-bottom:1px solid #e8ecec; }
    .chart-row:last-child { border-bottom:0; }
    .chart-name { font-weight:650; overflow-wrap:anywhere; }
    .chart-meta { color:var(--muted); font-size:12px; margin-top:4px; }
    .chart-actions { display:flex; gap:6px; flex-wrap:wrap; justify-content:end; }
    .chart-actions a,.chart-actions button { min-height:32px; padding:6px 9px; border:1px solid #c5ced1; border-radius:4px; background:#fff; text-decoration:none; font-size:12px; }
    .history { padding:4px 18px 14px; }
    .history-item { display:grid; grid-template-columns:150px 90px 1fr; gap:14px; padding:13px 0; border-bottom:1px solid #e8ecec; font-size:13px; }
    .history-item:last-child { border-bottom:0; }
    .history-time { color:var(--muted); font-size:12px; }
    .history-op { font-weight:700; color:var(--accent-dark); }
    .empty-state { padding:30px 18px; color:var(--muted); text-align:center; }
    dialog { width:min(1200px,calc(100% - 30px)); height:min(820px,calc(100% - 30px)); border:0; padding:0; box-shadow:0 18px 60px #0005; }
    dialog::backdrop { background:#1119; }
    .dialog-head { height:50px; display:flex; justify-content:space-between; align-items:center; padding:0 14px; background:#172226; color:#fff; }
    .dialog-head button { min-height:32px; border-color:#536268; background:#26353a; color:#fff; }
    iframe { display:block; width:100%; height:calc(100% - 50px); border:0; }
    /* Risk control dashboard theme */
    :root { --ink:#dbeafe; --muted:#7895b8; --line:#173e6b; --soft:#0a1d38; --paper:#081a32; --accent:#168cff; --accent-dark:#0d6fd8; --warn:#f2ab35; --danger:#ff6472; --good:#34d399; }
    html { color-scheme:dark; }
    body { min-height:100vh; background:radial-gradient(circle at 72% -12%,#0b3970 0,transparent 36%),linear-gradient(145deg,#020b18 0%,#06152b 52%,#020a16 100%); color:var(--ink); }
    body::before { content:""; position:fixed; inset:0; pointer-events:none; opacity:.22; background-image:linear-gradient(#148cff12 1px,transparent 1px),linear-gradient(90deg,#148cff12 1px,transparent 1px); background-size:38px 38px; mask-image:linear-gradient(to bottom,#000,transparent 90%); }
    .topbar { height:64px; position:relative; z-index:2; display:grid; grid-template-columns:minmax(170px,1fr) auto minmax(170px,1fr); align-items:center; padding:0 28px; background:#030d1eeb; border-bottom:1px solid #176fbd; box-shadow:0 4px 24px #006ed933; }
    .brand { display:flex; align-items:center; gap:10px; color:#f2f8ff; text-shadow:0 0 12px #168cff77; }
    .brand::before { content:""; width:25px; height:25px; border-radius:50%; border:4px solid #168cff; border-right-color:#22d3ee; box-shadow:0 0 15px #168cff99; }
    .back { justify-self:start; color:#b9d8fa; border-color:#245586; background:#081b36; }
    .back:hover,.back:focus-visible { color:#fff; border-color:#168cff; background:#0b3767; box-shadow:0 0 12px #168cff55; }
    .detail-account-search { position:relative; display:grid; grid-template-columns:minmax(150px,210px) auto; gap:7px; justify-self:end; align-items:center; }
    .detail-account-search input { min-height:34px; padding:6px 9px; font-size:13px; }
    .detail-account-search button { min-height:34px; padding:6px 11px; font-size:13px; }
    .detail-account-search-status { position:absolute; top:calc(100% + 5px); right:0; width:max-content; max-width:300px; padding:4px 7px; border:1px solid #245586; border-radius:4px; background:#06162b; color:#9bb7d5; font-size:12px; box-shadow:0 6px 16px #0008; }
    .detail-account-search-status:empty { display:none; }
    main { position:relative; z-index:1; width:min(1540px,calc(100% - 32px)); }
    .account-head,.quick-mark,.toolbar,.section { background:linear-gradient(145deg,#081b35ee,#06162cee); border-color:#194a7c; box-shadow:0 14px 34px #0007,inset 0 1px 0 #2a73ad33; }
    .account-head { overflow:hidden; border-top:1px solid #238ee8; border-radius:10px 10px 0 0; }
    .account-head::before { content:"账户风险画像"; align-self:flex-start; color:#4db9ff; border:1px solid #1c69a3; background:#09284a; border-radius:5px; padding:5px 9px; font-size:11px; letter-spacing:1px; }
    .account-head > div:first-child { flex:1; }
    .account-id { color:#f2f8ff; font-size:32px; text-shadow:0 0 16px #168cff88; }
    .badge { color:#9bb7d5; background:#102942; border:1px solid #285276; }
    .badge.marked { color:#7ff3d0; background:#0c3a36; border-color:#1f7e68; }
    .badge.empty { color:#ffd38c; background:#3a2a0e; border-color:#8d6520; }
    .badge.ea { color:#7ce7ff; background:#073449; border-color:#168ab2; box-shadow:0 0 10px #22c8ff33; }
    .badge.action { color:#a9d8ff; background:#0b2848; border-color:#285f8e; box-shadow:none; }
    .head-meta,.section-sub,.status,.history-time,.chart-meta,.job-text,.order-summary-meta,.risk-note { color:var(--muted); }
    .quick-mark { border-top:0; box-shadow:inset 0 1px 0 #164b7b; }
    .toolbar { border-top:1px solid #133b62; border-radius:0 0 10px 10px; box-shadow:0 14px 34px #0007; }
    label,.quick-label { color:#7999ba; }
    input,select,textarea { color:#dcecff; background:#06162b; border-color:#28527b; }
    input::placeholder,textarea::placeholder { color:#537294; }
    input:focus,select:focus,textarea:focus { outline:2px solid #168cff55; border-color:#168cff; box-shadow:0 0 16px #168cff33; }
    option { background:#07172c; color:#dbeafe; }
    button { color:#cce4ff; background:#0a2342; border-color:#2b5c8c; }
    button:hover:not(:disabled),button:focus-visible { color:#fff; border-color:#168cff; background:#0d3764; box-shadow:0 0 14px #168cff55; }
    button.primary { background:linear-gradient(135deg,#0875df,#159bff); border-color:#29a6ff; box-shadow:0 0 16px #168cff44; }
    button.primary:hover { background:linear-gradient(135deg,#168cff,#22b1ff); }
    .action-btn { color:#a9d8ff; background:#0b2848; border-color:#285f8e; }
    .action-btn:hover:not(:disabled) { color:#fff; background:#0d3764; border-color:#2faeff; }
    .action-btn.active { color:#fff; border-color:#39b5ff; background:linear-gradient(135deg,#0e72cf,#159bff); box-shadow:0 0 15px #168cff66,inset 0 1px 0 #8bd6ff55; }
    .quick-status select { color:#71d0ff; background:#082643; border-color:#2faeff; box-shadow:0 0 12px #168cff33; }
    .batch-toggle { color:#77d5ff; background:#082643; border-color:#277cb6; }
    .batch-toggle:hover { color:#fff; background:#0b355d; border-color:#2faeff; box-shadow:0 0 12px #168cff44; }
    .action-manager { border-color:#1f659c; background:#071b32; box-shadow:inset 0 1px 0 #2f78ad33; }
    .manage-action-chip { color:#9bdcff; background:#092744; border-color:#236ca7; }
    .manage-action-chip button { color:#ff8a94; background:#0b2949; border-left-color:#236ca7; }
    .manage-action-chip button:hover,.manage-action-chip button:focus-visible { color:#fff; background:#a72f3a; border-color:#ff6472; box-shadow:none; }
    .manage-action-chip.protected { color:#7895b8; background:#0a2039; }
    .section { overflow:hidden; border-radius:10px; }
    .section-head { min-height:50px; border-color:#17446f; background:linear-gradient(90deg,#0a2240cc,#07182e66); }
    .section-head h2 { color:#e9f5ff; font-size:16px; text-shadow:0 0 14px #168cff55; }
    .section-head h2::before { content:""; display:inline-block; width:3px; height:15px; margin-right:9px; vertical-align:-2px; border-radius:2px; background:#20a4ff; box-shadow:0 0 10px #20a4ffaa; }
    #financeMetrics { grid-template-columns:repeat(6,minmax(0,1fr)); gap:10px; padding:12px; }
    #frequencyMetrics { grid-template-columns:repeat(4,minmax(0,1fr)); gap:10px; padding:12px; }
    .risk-stat { min-height:88px; position:relative; overflow:hidden; border:1px solid #174674!important; border-radius:8px; padding:15px 16px; background:linear-gradient(145deg,#0a213f,#07172d); box-shadow:0 8px 20px #0004,inset 0 1px 0 #3c8dcc22; }
    .risk-stat::after { content:""; position:absolute; right:-20px; top:-24px; width:75px; height:75px; border-radius:50%; background:#168cff10; border:1px solid #168cff1f; }
    .risk-stat span { color:#7595b8; }
    .risk-stat b { position:relative; z-index:1; color:#edf7ff; font-size:19px; text-shadow:0 0 12px #168cff33; }
    .risk-stat b.positive,.metric b.positive,.positive { color:var(--good)!important; }
    .risk-stat b.negative,.metric b.negative,.negative { color:var(--danger)!important; }
    .risk-note { border-color:#12385f; background:#06162b99; }
    .metric-groups { gap:1px; background:#163c63; }
    .metric-group { border:0!important; background:#071a31; }
    .metric-group h3 { color:#55b9f8; }
    .metric { min-width:0; border-left:2px solid #164a77; padding-left:9px; }
    .metric span { color:#6f91b6; }
    .metric b { color:#ddecfb; }
    .split { gap:16px; }
    .table-wrap { scrollbar-color:#1d5d96 #061426; }
    table { color:#bed4ec; }
    th { color:#7598be; background:#0a2240; border-color:#17446f; }
    td { border-color:#102f51; }
    tbody tr:hover td { background:#0b294c; }
    .risk-table tr.total td { color:#e8f5ff; background:#0b2747; }
    .risk-table tr.current-account td { color:#eaf8ff; background:#0c3659; border-top-color:#2da9ff; border-bottom-color:#2da9ff; }
    .risk-table tr.current-account td:first-child { border-left-color:#22b1ff; box-shadow:inset 4px 0 0 #22b1ff; }
    .account-link { color:#55c3ff; background:#0b2a4b; border-color:#236ca7; }
    .account-link:hover,.account-link:focus-visible { color:#fff; background:#147cda; border-color:#3cb7ff; box-shadow:0 0 12px #168cff55; }
    .source-list,.chart-list,.history { background:#06162b55; }
    .source-row,.chart-row,.history-item { border-color:#102f51; }
    .history-op { color:#55c3ff; }
    .chart-actions a,.chart-actions button { color:#7dd1ff; background:#0a294a; border-color:#236ca7; }
    .chart-actions a:hover,.chart-actions a:focus-visible,.chart-actions button:hover,.chart-actions button:focus-visible { color:#fff; background:#147cda; border-color:#3cb7ff; box-shadow:0 0 12px #168cff55; }
    .progress { background:#071426; border:1px solid #164772; border-radius:999px; }
    .progress > div { background:linear-gradient(90deg,#0875df,#20c5ff); box-shadow:0 0 12px #20a4ff; }
    .order-details > summary { color:#dcecff; background:linear-gradient(90deg,#0a2240,#07182e); border-color:#194a7c; }
    .order-details > summary::before { content:"＋"; display:inline-grid; place-items:center; flex:0 0 26px; width:26px; height:26px; color:#62c7ff; border:1px solid #277cb6; border-radius:5px; background:#082643; }
    .order-details[open] > summary::before { content:"−"; }
    .order-details > summary:hover,.order-details > summary:focus-visible { color:#fff; background:#0b2d55; border-color:#2da9ff; box-shadow:inset 0 0 18px #168cff22; }
    .order-details[open] > summary { border-bottom-color:#17446f; }
    .order-table { min-width:1760px; }
    .order-table tbody tr:nth-child(even) td { background:#071a31; }
    .order-table tbody tr:hover td { background:#0d3158; }
    .reason-badge { display:inline-block; min-width:56px; padding:3px 7px; color:#78d3ff; background:#092a4a; border:1px solid #236fa8; border-radius:999px; text-align:center; font-size:12px; font-weight:700; }
    .order-comment { max-width:300px; min-width:180px; overflow:hidden; text-overflow:ellipsis; color:#d5e8fa; }
    .order-type.buy { color:var(--good); } .order-type.sell { color:var(--danger); }
    .order-pager { border-color:#163b62; background:#06172c; }
    dialog { background:#06162b; border:1px solid #237fc2; border-radius:10px; box-shadow:0 18px 70px #000c,0 0 30px #168cff44; }
    dialog::backdrop { background:#010711dd; backdrop-filter:blur(3px); }
    .account-source-dialog { width:min(620px,calc(100% - 30px)); height:auto; max-height:calc(100% - 30px); padding:18px; }
    .account-source-list { display:grid; gap:9px; }
    .account-source-option { display:flex; align-items:center; justify-content:space-between; gap:14px; width:100%; padding:13px 14px; border:1px solid #245586; border-radius:6px; background:#071f38; color:#dcecf8; text-align:left; cursor:pointer; }
    .account-source-option:hover,.account-source-option:focus-visible { border-color:#1ab5fb; background:#0a3152; }
    .account-source-option small { display:block; margin-top:4px; color:#819db5; font-size:11px; }
    .account-source-option strong { color:#58d2f5; white-space:nowrap; }
    .dialog-head { background:#071a31; border-bottom:1px solid #1d6097; }
    .dialog-head button,.dialog-head a { color:#b9d8fa; border:1px solid #245586; border-radius:5px; background:#081b36; text-decoration:none; padding:7px 10px; }
    .dialog-head-actions { display:flex; align-items:center; gap:8px; }
    .empty-state { color:#6587ab; }
    .dashboard-grid,.lower-grid { display:block; }
    .dashboard-column { min-width:0; }
    .viz-stack { display:grid; gap:12px; padding:12px; }
    .viz-panel { min-width:0; position:relative; padding:14px; border:1px solid #174674; border-radius:8px; background:linear-gradient(145deg,#091f3a,#06162b); box-shadow:inset 0 1px 0 #3c8dcc22; }
    .viz-panel-head { min-height:28px; display:flex; align-items:start; justify-content:space-between; gap:12px; margin-bottom:8px; }
    .viz-panel-head b { color:#ddecfb; font-size:13px; }
    .viz-panel-head span { color:#6f91b6; font-size:11px; text-align:right; }
    .viz-canvas { min-height:210px; display:grid; place-items:center; overflow:hidden; }
    .viz-canvas.compact { min-height:170px; }
    .viz-svg { display:block; width:100%; height:auto; overflow:visible; }
    .viz-grid-line { stroke:#19446e; stroke-width:1; stroke-dasharray:4 5; }
    .viz-zero-line { stroke:#6687aa; stroke-width:1; stroke-dasharray:3 4; }
    .viz-axis-text { fill:#6f91b6; font-size:10px; }
    .viz-note { color:#6f91b6; font-size:11px; line-height:1.6; }
    .viz-empty { color:#6587ab; font-size:12px; text-align:center; }
    .structure-grid { display:grid; grid-template-columns:repeat(2,minmax(0,1fr)); gap:12px; padding:12px; }
    .donut-wrap { min-height:190px; display:grid; grid-template-columns:150px minmax(0,1fr); align-items:center; gap:12px; }
    .donut-chart { width:142px; height:142px; margin:auto; }
    .donut-center-main { fill:#eff8ff; font-size:20px; font-weight:800; text-anchor:middle; }
    .donut-center-sub { fill:#7895b8; font-size:10px; text-anchor:middle; }
    .legend-list { display:grid; gap:8px; }
    .legend-item { display:grid; grid-template-columns:9px minmax(0,1fr) auto; gap:8px; align-items:center; color:#9db9d6; font-size:12px; }
    .legend-dot { width:9px; height:9px; border-radius:50%; box-shadow:0 0 8px currentColor; }
    .bar-list { display:grid; gap:9px; width:100%; }
    .bar-row { display:grid; grid-template-columns:minmax(72px,auto) minmax(90px,1fr) auto; gap:9px; align-items:center; font-size:12px; }
    .bar-label { color:#9db9d6; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }
    .bar-track { height:9px; overflow:hidden; border:1px solid #174674; border-radius:999px; background:#041225; }
    .bar-fill { height:100%; min-width:2px; border-radius:999px; box-shadow:0 0 10px currentColor; }
    .bar-value { min-width:72px; text-align:right; color:#dcecff; font-variant-numeric:tabular-nums; }
    .ip-summary { display:flex; flex-wrap:wrap; gap:7px; padding:12px 12px 0; }
    .ip-chip { padding:5px 8px; border:1px solid #236ca7; border-radius:999px; background:#0b2a4b; color:#86d7ff; font-size:11px; }
    .ip-list { display:grid; gap:9px; padding:12px; }
    .ip-row { display:grid; grid-template-columns:minmax(120px,.7fr) minmax(140px,1fr) minmax(180px,1.35fr); gap:12px; padding:12px; border:1px solid #16436d; border-radius:7px; background:#071a31; }
    .ip-address { color:#72d0ff; font-size:15px; font-weight:750; font-family:Consolas,"Segoe UI",sans-serif; }
    .ip-meta,.ip-location { color:#8caac8; font-size:11px; line-height:1.65; }
    .ip-location b { display:block; color:#dcecff; font-size:12px; }
    .section-head-actions { display:flex; align-items:center; justify-content:end; gap:8px; }
    .automation-grid { display:grid; grid-template-columns:1fr; gap:12px; padding:12px; }
    .automation-panel { min-width:0; overflow:hidden; border:1px solid #174674; border-radius:8px; background:linear-gradient(145deg,#091f3a,#06162b); }
    .automation-panel-head { display:flex; align-items:start; justify-content:space-between; gap:10px; padding:12px 14px; border-bottom:1px solid #174674; }
    .automation-panel-head b { color:#ddecfb; }
    .automation-panel-head span { color:#7895b8; font-size:11px; text-align:right; }
    .automation-summary { display:grid; grid-template-columns:repeat(3,minmax(0,1fr)); gap:8px; padding:10px; }
    .automation-summary .risk-stat { min-height:66px; padding:10px 11px; }
    .automation-summary .risk-stat b { font-size:16px; }
    .automation-table { min-width:720px; }
    .automation-note { padding:0 12px 10px; color:#7895b8; font-size:11px; line-height:1.55; }
    .section-head-actions button { min-height:30px; padding:5px 9px; font-size:11px; }
    .ea-query-entry { color:#d8f8ff; border-color:#2687a8; background:#0a4058; }
    .ea-query-entry:hover,.ea-query-entry:focus-visible { color:#fff; border-color:#51c9ef!important; background:#0d5874!important; box-shadow:0 0 14px #20bce944!important; }
    .relationship-entry { color:#e6d9ff; border-color:#7056a6; background:#30214f; }
    .relationship-entry:hover,.relationship-entry:focus-visible { color:#fff; border-color:#ac8fff!important; background:#4b3479!important; box-shadow:0 0 14px #9d76ff55!important; }
    .toxic-entry { color:#fff; border-color:#ff8a4c; background:linear-gradient(135deg,#b74218,#e85d24); box-shadow:0 0 15px #ff6a2f44; }
    .toxic-entry:hover,.toxic-entry:focus-visible { border-color:#ffb083!important; background:linear-gradient(135deg,#d9531f,#ff7837)!important; box-shadow:0 0 18px #ff6a2f66!important; }
    .historical-funds-entry { color:#dcf9ff; border-color:#287b9e; background:#073c57; box-shadow:0 0 14px #168cbb38; }
    .historical-funds-entry:hover,.historical-funds-entry:focus-visible { color:#fff; border-color:#5fd3f5!important; background:#07536f!important; box-shadow:0 0 18px #28c5ef55!important; }
    .toxic-dialog { width:min(1080px,calc(100% - 30px)); height:min(850px,calc(100% - 30px)); }
    .toxic-body { height:calc(100% - 50px); overflow:auto; padding:16px; }
    .historical-funds-dialog { width:min(1280px,calc(100% - 30px)); }
    .funds-summary { display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); margin:0 0 14px; border:1px solid #214f78; border-radius:7px; overflow:hidden; }
    .funds-summary div { min-height:68px; padding:10px 12px; border-left:1px solid #214f78; background:#071a31; }
    .funds-summary div:nth-child(4n+1) { border-left:0; }
    .funds-summary span,.funds-summary b { display:block; }
    .funds-summary span { color:#7895b8; font-size:11px; }
    .funds-summary b { margin-top:6px; color:#eaf6ff; font-size:16px; }
    .funds-chart { min-height:260px; padding:12px; border:1px solid #214f78; border-radius:7px; background:#06162b; }
    .funds-chart-head { display:flex; align-items:flex-end; justify-content:space-between; gap:12px; margin-bottom:8px; }
    .funds-chart-head b { color:#eaf6ff; font-size:13px; }
    .funds-chart-head span,.funds-note { color:#7895b8; font-size:11px; line-height:1.6; }
    .funds-chart svg { display:block; width:100%; height:220px; overflow:visible; }
    .funds-chart-grid { stroke:#16436c; stroke-width:1; }
    .funds-chart-balance { fill:none; stroke:#20b9ff; stroke-width:2.5; vector-effect:non-scaling-stroke; }
    .funds-chart-credit { fill:none; stroke:#e5af48; stroke-width:2; vector-effect:non-scaling-stroke; }
    .funds-chart-liquidation { fill:#ff5d66; stroke:#ffd4d7; stroke-width:1.5; vector-effect:non-scaling-stroke; cursor:pointer; }
    .funds-chart-liquidation:hover,.funds-chart-liquidation:focus { fill:#ff8790; stroke:#fff; outline:none; }
    .funds-liquidations { display:flex; align-items:center; gap:8px; flex-wrap:wrap; margin:10px 0 0; padding:10px 12px; border:1px solid #7c3945; border-radius:6px; background:#28131b; }
    .funds-liquidations b { color:#ffd9dc; font-size:11px; }
    .funds-liquidation-jump { min-height:28px; padding:4px 8px; color:#ffd9dc; border-color:#a94958; background:#4a1c27; font-size:11px; }
    .funds-liquidation-jump:hover,.funds-liquidation-jump:focus-visible { color:#fff; border-color:#ff7a85!important; background:#682430!important; }
    .funds-events-head { display:flex; align-items:end; justify-content:space-between; gap:12px; margin:16px 0 8px; }
    .funds-events-head b { color:#eaf6ff; font-size:13px; }
    .funds-event-pager { display:flex; justify-content:flex-end; align-items:center; gap:8px; margin:10px 0 0; color:#7895b8; font-size:11px; }
    .funds-events-table { min-width:1180px; }
    .funds-events-table td { vertical-align:top; }
    .funds-event-focus td { box-shadow:inset 0 2px 0 #ff8a93,inset 0 -2px 0 #ff8a93; }
    .funds-kind { display:inline-flex; padding:3px 6px; border:1px solid #376486; border-radius:4px; color:#bcd7ef; font-size:10px; white-space:nowrap; }
    .funds-kind.deposit,.funds-kind.bonus_grant,.funds-kind.negative_balance_clear { color:#70dfb1; border-color:#21835e; background:#0b322b; }
    .funds-kind.withdrawal,.funds-kind.bonus_remove { color:#ffc08a; border-color:#92572e; background:#38240f; }
    .funds-kind.internal_transfer { color:#d1c4ff; border-color:#6955a0; background:#282044; }
    .funds-liquidation-label { display:block; margin-top:4px; color:#ffafb5; font-size:10px; white-space:nowrap; }
    .funds-event-liquidation td { background:#2b151d; }
    .copy-panel { margin-bottom:16px; overflow:hidden; border:1px solid #234e79; border-radius:7px; background:#071a31; }
    .copy-panel:last-child { margin-bottom:0; }
    .copy-panel-title { display:flex; align-items:center; justify-content:space-between; gap:10px; padding:11px 16px; border-bottom:1px solid #234e79; color:#dbeafe; background:#092744; }
    .copy-panel .risk-note { border-top:0; border-bottom:1px solid #12385f; }
    .ea-range-controls { display:flex; align-items:end; flex-wrap:wrap; gap:10px; padding:12px 16px; border-bottom:1px solid #12385f; background:#06162b; }
    .ea-range-controls label { display:grid; gap:5px; font-size:12px; }
    .ea-range-controls input { min-height:34px; }
    .ea-range-controls button { min-height:34px; }
    .ea-range-hint { flex:1 1 180px; padding-bottom:8px; color:var(--muted); font-size:12px; }
    .copy-group-block { margin:12px; overflow:hidden; border:1px solid #1b4f7d; border-radius:6px; background:#06162b; }
    .copy-group-head { display:flex; align-items:center; justify-content:space-between; gap:12px; padding:11px 14px; border-bottom:1px solid #1b4f7d; }
    .copy-group-head span { color:var(--muted); font-size:12px; }
    .ea-group-block > summary { cursor:pointer; list-style:none; }
    .ea-group-block > summary::-webkit-details-marker { display:none; }
    .ea-group-block > summary::before { content:""; width:8px; height:8px; flex:0 0 auto; border-right:2px solid #59b8ee; border-bottom:2px solid #59b8ee; transform:rotate(-45deg); transition:transform .15s ease; }
    .ea-group-block[open] > summary::before { transform:rotate(45deg); }
    .ea-group-block:not([open]) > summary { border-bottom:0; }
    .ea-group-block > summary > div { flex:1; min-width:0; }
    .copy-group-table { min-width:1380px; }
    .copy-master-block { margin:12px; overflow:hidden; border:1px solid #1b5b88; border-radius:6px; background:#06162b; }
    .copy-master-head { display:flex; align-items:center; justify-content:space-between; gap:12px; padding:12px 14px; border-bottom:1px solid #1b5b88; background:#08233d; }
    .copy-master-head b { color:#e5f4ff; font-size:15px; }
    .copy-master-head span,.copy-master-head small { color:var(--muted); font-size:11px; }
    .copy-follower-table { min-width:1700px; }
    .copy-source-orders { margin:10px 12px 12px; border:1px solid #163e64; border-radius:5px; }
    .copy-source-orders summary { padding:8px 10px; cursor:pointer; color:#63b8ef; font-size:11px; }
    .copy-source-orders table { min-width:780px; }
    .relationship-network-dialog { width:min(1320px,calc(100% - 30px)); height:min(880px,calc(100% - 30px)); }
    .relationship-network-body { height:calc(100% - 50px); display:grid; grid-template-rows:auto minmax(0,1fr); overflow:hidden; }
    .relationship-network-controls { display:flex; align-items:center; justify-content:space-between; gap:12px; flex-wrap:wrap; padding:11px 14px; border-bottom:1px solid #174674; background:#06162b; }
    .relationship-network-filters { display:flex; align-items:center; gap:8px 12px; flex-wrap:wrap; }
    .relationship-network-filters label { display:flex; align-items:center; gap:5px; color:#abc7e4; font-size:11px; cursor:pointer; }
    .relationship-network-filters input { width:auto; min-height:auto; margin:0; accent-color:#8066bc; }
    .relationship-network-layout { min-height:0; display:grid; grid-template-columns:minmax(0,1.7fr) minmax(300px,.7fr); }
    .relationship-network-graph-wrap { min-width:0; min-height:0; display:grid; grid-template-rows:minmax(0,1fr) auto; padding:14px; background:#041225; }
    #relationshipNetworkGraph { display:block; width:100%; min-height:440px; height:auto; aspect-ratio:1000/620; touch-action:none; user-select:none; cursor:grab; contain:layout paint; }
    #relationshipNetworkGraph.dragging { cursor:grabbing; }
    .relationship-network-legend { display:flex; justify-content:center; gap:10px 16px; flex-wrap:wrap; padding-top:7px; color:#7895b8; font-size:11px; }
    .relationship-network-legend span { display:inline-flex; align-items:center; gap:5px; }
    .relationship-network-legend i { display:inline-block; width:9px; height:9px; border:1px solid #315474; }
    .relationship-network-legend .account { border-radius:50%; background:#298feb; }
    .relationship-network-legend .clue { border-radius:999px; background:#d36aac; }
    .relationship-network-legend .group { border-radius:2px; background:#8066bc; }
    .relationship-network-detail { min-width:0; overflow:auto; padding:16px; border-left:1px solid #174674; background:#06162b; }
    .relationship-detail-kind { color:#7895b8; font-size:11px; }
    .relationship-detail-title { margin:5px 0 10px; color:#e9f5ff; font-size:17px; }
    .relationship-detail-meta { display:grid; grid-template-columns:1fr 1fr; gap:8px; margin-bottom:14px; }
    .relationship-detail-meta div { padding:8px; border:1px solid #174674; border-radius:5px; background:#071a31; }
    .relationship-detail-meta span,.relationship-detail-meta b { display:block; }
    .relationship-detail-meta span { color:#7895b8; font-size:10px; }
    .relationship-detail-meta b { margin-top:3px; color:#e4f2ff; font-size:12px; word-break:break-word; }
    .relationship-evidence-title { margin-top:15px; padding-top:12px; border-top:1px solid #174674; color:#7895b8; font-size:11px; }
    .relationship-evidence { display:grid; gap:8px; margin-top:8px; }
    .relationship-evidence-row { padding:9px 10px; border:1px solid #174674; border-radius:5px; color:#bfd6ec; background:#071a31; font-size:12px; line-height:1.55; }
    .relationship-limitations { margin:0; padding:0 0 0 18px; color:#7895b8; font-size:11px; line-height:1.65; }
    .relationship-limitations li { margin:4px 0; }
    .relationship-network-graph .relation-edge { fill:none; stroke:#315474; stroke-width:2; }
    .relationship-network-graph .relation-edge.selected { stroke:#b69cff; stroke-width:4; }
    .relationship-network-graph .relation-edge-hit { fill:none; stroke:transparent; stroke-width:18; cursor:pointer; }
    .relationship-network-graph .relation-edge-label { pointer-events:none; }
    .relationship-network-graph .relation-edge-label rect { fill:#0a2440; stroke:#4f83ad; stroke-width:1; }
    .relationship-network-graph .relation-edge-label text { fill:#e3f3ff; font-size:10px; font-weight:700; text-anchor:middle; dominant-baseline:middle; }
    .relationship-network-graph .relation-edge-label.selected rect { fill:#352d5c; stroke:#c4aaff; }
    .relationship-network-graph .relation-node { cursor:grab; }
    .relationship-network-graph .relation-node:active { cursor:grabbing; }
    .relationship-network-graph .relation-node-shape { stroke:#315474; stroke-width:2; }
    .relationship-network-graph .relation-node.selected .relation-node-shape { stroke:#d7c7ff; stroke-width:4; }
    .relationship-network-graph .relation-node.account .relation-node-shape { fill:#1679cf; }
    .relationship-network-graph .relation-node.subject .relation-node-shape { fill:#168cff; stroke:#a6ddff; }
    .relationship-network-graph .relation-node.ip .relation-node-shape { fill:#bb5d99; }
    .relationship-network-graph .relation-node.ea_feature .relation-node-shape { fill:#9c4e8a; }
    .relationship-network-graph .relation-node.copy_group .relation-node-shape { fill:#7257aa; }
    .relationship-network-graph .relation-node.rebate .relation-node-shape { fill:#ba7943; }
    .relationship-network-graph .relation-node-label { fill:#eaf7ff; font-size:14px; text-anchor:middle; pointer-events:none; }
    .relationship-network-graph .relation-node-sub { fill:#9ab6d0; font-size:10px; text-anchor:middle; pointer-events:none; }
    .relationship-network-graph .relation-node-expand { fill:#d6c6ff; font-size:9px; text-anchor:middle; pointer-events:none; }
    .relationship-network-graph .relation-empty { fill:#7895b8; font-size:14px; text-anchor:middle; }
    .toxic-mode-row { display:flex; align-items:center; gap:10px; flex-wrap:wrap; padding:12px; border:1px solid #234e79; border-radius:7px; background:#071a31; }
    .toxic-mode { display:flex; align-items:center; gap:7px; min-height:34px; padding:6px 10px; border:1px solid #285f8e; border-radius:5px; color:#acd7ff; background:#092744; cursor:pointer; }
    .toxic-mode:has(input:checked) { color:#fff; border-color:#39b5ff; background:#0d5d9e; box-shadow:0 0 12px #168cff55; }
    .toxic-mode input,.toxic-check input { width:auto; min-height:auto; margin:0; accent-color:#168cff; }
    .toxic-mode-note { flex:1; min-width:240px; color:#7895b8; font-size:12px; text-align:right; }
    .toxic-selector { display:grid; grid-template-columns:repeat(3,minmax(0,1fr)); gap:8px; margin-top:12px; }
    .toxic-check { min-height:42px; display:flex; align-items:center; gap:8px; padding:8px 10px; border:1px solid #234e79; border-radius:6px; color:#bcd7ee; background:#071b32; cursor:pointer; }
    .toxic-check:hover { border-color:#2da9ff; background:#0a294a; }
    .toxic-selector.disabled { opacity:.45; pointer-events:none; }
    .toxic-controls-footer { display:flex; align-items:center; justify-content:space-between; gap:12px; margin-top:12px; }
    .toxic-results { display:grid; gap:10px; margin-top:14px; }
    .toxic-result { display:grid; grid-template-columns:92px minmax(150px,.8fr) minmax(320px,2fr) auto; gap:12px; align-items:start; padding:14px 12px; border:1px solid #174674; border-left:4px solid #477399; border-radius:7px; background:#071a31; }
    .toxic-result.warning { border-left-color:#f2ab35; } .toxic-result.high { border-left-color:#ff7c45; } .toxic-result.critical { border-left-color:#ff4758; box-shadow:inset 0 0 18px #ff475811; }
    .toxic-score { font-size:28px; font-weight:800; color:#edf7ff; line-height:1; }
    .toxic-score small { display:block; margin-top:5px; color:#7895b8; font-size:10px; font-weight:500; }
    .toxic-result-title b { display:block; color:#dcecff; }
    .toxic-result-title span,.toxic-result-summary,.toxic-limit { color:#7895b8; font-size:11px; line-height:1.55; }
    .toxic-result-summary b { color:#a9c7e5; }
    .toxic-analysis { display:grid; gap:8px; }
    .toxic-analysis-row { display:grid; grid-template-columns:78px minmax(0,1fr); gap:8px; color:#a9c7e5; font-size:12px; line-height:1.65; }
    .toxic-analysis-row b { color:#eef7ff; }
    .toxic-chain { display:grid; gap:10px; }
    .toxic-chain-head { color:#eef7ff; font-size:13px; line-height:1.6; }
    .toxic-chain-reasoning { color:#c5d9eb; font-size:12px; line-height:1.7; }
    .toxic-chain-facts { display:grid; gap:7px; border-left:2px solid #2b658d; padding-left:11px; }
    .toxic-chain-fact { display:grid; grid-template-columns:90px minmax(0,1fr); gap:8px; color:#a9c7e5; font-size:12px; line-height:1.65; }
    .toxic-chain-fact b { color:#eef7ff; }
    .toxic-chain-fact[data-strength="strong"] b { color:#7ee0bf; }
    .toxic-chain-fact[data-strength="missing"] b { color:#e6b56b; }
    .toxic-chain-section { border-top:1px solid #1c4268; padding-top:8px; }
    .toxic-chain-section b { color:#eef7ff; font-size:11px; }
    .toxic-chain-section ul { margin:5px 0 0 17px; padding:0; color:#a9c7e5; font-size:11px; line-height:1.65; }
    .toxic-accomplices { margin-top:10px; padding:9px 10px; border:1px solid #214d76; border-radius:6px; background:#081f39; }
    .toxic-accomplices-title { color:#eef7ff; font-size:12px; font-weight:700; }
    .toxic-accomplices-note { margin-top:4px; color:#7895b8; font-size:11px; }
    .toxic-accomplice-list { display:grid; grid-template-columns:repeat(auto-fit,minmax(190px,1fr)); gap:6px; margin-top:7px; }
    .toxic-accomplice { padding:7px 8px; border:1px solid #1c4268; border-radius:5px; color:#9dbbd8; font-size:11px; line-height:1.5; }
    .toxic-accomplice a { color:#63b8ef; font-weight:700; }
    .toxic-sync-comparison { padding:14px; border:1px solid #245577; border-radius:7px; background:#06172a; }
    .toxic-sync-head,.toxic-sync-detail-head { display:flex; align-items:flex-end; justify-content:space-between; gap:14px; }
    .toxic-sync-head h3,.toxic-sync-detail-head h3 { margin:0; color:#e5f2ff; font-size:14px; }
    .toxic-sync-head small,.toxic-sync-detail-head small { display:block; margin-top:4px; color:#7895b8; font-size:11px; }
    .toxic-sync-kpis { display:grid; grid-template-columns:repeat(6,minmax(105px,1fr)); margin-top:12px; border-top:1px solid #1d4667; border-bottom:1px solid #1d4667; }
    .toxic-sync-kpis div { padding:9px; border-left:1px solid #1d4667; }
    .toxic-sync-kpis div:first-child { border-left:0; }
    .toxic-sync-kpis span,.toxic-sync-kpis b { display:block; }
    .toxic-sync-kpis span { color:#7895b8; font-size:10px; }
    .toxic-sync-kpis b { margin-top:4px; color:#e4f2ff; font-size:15px; }
    .toxic-sync-table-wrap { margin-top:10px; max-height:250px; overflow:auto; }
    .toxic-sync-detail-head { margin-top:15px; }
    .toxic-sync-detail-head label { display:flex; flex-direction:column; gap:4px; width:230px; color:#8ba8c2; font-size:11px; }
    .toxic-sync-detail-wrap { margin-top:8px; max-height:480px; overflow:auto; }
    .toxic-sync-detail-table { min-width:1320px; }
    .toxic-sync-detail-table td { vertical-align:top; }
    .toxic-sync-detail-table td>b,.toxic-sync-detail-table td>a,.toxic-sync-detail-table td>small { display:block; }
    .toxic-sync-detail-table td>small { margin-top:3px; color:#7895b8; white-space:nowrap; }
    .toxic-sync-state { display:inline-flex; min-width:76px; justify-content:center; padding:4px 6px; border:1px solid #536b7f; border-radius:4px; font-size:10px; font-weight:700; }
    .toxic-sync-state.full { color:#76e0b5; border-color:#238866; background:#0b362c; }
    .toxic-sync-state.open-only { color:#ffc779; border-color:#9c6929; background:#38280e; }
    .toxic-tech { margin-top:10px; color:#7895b8; font-size:11px; }
    .toxic-tech summary { width:max-content; cursor:pointer; color:#63b8ef; }
    .toxic-tech-body { margin-top:7px; line-height:1.65; }
    .toxic-result-actions { display:flex; gap:6px; align-items:center; }
    .toxic-result-actions button { min-height:31px; padding:5px 9px; font-size:11px; }
    .toxic-empty { padding:36px 16px; color:#7895b8; text-align:center; border:1px dashed #28527b; border-radius:7px; }
    @media (min-width:951px) {
      .dashboard-grid { display:grid; grid-template-columns:minmax(0,7fr) minmax(420px,5fr); gap:16px; align-items:start; }
      .dashboard-column { display:grid; gap:16px; align-content:start; }
      .dashboard-column > .section { margin-top:0; }
      .dashboard-grid { margin-top:18px; }
      .lower-grid { display:grid; grid-template-columns:minmax(0,7fr) minmax(420px,5fr); gap:16px; align-items:start; }
      .lower-grid > .section,.lower-grid > .dashboard-column > .section { margin-top:18px; }
      .dashboard-grid #financeMetrics { grid-template-columns:repeat(3,minmax(0,1fr)); }
      .dashboard-grid #frequencyMetrics { grid-template-columns:repeat(4,minmax(0,1fr)); }
      .dashboard-grid .metric-groups { grid-template-columns:1fr; }
    }
    @media (max-width:950px) {
      .topbar { padding:0 14px; } main { width:calc(100% - 20px); margin-top:10px; }
      .account-head { align-items:flex-start; padding:17px 14px; } .head-meta { text-align:left; }
      .toolbar { grid-template-columns:1fr 1fr; } .metric-groups,.split { grid-template-columns:1fr; }
      .automation-grid { grid-template-columns:1fr; }
      .quick-mark-main { grid-template-columns:1fr 180px; } .quick-mark-main #saveBtn { grid-column:1/-1; }
      .quick-fields { grid-template-columns:1fr 1fr; }
      .metric-group { border-right:0; } .metric-group:nth-last-child(2) { border-bottom:1px solid var(--line); }
      .metrics,.risk-summary,#financeMetrics,#frequencyMetrics { grid-template-columns:repeat(2,1fr); } .risk-stat:nth-child(4n) { border-right:1px solid var(--line); } .risk-stat:nth-child(2n) { border-right:0; } .risk-stat:nth-last-child(-n+4) { border-bottom:1px solid var(--line); } .risk-stat:nth-last-child(-n+2) { border-bottom:0; }
      .form-grid { grid-template-columns:1fr; } .form-grid .wide { grid-column:auto; }
      .chart-product-grid { grid-template-columns:repeat(2,minmax(0,1fr)); }
      .toxic-sync-kpis { grid-template-columns:repeat(3,1fr); }
      .toxic-sync-kpis div:nth-child(4) { border-left:0; }
      .funds-summary { grid-template-columns:repeat(2,minmax(0,1fr)); }
      .funds-summary div:nth-child(odd) { border-left:0; }
    }
    @media (max-width:620px) {
      .topbar { height:auto; min-height:64px; grid-template-columns:auto minmax(0,1fr); gap:7px 12px; padding:8px 12px 10px; }
      .brand { justify-self:center; font-size:14px; }
      .brand::before { width:21px; height:21px; border-width:3px; }
      .detail-account-search { grid-column:1/-1; grid-template-columns:minmax(0,1fr) auto; width:100%; }
      .detail-account-search input { min-width:0; }
      .detail-account-search-status { max-width:calc(100vw - 24px); }
      .account-head { flex-direction:column; } .toolbar { grid-template-columns:1fr; }
      .quick-mark-main,.quick-fields { grid-template-columns:1fr; }
      .quick-mark-main #saveBtn { grid-column:auto; }
      #financeMetrics,#frequencyMetrics { grid-template-columns:1fr; }
      .automation-summary { grid-template-columns:1fr; }
      .chart-tools { grid-template-columns:1fr; } .chart-row { grid-template-columns:1fr; } .chart-actions { justify-content:start; }
      .history-item { grid-template-columns:1fr; gap:5px; } .source-row { grid-template-columns:1fr 1fr; }
      .toxic-selector { grid-template-columns:1fr; } .toxic-result { grid-template-columns:72px 1fr; } .toxic-result-summary,.toxic-result-actions { grid-column:1/-1; } .toxic-chain-fact { grid-template-columns:1fr; gap:2px; }
      .toxic-sync-kpis { grid-template-columns:repeat(2,1fr); } .toxic-sync-kpis div:nth-child(odd) { border-left:0; }
      .funds-summary { grid-template-columns:1fr; } .funds-summary div { border-left:0; border-top:1px solid #214f78; } .funds-summary div:first-child { border-top:0; }
      .relationship-network-layout { grid-template-columns:1fr; } .relationship-network-detail { border-left:0; border-top:1px solid #174674; max-height:360px; } #relationshipNetworkGraph { min-height:540px; }
    }
  </style>
</head>
<body>
  <header class="topbar">
    <a class="back" href="/">← 返回台账</a>
    <div class="brand">账号详情</div>
    <form class="detail-account-search" id="detailAccountSearchForm" role="search">
      <input id="detailAccountSearch" inputmode="numeric" autocomplete="off" placeholder="搜索账号" aria-label="搜索账号" />
      <button type="submit" id="detailAccountSearchBtn">查询</button>
      <span class="detail-account-search-status" id="detailAccountSearchStatus" role="status" aria-live="polite"></span>
    </form>
  </header>
  <main>
    <section class="account-head">
      <div><div class="account-id" id="accountId"></div><div class="badges" id="badges"></div></div>
      <div class="head-meta" id="headMeta">正在读取最新订单...</div>
    </section>
    <section class="quick-mark">
      <div class="quick-mark-main">
        <div><div class="quick-action-head"><div class="quick-label">快捷标记</div><button type="button" class="manage-actions-btn" id="manageActionsBtn">管理快捷标记</button></div><div class="action-grid" id="actions"></div><div class="action-manager" id="actionManager" hidden><div class="action-add-row"><input id="newQuickAction" maxlength="40" placeholder="输入新的快捷标记" /><button type="button" id="addQuickActionBtn">添加</button><span class="status" id="actionManageStatus"></span></div><div class="action-manage-list" id="actionManageList"></div></div></div>
        <label class="quick-status">本地状态（选择即保存）<select id="status" disabled><option>读取本地台账...</option></select></label>
        <button class="primary" id="saveBtn">保存全部</button>
      </div>
      <div class="quick-fields">
        <label id="customActionLabel" style="display:none">自定义动作<input id="customAction" /></label>
        <label>当前分组<input id="group" /></label>
        <label>风险标签<input id="tags" /></label>
        <label>风险备注<input id="note" /></label>
        <label>处理人 / 来源<input id="owner" /></label>
        <label class="batch-toggle disabled" id="batchSameNameLabel"><input type="checkbox" id="batchSameName" disabled /><span id="batchSameNameText">同名账户加载中...</span></label>
      </div>
      <div class="quick-footer"><span class="status" id="markState">正在读取本地台账...</span><span class="status" id="saveStatus"></span></div>
    </section>
    <section class="toolbar">
      <label>平台<select id="platform"><option value="">全部平台</option></select></label>
      <label>服务器<select id="server"><option value="">全部服务器</option></select></label>
      <label>品种<select id="symbol"><option value="">全部品种</option></select></label>
      <button id="refreshBtn" class="primary">刷新指标</button>
      <button id="copyOriginBtn" type="button" hidden>跟单查询</button>
      <button id="eaCommentBtn" class="ea-query-entry" type="button" hidden>EA 查询</button>
      <button id="relationshipNetworkBtn" class="relationship-entry" type="button" hidden>关系网络</button>
      <button id="toxicBtn" class="toxic-entry" type="button">Toxic 检测</button>
      <button id="historicalFundsBtn" class="historical-funds-entry" type="button">历史资金回溯</button>
    </section>

    <section class="section" id="sameNameSection">
      <div class="section-head"><h2>同名账户</h2><div class="section-sub" id="sameNameStatus"></div></div>
      <div class="table-wrap"><table class="risk-table"><thead><tr><th>服务器</th><th>账号</th><th>数据库状态</th><th>本地标记</th><th>账户余额</th><th>净值</th><th>净入金</th><th>持仓盈亏</th><th>平仓盈亏+手续费+利息</th><th>清零+补偿+奖励</th><th>返佣</th><th>综合盈利</th><th>最高持仓量</th></tr></thead><tbody id="sameNameRows"></tbody></table></div>
    </section>

    <div class="dashboard-grid">
      <div class="dashboard-column">
        <section class="section" id="pnlVisualizationSection">
          <div class="section-head"><h2>盈亏趋势</h2><div class="section-sub" id="visualStatus"></div></div>
          <div class="viz-stack">
            <div class="viz-panel"><div class="viz-panel-head"><b>累计净盈亏</b><span id="pnlSummary"></span></div><div class="viz-canvas" id="pnlChart"><div class="viz-empty">正在读取订单...</div></div></div>
            <div class="viz-panel"><div class="viz-panel-head"><b>每日净盈亏</b><span>最近 30 个交易日</span></div><div class="viz-canvas compact" id="dailyChart"></div></div>
          </div>
        </section>

        <section class="section" id="frequencySection">
          <div class="section-head"><h2>高频与持仓分析</h2><div class="section-sub" id="frequencyStatus"></div></div>
          <div class="risk-summary" id="frequencyMetrics"></div>
          <div class="viz-panel" style="margin:0 12px 12px"><div class="viz-panel-head"><b>持仓时长分布</b><span>订单数 / 胜率 / 净盈亏</span></div><div class="viz-canvas compact" id="holdingChart"></div></div>
          <div class="table-wrap"><table class="risk-table"><thead><tr><th>持仓时段</th><th>订单数量</th><th>胜率</th><th>总盈利</th><th>总手数</th><th>平均每手盈利</th><th>平均每单盈利</th><th>盈利占比</th><th>平均交易手数</th></tr></thead><tbody id="frequencyRows"></tbody></table></div>
        </section>

        <section class="section" id="automationSection">
          <div class="section-head"><h2>跟单 / EA 分析</h2><div class="section-sub" id="automationStatus">异步分析中...</div></div>
          <div class="automation-grid">
            <div class="automation-panel">
              <div class="automation-panel-head"><b>跟单分析</b><span>来源账号 · 订单/手数占比 · 盈亏</span></div>
              <div class="automation-summary" id="copyAutomationSummary"></div>
              <div class="table-wrap"><table class="risk-table automation-table"><thead><tr><th>来源账号</th><th>订单</th><th>订单占全部</th><th>手数</th><th>手数占全部</th><th>毛盈亏</th><th>净盈亏</th></tr></thead><tbody id="copyAutomationRows"></tbody></table></div>
              <div class="automation-note" id="copyAutomationNote"></div>
            </div>
            <div class="automation-panel">
              <div class="automation-panel-head"><b>EA 分析</b><span>ExpertID / Magic · 订单/手数占比 · 盈亏</span></div>
              <div class="automation-summary" id="eaAutomationSummary"></div>
              <div class="table-wrap"><table class="risk-table automation-table"><thead><tr><th>ExpertID / Magic</th><th>平台 / 服务器</th><th>订单</th><th>订单占全部</th><th>手数</th><th>手数占全部</th><th>毛盈亏</th><th>净盈亏</th></tr></thead><tbody id="eaAutomationRows"></tbody></table></div>
              <div class="automation-note" id="eaAutomationNote"></div>
            </div>
          </div>
        </section>

        <section class="section">
          <div class="section-head"><h2>交易指标</h2><div class="section-sub" id="metricStatus"></div></div>
          <div class="metric-groups" id="metricGroups"></div>
        </section>
      </div>

      <div class="dashboard-column">
        <section class="section" id="financeSection">
          <div class="section-head"><h2>账户资金情况</h2><div class="section-sub" id="financeStatus"></div></div>
          <div class="risk-summary" id="financeMetrics"></div>
          <div class="risk-note">平仓盈亏已扣除交易手续费、利息和税费；返佣来自 CRM 返佣明细，不计入交易手续费。</div>
        </section>

        <section class="section" id="structureSection">
          <div class="section-head"><h2>胜负与费用结构</h2><div class="section-sub">净值口径</div></div>
          <div class="structure-grid">
            <div class="viz-panel"><div class="viz-panel-head"><b>订单胜负</b><span>盈利 / 亏损 / 持平</span></div><div class="donut-wrap"><div id="outcomeDonut"></div><div class="legend-list" id="outcomeLegend"></div></div></div>
            <div class="viz-panel"><div class="viz-panel-head"><b>费用构成</b><span>手续费 / Swap / 税费</span></div><div class="viz-canvas compact" id="feeBars"></div></div>
          </div>
        </section>

        <section class="section" id="symbolSection">
          <div class="section-head"><h2>品种表现</h2><div class="section-sub" id="symbolCount"></div></div>
          <div class="viz-panel" style="margin:12px"><div class="viz-panel-head"><b>品种净盈亏</b><span>按绝对盈亏排序</span></div><div class="viz-canvas" id="symbolBars"></div></div>
          <div class="table-wrap"><table><thead><tr><th>品种</th><th>订单</th><th>手数</th><th>净盈亏</th><th>净胜率</th></tr></thead><tbody id="symbolRows"></tbody></table></div>
        </section>

        <section class="section" id="ipSection">
          <div class="section-head"><h2>登录 IP 来源</h2><div class="section-head-actions"><span class="section-sub" id="ipStatus">独立加载中...</span><button type="button" id="refreshIpBtn">刷新 IP</button></div></div>
          <div class="risk-note" id="ipCoverage">数据库仅提供最后登录 IP，本地历史自功能上线后开始累计。</div>
          <div class="ip-summary" id="ipSummary"></div>
          <div class="ip-list" id="ipRows"><div class="empty-state">正在查询登录 IP...</div></div>
        </section>
      </div>
    </div>

    <div class="lower-grid">
      <section class="section">
        <div class="section-head"><h2>交易图表</h2><div class="section-sub" id="chartCount"></div></div>
        <div class="chart-product-grid" id="chartProducts"><div class="empty-state">正在解析品种...</div></div>
        <div class="chart-tools"><div><div class="chart-filter-grid"><label>图表品种<select id="chartSymbol"><option value="">全部品种</option></select></label><label>开始时间<input type="datetime-local" id="chartStart" /></label><label>结束时间<input type="datetime-local" id="chartEnd" /></label></div><div class="chart-filter-grid"><label><input type="checkbox" id="includeTimeline" /> 包含资金与 Credit 回放</label><label><input type="checkbox" id="refreshTimelineCache" disabled /> 刷新全量资金缓存</label></div><div class="muted">时间留空即按所选账号和品种的全量历史生成；勾选回放后首次只读读取完整资金历史并缓存，后续生成复用缓存。</div><div class="progress"><div id="jobProgress"></div></div><div class="job-text" id="jobText"></div></div><button class="primary" id="generateBtn">生成 K 线图</button></div>
        <div class="chart-list" id="charts"></div>
      </section>
      <div class="dashboard-column">
        <section class="section">
          <div class="section-head"><h2>数据来源</h2><div class="section-sub" id="dbSource"></div></div>
          <div class="source-list" id="sourceRows"></div>
        </section>
        <section class="section">
          <div class="section-head"><h2>历史记录</h2><div class="section-sub" id="historyCount"></div></div>
          <div class="history" id="history"></div>
        </section>
      </div>
    </div>

    <section class="section" id="inlineKlineSection" hidden>
      <div class="section-head"><h2>交易 K 线</h2><div class="section-sub" id="inlineKlineStatus">正在读取最近 300 笔订单与 M1 报价...</div></div>
      <iframe class="inline-kline-frame" id="inlineKlineFrame" title="账户交易 K 线" sandbox="allow-scripts"></iframe>
      <div class="inline-kline-note">此处直接由账户详情服务只读加载，不创建 K 线任务；下方“生成 K 线图”功能仍可用于全量、筛选或资金回放。</div>
    </section>
    <details class="section order-details" id="orderDetails">
      <summary><span>所有订单</span><span class="order-summary-meta" id="orderSummary">展开后加载</span></summary>
      <div class="table-wrap"><table class="order-table"><thead><tr><th>订单号</th><th>平台 / 服务器</th><th>品种</th><th>方向</th><th>原因</th><th>注释 / EA</th><th>开仓时间</th><th>平仓时间</th><th>持仓时间</th><th>手数</th><th>毛盈亏</th><th>手续费</th><th>利息</th><th>税费</th><th>净盈亏</th><th>币种</th></tr></thead><tbody id="orderRows"><tr><td colspan="16"><div class="empty-state">展开后读取该账号全部订单</div></td></tr></tbody></table></div>
      <div class="order-pager"><span class="status" id="orderPageStatus"></span><button type="button" id="orderPrev">上一页</button><button type="button" id="orderNext">下一页</button></div>
    </details>
  </main>
  <dialog id="previewDialog"><div class="dialog-head"><b>AI 交易复核图表</b><div class="dialog-head-actions"><a id="previewOpen" href="#" target="_blank" rel="noopener">新窗口打开</a><button id="closePreview">关闭</button></div></div><iframe id="previewFrame" title="AI 交易复核图表"></iframe></dialog>
  <dialog id="accountSourceDialog" class="account-source-dialog"><div class="dialog-head"><b>选择平台 / 服务器</b><button id="closeAccountSource" type="button">关闭</button></div><p id="accountSourceHint" class="status">该账号存在多个交易来源，请选择要查看的详细数据。</p><div id="accountSourceList" class="account-source-list"></div></dialog>
  <dialog id="toxicDialog" class="toxic-dialog">
    <div class="dialog-head"><b>Toxic 专项检测</b><button id="closeToxic" type="button">关闭</button></div>
    <div class="toxic-body">
      <div class="toxic-mode-row">
        <label class="toxic-mode"><input type="radio" name="toxicMode" value="selected" checked />人工指定项目</label>
        <label class="toxic-mode"><input type="radio" name="toxicMode" value="screen" />只做初筛</label>
        <span class="toxic-mode-note" id="toxicModeNote">仅运行勾选的专项算法；需要 Tick 的项目会按需读取候选时段报价。</span>
      </div>
      <div class="toxic-selector" id="toxicSelector"></div>
      <div class="toxic-controls-footer">
        <span class="status" id="toxicStatus">尚未运行检测</span>
        <button class="primary" id="startToxic" type="button">开始检测</button>
      </div>
      <div class="progress"><div id="toxicProgress"></div></div>
      <div class="toxic-results" id="toxicResults"><div class="toxic-empty">人工选择项目，或先运行初筛后从结果中继续深度检测。</div></div>
    </div>
  </dialog>
  <dialog id="historicalFundsDialog" class="toxic-dialog historical-funds-dialog">
    <div class="dialog-head"><b>历史资金回溯</b><div class="dialog-head-actions"><span class="status" id="historicalFundsStatus"></span><button id="closeHistoricalFunds" type="button">关闭</button></div></div>
    <div class="toxic-body">
      <div class="funds-note" id="historicalFundsNote">按当前选定的平台和服务器读取完整的只读资金、订单和日快照数据。此页面只复现事实，不自动判断赔付。</div>
      <div class="funds-summary" id="historicalFundsSummary"></div>
      <div class="funds-chart" id="historicalFundsChart"><div class="empty-state">等待查询...</div></div>
      <div class="funds-liquidations" id="historicalFundsLiquidations" hidden></div>
      <div class="funds-events-head"><b>资金与订单事件</b><span class="funds-note" id="historicalFundsEventNote"></span></div>
      <div class="table-wrap"><table class="risk-table funds-events-table"><thead><tr><th>时间</th><th>事件</th><th>订单 / Deal</th><th>品种</th><th>备注</th><th>余额变化</th><th>Credit变化</th><th>实现盈亏</th><th>事件后余额</th><th>事件后Credit</th><th>权益状态</th></tr></thead><tbody id="historicalFundsEvents"><tr><td colspan="11"><div class="empty-state">尚未查询</div></td></tr></tbody></table></div>
      <div class="funds-event-pager"><span id="historicalFundsPageStatus"></span><button id="historicalFundsPrev" type="button">上一页</button><button id="historicalFundsNext" type="button">下一页</button></div>
    </div>
  </dialog>
  <dialog id="copyDialog" class="toxic-dialog">
    <div class="dialog-head"><b>跟单查询</b><div class="dialog-head-actions"><span class="status" id="copyExportStatus"></span><button id="exportCopyReportBtn" type="button">导出 Excel</button><button id="closeCopyDialog" type="button">关闭</button></div></div>
    <div class="toxic-body">
      <div class="ea-range-controls">
        <label>开始时间<input id="copyOriginStart" type="datetime-local" /></label>
        <label>结束时间<input id="copyOriginEnd" type="datetime-local" /></label>
        <button id="applyCopyOriginRange" type="button">查询</button>
        <span class="ea-range-hint">按开仓时间筛选 CPT 和 Signal；留空为全历史</span>
      </div>
      <section class="copy-panel">
        <div class="copy-panel-title"><b>CPT 单主与跟单人员收益</b><span class="section-sub">按源订单定位单主，并汇总所有跟单账号的实际收益</span></div>
        <div class="risk-note" id="copyOriginStatus">等待查询...</div>
        <div id="copyOriginRows"><div class="empty-state">尚未查询 CPT 单主</div></div>
      </section>
      <section class="copy-panel">
        <div class="copy-panel-title"><b>Signal 整组收益</b><span class="section-sub">按 users.comment 中的 Signal #… IN 识别同组账户</span></div>
        <div class="risk-note" id="copyGroupStatus">等待查询...</div>
        <div id="copyGroupResults"><div class="empty-state">尚未查询 Signal 跟单组</div></div>
      </section>
    </div>
  </dialog>
  <dialog id="eaCommentDialog" class="toxic-dialog">
    <div class="dialog-head"><b>EA 查询</b><div class="dialog-head-actions"><span class="status" id="eaExportStatus"></span><button id="exportEaReportBtn" type="button">导出 Excel</button><button id="closeEaCommentDialog" type="button">关闭</button></div></div>
    <div class="toxic-body">
      <section class="copy-panel">
        <div class="copy-panel-title"><b>相同 Comment 的 EA 账户收益</b><span class="section-sub">跨服务器按 Comment 匹配；同服务器同时核对 ExpertID / MAGIC，并逐账户列出匹配线索</span></div>
        <div class="ea-range-controls">
          <label>开始时间<input id="eaCommentStart" type="datetime-local" /></label>
          <label>结束时间<input id="eaCommentEnd" type="datetime-local" /></label>
          <button id="applyEaCommentRange" type="button">查询</button>
          <span class="ea-range-hint">按开仓时间筛选；留空为全历史</span>
        </div>
        <div class="risk-note" id="eaCommentStatus">等待查询...</div>
        <div id="eaCommentResults"><div class="empty-state">尚未查询 EA comment</div></div>
      </section>
    </div>
  </dialog>
  <dialog id="relationshipNetworkDialog" class="relationship-network-dialog">
    <div class="dialog-head"><b>账户关系网络</b><div class="dialog-head-actions"><button id="resetRelationshipNetworkBtn" type="button">恢复视图</button><button id="closeRelationshipNetworkDialog" type="button">关闭</button></div></div>
    <div class="relationship-network-body">
      <div class="relationship-network-controls"><div class="relationship-network-filters" id="relationshipNetworkFilters"></div><span class="status" id="relationshipNetworkStatus">点击关系网络后生成事实证据。</span></div>
      <div class="relationship-network-layout">
        <div class="relationship-network-graph-wrap">
          <canvas id="relationshipNetworkGraph" class="relationship-network-graph" width="1000" height="620" role="img" aria-label="账户关系证据网络">账户关系证据网络</canvas>
          <div class="relationship-network-legend"><span><i class="account"></i>交易账户</span><span><i class="clue"></i>数据线索</span><span><i class="group"></i>聚合关系</span></div>
        </div>
        <aside class="relationship-network-detail" aria-live="polite"><div class="relationship-detail-kind" id="relationshipDetailKind">尚未选择</div><h3 class="relationship-detail-title" id="relationshipDetailTitle">点击节点或连线查看事实依据</h3><div class="relationship-detail-meta" id="relationshipDetailMeta"></div><div class="relationship-evidence-title">可核验证据</div><div class="relationship-evidence" id="relationshipEvidence"></div><div class="relationship-evidence-title">范围与限制</div><ul class="relationship-limitations" id="relationshipLimitations"></ul></aside>
      </div>
    </div>
  </dialog>
  <script>
    const LOGIN=__ACCOUNT_LOGIN_JSON__;
    const $=id=>document.getElementById(id);
    const initialParams=new URLSearchParams(location.search);
    const state={detail:null,initialFilters:{platform:initialParams.get('platform')||'',server:initialParams.get('server')||'',symbol:initialParams.get('symbol')||''},riskRequest:0,automationRequest:0,selectedAction:'',jobTimer:null,inlineKlineKey:null,ledgerLoaded:false,formDirty:false,chartFiltersInitialized:false,chartSymbols:[],quickActions:[],protectedActions:['自定义'],sameNameAccounts:[LOGIN],orders:{page:1,pages:1,loading:false,loaded:false},ips:{loading:false},dialogCache:{copy:new Map(),ea:new Map(),relationship:new Map()},relationshipNetwork:null,historicalFunds:{loading:false,data:null,page:1,pageSize:100},toxic:{jobTimer:null,running:false,lastResult:null},accountLookupMatches:[]};
    const TOXIC_TYPES=[
      {id:'market_pushing',label:'推盘',tick:true},{id:'quote_latency_arbitrage',label:'报价延迟套利',tick:true},{id:'cross_platform_spread_arbitrage',label:'跨平台点差套利',tick:true},
      {id:'rebate_churning',label:'刷返佣'},{id:'bonus_arbitrage',label:'赠金套利'},{id:'short_close_trading',label:'短平交易'},
      {id:'internal_lock_arbitrage',label:'平台内多账户对锁'},{id:'high_leverage_lock_arbitrage',label:'高杠杆锁仓套利'},
      {id:'weekend_gap_trading',label:'周末跳空交易'},{id:'open_betting',label:'赌开盘'},{id:'news_event_betting',label:'新闻 / 高波动赌博',tick:true}
    ];
    const esc=value=>String(value??'').replace(/[&<>"']/g,c=>({"&":"&amp;","<":"&lt;",">":"&gt;",'"':"&quot;","'":"&#39;"}[c]));
    const num=(value,digits=2)=>new Intl.NumberFormat('zh-CN',{maximumFractionDigits:digits}).format(Number(value||0));
    const money=value=>`${Number(value||0)>0?'+':''}${num(value,2)}`;
    const amount=value=>new Intl.NumberFormat('zh-CN',{minimumFractionDigits:2,maximumFractionDigits:2}).format(Number(value||0));
    const pct=value=>value===null||value===undefined?'-':`${num(value,1)}%`;
    const duration=value=>{value=Number(value||0);if(!value)return '-';if(value<60)return `${num(value,0)} 秒`;if(value<3600)return `${num(value/60,1)} 分`;if(value<86400)return `${num(value/3600,1)} 小时`;return `${num(value/86400,1)} 天`;};
    const datetimeLocal=value=>String(value||'').replace(' ','T').slice(0,16);
    const databaseTime=(value,end=false)=>value?`${String(value).replace('T',' ').slice(0,16)}:${end?'59':'00'}`:'';
    async function json(url,options={}){if(options.body&&!options.headers)options.headers={'Content-Type':'application/json; charset=utf-8'};const res=await fetch(url,options);const data=await res.json();if(!res.ok||!data.ok)throw new Error(data.error||`HTTP ${res.status}`);return data;}
    function optionList(select,items,placeholder){const current=select.value;select.innerHTML=`<option value="">${esc(placeholder)}</option>`+items.map(item=>{const value=typeof item==='string'?item:item.value;const label=typeof item==='string'?item:item.label;return `<option value="${esc(value)}">${esc(label)}</option>`;}).join('');select.value=current;}
    function metric(label,value,cls=''){return `<div class="metric"><span>${label}</span><b class="${cls}">${value}</b></div>`;}
    function group(title,items){return `<div class="metric-group"><h3>${title}</h3><div class="metrics">${items.join('')}</div></div>`;}
    function profitClass(value){return Number(value)>0?'positive':Number(value)<0?'negative':'';}
    function renderMetrics(db){
      const m=db.metrics||{}, meta=db.accountMeta||{}, unit=meta.displayCurrency?` (${meta.displayCurrency})`:'';
      if(!db.exists){$("metricGroups").innerHTML=`<div class="empty-state" style="grid-column:1/-1">${esc(db.error||'账户暂未做单')}</div>`;$("symbolRows").innerHTML='<tr><td colspan="5"><div class="empty-state">暂无订单</div></td></tr>';$("sourceRows").innerHTML='<div class="empty-state">暂无来源</div>';return;}
      $("metricGroups").innerHTML=[
        group('交易概览',[metric('订单数',num(m.orderCount,0)),metric('可画图',num(m.chartableOrderCount,0)),metric('品种数',num(m.symbolCount,0)),metric('活跃天数',num(m.activeDays,0))]),
        group('盈亏表现',[metric(`平仓净盈亏${unit}`,money(m.netProfit),profitClass(m.netProfit)),metric(`平仓毛盈亏${unit}`,money(m.grossProfit),profitClass(m.grossProfit)),metric('净胜率',pct(m.winRate)),metric('净盈利 / 净亏损',`${num(m.winningOrders,0)} / ${num(m.losingOrders,0)}`),metric(`平均 / 中位净盈亏${unit}`,`${money(m.averageProfit)} / ${money(m.medianProfit)}`)]),
        group('持仓行为',[metric('平均持仓',duration(m.averageHoldingSeconds)),metric('中位持仓',duration(m.medianHoldingSeconds)),metric('5 分钟内',pct(m.shortHoldingRatio)),metric('1 分钟内',pct(m.oneMinuteHoldingRatio))]),
        group('手数行为',[metric('总手数',num(m.totalVolume,4)),metric('平均手数',num(m.averageVolume,4)),metric('最大手数',num(m.maxVolume,4)),metric('日均订单',num(m.ordersPerActiveDay,1))]),
        group('短线 / 高频特征',[metric('分钟峰值',num(m.maxOrdersInOneMinute,0)),metric('平均下单间隔',duration(m.averageOrderGapSeconds)),metric('短持仓比例',pct(m.shortHoldingRatio)),metric('活跃交易日',num(m.activeDays,0))]),
        group('佣金 / 费用占比',[metric(`手续费${unit}`,money(m.commissionTotal),profitClass(m.commissionTotal)),metric(`利息 / Swap${unit}`,money(m.swapTotal),profitClass(m.swapTotal)),metric(`税费${unit}`,money(m.taxesTotal),profitClass(m.taxesTotal)),metric('费用 / 毛盈亏',pct(m.feeToProfitRatio))])
      ].join('');
      $("symbolCount").textContent=`${m.bySymbol?.length||0} 个品种`;
      $("symbolRows").innerHTML=(m.bySymbol||[]).map(row=>`<tr><td><b>${esc(row.symbol)}</b></td><td>${num(row.orders,0)}</td><td>${num(row.volume,4)}</td><td class="${profitClass(row.profit)}">${money(row.profit)}</td><td>${pct(row.winRate)}</td></tr>`).join('')||'<tr><td colspan="5"><div class="empty-state">暂无品种统计</div></td></tr>';
      $("sourceRows").innerHTML=(m.bySource||[]).map(row=>`<div class="source-row"><b>${esc(row.platform)}${row.currency?` · ${esc(row.currency)}`:''}</b><span>${esc(row.server)}</span><span>${num(row.orders,0)} 单</span><span class="${profitClass(row.profit)}">${money(row.profit)}</span></div>`).join('')||'<div class="empty-state">暂无来源</div>';
    }
    function svgEmpty(text){return `<div class="viz-empty">${esc(text)}</div>`;}
    function lineChart(points){
      if(!points?.length)return svgEmpty('暂无可绘制的盈亏数据');
      const width=720,height=245,left=55,right=18,top=18,bottom=31,values=points.map(row=>Number(row.value||0)),low=Math.min(0,...values),high=Math.max(0,...values),range=(high-low)||1,innerW=width-left-right,innerH=height-top-bottom;
      const x=index=>left+(points.length===1?innerW/2:index/(points.length-1)*innerW),y=value=>top+(high-Number(value||0))/range*innerH;
      const coords=points.map((row,index)=>`${x(index).toFixed(1)},${y(row.value).toFixed(1)}`).join(' '),zero=y(0),grid=[0,.25,.5,.75,1].map(t=>{const value=high-range*t,yy=y(value);return `<line class="viz-grid-line" x1="${left}" y1="${yy}" x2="${width-right}" y2="${yy}"/><text class="viz-axis-text" x="${left-7}" y="${yy+3}" text-anchor="end">${esc(num(value,0))}</text>`;}).join('');
      const markers=points.map((row,index)=>{if(points.length>36&&index%Math.ceil(points.length/30)!==0&&index!==points.length-1)return '';return `<circle cx="${x(index)}" cy="${y(row.value)}" r="3" fill="${Number(row.change)>=0?'#34d399':'#ff6472'}"><title>${esc(row.time||`第 ${row.index} 笔`)} · 变动 ${money(row.change)} · 累计 ${money(row.value)}${row.adjustment?' · 费用调整':''}</title></circle>`;}).join('');
      return `<svg class="viz-svg" viewBox="0 0 ${width} ${height}" role="img" aria-label="累计净盈亏曲线"><defs><linearGradient id="pnlArea" x1="0" y1="0" x2="0" y2="1"><stop offset="0" stop-color="#168cff" stop-opacity=".42"/><stop offset="1" stop-color="#168cff" stop-opacity=".02"/></linearGradient></defs>${grid}<line class="viz-zero-line" x1="${left}" y1="${zero}" x2="${width-right}" y2="${zero}"/><polygon points="${left},${zero} ${coords} ${width-right},${zero}" fill="url(#pnlArea)"/><polyline points="${coords}" fill="none" stroke="#28a9ff" stroke-width="2.5" stroke-linejoin="round" stroke-linecap="round" filter="drop-shadow(0 0 5px #168cff88)"/>${markers}<text class="viz-axis-text" x="${left}" y="${height-8}">${esc(points[0].time||'开始')}</text><text class="viz-axis-text" x="${width-right}" y="${height-8}" text-anchor="end">${esc(points[points.length-1].time||'结束')}</text></svg>`;
    }
    function dailyBarChart(rows){
      rows=(rows||[]).slice(-30);if(!rows.length)return svgEmpty('暂无每日盈亏数据');
      const width=720,height=185,left=42,right=12,top=12,bottom=32,innerW=width-left-right,innerH=height-top-bottom,maxAbs=Math.max(1,...rows.map(row=>Math.abs(Number(row.profit||0)))),zero=top+innerH/2,slot=innerW/rows.length,barW=Math.max(3,slot*.62);
      const bars=rows.map((row,index)=>{const value=Number(row.profit||0),barH=Math.abs(value)/maxAbs*(innerH/2-4),yy=value>=0?zero-barH:zero,color=value>=0?'#34d399':'#ff6472',label=(index===0||index===rows.length-1||index%Math.max(1,Math.ceil(rows.length/6))===0)?`<text class="viz-axis-text" x="${left+slot*index+slot/2}" y="${height-9}" text-anchor="middle">${esc(String(row.date||'').slice(5))}</text>`:'';return `<rect x="${left+slot*index+(slot-barW)/2}" y="${yy}" width="${barW}" height="${Math.max(1,barH)}" rx="2" fill="${color}"><title>${esc(row.date)} · ${money(value)}</title></rect>${label}`;}).join('');
      return `<svg class="viz-svg" viewBox="0 0 ${width} ${height}" role="img" aria-label="每日净盈亏柱图"><line class="viz-zero-line" x1="${left}" y1="${zero}" x2="${width-right}" y2="${zero}"/>${bars}<text class="viz-axis-text" x="${left-6}" y="${top+4}" text-anchor="end">${esc(num(maxAbs,0))}</text><text class="viz-axis-text" x="${left-6}" y="${top+innerH}" text-anchor="end">-${esc(num(maxAbs,0))}</text></svg>`;
    }
    function donutChart(outcomes){
      const items=[{label:'盈利',value:Number(outcomes?.winning||0),color:'#34d399'},{label:'亏损',value:Number(outcomes?.losing||0),color:'#ff6472'},{label:'持平',value:Number(outcomes?.breakeven||0),color:'#f2ab35'}],total=items.reduce((sum,row)=>sum+row.value,0),radius=54,circ=2*Math.PI*radius;let offset=0;
      const rings=total?items.map(row=>{const length=row.value/total*circ,part=`<circle cx="75" cy="75" r="${radius}" fill="none" stroke="${row.color}" stroke-width="16" stroke-dasharray="${length} ${circ-length}" stroke-dashoffset="${-offset}" transform="rotate(-90 75 75)"><title>${row.label} ${row.value} 笔</title></circle>`;offset+=length;return part;}).join(''):`<circle cx="75" cy="75" r="${radius}" fill="none" stroke="#173e6b" stroke-width="16"/>`;
      $("outcomeDonut").innerHTML=`<svg class="donut-chart" viewBox="0 0 150 150">${rings}<text class="donut-center-main" x="75" y="73">${num(total,0)}</text><text class="donut-center-sub" x="75" y="91">订单</text></svg>`;
      $("outcomeLegend").innerHTML=items.map(row=>`<div class="legend-item"><i class="legend-dot" style="color:${row.color};background:${row.color}"></i><span>${row.label}</span><b>${num(row.value,0)} · ${total?pct(row.value/total*100):'0%'}</b></div>`).join('');
    }
    function barList(rows,labelKey,valueKey,valueFormatter=money){
      if(!rows?.length)return svgEmpty('暂无数据');const max=Math.max(1,...rows.map(row=>Math.abs(Number(row[valueKey]||0))));
      return `<div class="bar-list">${rows.map(row=>{const value=Number(row[valueKey]||0),color=value>=0?'#34d399':'#ff6472',width=Math.max(1,Math.abs(value)/max*100);return `<div class="bar-row"><span class="bar-label" title="${esc(row[labelKey])}">${esc(row[labelKey])}</span><span class="bar-track"><i class="bar-fill" style="display:block;width:${width}%;color:${color};background:${color}"></i></span><b class="bar-value ${profitClass(value)}">${valueFormatter(value)}</b></div>`;}).join('')}</div>`;
    }
    function holdingBars(rows){
      if(!rows?.length)return svgEmpty('暂无持仓数据');const max=Math.max(1,...rows.map(row=>Number(row.orders||0)));
      return `<div class="bar-list">${rows.map(row=>`<div class="bar-row"><span class="bar-label">${esc(row.label)}</span><span class="bar-track"><i class="bar-fill" style="display:block;width:${Math.max(1,Number(row.orders||0)/max*100)}%;color:#28a9ff;background:#28a9ff"></i></span><b class="bar-value">${num(row.orders,0)} 单 · ${pct(row.winRate)} · <span class="${profitClass(row.grossProfit)}">${money(row.grossProfit)}</span></b></div>`).join('')}</div>`;
    }
    function renderVisualizations(db){
      const v=db.visualizations||{},meta=db.accountMeta||{},unit=meta.displayCurrency||'';
      $("visualStatus").textContent=db.exists?`${num(db.metrics?.orderCount,0)} 笔订单 · ${unit}`:'暂无订单';
      $("pnlSummary").innerHTML=`净盈亏 <b class="${profitClass(v.netTotal)}">${money(v.netTotal)}</b> · 最大回撤 <b class="negative">-${num(v.maxDrawdown,2)}</b>`;
      $("pnlChart").innerHTML=lineChart(v.pnlSeries||[]);$("dailyChart").innerHTML=dailyBarChart(v.dailyPnl||[]);donutChart(v.outcomes||{});
      $("feeBars").innerHTML=barList(v.feeBreakdown||[],'label','value',money);$("symbolBars").innerHTML=barList(v.symbolPerformance||[],'symbol','profit',money);$("holdingChart").innerHTML=holdingBars(v.holdingBuckets||[]);
    }
    function renderIps(data){
      const rows=data.records||[],missing=(data.sources||[]).filter(row=>row.accountExists&&!row.available),countries=new Set(rows.map(row=>row.geo?.country).filter(Boolean)),isps=new Set(rows.map(row=>row.geo?.isp).filter(Boolean));
      $("ipStatus").textContent=`${rows.length} 个已知 IP${missing.length?` · ${missing.length} 个来源未导出 IP`:''} · ${esc(data.refreshedAt||'-')}`;$("ipCoverage").textContent=data.coverage?.notice||'数据库仅提供最后登录 IP，本地历史自功能上线后开始累计。';
      $("ipSummary").innerHTML=[`${rows.length} 个历史 IP`,`${countries.size} 个国家/地区`,`${isps.size} 个运营商`].map(value=>`<span class="ip-chip">${esc(value)}</span>`).join('');
      const knownHtml=rows.map(row=>{const geo=row.geo||{},location=[geo.country,geo.region,geo.city].filter(Boolean).join(' · ')||({private:'内网或保留地址',invalid:'IP 格式无效',unavailable:'归属地查询暂不可用'}[geo.status]||'归属地未知'),network=[geo.isp,geo.asn?`AS${geo.asn}`:''].filter(Boolean).join(' · ')||'-';return `<div class="ip-row"><div><div class="ip-address">${esc(row.ip)}</div><div class="ip-meta">${esc([row.platform,row.server].filter(Boolean).join(' / '))}</div></div><div class="ip-meta">数据库最后访问<br><b>${esc(row.lastAccessAt||'-')}</b><br>本地发现 ${esc(row.firstSeenAt||'-')}<br>最后观察 ${esc(row.lastSeenAt||'-')}</div><div class="ip-location"><b>${esc(location)}</b>${esc(network)}</div></div>`;}).join('');
      const missingHtml=missing.map(row=>`<div class="ip-row"><div><div class="ip-address">IP 未导出</div><div class="ip-meta">${esc([row.platform,row.server].filter(Boolean).join(' / '))}</div></div><div class="ip-meta">数据库最后登录<br><b>${esc(row.lastAccessAt||'-')}</b></div><div class="ip-location"><b>当前无法获取</b>${esc(row.reason||'数据源未提供 IP 字段')}</div></div>`).join('');
      $("ipRows").innerHTML=knownHtml+missingHtml||'<div class="empty-state">数据库暂无登录 IP 记录</div>';
    }
    async function loadIps(){if(state.ips.loading)return;state.ips.loading=true;$("refreshIpBtn").disabled=true;$("ipStatus").textContent='正在查询 IP 与归属地...';try{renderIps(await json(`/api/accounts/by-login/${encodeURIComponent(LOGIN)}/login-ips`));}catch(err){$("ipStatus").textContent='查询失败';$("ipRows").innerHTML=`<div class="empty-state">${esc(err.message)}</div>`;}finally{state.ips.loading=false;$("refreshIpBtn").disabled=false;}}
    function automationSummary(section){
      const labels=[['订单',num(section.orders,0)],['订单占全部',pct(section.orderRatio)],['手数占全部',pct(section.volumeRatio)],['毛盈亏',money(section.grossProfit)],['净盈亏',money(section.netProfit)],['手数',num(section.volume,4)]];
      return labels.map(([label,value])=>riskStat(label,value,label==='毛盈亏'||label==='净盈亏'?section[label==='毛盈亏'?'grossProfit':'netProfit']:null)).join('');
    }
    function renderAutomation(data){
      const copy=data.copy||{},ea=data.ea||{};
      $("automationStatus").textContent=`${num(data.totalOrders,0)} 笔订单 · ${num(data.totalVolume,4)} 手 · ${esc(data.refreshedAt||'-')}`;
      $("copyAutomationSummary").innerHTML=automationSummary(copy);
      $("eaAutomationSummary").innerHTML=automationSummary(ea);
      const copyRows=copy.origins||[];
      $("copyAutomationRows").innerHTML=copyRows.map(row=>{
        const account=row.unresolved?esc(row.account):`<a class="account-link" href="/account/${encodeURIComponent(row.account)}?platform=${encodeURIComponent(row.platform||'')}&server=${encodeURIComponent(row.server||'')}">${esc(row.account)}</a>`;
        return `<tr><td>${account}${row.server?`<small class="muted">${esc(row.server)}</small>`:''}</td><td>${num(row.orders,0)}</td><td>${pct(row.orderRatio)}</td><td>${num(row.volume,4)}</td><td>${pct(row.volumeRatio)}</td><td class="${profitClass(row.grossProfit)}">${money(row.grossProfit)}</td><td class="${profitClass(row.netProfit)}"><b>${money(row.netProfit)}</b></td></tr>`;
      }).join('')||'<tr><td colspan="7"><div class="empty-state">未识别到跟单订单</div></td></tr>';
      const hasSignalSource=copyRows.some(row=>row.sourceType==='signal');
      $("copyAutomationNote").textContent=copy.detected?(copy.errors?.length?`部分来源查询失败：${copy.errors.join('；')}`:(hasSignalSource?'CPT 跟单按订单号反查发起账号；Signal 跟单当前只能确认 Signal 标识，源账户未在导出数据中提供。':'来源按 CPT 订单号归属；无法反查的同步订单会列为“来源未解析”。')):'当前筛选范围内没有识别到跟单订单。';
      const eaRows=ea.groups||[];
      $("eaAutomationRows").innerHTML=eaRows.map(row=>`<tr><td><b>${esc(row.expertId)}</b>${row.symbols?.length?`<small class="muted">${esc(row.symbols.join('、'))}</small>`:''}</td><td>${esc([row.platform,row.server].filter(Boolean).join(' / ')||'-')}</td><td>${num(row.orders,0)}</td><td>${pct(row.orderRatio)}</td><td>${num(row.volume,4)}</td><td>${pct(row.volumeRatio)}</td><td class="${profitClass(row.grossProfit)}">${money(row.grossProfit)}</td><td class="${profitClass(row.netProfit)}"><b>${money(row.netProfit)}</b></td></tr>`).join('')||'<tr><td colspan="8"><div class="empty-state">未识别到 EA 订单</div></td></tr>';
      $("eaAutomationNote").textContent=ea.detected?'按 ExpertID / Magic 分组；未导出的 ID 统一归入“EA（未标记ID）”。':'当前筛选范围内没有识别到 EA 订单。';
    }
    async function loadAutomation(q){const request=++state.automationRequest;try{const data=await json(`/api/accounts/by-login/${encodeURIComponent(LOGIN)}/automation-analysis?${q}`);if(request!==state.automationRequest)return;renderAutomation(data);}catch(err){if(request!==state.automationRequest)return;$("automationStatus").textContent='分析失败';$("copyAutomationRows").innerHTML=`<tr><td colspan="7"><div class="empty-state">${esc(err.message)}</div></td></tr>`;$("eaAutomationRows").innerHTML=`<tr><td colspan="8"><div class="empty-state">${esc(err.message)}</div></td></tr>`;}}
    function riskStat(label,value,rawValue=null){return `<div class="risk-stat"><span>${esc(label)}</span><b class="${rawValue===null?'':profitClass(rawValue)}">${value}</b></div>`;}
    function renderRiskPanels(panels){
      panels=panels||{};
      if(!panels.available){
        const reason=esc(panels.reason||'当前账号暂无风控面板数据');
        $("financeStatus").textContent=$("frequencyStatus").textContent=$("sameNameStatus").textContent='不可用';
        $("financeMetrics").innerHTML=`<div class="empty-state" style="grid-column:1/-1">${reason}</div>`;
        $("frequencyMetrics").innerHTML=`<div class="empty-state" style="grid-column:1/-1">${reason}</div>`;
        $("frequencyRows").innerHTML='<tr><td colspan="9"><div class="empty-state">暂无高频分析</div></td></tr>';
        $("sameNameRows").innerHTML='<tr><td colspan="13"><div class="empty-state">暂无同名账户</div></td></tr>';
        updateBatchControl([]);
        return;
      }
      const f=panels.finance||{}, unit=f.displayCurrency?` (${f.displayCurrency})`:'';
      $("financeStatus").textContent=`${panels.server||''}${f.currency?` · ${f.currency}`:''}`;
      $("financeMetrics").innerHTML=[
        riskStat(`综合盈利${unit}`,amount(f.comprehensiveProfit),f.comprehensiveProfit),riskStat(`净入金${unit}`,amount(f.netDeposit),f.netDeposit),riskStat(`负值清零${unit}`,amount(f.negativeBalanceClear),f.negativeBalanceClear),riskStat(`利息${unit}`,amount(f.interest),f.interest),
        riskStat(`平仓盈亏${unit}`,amount(f.closedNetProfit),f.closedNetProfit),riskStat('综合盈利率',f.comprehensiveProfitRate===null||f.comprehensiveProfitRate===undefined?'-':amount(f.comprehensiveProfitRate)),riskStat(`手续费${unit}`,amount(f.tradingFees),f.tradingFees),riskStat('爆仓率',f.liquidationRate===null||f.liquidationRate===undefined?'暂无数据':pct(f.liquidationRate)),
        riskStat('爆仓金额比',f.liquidationAmountRatio===null||f.liquidationAmountRatio===undefined?'暂无数据':pct(f.liquidationAmountRatio)),riskStat(`补偿${unit}`,amount(f.compensation),f.compensation),riskStat(`返佣${unit}`,amount(f.rebate),f.rebate),riskStat(`奖励${unit}`,amount(f.reward),f.reward)
      ].join('');

      const h=panels.highFrequency||{}, buckets=h.buckets||[];
      $("frequencyStatus").textContent=`${num(h.orderCount,0)} 笔平仓订单`;
      $("frequencyMetrics").innerHTML=[riskStat('平均持仓时间',`${num(h.averageHoldingMinutes,0)} 分钟`),riskStat('盈利单平均持仓时间',`${num(h.winningAverageHoldingMinutes,2)} 分钟`),riskStat('亏损单平均持仓时间',`${num(h.losingAverageHoldingMinutes,2)} 分钟`),riskStat('高频订单占比',`${num(h.highFrequencyOrderRatio,0)}%`)].join('');
      $("frequencyRows").innerHTML=buckets.map(row=>`<tr><td>${esc(row.label)}</td><td>${num(row.orders,0)}</td><td>${num(row.winRate,0)}%</td><td class="${profitClass(row.grossProfit)}">${amount(row.grossProfit)}</td><td>${amount(row.volume)}</td><td class="${profitClass(row.averageProfitPerLot)}">${amount(row.averageProfitPerLot)}</td><td class="${profitClass(row.averageProfitPerOrder)}">${amount(row.averageProfitPerOrder)}</td><td>${num(row.profitShare,0)}%</td><td>${amount(row.averageVolume)}</td></tr>`).join('')||'<tr><td colspan="9"><div class="empty-state">暂无高频分析</div></td></tr>';

      const rows=panels.sameName||[], totals=panels.sameNameTotals||{};
      $("sameNameStatus").textContent=`${rows.length} 个关联账号 · 仅展示账号和交易数据`;
      const sameRow=row=>`<tr class="${String(row.account)===String(LOGIN)?'current-account':''}"><td>${esc([row.platform,row.server].filter(Boolean).join(' · '))}</td><td><a class="account-link" href="/account/${encodeURIComponent(row.account)}">${esc(row.account)}</a>${row.currency?` · ${esc(row.currency)}`:''}</td><td><b>${esc(row.databaseStatus||'-')}</b></td><td>${esc(row.localStatus||'-')}</td><td>${amount(row.balance)}</td><td>${amount(row.equity)}</td><td>${amount(row.netDeposit)}</td><td class="${profitClass(row.holdingProfit)}">${amount(row.holdingProfit)}</td><td class="${profitClass(row.closedNetProfit)}">${amount(row.closedNetProfit)}</td><td class="${profitClass(row.adjustments)}">${amount(row.adjustments)}</td><td>${amount(row.rebate)}</td><td class="${profitClass(row.comprehensiveProfit)}">${amount(row.comprehensiveProfit)}</td><td>${amount(row.highestHoldingVolume)}</td></tr>`;
      const totalRow=rows.length?`<tr class="total"><td></td><td>合计</td><td>-</td><td>-</td><td>${amount(totals.balance)}</td><td>${amount(totals.equity)}</td><td>${amount(totals.netDeposit)}</td><td>${amount(totals.holdingProfit)}</td><td>${amount(totals.closedNetProfit)}</td><td>${amount(totals.adjustments)}</td><td>${amount(totals.rebate)}</td><td>${amount(totals.comprehensiveProfit)}</td><td>-</td></tr>`:'';
      $("sameNameRows").innerHTML=rows.map(sameRow).join('')+totalRow||'<tr><td colspan="13"><div class="empty-state">暂无同名账户</div></td></tr>';
      updateBatchControl(rows);
    }
    function updateBatchControl(rows){
      const accounts=[...new Set((rows||[]).map(row=>String(row.account||'')).filter(Boolean))];if(!accounts.includes(String(LOGIN)))accounts.unshift(String(LOGIN));state.sameNameAccounts=accounts;
      const enabled=accounts.length>1;$("batchSameName").disabled=!enabled;if(!enabled)$("batchSameName").checked=false;$("batchSameNameLabel").classList.toggle('disabled',!enabled);$("batchSameNameText").textContent=enabled?`同时保存到 ${accounts.length} 个同名账户（仅本地台账）`:'没有其他同名账户';
    }
    function selectAction(value,dirty=false){state.selectedAction=value||'';if(dirty)state.formDirty=true;document.querySelectorAll('.action-btn').forEach(btn=>btn.classList.toggle('active',btn.dataset.action===state.selectedAction));const custom=state.selectedAction==='自定义';$("customActionLabel").style.display=custom?'block':'none';}
    function renderActionButtons(){
      $("actions").innerHTML=state.quickActions.map(value=>`<button type="button" class="action-btn" data-action="${esc(value)}">${esc(value)}</button>`).join('');
      document.querySelectorAll('.action-btn').forEach(btn=>btn.addEventListener('click',()=>selectAction(btn.dataset.action,true)));
    }
    function renderActionManager(){
      const protectedSet=new Set(state.protectedActions||[]);
      $("actionManageList").innerHTML=state.quickActions.map(value=>protectedSet.has(value)?`<span class="manage-action-chip protected"><span>${esc(value)}</span></span>`:`<span class="manage-action-chip"><span>${esc(value)}</span><button type="button" data-delete-action="${esc(value)}" title="删除 ${esc(value)}" aria-label="删除 ${esc(value)}">×</button></span>`).join('');
      document.querySelectorAll('[data-delete-action]').forEach(btn=>btn.addEventListener('click',()=>deleteQuickAction(btn.dataset.deleteAction)));
    }
    function applyQuickActions(actions){
      const selected=state.selectedAction,customValue=$("customAction").value;state.quickActions=[...new Set((actions||[]).filter(Boolean))];renderActionButtons();renderActionManager();
      if(state.quickActions.includes(selected))selectAction(selected,false);else if(selected&&selected!=='自定义'){$("customAction").value=selected;selectAction('自定义',false);}else{$("customAction").value=customValue;selectAction(selected,false);}
    }
    function renderActions(actions,record){
      state.quickActions=[...new Set((actions||[]).filter(Boolean))];renderActionButtons();renderActionManager();const action=record?.['建议动作']||'';selectAction(state.quickActions.includes(action)?action:(action?'自定义':''),false);if(action&&!state.quickActions.includes(action))$("customAction").value=action;
    }
    async function addQuickAction(){const action=$("newQuickAction").value.trim();if(!action){$("actionManageStatus").textContent='请输入快捷标记';return;}$("addQuickActionBtn").disabled=true;$("actionManageStatus").textContent='正在添加...';try{const data=await json('/api/quick-actions',{method:'POST',body:JSON.stringify({action})});state.protectedActions=data.protected||['自定义'];applyQuickActions(data.actions);$("newQuickAction").value='';$("actionManageStatus").textContent=`已添加 ${action}`;}catch(err){$("actionManageStatus").textContent=err.message;}finally{$("addQuickActionBtn").disabled=false;}}
    async function deleteQuickAction(action){$("actionManageStatus").textContent=`正在删除 ${action}...`;try{const data=await json(`/api/quick-actions/${encodeURIComponent(action)}`,{method:'DELETE'});state.protectedActions=data.protected||['自定义'];applyQuickActions(data.actions);$("actionManageStatus").textContent=`已删除 ${action}；历史记录不受影响`;}catch(err){$("actionManageStatus").textContent=err.message;}}
    function toggleActionManager(){const panel=$("actionManager"),opening=panel.hidden;panel.hidden=!opening;$("manageActionsBtn").textContent=opening?'完成管理':'管理快捷标记';if(opening)$("newQuickAction").focus();}
    function renderRecord(detail,populate=true){
      const r=detail.record||{};$("markState").textContent=detail.marked?`记录 ${r['记录ID']||''}`:'尚未加入台账';
      if(populate){$("group").value=r['当前分组']||'';$("tags").value=r['风险标签']||'';$("note").value=r['风险/问题备注']||'';$("owner").value=r['处理人/来源']||'';$("status").innerHTML=(detail.statuses||[]).map(v=>`<option>${esc(v)}</option>`).join('');$("status").value=r['状态']||detail.statuses?.[0]||'';renderActions(detail.actions,r);state.formDirty=false;}
      $("status").disabled=false;state.ledgerLoaded=true;
    }
    async function loadLedger(){try{const detail=await json(`/api/accounts/by-login/${encodeURIComponent(LOGIN)}/ledger`);renderRecord(detail,!state.formDirty);}catch(err){$("markState").textContent=err.message;$("status").disabled=true;}}
    function chartRow(chart){return `<div class="chart-row"><div><div class="chart-name">${esc(chart.name)}</div><div class="chart-meta">${esc(chart.start||'-')} 至 ${esc(chart.end||'-')} · ${num(chart.size,0)} bytes</div></div><div class="chart-actions"><a href="${esc(chart.url)}" target="_blank" rel="noopener">打开 AI 图表</a><button type="button" data-preview="${esc(chart.url)}">预览</button><button type="button" data-copy="${esc(chart.url)}">复制链接</button></div></div>`;}
    function bindCharts(){document.querySelectorAll('[data-preview]').forEach(btn=>btn.addEventListener('click',()=>preview(btn.dataset.preview)));document.querySelectorAll('[data-copy]').forEach(btn=>btn.addEventListener('click',async()=>{try{await navigator.clipboard.writeText(btn.dataset.copy);btn.textContent='已复制';setTimeout(()=>btn.textContent='复制链接',1000);}catch{prompt('复制图表链接',btn.dataset.copy);}}));}
    function renderCharts(charts){$("chartCount").textContent=`${charts.length} 个图表`;$("charts").innerHTML=charts.length?charts.map(chartRow).join(''):'<div class="empty-state">暂无图表</div>';bindCharts();}
    function chartRange(symbol){return symbol?state.chartSymbols.find(row=>row.symbol===symbol):{firstTime:state.detail?.database?.metrics?.firstTradeTime,lastTime:state.detail?.database?.metrics?.lastTradeTime};}
    function selectChartSymbol(symbol,updateRange=true){
      $("chartSymbol").value=symbol||'';document.querySelectorAll('[data-chart-product]').forEach(btn=>btn.classList.toggle('active',btn.dataset.chartProduct===(symbol||'')));
      if(updateRange){const row=chartRange(symbol);if(row){$("chartStart").value=datetimeLocal(row.firstTime);$("chartEnd").value=datetimeLocal(row.lastTime);}}
    }
    function renderChartProducts(db){
      const m=db.metrics||{},rows=m.bySymbol||[],current=rows.some(row=>row.symbol===$("chartSymbol").value)?$("chartSymbol").value:'';state.chartSymbols=rows;
      $("chartSymbol").innerHTML='<option value="">全部品种</option>'+rows.map(row=>`<option value="${esc(row.symbol)}">${esc(row.symbol)} · ${num(row.orders,0)} 单 · ${money(row.profit)}</option>`).join('');
      const cards=[{symbol:'',orders:m.orderCount||0,volume:m.totalVolume||0,profit:m.netProfit||0,firstTime:m.firstTradeTime,lastTime:m.lastTradeTime},...rows];
      $("chartProducts").innerHTML=cards.length?cards.map(row=>`<button type="button" class="chart-product" data-chart-product="${esc(row.symbol)}"><b class="chart-product-name">${esc(row.symbol||'全部品种')}</b><strong class="chart-product-profit ${profitClass(row.profit)}">${money(row.profit)}</strong><span class="chart-product-meta">${num(row.orders,0)} 单 · ${num(row.volume,4)} 手</span><span class="chart-product-time">${esc(row.firstTime||'-')} 至 ${esc(row.lastTime||'-')}</span></button>`).join(''):'<div class="empty-state">暂无可画图品种</div>';
      document.querySelectorAll('[data-chart-product]').forEach(btn=>btn.addEventListener('click',()=>selectChartSymbol(btn.dataset.chartProduct,true)));
      selectChartSymbol(current,false);
      if(!state.chartFiltersInitialized){selectChartSymbol(current,true);state.chartFiltersInitialized=true;}
    }
    function renderHistory(rows){$("historyCount").textContent=`${rows.length} 条记录`;$("history").innerHTML=rows.length?rows.map(row=>`<div class="history-item"><div class="history-time">${esc(row['修改时间']||'-')}</div><div class="history-op">${esc(row['操作']||'修改')}</div><div><b>${esc(row['修改字段']||'-')}</b><div class="history-time">${esc(row['处理人/来源']||'')}</div></div></div>`).join(''):'<div class="empty-state">暂无历史记录</div>';}
    function orderRow(row){const direction=row.type==='buy'?'买入':row.type==='sell'?'卖出':row.type||'-',comment=[row.comment,row.expertId?`EA #${row.expertId}`:''].filter(Boolean).join(' · ')||'-';return `<tr><td>${esc(row.ticket||'-')}</td><td>${esc([row.platform,row.server].filter(Boolean).join(' / ')||'-')}</td><td><b>${esc(row.symbol||'-')}</b></td><td class="order-type ${esc(row.type||'')}">${esc(direction)}</td><td><span class="reason-badge">${esc(row.reason||'-')}</span></td><td><div class="order-comment" title="${esc(comment)}">${row.isCopyTrade?'<span class="reason-badge">跟单</span> ':''}${esc(comment)}</div></td><td>${esc(row.openTime||'-')}</td><td>${esc(row.closeTime||'-')}</td><td>${duration(row.holdingSeconds)}</td><td>${num(row.volume,4)}</td><td class="${profitClass(row.profit)}">${amount(row.profit)}</td><td class="${profitClass(row.commission)}">${amount(row.commission)}</td><td class="${profitClass(row.swap)}">${amount(row.swap)}</td><td class="${profitClass(row.taxes)}">${amount(row.taxes)}</td><td class="${profitClass(row.netProfit)}"><b>${amount(row.netProfit)}</b></td><td>${esc(row.currency||row.displayCurrency||'-')}</td></tr>`;}
    function renderOrders(data){
      const rows=data.orders||[];state.orders.page=data.page||1;state.orders.pages=data.pages||1;state.orders.loaded=true;
      $("orderSummary").textContent=`${num(data.total,0)} 笔订单${data.truncated?' · 最多展示 50,000 笔':''}`;
      $("orderPageStatus").textContent=`第 ${state.orders.page} / ${state.orders.pages} 页 · 每页 ${data.pageSize} 笔`;
      $("orderRows").innerHTML=rows.length?rows.map(orderRow).join(''):'<tr><td colspan="16"><div class="empty-state">暂无订单</div></td></tr>';
      $("orderPrev").disabled=state.orders.page<=1;$("orderNext").disabled=state.orders.page>=state.orders.pages;
    }
    async function loadOrders(page=1){
      if(state.orders.loading)return;state.orders.loading=true;$("orderSummary").textContent='正在读取订单...';$("orderPrev").disabled=true;$("orderNext").disabled=true;
      const current=filters(),q=new URLSearchParams({page:String(page),pageSize:'100',platform:current.platform||'',server:current.server||''});
      try{renderOrders(await json(`/api/accounts/by-login/${encodeURIComponent(LOGIN)}/orders?${q}`));}catch(err){$("orderSummary").textContent=err.message;$("orderRows").innerHTML=`<tr><td colspan="16"><div class="empty-state">${esc(err.message)}</div></td></tr>`;}finally{state.orders.loading=false;}
    }
    function render(detail,keepFilters=false){
      state.detail=detail;const db=detail.database||{},source=db.latestSource||{},meta=db.accountMeta||{};$("accountId").textContent=detail.account;document.title=`账号 ${detail.account} · 风控台账`;
      const currencyText=meta.isCentAccount?'USC 美分账户 · 金额已按 USD 折算':(meta.currency==='USD'?'USD 美元账户':(meta.currency?`${meta.currency} 账户`:'币种未识别 · 金额未缩放'));
      const hasEa=Boolean(db.allMetrics?.hasEaTrades||db.metrics?.hasEaTrades),hasCopy=Boolean(db.allMetrics?.hasCopyTrades||db.metrics?.hasCopyTrades);$("badges").innerHTML=`<span class="badge ${detail.marked?'marked':''}">${detail.marked?'已标记':'未标记'}</span><span class="badge ${db.exists?'':'empty'}">${db.exists?'数据库有订单':'账户暂未做单'}</span><span class="badge">${esc(currencyText)}</span>${hasEa?'<span class="badge ea">EA</span>':''}${hasCopy?'<span class="badge marked">跟单</span>':''}${detail.record?.['建议动作']?`<span class="badge action">${esc(detail.record['建议动作'])}</span>`:''}`;$("copyOriginBtn").hidden=!db.exists;$("eaCommentBtn").hidden=!db.exists;$("relationshipNetworkBtn").hidden=!db.exists;
      $("headMeta").innerHTML=`${esc([source.platform,source.server].filter(Boolean).join(' / ')||'未识别平台')}<br>最近交易 ${esc(db.lastTime||'-')}<br>刷新 ${esc(db.refreshedAt||'-')}`;
      $("metricStatus").textContent=db.exists?`${db.orderCount} 条订单 · ${db.firstTime||'-'} 至 ${db.lastTime||'-'}`:(db.error||'账户暂未做单');$("dbSource").textContent=db.dbSource?db.dbSource.toUpperCase():'';
      if(!keepFilters){optionList($("platform"),db.platforms||[],'全部平台');optionList($("server"),db.servers||[],'全部服务器');optionList($("symbol"),db.symbols||[],'全部品种');if(state.initialFilters){for(const key of ['platform','server','symbol'])if(state.initialFilters[key])$(key).value=state.initialFilters[key];state.initialFilters=null;}}
      renderChartProducts(db);
      renderRiskPanels(db.riskPanels);renderMetrics(db);renderVisualizations(db);if(!state.ledgerLoaded)renderRecord(detail,!state.formDirty);renderCharts(detail.charts||[]);renderHistory(detail.history||[]);$("generateBtn").disabled=!db.exists;$("toxicBtn").disabled=!db.exists;$("historicalFundsBtn").disabled=!db.exists;
    }
    function filters(){return state.initialFilters||{platform:$("platform").value,server:$("server").value,symbol:$("symbol").value};}
    async function openAccountFromDetailSearch(event){
      event.preventDefault();const input=$("detailAccountSearch"),button=$("detailAccountSearchBtn"),status=$("detailAccountSearchStatus"),account=input.value.trim();
      if(!/^\d+$/.test(account)){status.textContent=account?'请输入数字账号':'请输入账号';input.focus();return;}
      button.disabled=true;status.textContent='正在定位账号...';
      try{
        const data=await json(`/api/account-lookup?account=${encodeURIComponent(account)}`),matches=(data.databases||[]).length?data.databases:[data.database].filter(Boolean);
        if(data.database?.queryFailed)throw new Error(data.database.error||'账号查询失败，请稍后重试');
        if(!matches.length)throw new Error('未找到该账号');
        if(matches.length>1){openAccountSourceDialog(account,matches);status.textContent=`找到 ${matches.length} 个平台/服务器，请选择`;return;}
        const current=filters(),selected=matches.find(item=>item.latestSource?.platform===current.platform&&item.latestSource?.server===current.server)||matches[0],source=selected.latestSource||{};
        const query=new URLSearchParams();if(source.platform)query.set('platform',source.platform);if(source.server)query.set('server',source.server);
        location.assign(`/account/${encodeURIComponent(account)}${query.size?`?${query}`:''}`);
      }catch(err){status.textContent=err.message||'账号查询失败';button.disabled=false;}
      finally{if(status.textContent==='正在定位账号...')button.disabled=false;}
    }
    function openAccountSourceDialog(account,matches){
      state.accountLookupMatches=matches;const dialog=$("accountSourceDialog"),list=$("accountSourceList");
      $("accountSourceHint").textContent=`账号 ${account} 同时存在多个平台/服务器，请选择要查看的详细数据。`;
      list.innerHTML=matches.map((item,index)=>{const source=item.latestSource||{};return `<button type="button" class="account-source-option" data-source-index="${index}"><span><b>${esc(source.platform||'-')} / ${esc(source.server||'-')}</b><small>${item.exists?`${num(item.orderCount,0)} 笔订单`:'账户暂未做单'}</small></span><strong>查看详情 →</strong></button>`;}).join('');
      list.querySelectorAll('[data-source-index]').forEach(button=>button.addEventListener('click',()=>{const item=state.accountLookupMatches[Number(button.dataset.sourceIndex)]||{},source=item.latestSource||{},query=new URLSearchParams();if(source.platform)query.set('platform',source.platform);if(source.server)query.set('server',source.server);dialog.close();location.assign(`/account/${encodeURIComponent(account)}${query.size?`?${query}`:''}`);}));
      dialog.showModal();
    }
    function automationDialogQuery(){const q=new URLSearchParams(filters());[...q.keys()].forEach(key=>{if(!q.get(key))q.delete(key)});q.sort();return q;}
    function copyOriginDialogQuery(){const q=automationDialogQuery(),start=databaseTime($("copyOriginStart").value),end=databaseTime($("copyOriginEnd").value,true);if(start&&end&&start>end)throw new Error('跟单开始时间不能晚于结束时间');if(start)q.set('start',start);if(end)q.set('end',end);q.sort();return q;}
    function eaCommentDialogQuery(){const q=automationDialogQuery(),start=databaseTime($("eaCommentStart").value),end=databaseTime($("eaCommentEnd").value,true);if(start&&end&&start>end)throw new Error('EA 开始时间不能晚于结束时间');if(start)q.set('start',start);if(end)q.set('end',end);q.sort();return q;}
    function clearAutomationDialogCache(){state.dialogCache.copy.clear();state.dialogCache.ea.clear();state.dialogCache.relationship.clear();state.relationshipNetwork=null;}
    const RELATIONSHIP_GRAPH_WIDTH=1000,RELATIONSHIP_GRAPH_HEIGHT=620,RELATIONSHIP_GRAPH_CACHE_SCALE=3;
    const relationshipGraph=$('relationshipNetworkGraph');
    let relationshipCanvasContext=null,relationshipDrag=null;
    function relationShort(value,limit=18){const text=String(value||'');return text.length>limit?`${text.slice(0,limit-1)}…`:text||'-';}
    function relationEntityMap(network){return new Map((network.data.entities||[]).map(entity=>[entity.id,entity]));}
    function relationChild(edge,network){return edge.source!==network.subjectId&&['ea_feature','copy_order','copy_group'].includes(edge.type);}
    function relationVisibleEdges(network){return (network.data.relationships||[]).filter(edge=>network.visibleTypes.has(edge.type)&&(!relationChild(edge,network)||network.expanded.has(edge.source)));}
    function relationVisibleEntities(network,edges){const ids=new Set([network.subjectId]);edges.forEach(edge=>{ids.add(edge.source);ids.add(edge.target);});return (network.data.entities||[]).filter(entity=>ids.has(entity.id));}
    function relationChildCounts(network){const counts=new Map();(network.data.relationships||[]).forEach(edge=>{if(relationChild(edge,network))counts.set(edge.source,(counts.get(edge.source)||0)+1);});return counts;}
    function relationLane(entity,edges,network){
      if(entity.id===network.subjectId)return 'subject';
      if(entity.type==='ip')return 'ip';
      if(entity.type==='rebate')return 'rebate';
      if(entity.type==='ea_feature')return 'ea';
      if(entity.type==='copy_group')return 'copyGroup';
      const inbound=edges.find(edge=>edge.target===entity.id);
      if(inbound?.type==='same_crm_user')return 'same';
      if(inbound?.type==='ea_feature')return 'ea';
      if(inbound?.type==='copy_order'||inbound?.type==='copy_group')return 'copy';
      return 'other';
    }
    function relationLayout(network,entities,edges){
      const lanes={same:{x:150,y:150},ip:{x:150,y:435},ea:{x:820,y:130},copy:{x:820,y:350},copyGroup:{x:790,y:510},rebate:{x:500,y:550},other:{x:500,y:90}};
      const groups=new Map();entities.forEach(entity=>{if(entity.id===network.subjectId)return;const lane=relationLane(entity,edges,network);if(!groups.has(lane))groups.set(lane,[]);groups.get(lane).push(entity);});
      const subject=entities.find(entity=>entity.id===network.subjectId);if(subject&&!network.positions.has(subject.id))network.positions.set(subject.id,{x:500,y:310});
      groups.forEach((items,lane)=>{const anchor=lanes[lane]||lanes.other,columns=items.length>3?2:1,rows=Math.ceil(items.length/columns);items.forEach((entity,index)=>{if(network.positions.has(entity.id))return;const column=index%columns,row=Math.floor(index/columns),x=anchor.x+(column-(columns-1)/2)*150,y=anchor.y+(row-(rows-1)/2)*82;network.positions.set(entity.id,{x:Math.max(72,Math.min(928,x)),y:Math.max(48,Math.min(572,y))});});});
    }
    function relationNodeSub(entity,childCount,expanded){
      if(childCount)return `${expanded?'收起':'展开'} ${childCount} 个成员`;
      return relationShort(entity.detail||[entity.platform,entity.server].filter(Boolean).join(' / '),20);
    }
    function relationRenderDetail(network){
      const entities=relationEntityMap(network),selected=network.selected||{kind:'node',id:network.subjectId},limits=(network.data.limitations||[]).filter(Boolean);
      $('relationshipLimitations').innerHTML=limits.map(item=>`<li>${esc(item)}</li>`).join('')||'<li>无额外限制。</li>';
      if(selected.kind==='edge'){
        const edge=(network.data.relationships||[]).find(item=>item.id===selected.id);if(!edge)return;
        const source=entities.get(edge.source)||{},target=entities.get(edge.target)||{};
        $('relationshipDetailKind').textContent=edge.typeLabel||'关系证据';$('relationshipDetailTitle').textContent=edge.label||'关系记录';
        $('relationshipDetailMeta').innerHTML=[['关系类别',edge.typeLabel||'-'],['关系对象',`${source.label||'-'} → ${target.label||'-'}`]].map(([label,value])=>`<div><span>${esc(label)}</span><b>${esc(value)}</b></div>`).join('');
        $('relationshipEvidence').innerHTML=(edge.evidence||[]).map((item,index)=>`<div class="relationship-evidence-row"><span class="muted">证据 ${index+1}</span><br>${esc(item)}</div>`).join('')||'<div class="relationship-evidence-row">该关系没有额外证据字段。</div>';
        return;
      }
      const entity=entities.get(selected.id);if(!entity)return;
      const connected=(network.data.relationships||[]).filter(edge=>edge.source===entity.id||edge.target===entity.id);
      $('relationshipDetailKind').textContent=entity.isSubject?'当前账户':({account:'交易账户',ip:'登录 IP',ea_feature:'EA / 路由特征',copy_group:'跟单组',rebate:'返佣记录'}[entity.type]||'实体');
      $('relationshipDetailTitle').textContent=entity.label||'-';
      $('relationshipDetailMeta').innerHTML=[['平台',entity.platform||'-'],['服务器',entity.server||'-'],['实体说明',entity.detail||'-'],['关联记录',`${connected.length} 条`]].map(([label,value])=>`<div><span>${esc(label)}</span><b>${esc(value)}</b></div>`).join('');
      const evidence=connected.flatMap(edge=>(edge.evidence||[]).map(item=>`${edge.label||edge.typeLabel}：${item}`));
      $('relationshipEvidence').innerHTML=evidence.slice(0,24).map((item,index)=>`<div class="relationship-evidence-row"><span class="muted">证据 ${index+1}</span><br>${esc(item)}</div>`).join('')||'<div class="relationship-evidence-row">该实体当前没有可显示的关系证据。</div>';
    }
    function relationEdgeLabel(edge){return relationShort(edge.typeLabel||edge.label||'关系证据',18);}
    function relationCanvasViewport(){const rect=relationshipGraph.getBoundingClientRect(),scale=Math.min(rect.width/RELATIONSHIP_GRAPH_WIDTH,rect.height/RELATIONSHIP_GRAPH_HEIGHT)||1;return {rect,scale,offsetX:(rect.width-RELATIONSHIP_GRAPH_WIDTH*scale)/2,offsetY:(rect.height-RELATIONSHIP_GRAPH_HEIGHT*scale)/2};}
    function relationCanvasContextFor(viewport){const dpr=Math.min(window.devicePixelRatio||1,2),width=Math.max(1,Math.round(viewport.rect.width*dpr)),height=Math.max(1,Math.round(viewport.rect.height*dpr));if(relationshipGraph.width!==width||relationshipGraph.height!==height){relationshipGraph.width=width;relationshipGraph.height=height;}relationshipCanvasContext=relationshipCanvasContext||relationshipGraph.getContext('2d',{alpha:false,desynchronized:true});if(relationshipCanvasContext)relationshipCanvasContext.setTransform(dpr,0,0,dpr,0,0);return relationshipCanvasContext;}
    function relationCanvasRawPoint(clientX,clientY,viewport){return {x:(clientX-viewport.rect.left-viewport.offsetX)/viewport.scale,y:(clientY-viewport.rect.top-viewport.offsetY)/viewport.scale};}
    function relationCanvasPoint(clientX,clientY,network,viewport){const raw=relationCanvasRawPoint(clientX,clientY,viewport);return {x:(raw.x-network.tx)/network.scale,y:(raw.y-network.ty)/network.scale};}
    function relationCanvasRoundedRect(context,x,y,width,height,radius){const r=Math.min(radius,width/2,height/2);context.beginPath();context.moveTo(x+r,y);context.arcTo(x+width,y,x+width,y+height,r);context.arcTo(x+width,y+height,x,y+height,r);context.arcTo(x,y+height,x,y,r);context.arcTo(x,y,x+width,y,r);context.closePath();}
    function relationCanvasNode(context,entity,point,childCount,network){const selected=network.selected?.kind==='node'&&network.selected.id===entity.id,isSubject=Boolean(entity.isSubject);context.save();context.translate(point.x,point.y);context.fillStyle=isSubject?'#168cff':({account:'#1679cf',ip:'#bb5d99',ea_feature:'#9c4e8a',copy_group:'#7257aa',rebate:'#ba7943'}[entity.type]||'#1679cf');context.strokeStyle=selected?'#d7c7ff':(isSubject?'#a6ddff':'#315474');context.lineWidth=selected?4:2;if(entity.type==='account'){context.beginPath();context.arc(0,0,isSubject?44:29,0,Math.PI*2);context.fill();context.stroke();}else if(entity.type==='rebate'){context.beginPath();context.moveTo(0,-30);context.lineTo(28,-15);context.lineTo(28,15);context.lineTo(0,30);context.lineTo(-28,15);context.lineTo(-28,-15);context.closePath();context.fill();context.stroke();}else{const width=entity.type==='ip'?118:132;relationCanvasRoundedRect(context,-width/2,-27,width,54,entity.type==='ip'?27:7);context.fill();context.stroke();}context.fillStyle='#eaf7ff';context.font='14px Segoe UI, Microsoft YaHei, sans-serif';context.textAlign='center';context.textBaseline='middle';context.fillText(relationShort(entity.label,16),0,-3);context.fillStyle='#9ab6d0';context.font='10px Segoe UI, Microsoft YaHei, sans-serif';context.fillText(relationNodeSub(entity,childCount,network.expanded.has(entity.id)),0,16);if(childCount){context.fillStyle='#d6c6ff';context.font='9px Segoe UI, Microsoft YaHei, sans-serif';context.fillText('点击切换成员',0,39);}context.restore();}
    function relationClearCanvas(){const viewport=relationCanvasViewport(),context=relationCanvasContextFor(viewport);if(context){context.fillStyle='#041225';context.fillRect(0,0,viewport.rect.width,viewport.rect.height);}}
    function relationWorldViewport(network,viewport){const left=(0-viewport.offsetX)/viewport.scale,top=(0-viewport.offsetY)/viewport.scale,right=(viewport.rect.width-viewport.offsetX)/viewport.scale,bottom=(viewport.rect.height-viewport.offsetY)/viewport.scale;return {left:(left-network.tx)/network.scale,top:(top-network.ty)/network.scale,right:(right-network.tx)/network.scale,bottom:(bottom-network.ty)/network.scale};}
    function relationBoundsIntersect(a,b){return a.left<=b.right&&a.right>=b.left&&a.top<=b.bottom&&a.bottom>=b.top;}
    function relationNodeBounds(entity,point){if(entity.type==='account'){const radius=entity.isSubject?44:29;return {left:point.x-radius,right:point.x+radius,top:point.y-radius,bottom:point.y+radius};}if(entity.type==='rebate')return {left:point.x-30,right:point.x+30,top:point.y-32,bottom:point.y+32};const width=entity.type==='ip'?118:132;return {left:point.x-width/2,right:point.x+width/2,top:point.y-33,bottom:point.y+44};}
    function relationEdgeIntersectsViewport(source,target,viewport){return relationBoundsIntersect({left:Math.min(source.x,target.x),right:Math.max(source.x,target.x),top:Math.min(source.y,target.y),bottom:Math.max(source.y,target.y)},viewport);}
    function relationEdgeLabelMetric(context,network,edge){const label=relationEdgeLabel(edge),cached=network.edgeLabelMetrics?.get(edge.id);if(cached?.label===label)return cached;context.save();context.font='700 10px Segoe UI, Microsoft YaHei, sans-serif';const metric={label,width:Math.max(74,Math.min(170,context.measureText(label).width+20))};context.restore();if(!network.edgeLabelMetrics)network.edgeLabelMetrics=new Map();network.edgeLabelMetrics.set(edge.id,metric);return metric;}
    function relationCanvasEdgeLabel(context,edge,source,target,index,selected,network,worldViewport){const dx=target.x-source.x,dy=target.y-source.y,distance=Math.hypot(dx,dy)||1,side=index%2?1:-1,offset=16+(index%3)*6,x=(source.x+target.x)/2+(-dy/distance)*side*offset,y=(source.y+target.y)/2+(dx/distance)*side*offset,metric=relationEdgeLabelMetric(context,network,edge);if(!relationBoundsIntersect({left:x-metric.width/2,right:x+metric.width/2,top:y-11,bottom:y+11},worldViewport))return;context.save();context.font='700 10px Segoe UI, Microsoft YaHei, sans-serif';relationCanvasRoundedRect(context,x-metric.width/2,y-11,metric.width,22,11);context.fillStyle=selected?'#352d5c':'#0a2440';context.fill();context.strokeStyle=selected?'#c4aaff':'#4f83ad';context.lineWidth=1;context.stroke();context.fillStyle='#e3f3ff';context.textAlign='center';context.textBaseline='middle';context.fillText(metric.label,x,y+1);context.restore();}
    function relationDrawEdge(context,edge,source,target,index,selected,network,worldViewport){if(!relationEdgeIntersectsViewport(source,target,worldViewport))return;context.beginPath();context.moveTo(source.x,source.y);context.lineTo(target.x,target.y);context.strokeStyle=selected?'#b69cff':'#315474';context.lineWidth=selected?4:2;context.stroke();relationCanvasEdgeLabel(context,edge,source,target,index,selected,network,worldViewport);}
    function relationSceneContext(network){const width=RELATIONSHIP_GRAPH_WIDTH*RELATIONSHIP_GRAPH_CACHE_SCALE,height=RELATIONSHIP_GRAPH_HEIGHT*RELATIONSHIP_GRAPH_CACHE_SCALE;if(!network.sceneCanvas)network.sceneCanvas=document.createElement('canvas');if(network.sceneCanvas.width!==width||network.sceneCanvas.height!==height){network.sceneCanvas.width=width;network.sceneCanvas.height=height;network.sceneContext=null;}network.sceneContext=network.sceneContext||network.sceneCanvas.getContext('2d',{alpha:true,desynchronized:true});if(network.sceneContext)network.sceneContext.setTransform(RELATIONSHIP_GRAPH_CACHE_SCALE,0,0,RELATIONSHIP_GRAPH_CACHE_SCALE,0,0);return network.sceneContext;}
    function relationRecordFrame(network,started,cacheBuildMs=0){const now=performance.now(),stats=network.frameStats||(network.frameStats={frames:0,totalMs:0,maxMs:0,cacheBuilds:0,cacheTotalMs:0,maxCacheMs:0,lastInputToPaintMs:0});const duration=now-started;stats.frames+=1;stats.totalMs+=duration;stats.maxMs=Math.max(stats.maxMs,duration);if(cacheBuildMs){stats.cacheBuilds+=1;stats.cacheTotalMs+=cacheBuildMs;stats.maxCacheMs=Math.max(stats.maxCacheMs,cacheBuildMs);}if(network.latestInputAt)stats.lastInputToPaintMs=now-network.latestInputAt;relationshipGraph.dataset.frameMs=duration.toFixed(2);relationshipGraph.dataset.cacheBuildMs=cacheBuildMs.toFixed(2);relationshipGraph.dataset.cacheBuilds=String(stats.cacheBuilds);relationshipGraph.dataset.inputToPaintMs=stats.lastInputToPaintMs.toFixed(2);window.__kdeskRelationshipPerf=stats;}
    function relationDrawSceneContent(context,network,worldViewport,excludedNodeId=''){const edges=network.renderedEdges||[],entities=network.renderedEntities||[];edges.forEach((edge,index)=>{if(edge.source===excludedNodeId||edge.target===excludedNodeId)return;const source=network.positions.get(edge.source),target=network.positions.get(edge.target);if(source&&target)relationDrawEdge(context,edge,source,target,index,network.selected?.kind==='edge'&&network.selected.id===edge.id,network,worldViewport);});entities.forEach(entity=>{if(entity.id===excludedNodeId)return;const point=network.positions.get(entity.id);if(point&&relationBoundsIntersect(relationNodeBounds(entity,point),worldViewport))relationCanvasNode(context,entity,point,network.childCounts?.get(entity.id)||0,network);});}
    function relationBuildSceneCache(network){const context=relationSceneContext(network);if(!context)return null;context.clearRect(0,0,RELATIONSHIP_GRAPH_WIDTH,RELATIONSHIP_GRAPH_HEIGHT);relationDrawSceneContent(context,network,{left:0,top:0,right:RELATIONSHIP_GRAPH_WIDTH,bottom:RELATIONSHIP_GRAPH_HEIGHT},network.dragNodeId||'');network.sceneDirty=false;return network.sceneCanvas;}
    function relationDrawDragOverlay(context,network,worldViewport){const nodeId=network.dragNodeId;if(!nodeId)return;const entity=(network.renderedEntities||[]).find(item=>item.id===nodeId),point=network.positions.get(nodeId);if(!entity||!point)return;(network.renderedEdges||[]).forEach((edge,index)=>{if(edge.source!==nodeId&&edge.target!==nodeId)return;const source=network.positions.get(edge.source),target=network.positions.get(edge.target);if(source&&target)relationDrawEdge(context,edge,source,target,index,network.selected?.kind==='edge'&&network.selected.id===edge.id,network,worldViewport);});if(relationBoundsIntersect(relationNodeBounds(entity,point),worldViewport))relationCanvasNode(context,entity,point,network.childCounts?.get(entity.id)||0,network);}
    function relationDrawGraph(network){const started=performance.now(),viewport=relationCanvasViewport(),context=relationCanvasContextFor(viewport);if(!context)return;context.fillStyle='#041225';context.fillRect(0,0,viewport.rect.width,viewport.rect.height);const edges=network.renderedEdges||[];if(!edges.length){context.fillStyle='#7895b8';context.font='14px Segoe UI, Microsoft YaHei, sans-serif';context.textAlign='center';context.fillText('当前筛选条件下没有可显示的关系记录',viewport.rect.width/2,viewport.rect.height/2);relationRecordFrame(network,started);return;}let cacheBuildMs=0;if(network.sceneDirty||!network.sceneCanvas){const cacheStarted=performance.now();relationBuildSceneCache(network);cacheBuildMs=performance.now()-cacheStarted;}const scene=network.sceneCanvas;if(scene){const worldViewport=relationWorldViewport(network,viewport);context.save();context.translate(viewport.offsetX+viewport.scale*network.tx,viewport.offsetY+viewport.scale*network.ty);context.scale(viewport.scale*network.scale,viewport.scale*network.scale);context.drawImage(scene,0,0,scene.width,scene.height,0,0,RELATIONSHIP_GRAPH_WIDTH,RELATIONSHIP_GRAPH_HEIGHT);relationDrawDragOverlay(context,network,worldViewport);context.restore();}relationRecordFrame(network,started,cacheBuildMs);}
    function relationSetSelection(network,selected){network.selected=selected;network.sceneDirty=true;relationRenderDetail(network);relationScheduleCanvas(network);}
    function relationApplyPendingInput(network){const drag=relationshipDrag;if(drag?.network===network&&drag.latest){const viewport=relationCanvasViewport();if(drag.kind==='node'){const point=relationCanvasPoint(drag.latest.x,drag.latest.y,network,viewport);network.positions.set(drag.entity.id,{x:Math.max(72,Math.min(928,point.x+drag.dx)),y:Math.max(48,Math.min(572,point.y+drag.dy))});}else if(drag.kind==='pan'){const raw=relationCanvasRawPoint(drag.latest.x,drag.latest.y,viewport);network.tx=drag.tx+(raw.x-drag.raw.x);network.ty=drag.ty+(raw.y-drag.raw.y);}drag.latest=null;}const wheel=network.pendingWheel;if(wheel){network.pendingWheel=null;const viewport=relationCanvasViewport(),raw=relationCanvasRawPoint(wheel.x,wheel.y,viewport),point={x:(raw.x-network.tx)/network.scale,y:(raw.y-network.ty)/network.scale},factor=Math.exp(-Math.max(-360,Math.min(360,wheel.delta))*0.00075),scale=Math.max(.7,Math.min(1.8,network.scale*factor));network.scale=scale;network.tx=raw.x-point.x*scale;network.ty=raw.y-point.y*scale;}}
    function relationScheduleCanvas(network){if(network.canvasFrame)return;network.canvasFrame=requestAnimationFrame(()=>{network.canvasFrame=0;relationApplyPendingInput(network);relationDrawGraph(network);});}
    function relationFlushCanvas(network){if(network.canvasFrame){cancelAnimationFrame(network.canvasFrame);network.canvasFrame=0;}relationApplyPendingInput(network);relationDrawGraph(network);}
    function relationPointToSegmentDistance(point,source,target){const dx=target.x-source.x,dy=target.y-source.y,length=dx*dx+dy*dy||1,t=Math.max(0,Math.min(1,((point.x-source.x)*dx+(point.y-source.y)*dy)/length)),x=source.x+t*dx,y=source.y+t*dy;return Math.hypot(point.x-x,point.y-y);}
    function relationNodeContains(entity,point,network){const center=network.positions.get(entity.id);if(!center)return false;const dx=point.x-center.x,dy=point.y-center.y;if(entity.type==='account')return Math.hypot(dx,dy)<=((entity.isSubject?44:29)+6);if(entity.type==='rebate')return Math.abs(dx)/34+Math.abs(dy)/36<=1;const width=entity.type==='ip'?118:132;return Math.abs(dx)<=width/2+6&&Math.abs(dy)<=33;}
    function relationHitTest(network,point){const entities=[...(network.renderedEntities||[])].reverse(),entity=entities.find(item=>relationNodeContains(item,point,network));if(entity)return {kind:'node',entity};const edge=[...(network.renderedEdges||[])].reverse().find(item=>{const source=network.positions.get(item.source),target=network.positions.get(item.target);return source&&target&&relationPointToSegmentDistance(point,source,target)<=10/network.scale;});return edge?{kind:'edge',edge}:null;}
    function relationActivateNode(network,entity){const childCount=network.childCounts?.get(entity.id)||0;if(childCount){network.expanded.has(entity.id)?network.expanded.delete(entity.id):network.expanded.add(entity.id);network.selected={kind:'node',id:entity.id};relationRenderGraph();}else relationSetSelection(network,{kind:'node',id:entity.id});}
    function relationRenderGraph(){const network=state.relationshipNetwork;if(!network)return;const edges=relationVisibleEdges(network),entities=relationVisibleEntities(network,edges);network.renderedEdges=edges;network.renderedEntities=entities;network.childCounts=relationChildCounts(network);relationLayout(network,entities,edges);network.sceneDirty=true;relationRenderDetail(network);relationDrawGraph(network);}
    function relationBindGraph(){
      relationshipGraph.addEventListener('pointerdown',event=>{const network=state.relationshipNetwork;if(!network)return;const viewport=relationCanvasViewport(),point=relationCanvasPoint(event.clientX,event.clientY,network,viewport),raw=relationCanvasRawPoint(event.clientX,event.clientY,viewport),hit=relationHitTest(network,point);relationshipDrag={network,pointer:event.pointerId,x:event.clientX,y:event.clientY,raw,tx:network.tx,ty:network.ty,moved:false,...(hit||{kind:'pan'})};if(hit?.kind==='node'){const position=network.positions.get(hit.entity.id);relationshipDrag.dx=position.x-point.x;relationshipDrag.dy=position.y-point.y;network.dragNodeId=hit.entity.id;network.sceneDirty=true;relationScheduleCanvas(network);}relationshipGraph.classList.add('dragging');relationshipGraph.setPointerCapture(event.pointerId);});
      relationshipGraph.addEventListener('pointermove',event=>{const network=state.relationshipNetwork,drag=relationshipDrag;if(!network||!drag||drag.network!==network||drag.pointer!==event.pointerId){if(network){const viewport=relationCanvasViewport(),point=relationCanvasPoint(event.clientX,event.clientY,network,viewport),hit=relationHitTest(network,point);relationshipGraph.style.cursor=hit?.kind==='edge'?'pointer':'grab';}return;}if(!drag.moved&&Math.hypot(event.clientX-drag.x,event.clientY-drag.y)>=4)drag.moved=true;if(drag.kind==='node'||drag.kind==='pan'){drag.latest={x:event.clientX,y:event.clientY};network.latestInputAt=performance.now();relationScheduleCanvas(network);}});
      function relationEndPointer(event,cancelled=false){const network=state.relationshipNetwork,drag=relationshipDrag;if(!network||!drag||drag.network!==network||drag.pointer!==event.pointerId)return;relationFlushCanvas(network);if(drag.kind==='node'){network.dragNodeId=null;network.sceneDirty=true;relationDrawGraph(network);}relationshipDrag=null;relationshipGraph.classList.remove('dragging');if(cancelled||drag.moved)return;if(drag.kind==='node')relationActivateNode(network,drag.entity);else if(drag.kind==='edge')relationSetSelection(network,{kind:'edge',id:drag.edge.id});else relationSetSelection(network,{kind:'node',id:network.subjectId});}
      relationshipGraph.addEventListener('pointerup',event=>relationEndPointer(event));relationshipGraph.addEventListener('pointercancel',event=>relationEndPointer(event,true));
      relationshipGraph.addEventListener('wheel',event=>{const network=state.relationshipNetwork;if(!network)return;event.preventDefault();const pending=network.pendingWheel||{delta:0,x:event.clientX,y:event.clientY};pending.delta+=event.deltaY;pending.x=event.clientX;pending.y=event.clientY;network.pendingWheel=pending;network.latestInputAt=performance.now();relationScheduleCanvas(network);},{passive:false});
      if(typeof ResizeObserver!=='undefined')new ResizeObserver(()=>{if(state.relationshipNetwork)relationScheduleCanvas(state.relationshipNetwork);}).observe(relationshipGraph);
    }
    function relationRenderFilters(network){
      $('relationshipNetworkFilters').innerHTML=(network.data.relationTypes||[]).map(type=>`<label><input type="checkbox" data-relation-type="${esc(type.id)}" ${network.visibleTypes.has(type.id)?'checked':''}/><span>${esc(type.label)}</span></label>`).join('');
      document.querySelectorAll('[data-relation-type]').forEach(input=>input.addEventListener('change',()=>{input.checked?network.visibleTypes.add(input.dataset.relationType):network.visibleTypes.delete(input.dataset.relationType);const current=network.selected;if(current?.kind==='edge'&&!relationVisibleEdges(network).some(edge=>edge.id===current.id))network.selected={kind:'node',id:network.subjectId};relationRenderGraph();}));
    }
    function renderRelationshipNetwork(data){
      const subject=(data.entities||[]).find(entity=>entity.isSubject)||data.entities?.[0],availableTypes=new Set((data.relationTypes||[]).map(item=>item.id));
      state.relationshipNetwork={data,subjectId:subject?.id||'',visibleTypes:availableTypes,expanded:new Set(),positions:new Map(),edgeLabelMetrics:new Map(),sceneCanvas:null,sceneContext:null,sceneDirty:true,dragNodeId:null,frameStats:{frames:0,totalMs:0,maxMs:0,cacheBuilds:0,cacheTotalMs:0,maxCacheMs:0,lastInputToPaintMs:0},selected:{kind:'node',id:subject?.id||''},scale:1,tx:0,ty:0};
      const summary=data.summary||{},failed=(data.coverage||[]).filter(item=>item.status==='failed');$('relationshipNetworkStatus').textContent=`实体 ${num(summary.entityCount,0)} 个 · 关系 ${num(summary.relationshipCount,0)} 条 · 证据 ${num(summary.evidenceCount,0)} 条${failed.length?` · ${failed.length} 个数据源未完成`:''}`;
      relationRenderFilters(state.relationshipNetwork);relationRenderGraph();
    }
    function openRelationshipNetwork(){
      const query=automationDialogQuery();query.set('account',LOGIN);location.assign(`/kuzu-risk?${query.toString()}`);
    }
    function resetRelationshipNetwork(){const network=state.relationshipNetwork;if(!network)return;network.visibleTypes=new Set((network.data.relationTypes||[]).map(item=>item.id));network.expanded.clear();network.positions.clear();network.dragNodeId=null;network.selected={kind:'node',id:network.subjectId};network.scale=1;network.tx=0;network.ty=0;relationRenderFilters(network);relationRenderGraph();}
    relationBindGraph();
    function renderCopyOrigins(data){
      const origins=data.origins||[];
      if(!data.detected){$("copyOriginStatus").textContent='当前账户订单备注中没有识别到 CPT 跟单订单号';$("copyOriginRows").innerHTML='<div class="empty-state">没有 CPT 来源数据</div>';return;}
      $("copyOriginStatus").textContent=origins.length?`定位到 ${origins.length} 个单主 · 当前账号跟单 ${num(data.copyOrders,0)} / 全部 ${num(data.totalOrders,0)} 单 · 已归属 ${num(data.mappedCopyOrders,0)} 单${data.unmappedCopyOrders?` · 未归属 ${num(data.unmappedCopyOrders,0)} 单`:''}`:`已识别 ${num(data.searchedOrders,0)} 个 CPT 订单号，但当前数据库未找到对应单主`;
      $("copyOriginRows").innerHTML=origins.length?origins.map((row,index)=>{
        const href=`/account/${encodeURIComponent(row.account)}?platform=${encodeURIComponent(row.platform||'')}&server=${encodeURIComponent(row.server||'')}`,followers=row.followers||[],summary=row.followerSummary||{},discovery=row.followerDiscovery||{},currency=summary.currency||'',sourceOrders=row.sourceOrders||[];
        const followerRows=followers.map(item=>{const followerHref=`/account/${encodeURIComponent(item.account)}?platform=${encodeURIComponent(item.platform||row.platform||'')}&server=${encodeURIComponent(item.server||row.server||'')}`;return `<tr class="${item.isCurrentAccount?'current-account':''}"><td><a class="account-link" href="${esc(followerHref)}">${esc(item.account)}</a>${item.isCurrentAccount?' <span class="badge marked">当前账号</span>':''}</td><td>${esc([item.platform,item.server].filter(Boolean).join(' / ')||'-')}</td><td>${num(item.matchedSourceOrders,0)}</td><td><b>${num(item.orders,0)}</b></td><td>${num(item.volume,4)}</td><td class="${profitClass(item.grossProfit)}">${money(item.grossProfit)}</td><td class="${profitClass(item.commission)}">${money(item.commission)}</td><td class="${profitClass(item.swap)}">${money(item.swap)}</td><td class="${profitClass(item.taxes)}">${money(item.taxes)}</td><td class="${profitClass(item.netProfit)}"><b>${money(item.netProfit)}</b></td><td>${esc(item.displayCurrency||item.currency||'-')}${item.isCentAccount?'（USC折算）':''}</td><td>${esc((item.symbols||[]).join('、')||'-')}</td><td>${esc(item.firstTime||'-')}<br><span class="muted">至 ${esc(item.lastTime||'-')}</span></td><td>${esc((item.tickets||[]).join('、')||'-')}</td></tr>`;}).join('');
        const sourceRows=sourceOrders.map(item=>`<tr><td>${esc(item.orderId||'-')}</td><td>${esc(item.ticket||'-')}</td><td>${esc(item.symbol||'-')}</td><td>${esc(item.time||'-')}</td></tr>`).join('');
        const scanNote=discovery.error?`跟单人员查询失败：${discovery.error}`:[discovery.sourceOrdersTruncated?'源订单查询未完整':'',discovery.candidateRowsTruncated?'候选订单查询未完整':''].filter(Boolean).join('；');
        return `<section class="copy-master-block"><div class="copy-master-head"><div><b>单主：<a class="account-link" href="${esc(href)}">${esc(row.account)}</a></b>${index===0?' <span class="badge marked">主要单主</span>':''}<br><small>${esc([row.platform,row.server].filter(Boolean).join(' / '))} · ${num(row.matchedOrders,0)} 个源订单 · ${esc((row.symbols||[]).join('、')||'-')}</small></div><span>${esc(row.firstTime||'-')} 至 ${esc(row.lastTime||'-')}</span></div><div class="risk-summary">${copyRiskStat('跟单人员',`${num(summary.accounts,0)} 人`)}${copyRiskStat('盈利 / 亏损人员',`${num(summary.profitableAccounts,0)} / ${num(summary.losingAccounts,0)}`)}${copyRiskStat('跟单订单',`${num(summary.orders,0)} 单`)}${copyRiskStat('跟单手数',num(summary.volume,4))}${copyRiskStat(`跟单毛盈亏${currency?` (${currency})`:''}`,money(summary.grossProfit),profitClass(summary.grossProfit))}${copyRiskStat('手续费 / Fee',money(summary.commission),profitClass(summary.commission))}${copyRiskStat('利息 / Swap',money(summary.swap),profitClass(summary.swap))}${copyRiskStat(`跟单净盈亏${currency?` (${currency})`:''}`,money(summary.netProfit),profitClass(summary.netProfit))}${copyRiskStat('当前账号跟单',`${num(row.orders,0)} 单 / ${num(row.volume,4)} 手`)}${copyRiskStat('当前账号净盈亏',money(row.netProfit),profitClass(row.netProfit))}</div>${scanNote?`<div class="risk-note">${esc(scanNote)}</div>`:''}<div class="table-wrap"><table class="risk-table copy-follower-table"><thead><tr><th>跟单账号</th><th>平台 / 服务器</th><th>匹配源单</th><th>跟单订单</th><th>手数</th><th>毛盈亏</th><th>手续费 / Fee</th><th>利息</th><th>税费</th><th>净盈亏</th><th>币种</th><th>品种</th><th>首次 / 最后</th><th>样例跟单单号</th></tr></thead><tbody>${followerRows||'<tr><td colspan="14"><div class="empty-state">没有发现其他跟单人员，或跟单记录超出当前数据范围</div></td></tr>'}</tbody></table></div><details class="copy-source-orders"><summary>查看单主源订单（${num(sourceOrders.length,0)} 单）</summary><div class="table-wrap"><table class="risk-table"><thead><tr><th>源订单号</th><th>单主订单 / Position</th><th>品种</th><th>时间</th></tr></thead><tbody>${sourceRows||'<tr><td colspan="4"><div class="empty-state">没有源订单明细</div></td></tr>'}</tbody></table></div></details></section>`;
      }).join(''):'<div class="empty-state">未找到单主；可能对应订单不在当前导出范围内</div>';
    }
    function copyRiskStat(label,value,cls=''){return `<div class="risk-stat"><span>${esc(label)}</span><b class="${esc(cls)}">${value}</b></div>`;}
    function renderCopyGroupProfit(data){
      const groups=data.groups||[];
      if(!data.detected||!groups.length){const errors=(data.errors||[]).join('；');$("copyGroupStatus").textContent=errors?`没有找到 Signal 跟单组；${errors}`:'当前账户的 users.comment 中没有识别到 Signal #… IN';$("copyGroupResults").innerHTML='<div class="empty-state">没有 Signal 整组收益数据</div>';return;}
      $("copyGroupStatus").textContent=`识别到 ${groups.length} 个 Signal 跟单组 · ${data.definition||''}`;
      $("copyGroupResults").innerHTML=groups.map(group=>{const t=group.totals||{},members=group.members||[],source=[group.platform,group.server].filter(Boolean).join(' / '),statusText=Object.entries(t.statusCounts||{}).map(([key,value])=>`${key} ${value}`).join('、')||'-',share=t.rebateShareOfClientLoss===null||t.rebateShareOfClientLoss===undefined?'-':`${num(Number(t.rebateShareOfClientLoss)*100,1)}%`,limitations=(group.limitations||[]).filter(Boolean),rows=members.map(row=>{const href=`/account/${encodeURIComponent(row.account)}?platform=${encodeURIComponent(group.platform||'')}&server=${encodeURIComponent(group.server||'')}`;return `<tr class="${String(row.account)===String(LOGIN)?'current-account':''}"><td><a class="account-link" href="${esc(href)}">${esc(row.account)}</a></td><td>${esc(row.status||'-')}</td><td>${num(row.closedOrders,0)}</td><td>${num(row.openOrders,0)}</td><td>${num(row.closedLots,4)}</td><td class="${profitClass(row.closedNetProfit)}">${money(row.closedNetProfit)}</td><td class="${profitClass(row.floatingNetProfit)}">${money(row.floatingNetProfit)}</td><td class="${profitClass(row.combinedNetProfit)}"><b>${money(row.combinedNetProfit)}</b></td><td>${money(row.rebate)}</td><td>${esc(row.currency||'-')}</td><td>${esc(row.firstClose||'-')}</td><td>${esc(row.lastClose||'-')}</td></tr>`;}).join('');return `<section class="copy-group-block"><div class="copy-group-head"><div><b>${esc(group.signalTag||'-')}</b><br><span>${esc(source)} · 当前账号 ${esc(group.account||LOGIN)}</span></div><span>${group.truncated?'账户列表达到安全上限，结果可能不完整':`${num(t.accounts,0)} 个账户`}</span></div><div class="risk-summary">${copyRiskStat('账户 / 盈利 / 亏损',`${num(t.accounts,0)} / ${num(t.profitableAccounts,0)} / ${num(t.losingAccounts,0)}`)}${copyRiskStat('平仓单 / 当前持仓',`${num(t.closedOrders,0)} / ${num(t.openOrders,0)}`)}${copyRiskStat('累计平仓手数',num(t.closedLots,4))}${copyRiskStat('Status 分布',esc(statusText))}${copyRiskStat('平仓净盈亏',money(t.closedNetProfit),profitClass(t.closedNetProfit))}${copyRiskStat('持仓浮动盈亏',money(t.floatingNetProfit),profitClass(t.floatingNetProfit))}${copyRiskStat('整组交易盈亏',money(t.combinedNetProfit),profitClass(t.combinedNetProfit))}${copyRiskStat('产生返佣',money(t.rebate))}${copyRiskStat('平均每手返佣',t.rebatePerLot===null||t.rebatePerLot===undefined?'-':money(t.rebatePerLot))}${copyRiskStat('返佣 / 客户亏损',share)}${copyRiskStat('B-book 粗算剩余',money(t.estimatedPlatformAfterRebate),profitClass(t.estimatedPlatformAfterRebate))}${copyRiskStat('数据口径','交易盈亏与返佣分列')}</div><div class="risk-note">B-book 粗算剩余 = 客户交易亏损 − 返佣；不含外部对冲、LP、奖金、出入金和其他成本。${limitations.length?` 限制：${esc(limitations.join('；'))}`:''}</div><div class="table-wrap"><table class="risk-table copy-group-table"><thead><tr><th>账户</th><th>Status</th><th>平仓单</th><th>持仓单</th><th>平仓手数</th><th>平仓净盈亏</th><th>浮动盈亏</th><th>综合交易盈亏</th><th>返佣</th><th>币种</th><th>首次平仓</th><th>最后平仓</th></tr></thead><tbody>${rows||'<tr><td colspan="12"><div class="empty-state">该组暂无交易数据</div></td></tr>'}</tbody></table></div></section>`;}).join('');
    }
    async function loadCopyOrigins(){
      let q;try{q=copyOriginDialogQuery();}catch(err){$("copyOriginStatus").textContent=err.message;$("copyOriginRows").innerHTML='<div class="empty-state">请修正时间范围后重新查询</div>';$("copyGroupStatus").textContent=err.message;$("copyGroupResults").innerHTML='<div class="empty-state">请修正时间范围后重新查询</div>';return;}
      $("copyOriginStatus").textContent='正在定位单主并汇总全部跟单人员收益...';$("copyOriginRows").innerHTML='<div class="empty-state">查询单主、源订单和跟单人员中...</div>';$("copyGroupStatus").textContent='正在读取 users.comment 并汇总整组收益...';$("copyGroupResults").innerHTML='<div class="empty-state">Signal 跟单组查询中...</div>';
      const cacheKey=q.toString();
      let request=state.dialogCache.copy.get(cacheKey);
      if(!request){
        request=Promise.allSettled([json(`/api/accounts/by-login/${encodeURIComponent(LOGIN)}/copy-origins?${q}`),json(`/api/accounts/by-login/${encodeURIComponent(LOGIN)}/copy-group-profit?${q}`)]);
        state.dialogCache.copy.set(cacheKey,request);
      }
      const [originResult,groupResult]=await request;
      if(originResult.status!=='fulfilled'||groupResult.status!=='fulfilled')state.dialogCache.copy.delete(cacheKey);
      if(originResult.status==='fulfilled')renderCopyOrigins(originResult.value);else{$("copyOriginStatus").textContent=originResult.reason?.message||'CPT 单主查询失败';$("copyOriginRows").innerHTML='<div class="empty-state">查询失败</div>';}
      if(groupResult.status==='fulfilled')renderCopyGroupProfit(groupResult.value);else{$("copyGroupStatus").textContent=groupResult.reason?.message||'Signal 整组收益查询失败';$("copyGroupResults").innerHTML='<div class="empty-state">查询失败</div>';}
    }
    async function openCopyOrigins(){$("copyDialog").showModal();await loadCopyOrigins();}
    function renderEaCommentGroups(data){
      const groups=data.groups||[],errors=(data.errors||[]).filter(Boolean),eaSummary=data.eaSummary||{},routeSummary=data.possibleCopyRouteSummary||{};
      if(!data.detected||!groups.length){$("eaCommentStatus").textContent=errors.length?`未找到可查询的 EA comment；${errors.join('；')}`:'当前筛选范围内没有识别到有效的 EA comment';$("eaCommentResults").innerHTML='<div class="empty-state">没有相同 comment 的 EA 收益数据</div>';return;}
      const accountSet=new Set(groups.flatMap(group=>(group.members||[]).map(row=>`${row.database||''}/${row.server||group.server}/${row.account}`)));
      $("eaCommentStatus").textContent=`EA ${num(eaSummary.groups,0)} 组 / ${num(eaSummary.accounts,0)} 个账户 / 净盈亏 ${money(eaSummary.netProfit)} · 可能是跟单路由 ${num(routeSummary.groups,0)} 组（不计入 EA） · 共覆盖 ${accountSet.size} 个服务器账户 · ${data.definition||''}${errors.length?` · 部分查询失败：${errors.join('；')}`:''}`;
      $("eaCommentResults").innerHTML=groups.map(group=>{
        const totals=group.totals||{},members=group.members||[],currency=totals.currency||'',isRoute=group.classification==='possible_copy_route',groupLabel=isRoute?'可能是跟单路由':'EA',limitations=[...(group.limitations||[]),group.truncated?'命中记录达到 50,000 条安全上限，结果可能不完整':''].filter(Boolean);
        const rows=members.map(row=>{const href=`/account/${encodeURIComponent(row.account)}?platform=${encodeURIComponent(row.platform||group.platform||'')}&server=${encodeURIComponent(row.server||group.server||'')}`,allExpertIds=row.expertIds||[],expertIds=allExpertIds.length>8?`${allExpertIds.slice(0,8).join('、')}…（共${allExpertIds.length}个）`:allExpertIds.join('、')||'-',matchClues=(row.matchClues||[row.matchClue]).filter(Boolean).join('；')||'-';return `<tr class="${row.isCurrentAccount?'current-account':''}"><td>${esc(row.database||'-')}</td><td>${esc(row.server||'-')}</td><td><a class="account-link" href="${esc(href)}">${esc(row.account)}</a>${row.isCurrentAccount?' <span class="badge marked">当前账号</span>':''}</td><td>${esc(expertIds)}</td><td>${esc(matchClues)}</td><td><b>${num(row.orders,0)}</b></td><td>${num(row.volume,4)}</td><td class="${profitClass(row.grossProfit)}">${money(row.grossProfit)}</td><td class="${profitClass(row.commission)}">${money(row.commission)}</td><td class="${profitClass(row.swap)}">${money(row.swap)}</td><td class="${profitClass(row.taxes)}">${money(row.taxes)}</td><td class="${profitClass(row.netProfit)}"><b>${money(row.netProfit)}</b></td><td>${esc(row.currency||'-')}${row.isCentAccount?'（USC折算）':''}</td><td>${esc((row.symbols||[]).join('、')||'-')}</td><td>${esc(row.firstTime||'-')}<br><span class="muted">至 ${esc(row.lastTime||'-')}</span></td><td>${esc((row.tickets||[]).join('、')||'-')}</td></tr>`;}).join('');
        const groupIdentifier=group.signatureType==='expert-sequence'?`共享完整 ExpertID ${num(group.expertSequence?.sharedAcrossAllAccounts,0)} 个`:`查询账号标识 ${String(group.expertId??'-')}`;
        return `<details class="copy-group-block ea-group-block" open><summary class="copy-group-head"><div><b>${esc(groupLabel)}：${esc(group.comment||'-')}</b>${isRoute?' <span class="badge">不计入 EA 汇总</span>':''}<br><span>${esc(group.classificationEvidence||'')} · ${esc(groupIdentifier)} · ${esc((group.servers||[group.server]).filter(Boolean).join(' / ')||'-')} · 当前账号 ${num(group.currentOrders,0)} 单 / ${num(group.currentVolume,4)} 手 / 净盈亏 ${money(group.currentNetProfit)}</span></div><span>${num(totals.accounts,0)} 个账户</span></summary><div class="risk-summary">${copyRiskStat('账户 / 盈利 / 亏损',`${num(totals.accounts,0)} / ${num(totals.profitableAccounts,0)} / ${num(totals.losingAccounts,0)}`)}${copyRiskStat('已平仓订单',`${num(totals.orders,0)} 单`)}${copyRiskStat('累计手数',num(totals.volume,4))}${copyRiskStat(`毛盈亏${currency?` (${currency})`:''}`,money(totals.grossProfit),profitClass(totals.grossProfit))}${copyRiskStat('手续费 / Fee',money(totals.commission),profitClass(totals.commission))}${copyRiskStat('利息 / Swap',money(totals.swap),profitClass(totals.swap))}${copyRiskStat('税费',money(totals.taxes),profitClass(totals.taxes))}${copyRiskStat(`净盈亏${currency?` (${currency})`:''}`,money(totals.netProfit),profitClass(totals.netProfit))}</div>${limitations.length?`<div class="risk-note">${esc(limitations.join('；'))}</div>`:''}<div class="table-wrap"><table class="risk-table copy-follower-table"><thead><tr><th>数据库</th><th>服务器</th><th>账户</th><th>EA标识</th><th>匹配线索</th><th>平仓订单</th><th>手数</th><th>毛盈亏</th><th>手续费 / Fee</th><th>利息</th><th>税费</th><th>净盈亏</th><th>币种</th><th>品种</th><th>首次 / 最后</th><th>样例订单号</th></tr></thead><tbody>${rows||'<tr><td colspan="16"><div class="empty-state">未找到使用此 Comment 的已平仓交易</div></td></tr>'}</tbody></table></div></details>`;
      }).join('');
    }
    async function loadEaCommentGroups(){
      let q;try{q=eaCommentDialogQuery();}catch(err){$("eaCommentStatus").textContent=err.message;$("eaCommentResults").innerHTML='<div class="empty-state">请修正时间范围后重新查询</div>';return;}
      $("eaCommentStatus").textContent='正在识别 EA comment 并汇总同备注账户收益...';$("eaCommentResults").innerHTML='<div class="empty-state">EA comment 查询中...</div>';
      const cacheKey=q.toString();
      let request=state.dialogCache.ea.get(cacheKey);
      if(!request){request=json(`/api/accounts/by-login/${encodeURIComponent(LOGIN)}/ea-comment-profit?${q}`);state.dialogCache.ea.set(cacheKey,request);}
      try{renderEaCommentGroups(await request);}catch(err){state.dialogCache.ea.delete(cacheKey);$("eaCommentStatus").textContent=err.message||'EA 查询失败';$("eaCommentResults").innerHTML='<div class="empty-state">查询失败</div>';}
    }
    async function openEaCommentGroups(){$("eaCommentDialog").showModal();await loadEaCommentGroups();}
    async function downloadAutomationReport(kind){
      const isCopy=kind==='copy',button=$(isCopy?'exportCopyReportBtn':'exportEaReportBtn'),status=$(isCopy?'copyExportStatus':'eaExportStatus');button.disabled=true;status.textContent='正在整理报表...';
      try{
        const q=kind==='ea'?eaCommentDialogQuery():copyOriginDialogQuery();
        const endpoint=isCopy?'copy-report.xlsx':'ea-report.xlsx',response=await fetch(`/api/accounts/by-login/${encodeURIComponent(LOGIN)}/${endpoint}?${q}`);
        if(!response.ok){let message=`HTTP ${response.status}`;try{const payload=await response.json();message=payload.detail||payload.error||message;}catch(_error){}throw new Error(message);}
        const blob=await response.blob(),disposition=response.headers.get('content-disposition')||'',matched=disposition.match(/filename="([^"]+)"/i),filename=matched?.[1]||`${isCopy?'copy_profit':'ea_profit'}_${LOGIN}.xlsx`,url=URL.createObjectURL(blob),anchor=document.createElement('a');
        anchor.href=url;anchor.download=filename;document.body.appendChild(anchor);anchor.click();anchor.remove();URL.revokeObjectURL(url);status.textContent='报表已导出';
      }catch(err){status.textContent=err.message||'报表导出失败';}finally{button.disabled=false;}
    }
    async function loadRiskPanels(q){const request=++state.riskRequest;try{const data=await json(`/api/accounts/by-login/${encodeURIComponent(LOGIN)}/risk-panels?${q}`);if(request!==state.riskRequest)return;if(state.detail?.database)state.detail.database.riskPanels=data.riskPanels;renderRiskPanels(data.riskPanels);}catch(err){if(request===state.riskRequest)renderRiskPanels({available:false,reason:err.message});}}
    async function load(keepFilters=false){$("refreshBtn").disabled=true;$("metricStatus").textContent='正在读取最新订单...';const q=new URLSearchParams(filters());[...q.keys()].forEach(k=>{if(!q.get(k))q.delete(k)});try{const detail=await json(`/api/accounts/by-login/${encodeURIComponent(LOGIN)}/detail?${q}`);if(detail.database?.requiresSourceSelection){openAccountSourceDialog(LOGIN,(detail.database.sourceCandidates||[]).map(item=>({exists:true,orderCount:item.orderCount,latestSource:{platform:item.platform,server:item.server}})));$("metricStatus").textContent='请先选择平台 / 服务器';return;}render(detail,keepFilters);loadInlineKline();loadRiskPanels(q);loadAutomation(q);}catch(err){$("metricStatus").textContent=err.message;}finally{$("refreshBtn").disabled=false;}}
    function targetAccounts(){return $("batchSameName").checked&&!$("batchSameName").disabled?state.sameNameAccounts:[String(LOGIN)];}
    async function markRequest(payload,accounts=targetAccounts()){const batch=accounts.length>1;return json(batch?'/api/accounts/mark-batch':'/api/accounts/mark',{method:'POST',body:JSON.stringify(batch?{...payload,accounts}:{...payload,account:LOGIN})});}
    function applyLocalAction(action,accounts){const panels=state.detail?.database?.riskPanels;if(!panels?.sameName)return;const targets=new Set(accounts.map(String));panels.sameName.forEach(row=>{if(targets.has(String(row.account)))row.localStatus=action;});renderRiskPanels(panels);}
    async function saveStatusOnly(){if($("status").disabled)return;const accounts=targetAccounts(),value=$("status").value;$("status").disabled=true;$("saveStatus").textContent=accounts.length>1?`正在保存 ${accounts.length} 个账号状态...`:'正在保存状态...';try{await markRequest({status:value},accounts);$("markState").textContent=accounts.length>1?`已批量记录 ${accounts.length} 个同名账户`:'已记录本地台账';$("saveStatus").textContent='状态已保存';}catch(err){$("saveStatus").textContent=err.message;}finally{$("status").disabled=false;}}
    async function save(){const action=state.selectedAction==='自定义'?($("customAction").value.trim()||'自定义'):state.selectedAction,accounts=targetAccounts();$("saveBtn").disabled=true;$("saveStatus").textContent=accounts.length>1?`正在保存 ${accounts.length} 个账号...`:'正在保存...';try{await markRequest({action,group:$("group").value.trim(),tags:$("tags").value.trim(),note:$("note").value.trim(),status:$("status").value,owner:$("owner").value.trim()},accounts);applyLocalAction(action,accounts);state.formDirty=false;$("markState").textContent=accounts.length>1?`已批量记录 ${accounts.length} 个同名账户`:'已记录本地台账';$("saveStatus").textContent=accounts.length>1?`已保存 ${accounts.length} 个账号`:'已保存';}catch(err){$("saveStatus").textContent=err.message;}finally{$("saveBtn").disabled=false;}}
    function preview(url){$("previewFrame").src=url;$("previewOpen").href=url;$("previewDialog").showModal();}
    async function loadInlineKline(){const db=state.detail?.database||{},source=db.latestSource||{},section=$("inlineKlineSection"),status=$("inlineKlineStatus"),frame=$("inlineKlineFrame");if(!db.exists||!source.platform||!source.server){section.hidden=true;return;}const current=filters(),query=new URLSearchParams({platform:current.platform||source.platform,server:current.server||source.server,recentOrders:'300'}),key=`${query.toString()}|${db.lastTime||db.orderCount||0}`;section.hidden=false;if(state.inlineKlineKey===key)return;state.inlineKlineKey=key;status.textContent='正在读取最近 300 笔订单与 M1 报价...';frame.removeAttribute('srcdoc');try{const response=await fetch(`/api/accounts/by-login/${encodeURIComponent(LOGIN)}/inline-kline?${query}`);if(!response.ok){let issue={};try{issue=await response.json();}catch(_){ }throw new Error(issue.error||issue.detail||`HTTP ${response.status}`);}frame.srcdoc=await response.text();status.textContent='最近 300 笔订单 · 已直接加载';}catch(err){frame.srcdoc='';status.textContent=`K 线暂时无法加载：${err.message}`;}}
    async function generate(){const start=databaseTime($("chartStart").value),end=databaseTime($("chartEnd").value,true);if(start&&end&&start>end){$("jobText").textContent='开始时间不能晚于结束时间';return;}$("generateBtn").disabled=true;$("jobText").textContent='正在提交生成任务...';$("jobProgress").style.width='3%';try{const current=filters(),includeTimeline=$("includeTimeline").checked,payload={account:LOGIN,platform:current.platform,server:current.server,symbol:$("chartSymbol").value,start,end,includeTimeline,refreshTimelineCache:includeTimeline&&$("refreshTimelineCache").checked};const data=await json('/api/kline/generate-from-db',{method:'POST',body:JSON.stringify(payload)});$("jobText").textContent=`已提交 · ${payload.symbol||'全部品种'} · ${start||'全量'} 至 ${end||'全量'}${includeTimeline?' · 含资金回放':''}`;poll(data.job.id);}catch(err){$("jobText").textContent=err.message;$("generateBtn").disabled=false;}}
    async function poll(id,options={}){clearTimeout(state.jobTimer);try{const data=await json(`/api/kline/jobs/${encodeURIComponent(id)}`),job=data.job||{};$("jobProgress").style.width=`${Math.max(0,Math.min(100,Number(job.percent||0)))}%`;$("jobText").textContent=[job.message,job.elapsedSeconds?`${job.elapsedSeconds}s`:''].filter(Boolean).join(' · ');if(job.status==='done'){if(job.chart){state.detail.charts=[job.chart,...(state.detail.charts||[]).filter(c=>c.name!==job.chart.name)];renderCharts(state.detail.charts);if(!options.auto)preview(job.chart.url);}$("generateBtn").disabled=false;return;}if(job.status==='failed'||job.status==='missing'){$("generateBtn").disabled=false;return;}state.jobTimer=setTimeout(()=>poll(id,options),1000);}catch(err){$("jobText").textContent=err.message;$("generateBtn").disabled=false;}}
    function toxicMode(){return document.querySelector('input[name="toxicMode"]:checked')?.value||'selected';}
    function renderToxicSelector(){
      $("toxicSelector").innerHTML=TOXIC_TYPES.map(item=>`<label class="toxic-check"><input type="checkbox" value="${esc(item.id)}" /><span>${esc(item.label)}${item.tick?' · Tick':''}</span></label>`).join('');
    }
    function updateToxicMode(){
      const screen=toxicMode()==='screen';$("toxicSelector").classList.toggle('disabled',screen);
      document.querySelectorAll('#toxicSelector input').forEach(input=>input.disabled=screen);
      $("toxicModeNote").textContent=screen?'对全部类型做轻量初筛；结果中可对单项继续深度检测。':'仅运行勾选的专项算法；需要 Tick 的项目会按需读取候选时段报价。';
      $("startToxic").textContent=screen?'开始初筛':'开始专项检测';
    }
    const fundsKinds={trade_open:'开仓',trade_close:'平仓',deposit:'外部入金',withdrawal:'外部出金',internal_transfer:'内部划转',bonus_grant:'Credit 增加',bonus_remove:'Credit 扣减',negative_balance_clear:'负余额清零',compensation:'补偿入账',cash_reversal:'资金冲正',adjustment:'账务调整',other_balance:'余额调整'};
    const fundsValue=value=>value===null||value===undefined||value===''?'—':amount(value);
    const validFundsValue=value=>value!==null&&value!==undefined&&value!==''&&Number.isFinite(Number(value));
    function fundsSeriesPoints(points,key,min,max,width,height){return points.filter(row=>validFundsValue(row[key])).map((row,index,rows)=>`${(rows.length===1?0:index/(rows.length-1))*width},${height-((Number(row[key])-min)/(max-min||1))*height}`).join(' ');}
    function jumpToHistoricalFundsEvent(value){
      const funds=state.historicalFunds,eventIndex=Number(value),events=funds.data?.events||[];
      if(!Number.isInteger(eventIndex)||eventIndex<0||eventIndex>=events.length)return;
      funds.page=Math.floor(eventIndex/funds.pageSize)+1;renderHistoricalFundsEvents();
      const target=$(`historicalFundsEvent-${eventIndex}`);if(!target)return;
      target.scrollIntoView({block:'center',behavior:'smooth'});target.classList.add('funds-event-focus');
      setTimeout(()=>target.classList.remove('funds-event-focus'),1800);
    }
    function bindLiquidationJumps(root){root.querySelectorAll('[data-liquidation-index]').forEach(node=>{const jump=()=>jumpToHistoricalFundsEvent(node.dataset.liquidationIndex);node.addEventListener('click',jump);node.addEventListener('keydown',event=>{if(event.key==='Enter'||event.key===' '){event.preventDefault();jump();}});});}
    function renderHistoricalFundsChart(data){
      const curve=(data.curve||[]).filter(row=>validFundsValue(row.balance)||validFundsValue(row.credit));
      if(!curve.length){$("historicalFundsChart").innerHTML='<div class="empty-state">没有可绘制的余额或 Credit 快照</div>';return;}
      const values=curve.flatMap(row=>[row.balance,row.credit]).filter(validFundsValue).map(Number),min=Math.min(...values),max=Math.max(...values),pad=Math.max((max-min)*.08,1),low=min-pad,high=max+pad,width=1000,height=190;
      const balance=fundsSeriesPoints(curve,'balance',low,high,width,height),credit=fundsSeriesPoints(curve,'credit',low,high,width,height);
      const markers=(data.liquidationPoints||[]).map(point=>{const index=curve.findIndex(row=>Number(row.eventIndex)===Number(point.eventIndex));if(index<0)return '';const row=curve[index],value=validFundsValue(row.balance)?Number(row.balance):Number(row.credit),x=(curve.length===1?0:index/(curve.length-1))*width,y=height-((value-low)/(high-low||1))*height,label=`${point.label||'爆仓标记'} · ${point.timestamp||'-'} · ${point.orderId||'无订单号'}`;return `<circle class="funds-chart-liquidation" data-liquidation-index="${Number(point.eventIndex)}" cx="${x.toFixed(2)}" cy="${y.toFixed(2)}" r="5" tabindex="0" role="button" aria-label="${esc(label)}，跳转到事件"><title>${esc(label)}</title></circle>`;}).join('');
      const chart=$("historicalFundsChart");chart.innerHTML=`<div class="funds-chart-head"><b>余额与 Credit 回放</b><span>蓝色：余额 · 金色：Credit · 红点：爆仓标记，可点击跳转</span></div><svg viewBox="0 0 ${width} 220" role="img" aria-label="历史余额与 Credit 曲线"><line class="funds-chart-grid" x1="0" y1="0" x2="${width}" y2="0"/><line class="funds-chart-grid" x1="0" y1="${height/2}" x2="${width}" y2="${height/2}"/><line class="funds-chart-grid" x1="0" y1="${height}" x2="${width}" y2="${height}"/><polyline class="funds-chart-balance" points="${balance}"/><polyline class="funds-chart-credit" points="${credit}"/>${markers}<text x="0" y="12" fill="#7895b8" font-size="11">最高 ${esc(amount(high))}</text><text x="0" y="214" fill="#7895b8" font-size="11">最低 ${esc(amount(low))}</text></svg>`;bindLiquidationJumps(chart);
    }
    function renderHistoricalFundsEvents(){
      const funds=state.historicalFunds,data=funds.data,events=data?.events||[],start=(funds.page-1)*funds.pageSize,rows=events.slice(start,start+funds.pageSize),pages=Math.max(1,Math.ceil(events.length/funds.pageSize));
      $("historicalFundsPageStatus").textContent=`第 ${funds.page} / ${pages} 页，共 ${num(events.length,0)} 条`;
      $("historicalFundsPrev").disabled=funds.page<=1;$("historicalFundsNext").disabled=funds.page>=pages;
      $("historicalFundsEvents").innerHTML=rows.length?rows.map(row=>`<tr id="historicalFundsEvent-${Number(row.eventIndex)}" class="${row.liquidation?'funds-event-liquidation':''}"><td>${esc(row.timestamp||'-')}</td><td><span class="funds-kind ${esc(row.kind||'other')}">${esc(fundsKinds[row.kind]||'其他')}</span>${row.liquidation?`<span class="funds-liquidation-label">${esc(row.liquidation.label||'爆仓标记')}</span>`:''}</td><td>${esc(row.orderId||'-')}</td><td>${esc(row.symbol||'-')}</td><td>${esc(row.comment||'-')}</td><td class="${profitClass(row.deltaBalance)}">${fundsValue(row.deltaBalance)}</td><td class="${profitClass(row.deltaCredit)}">${fundsValue(row.deltaCredit)}</td><td class="${profitClass(row.realizedPnl)}">${fundsValue(row.realizedPnl)}</td><td>${fundsValue(row.balance)}</td><td>${fundsValue(row.credit)}</td><td>${row.equityStatus==='authoritative_daily'?'日快照':row.equityStatus==='missing_intraday_snapshot'?'无盘中快照':row.equityStatus==='before_first_anchor'?'首个快照前':'—'}</td></tr>`).join(''):'<tr><td colspan="11"><div class="empty-state">没有可回放的事件</div></td></tr>';
    }
    function renderHistoricalFunds(data){
      const summary=data.summary||{},currency=summary.currency||data.currency||'USD',items=[['外部入金',summary.externalDeposit],['外部出金',summary.externalWithdrawal],['外部净入金',summary.externalNetDeposit],['内部划转',summary.internalTransfer],['Credit 增加',summary.bonusGranted],['Credit 扣减',summary.bonusRemoved],['负余额清零',summary.negativeBalanceCleared],['爆仓标记',summary.liquidationCount]];
      $("historicalFundsSummary").innerHTML=items.map(([label,value],index)=>`<div><span>${esc(label)}${index===7?'':' ('+esc(currency)+')'}</span><b class="${index===7?'':profitClass(index===1||index===5?-Number(value||0):value)}">${index===7?num(value,0):fundsValue(value)}</b></div>`).join('');
      const anchorNote=data.coverage?.dailyAnchorsAvailable===false?`${data.coverage?.dailyAnchorReason||'历史日快照不可用。'} `:'';
      $("historicalFundsNote").textContent=`来源：${data.platform||'-'} / ${data.server||'-'}。完整读取 ${num(data.coverage?.eventRows,0)} 条订单与资金事件、${num(data.coverage?.dailyAnchors,0)} 条日快照；${anchorNote}余额和 Credit 按事件回放，权益不补造。`;
      const equityNote=summary.equityCoverage==='daily_anchors_only'?'权益只在日快照处有事实值；':summary.equityCoverage==='current_anchor_only'?'仅当前账户快照有权益事实值；':'未取得权益快照；';
      $("historicalFundsEventNote").textContent=`${equityNote}外部入出金、内部划转与 Credit 分开列示。`;
      const liquidations=data.liquidationPoints||[],liquidationPanel=$("historicalFundsLiquidations");liquidationPanel.hidden=!liquidations.length;liquidationPanel.innerHTML=liquidations.length?`<b>爆仓点位 ${num(liquidations.length,0)} 个</b>${liquidations.map(point=>`<button class="funds-liquidation-jump" type="button" data-liquidation-index="${Number(point.eventIndex)}">${esc(point.label||'爆仓标记')} · ${esc(point.timestamp||'-')} · ${esc(point.orderId||'无订单号')}</button>`).join('')}`:'';bindLiquidationJumps(liquidationPanel);
      renderHistoricalFundsChart(data);renderHistoricalFundsEvents();
    }
    async function openHistoricalFunds(){
      const funds=state.historicalFunds;if(funds.loading)return;funds.loading=true;funds.data=null;funds.page=1;$("historicalFundsStatus").textContent='正在读取完整历史...';$("historicalFundsSummary").innerHTML='';$("historicalFundsChart").innerHTML='<div class="empty-state">正在读取订单、资金流水和日快照...</div>';$("historicalFundsEvents").innerHTML='<tr><td colspan="11"><div class="empty-state">正在读取...</div></td></tr>';$("historicalFundsDialog").showModal();
      try{const current=filters(),query=new URLSearchParams({platform:current.platform||'',server:current.server||''});const data=await json(`/api/accounts/by-login/${encodeURIComponent(LOGIN)}/historical-funds?${query}`);if(!data.available)throw new Error(data.reason||'未找到历史资金数据');funds.data=data;renderHistoricalFunds(data);$("historicalFundsStatus").textContent=`已回放 ${num(data.summary?.eventCount,0)} 条事件`;}
      catch(err){$("historicalFundsStatus").textContent=err.message;$("historicalFundsChart").innerHTML=`<div class="empty-state">${esc(err.message)}</div>`;}
      finally{funds.loading=false;}
    }
    function openToxic(){if(!$("toxicSelector").children.length)renderToxicSelector();updateToxicMode();$("toxicDialog").showModal();}
    function selectedToxicTypes(){return [...document.querySelectorAll('#toxicSelector input:checked')].map(input=>input.value);}
    function toxicResultClass(result){return Number(result.score)>=90?'critical':Number(result.score)>=75?'high':Number(result.score)>=60?'warning':'';}
    function toxicSyncDelta(value){const parsed=Number(value);return Number.isFinite(parsed)?`${parsed>0?'+':''}${parsed.toFixed(3).replace(/\.000$/,'')} 秒`:'-';}
    function toxicSyncDirection(value){return value==='buy'?'买入':value==='sell'?'卖出':value||'-';}
    function renderToxicSyncRows(sync,peer=''){
      const rows=(sync?.comparisonRows||[]).filter(row=>!peer||String(row.peerAccount)===String(peer)).slice(0,200);
      if(!rows.length)return '<tr><td colspan="7"><div class="empty-state">该关联账户暂无可展示的订单对</div></td></tr>';
      return rows.map(row=>{const href=`/account/${encodeURIComponent(row.peerAccount||'')}?platform=${encodeURIComponent(row.peerPlatform||'')}&server=${encodeURIComponent(row.peerServer||'')}`;return `<tr><td><span class="toxic-sync-state ${row.closeSynchronized?'full':'open-only'}">${row.closeSynchronized?'开平仓同步':'仅开仓同步'}</span></td><td><b>${esc(row.targetAccount||LOGIN)}</b><small>#${esc(row.targetTicket||'-')}</small></td><td><a href="${esc(href)}" target="_blank">${esc(row.peerAccount||'-')}</a><small>#${esc(row.peerTicket||'-')} · ${esc(row.peerServer||'-')}</small></td><td><b>${esc(row.targetSymbol||'-')}</b><small>${esc(toxicSyncDirection(row.targetDirection))}</small></td><td><b>${num(row.targetVolume,2)} / ${num(row.peerVolume,2)}</b><small>主体 / 关联</small></td><td><b>${esc(toxicSyncDelta(row.openDeltaSeconds))}</b><small>${esc(row.targetOpened||'-')}</small><small>${esc(row.peerOpened||'-')}</small></td><td><b class="${row.closeSynchronized?'positive':'negative'}">${esc(toxicSyncDelta(row.closeDeltaSeconds))}</b><small>${esc(row.targetClosed||'-')}</small><small>${esc(row.peerClosed||'-')}</small></td></tr>`;}).join('');
    }
    function renderToxicSyncComparison(result){
      const sync=result?.pushSync;if(!sync?.available)return '';
      const peers=sync.suspectedAccounts||[],options=peers.map(item=>`<option value="${esc(item.account||'')}">${esc(item.account||'-')} · ${num(item.matches,0)} 单</option>`).join('');
      const peerRows=peers.map(item=>{const href=`/account/${encodeURIComponent(item.account||'')}?platform=${encodeURIComponent(item.platform||'')}&server=${encodeURIComponent(item.server||'')}`;return `<tr><td><a href="${esc(href)}" target="_blank">${esc(item.account||'-')}</a></td><td>${esc(item.server||'-')}</td><td>${num(item.matches,0)} 单</td><td>${num(item.closeMatches,0)} 单</td><td>${num(item.matchRatio,1)}%</td><td>${num(item.closeMatchRatio,1)}%</td></tr>`;}).join('');
      return `<div class="toxic-sync-comparison"><div class="toxic-sync-head"><div><h3>同步订单逐笔对比</h3><small>同品种、同方向且开仓相差不超过2秒；平仓时间差单独判断</small></div><span>抽样 ${num(sync.sampledOrders,0)} 单</span></div><div class="toxic-sync-kpis"><div><span>任意开仓匹配</span><b>${num(sync.matchedRatio,1)}%</b></div><div><span>反复账户协调开仓</span><b>${num(sync.coordinatedMatchedRatio,1)}%</b></div><div><span>协调手数覆盖</span><b>${num(sync.coordinatedVolumeRatio,1)}%</b></div><div><span>协调平仓</span><b>${num(sync.coordinatedCloseRatio,1)}%</b></div><div><span>反复关联账户</span><b>${num(sync.recurringPeerAccounts,0)}</b></div><div><span>重复门槛</span><b>${num(sync.recurringMinMatches,0)} 单</b></div></div><div class="toxic-sync-table-wrap"><table><thead><tr><th>关联账户</th><th>服务器</th><th>同步开仓</th><th>同步平仓</th><th>开仓覆盖</th><th>平仓覆盖</th></tr></thead><tbody>${peerRows||'<tr><td colspan="6"><div class="empty-state">没有达到重复门槛的关联账户</div></td></tr>'}</tbody></table></div><div class="toxic-sync-detail-head"><div><h3>相似订单明细</h3><small>共 ${num(sync.comparisonTotal,0)} 组，页面最多显示200组</small></div><label>关联账户<select id="toxicSyncPeerFilter"><option value="">全部账户</option>${options}</select></label></div><div class="toxic-sync-detail-wrap"><table class="toxic-sync-detail-table"><thead><tr><th>同步结论</th><th>主体订单</th><th>关联订单</th><th>品种 / 方向</th><th>手数对比</th><th>开仓时间对比</th><th>平仓时间对比</th></tr></thead><tbody id="toxicSyncRows">${renderToxicSyncRows(sync)}</tbody></table></div>${sync.comparisonTruncated?`<div class="toxic-accomplices-note">对照记录较多，后端仅保留前 ${num(sync.comparisonLimit,0)} 组。</div>`:''}</div>`;
    }
    function renderEvidenceChain(row){
      const chain=row.evidenceChain||{};
      if(!chain.headline)return '';
      const facts=(chain.facts||[]).map(item=>`<div class="toxic-chain-fact" data-strength="${esc(item.strength||'partial')}"><b>${esc(item.title||'-')}</b><span>${esc(item.text||'-')}</span></div>`).join('');
      const list=(items)=>items&&items.length?`<ul>${items.map(item=>`<li>${esc(item)}</li>`).join('')}</ul>`:'';
      const basis=chain.scoreBasis||{};
      const basisText=`结构 ${num(basis.structure,1)}/40 · Tick ${num(basis.tick,1)}/35 · 协同 ${num(basis.coordination,1)}/15 · 辅助 ${num(basis.context,1)}/10 · 一致性加分 ${num(basis.consistencyBonus,1)} · 反证扣分 ${num(basis.counterevidence,1)}`;
      return `<div class="toxic-chain"><div class="toxic-chain-head">${esc(chain.headline)}</div><div class="toxic-chain-reasoning">${esc(chain.reasoning||'')}</div><div class="toxic-chain-facts">${facts}</div><div class="toxic-chain-section"><b>已观察到的风险</b>${list(chain.observedRisks||chain.riskPoints)}</div><div class="toxic-chain-section"><b>反证与替代解释</b>${list(chain.counterpoints)}</div><div class="toxic-chain-section"><b>仍不能确定</b>${list(chain.uncertainties)}</div><div class="toxic-chain-section"><b>下一步复核</b>${list(chain.nextChecks)}</div><div class="toxic-chain-reasoning">评分依据：${esc(basisText)}</div></div>`;
    }
    function renderPositionRiskEvidence(row){
      if(!['weekend_gap_trading','open_betting'].includes(row.type))return '';
      const event=row?.evidence?.bestEvent||{},coverage=event.peerSearchCoverage||{};
      if(!Object.keys(event).length)return '';
      const direction=value=>value==='buy'?'买入':value==='sell'?'卖出':value||'-';
      const matchRows=(rows,empty)=>rows?.length?rows.slice(0,100).map(item=>`<tr><td><b>${esc(item.targetOrderId||'-')}</b><small>仓位 ${esc(item.targetPositionId||'-')}</small></td><td><b>${esc(item.account||'-')}</b><small>${esc(item.platform||'-')} / ${esc(item.server||'-')}</small><small>${esc(item.database||'-')}</small></td><td><b>${esc(item.orderId||'-')}</b><small>仓位 ${esc(item.positionId||'-')} · 成交 ${esc(item.dealId||'-')}</small></td><td><b>${esc(item.symbol||'-')}</b><small>主体 ${esc(direction(item.targetDirection))} ${num(item.targetVolume,2)} 手 / 关联 ${esc(direction(item.direction))} ${num(item.volume,2)} 手</small><small>手数相似度 ${num(Number(item.lotSimilarity||0)*100,1)}%</small></td><td><b>差 ${num(item.openDeltaSeconds,3)} 秒</b><small>主体 ${esc(item.targetOpenTime||'-')}</small><small>关联 ${esc(item.openTime||'-')}</small></td><td><b>差 ${num(item.closeDeltaSeconds,3)} 秒</b><small>主体 ${esc(item.targetCloseTime||'-')}</small><small>关联 ${esc(item.closeTime||'-')}</small></td></tr>`).join(''):`<tr><td colspan="6"><div class="empty-state">${esc(empty)}</div></td></tr>`;
      const heavyRows=(event.heavyOrders||[]).slice(0,100).map(item=>`<tr><td><b>${esc(item.orderId||'-')}</b><small>仓位 ${esc(item.positionId||'-')} · 成交 ${esc(item.dealId||'-')}</small></td><td>${esc(item.symbol||'-')} / ${esc(direction(item.direction))}</td><td>${num(item.volume,2)} 手</td><td>${esc(item.openTime||'-')}</td><td>${esc(item.closeTime||'未平仓')}</td><td>${num(item.positionNetProfit,2)}</td></tr>`).join('');
      const failureRows=(coverage.failures||[]).map(item=>`<li>${esc(item.platform||'-')} / ${esc(item.server||'-')} / ${esc(item.database||'-')}：${esc(item.reason||'-')}</li>`).join('');
      const scannedSources=(coverage.scannedSources||[]).map(item=>`${esc(item.platform||'-')} / ${esc(item.server||'-')} / ${esc(item.database||'-')}`).join('；');
      const skipped=(coverage.skippedTargetOrders||[]).map(item=>`${esc(item.orderId||'未知订单')}：${esc(item.reason||'-')}`).join('；');
      return `<div class="toxic-sync-comparison"><div class="toxic-sync-head"><div><h3>重仓与全平台同步开平仓证据</h3><small>${esc(coverage.scope||'AC/DBG 全平台 MT4 + MT5')} · 开仓和平仓均需在 ±${num(coverage.toleranceSeconds??5,0)} 秒内</small></div><span>${esc(coverage.status||'数据不足')} · ${num(coverage.scannedSourceCount,0)}/${num(coverage.physicalSourceTotal,0)} 库</span></div><div class="toxic-sync-kpis"><div><span>峰值仓位</span><b>${num(event.peakLots,2)} 手 / ${num(event.peakOrderCount||event.orderCount,0)} 单</b></div><div><span>账户杠杆</span><b>1:${num(event.leverage,0)}</b></div><div><span>预计占用保证金</span><b>${num(event.estimatedMargin,2)}</b></div><div><span>保证金占权益</span><b>${num(Number(event.marginRatio||0)*100,1)}%（越高越满）</b></div><div><span>估算保证金水平</span><b>${num(event.estimatedMarginLevel,1)}%（越低越满）</b></div><div><span>穿仓判断</span><b>${esc(event.penetrationStatus||'数据不足')}</b></div></div>${scannedSources?`<div class="toxic-accomplices-note">已完成：${scannedSources}</div>`:''}${(event.penetrationDataGaps||[]).length?`<div class="toxic-accomplices-note">数据不足原因：${esc(event.penetrationDataGaps.join('；'))}</div>`:''}${skipped?`<div class="toxic-accomplices-note">未参与同步匹配：${skipped}</div>`:''}${failureRows?`<div class="toxic-chain-section"><b>未完成的数据源</b><ul>${failureRows}</ul></div>`:''}${event.peerMatchesTruncated?`<div class="toxic-accomplices-note">订单对较多，接口最多保留 ${num(event.peerMatchDetailLimit||500,0)} 组明细，本页每类最多显示100组；账号总数和订单对总数仍按完整结果计算。</div>`:''}<div class="toxic-sync-detail-head"><div><h3>重仓开赌订单</h3><small>${num((event.heavyOrders||[]).length,0)} 单</small></div></div><div class="toxic-sync-detail-wrap"><table class="toxic-sync-detail-table"><thead><tr><th>订单 / 仓位</th><th>品种 / 方向</th><th>手数</th><th>开仓时间</th><th>平仓时间</th><th>仓位净盈亏</th></tr></thead><tbody>${heavyRows||'<tr><td colspan="6"><div class="empty-state">没有订单级明细</div></td></tr>'}</tbody></table></div><div class="toxic-sync-detail-head"><div><h3>同步同向订单</h3><small>共 ${num(event.sameDirectionMatchTotal??(event.sameDirectionMatches||[]).length,0)} 组</small></div></div><div class="toxic-sync-detail-wrap"><table class="toxic-sync-detail-table"><thead><tr><th>目标订单</th><th>同行账号 / 来源</th><th>同行订单 / 仓位</th><th>品种 / 方向 / 手数</th><th>同步开仓</th><th>同步平仓</th></tr></thead><tbody>${matchRows(event.sameDirectionMatches,'未找到同步同向开平仓订单')}</tbody></table></div><div class="toxic-sync-detail-head"><div><h3>同步反向疑似对锁订单</h3><small>共 ${num(event.oppositeDirectionMatchTotal??(event.oppositeDirectionMatches||[]).length,0)} 组</small></div></div><div class="toxic-sync-detail-wrap"><table class="toxic-sync-detail-table"><thead><tr><th>目标订单</th><th>反向账号 / 来源</th><th>反向订单 / 仓位</th><th>品种 / 方向 / 手数</th><th>同步开仓</th><th>同步平仓</th></tr></thead><tbody>${matchRows(event.oppositeDirectionMatches,'未找到同步反向开平仓订单')}</tbody></table></div></div>`;
    }
    function renderInternalLockEvidence(row){
      if(row.type!=='internal_lock_arbitrage')return '';
      const query=row?.evidence?.hedgeQuery||{},coverage=query.coverage||{};
      if(!Object.keys(query).length)return '';
      const direction=value=>value==='buy'?'买入':value==='sell'?'卖出':value||'-';
      const accountRows=(query.accounts||[]).map(item=>{const href=`/account/${encodeURIComponent(item.account||'')}?platform=${encodeURIComponent(item.platform||'')}&server=${encodeURIComponent(item.server||'')}`;return `<tr><td><a href="${esc(href)}" target="_blank">${esc(item.account||'-')}</a></td><td>${esc(item.platform||'-')} / ${esc(item.server||'-')}<small>${esc(item.database||'-')}</small></td><td>${num(item.matchCount,0)}</td><td>${num(item.targetLots,2)} / ${num(item.peerLots,2)} 手</td></tr>`;}).join('');
      const matchRows=(query.matches||[]).slice(0,100).map(item=>`<tr><td><b>${esc(item.targetOrderId||'-')}</b><small>仓位 ${esc(item.targetPositionId||'-')}</small></td><td><a href="/account/${encodeURIComponent(item.account||'')}?platform=${encodeURIComponent(item.platform||'')}&server=${encodeURIComponent(item.server||'')}" target="_blank">${esc(item.account||'-')}</a><small>${esc(item.platform||'-')} / ${esc(item.server||'-')}</small><small>${esc(item.database||'-')}</small></td><td><b>${esc(item.orderId||'-')}</b><small>仓位 ${esc(item.positionId||'-')} · 成交 ${esc(item.dealId||'-')}</small></td><td><b>${esc(item.symbol||'-')}</b><small>主体 ${esc(direction(item.targetDirection))} ${num(item.targetVolume,2)} 手 / 对方 ${esc(direction(item.direction))} ${num(item.volume,2)} 手</small><small>手数相似度 ${num(Number(item.lotSimilarity||0)*100,1)}%</small></td><td><b>差 ${num(item.openDeltaSeconds,3)} 秒</b><small>主体 ${esc(item.targetOpenTime||'-')}</small><small>对方 ${esc(item.openTime||'-')}</small></td><td><b>差 ${num(item.closeDeltaSeconds,3)} 秒</b><small>主体 ${esc(item.targetCloseTime||'-')}</small><small>对方 ${esc(item.closeTime||'-')}</small></td></tr>`).join('');
      const failures=(coverage.failures||[]).map(item=>`<li>${esc(item.platform||'-')} / ${esc(item.server||'-')} / ${esc(item.database||'-')}：${esc(item.reason||'-')}</li>`).join('');
      return `<div class="toxic-sync-comparison"><div class="toxic-sync-head"><div><h3>平台内多账户对锁查询</h3><small>${esc(query.rule||'同品种、反方向、手数相似度至少80%，且双方开仓和平仓时间差都不超过5秒')} · ${esc(query.scope||'AC/DBG 全平台 MT4 + MT5')}</small></div><span>${esc(coverage.status||'数据不足')} · ${num(coverage.scannedSourceCount,0)}/${num(coverage.physicalSourceTotal,0)} 库</span></div><div class="toxic-sync-kpis"><div><span>已检查目标订单</span><b>${num(query.targetOrderCount,0)}</b></div><div><span>疑似对锁账号</span><b>${num(query.accountCount,0)}</b></div><div><span>反向同步订单对</span><b>${num(query.matchTotal,0)}</b></div><div><span>对锁手数门槛</span><b>${num(Number(query.lotSimilarityThreshold||0.8)*100,0)}%</b></div><div><span>未平仓未参与</span><b>${num(query.openPositionCount,0)}</b></div></div>${failures?`<div class="toxic-chain-section"><b>未完成的数据源</b><ul>${failures}</ul></div>`:''}${query.detailsTruncated?`<div class="toxic-accomplices-note">账号与订单对总数按完整结果计算；订单明细最多保留 ${num(query.detailLimit,0)} 组，本页显示前100组。</div>`:''}<div class="toxic-sync-detail-head"><div><h3>疑似对锁账号</h3><small>${num(query.accountCount,0)} 个</small></div></div><div class="toxic-sync-detail-wrap"><table class="toxic-sync-detail-table"><thead><tr><th>账号</th><th>平台 / 服务器 / 数据库</th><th>匹配订单对</th><th>主体 / 对方手数合计</th></tr></thead><tbody>${accountRows||'<tr><td colspan="4"><div class="empty-state">未发现同时满足开平仓同步与手数接近条件的反向账户</div></td></tr>'}</tbody></table></div><div class="toxic-sync-detail-head"><div><h3>反向同步开平仓订单</h3><small>共 ${num(query.matchTotal,0)} 组</small></div></div><div class="toxic-sync-detail-wrap"><table class="toxic-sync-detail-table"><thead><tr><th>主体订单</th><th>对锁账号 / 来源</th><th>对方订单 / 仓位</th><th>品种 / 方向 / 手数</th><th>同步开仓</th><th>同步平仓</th></tr></thead><tbody>${matchRows||'<tr><td colspan="6"><div class="empty-state">未找到同时满足手数相似度至少80%的反向同步开平仓订单</div></td></tr>'}</tbody></table></div><div class="toxic-accomplices-note">反向同步只表示疑似对锁，不能单独证明双方存在套利关系。</div></div>`;
    }
    function renderToxicResults(result){
      state.toxic.lastResult=result;const rows=result?.results||[];
      if(!rows.length){$("toxicResults").innerHTML='<div class="toxic-empty">没有检测结果</div>';return;}
      $("toxicResults").innerHTML=rows.map(row=>{
        const metrics=(row.metrics||[]).map(item=>`<b>${esc(item.label)}：${esc(item.value)}</b>`).join(' · ');
        const triggers=(row.triggeredRules||[]).map(item=>`<b>${esc(item)}</b>`).join('；');
        const limits=(row.limitations||[]).join('；');
        const analysis=(row.analysis||[]).map(item=>`<div class="toxic-analysis-row"><b>${esc(item.title)}</b><span>${esc(item.text)}</span></div>`).join('');
        const accomplices=row.type==='market_pushing'?(row.suspectedAccomplices||[]):[];
        const accompliceList=accomplices.map(item=>{const href=`/account/${encodeURIComponent(item.account||'')}?platform=${encodeURIComponent(item.platform||'')}&server=${encodeURIComponent(item.server||'')}`;return `<div class="toxic-accomplice"><a href="${esc(href)}" target="_blank">${esc(item.account||'-')}</a><br>${esc(item.server||item.platform||'-')} · 开仓同步 ${num(item.matches,0)} 次 (${num(item.matchRatio,1)}%) · 平仓同步 ${num(item.closeMatches,0)} 次</div>`;}).join('');
        const accompliceHtml=row.type==='market_pushing'?`<div class="toxic-accomplices"><div class="toxic-accomplices-title">疑似同伙账户</div>${accompliceList?`<div class="toxic-accomplice-list">${accompliceList}</div>`:`<div class="toxic-accomplices-note">当前未找到达到反复匹配门槛的协同账户；这不会降低本账户自身的推盘嫌疑。</div>`}</div>`:'';
        const chain=renderEvidenceChain(row);
        const readable=(chain|| (analysis?`<div class="toxic-analysis">${analysis}</div>`:`<div>${esc(row.summary||'')}</div>`))+accompliceHtml+renderInternalLockEvidence(row)+renderPositionRiskEvidence(row);
        const tech=(metrics||triggers||limits)?`<details class="toxic-tech"><summary>查看技术明细</summary><div class="toxic-tech-body">${metrics}${triggers?`<br><span>触发规则：</span>${triggers}`:''}${limits?`<div class="toxic-limit">模型限制：${esc(limits)}</div>`:''}</div></details>`:'';
        const queryOnly=row.type==='internal_lock_arbitrage',query=row?.evidence?.hedgeQuery||{};
        const scoreValue=queryOnly?num(query.accountCount,0):num(row.score,1),scoreLabel=queryOnly?'疑似账号':esc(row.level);
        return `<div class="toxic-result ${toxicResultClass(row)}"><div class="toxic-score">${scoreValue}<small>${scoreLabel}</small></div><div class="toxic-result-title"><b>${esc(row.label)}</b><span>${queryOnly?'专项查询':(row.stage==='deep'?'深度检测':'初筛')} · 置信度 ${num(row.confidence,0)}%</span></div><div class="toxic-result-summary">${readable}${tech}</div><div class="toxic-result-actions">${row.stage!=='deep'?`<button type="button" data-toxic-deep="${esc(row.type)}">深度检测</button>`:''}</div></div>`;
      }).join('')+renderToxicSyncComparison(result);
      document.querySelectorAll('[data-toxic-deep]').forEach(btn=>btn.addEventListener('click',()=>deepToxic(btn.dataset.toxicDeep)));
      const syncFilter=$("toxicSyncPeerFilter");if(syncFilter)syncFilter.addEventListener('change',()=>{$("toxicSyncRows").innerHTML=renderToxicSyncRows(result.pushSync,syncFilter.value);});
    }
    async function startToxic(modeOverride=null,typesOverride=null){
      if(state.toxic.running)return;const mode=modeOverride||toxicMode(),types=typesOverride||selectedToxicTypes();
      if(mode==='selected'&&!types.length){$("toxicStatus").textContent='请至少勾选一个检测项目';return;}
      state.toxic.running=true;$("startToxic").disabled=true;$("toxicBtn").disabled=true;$("toxicProgress").style.width='3%';$("toxicStatus").textContent='正在提交检测任务...';
      try{const current=filters(),payload={mode,types,...current};const data=await json(`/api/accounts/by-login/${encodeURIComponent(LOGIN)}/toxic-checks`,{method:'POST',body:JSON.stringify(payload)});pollToxic(data.job.id);}catch(err){$("toxicStatus").textContent=err.message;state.toxic.running=false;$("startToxic").disabled=false;$("toxicBtn").disabled=false;}
    }
    function deepToxic(type){
      document.querySelector('input[name="toxicMode"][value="selected"]').checked=true;updateToxicMode();
      document.querySelectorAll('#toxicSelector input').forEach(input=>input.checked=input.value===type);startToxic('selected',[type]);
    }
    async function pollToxic(id){
      clearTimeout(state.toxic.jobTimer);
      try{const data=await json(`/api/toxic/jobs/${encodeURIComponent(id)}`),job=data.job||{};$("toxicProgress").style.width=`${Math.max(0,Math.min(100,Number(job.percent||0)))}%`;$("toxicStatus").textContent=[job.message,job.elapsedSeconds?`${job.elapsedSeconds}s`:'',job.cached?'缓存':''].filter(Boolean).join(' · ');
        if(job.status==='done'){renderToxicResults(job.result);state.toxic.running=false;$("startToxic").disabled=false;$("toxicBtn").disabled=false;return;}
        if(job.status==='failed'||job.status==='missing'){state.toxic.running=false;$("startToxic").disabled=false;$("toxicBtn").disabled=false;return;}
        state.toxic.jobTimer=setTimeout(()=>pollToxic(id),800);
      }catch(err){$("toxicStatus").textContent=err.message;state.toxic.running=false;$("startToxic").disabled=false;$("toxicBtn").disabled=false;}
    }
    $("accountId").textContent=LOGIN;renderToxicSelector();$("detailAccountSearchForm").addEventListener('submit',openAccountFromDetailSearch);$("refreshBtn").addEventListener('click',()=>{clearAutomationDialogCache();load(true);});$("refreshIpBtn").addEventListener('click',loadIps);$("copyOriginBtn").addEventListener('click',openCopyOrigins);$("applyCopyOriginRange").addEventListener('click',loadCopyOrigins);$("exportCopyReportBtn").addEventListener('click',()=>downloadAutomationReport('copy'));$("closeCopyDialog").addEventListener('click',()=>$("copyDialog").close());$("eaCommentBtn").addEventListener('click',openEaCommentGroups);$("applyEaCommentRange").addEventListener('click',loadEaCommentGroups);$("exportEaReportBtn").addEventListener('click',()=>downloadAutomationReport('ea'));$("closeEaCommentDialog").addEventListener('click',()=>$("eaCommentDialog").close());$("relationshipNetworkBtn").addEventListener('click',openRelationshipNetwork);$("resetRelationshipNetworkBtn").addEventListener('click',resetRelationshipNetwork);$("closeRelationshipNetworkDialog").addEventListener('click',()=>$("relationshipNetworkDialog").close());$("manageActionsBtn").addEventListener('click',toggleActionManager);$("addQuickActionBtn").addEventListener('click',addQuickAction);$("newQuickAction").addEventListener('keydown',event=>{if(event.key==='Enter'){event.preventDefault();addQuickAction();}});$("saveBtn").addEventListener('click',save);$("status").addEventListener('change',saveStatusOnly);$("chartSymbol").addEventListener('change',event=>selectChartSymbol(event.target.value,true));$("includeTimeline").addEventListener('change',event=>{$("refreshTimelineCache").disabled=!event.target.checked;if(!event.target.checked)$("refreshTimelineCache").checked=false;});$("generateBtn").addEventListener('click',generate);$("toxicBtn").addEventListener('click',openToxic);$("historicalFundsBtn").addEventListener('click',openHistoricalFunds);$("historicalFundsPrev").addEventListener('click',()=>{if(state.historicalFunds.page>1){state.historicalFunds.page--;renderHistoricalFundsEvents();}});$("historicalFundsNext").addEventListener('click',()=>{const pages=Math.ceil((state.historicalFunds.data?.events?.length||0)/state.historicalFunds.pageSize);if(state.historicalFunds.page<pages){state.historicalFunds.page++;renderHistoricalFundsEvents();}});$("closeHistoricalFunds").addEventListener('click',()=>$("historicalFundsDialog").close());$("startToxic").addEventListener('click',()=>startToxic());document.querySelectorAll('input[name="toxicMode"]').forEach(input=>input.addEventListener('change',updateToxicMode));$("closeToxic").addEventListener('click',()=>$("toxicDialog").close());$("closePreview").addEventListener('click',()=>{$("previewDialog").close();$("previewFrame").src='about:blank';});
    $("closeAccountSource").addEventListener('click',()=>$("accountSourceDialog").close());
    ["customAction","group","tags","note","owner"].forEach(id=>$(id).addEventListener('input',()=>{state.formDirty=true;}));
    $("orderDetails").addEventListener('toggle',()=>{if($("orderDetails").open&&!state.orders.loaded)loadOrders(1);});$("orderPrev").addEventListener('click',()=>loadOrders(state.orders.page-1));$("orderNext").addEventListener('click',()=>loadOrders(state.orders.page+1));
    loadLedger();load().catch(err=>{$("metricStatus").textContent=err.message;});loadIps();
  </script>
</body>
</html>"""


def main() -> None:
    init_workbook()
    normalize_initial_record_times()
    server = ThreadingHTTPServer((HOST, PORT), Handler)
    print(f"问题账户台账网页已启动: http://{HOST}:{PORT}")
    print(f"Excel: {WORKBOOK_PATH}")
    server.serve_forever()


if __name__ == "__main__":
    main()
